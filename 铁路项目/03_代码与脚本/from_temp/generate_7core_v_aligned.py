import docx
from docx.shared import Pt
from docx.oxml.ns import qn
import os

output_dir = r"D:\铁路线路智能检测机器人\04-项目文档\设计文档"
os.makedirs(output_dir, exist_ok=True)
doc = docx.Document()

doc.styles['Normal'].font.name = u'微软雅黑'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
doc.styles['Normal'].font.size = Pt(10.5)

def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = u'微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    return h

add_heading(doc, '铁路线路智能检测机器人', 0)
add_heading(doc, '7大核心工作板块架构方案 (V2.1 参数一致性对齐版)', 1)
doc.add_paragraph('本文档已与《硬件V2.1》、《软件V2.1》及《总体方案V2.1》完成底层参数的100%交叉对齐。确保所有物理边界（10网口、10km/h、120mm轮径、刚性底盘）在各板块逻辑闭环。')

add_heading(doc, '1. 硬件控制与物理感知 (纯物理基石)', 2)
table1 = doc.add_table(rows=1, cols=3, style='Table Grid')
hdr1 = table1.rows[0].cells
hdr1[0].text, hdr1[1].text, hdr1[2].text = '核心子项', '设计规范', '一致性对齐参数'
data1 = [
    ('网口物理拓扑', '10网口硬隔离，LAN1(EtherCAT), LAN2(SIM路由器), PCIe-A/B(相机激光)', '强制开启 MTU 9014 巨型帧'),
    ('光学与同步', '遮光罩+防尘柔性裙边屏蔽环境光；相机参数锁死，I/O触发频闪LED同步', '曝光时间限制 ≤ 200μs (匹配10km/h)'),
    ('底层测距锚定', '读取4个伺服电机绝对值编码器，算法剔除打滑轮求有效均值', '基于 120mm 轮径及减速比计算脉冲当量')
]
for item in data1:
    r = table1.add_row().cells
    r[0].text, r[1].text, r[2].text = item

add_heading(doc, '2. 软件平台与数据流转 (C#高并发管道)', 2)
table2 = doc.add_table(rows=1, cols=3, style='Table Grid')
hdr2 = table2.rows[0].cells
hdr2[0].text, hdr2[1].text, hdr2[2].text = '核心子项', '设计规范', '一致性对齐参数'
data2 = [
    ('CPU内核隔离', 'Windows系统底层将EtherCAT线程绑定至独立CPU核心，防中断风暴', '隔离 GigE 网卡产生的高频 DPC 中断'),
    ('零拷贝与数据对齐', '非托管内存池防GC卡顿；以里程脉冲 (Frame ID) 为主键建立缓冲字典', '同一 Frame ID 凑齐出栈，超时500ms丢弃残帧'),
    ('软件硬件联动看门狗', 'C#探针侦测网络或设备死机，调用底层指令硬重启', '通过 RS485/CAN 向 PMB 电源板发送断电指令')
]
for item in data2:
    r = table2.add_row().cells
    r[0].text, r[1].text, r[2].text = item

add_heading(doc, '3. 核心算法与数据处理 (绝对解耦)', 2)
table3 = doc.add_table(rows=1, cols=3, style='Table Grid')
hdr3 = table3.rows[0].cells
hdr3[0].text, hdr3[1].text, hdr3[2].text = '核心子项', '设计规范', '一致性对齐参数'
data3 = [
    ('轨距/水平通道 (独立)', '8个测距仪相加算轨距；陀螺仪Roll角算超高水平。绝不互相干涉。', '刚性底盘下，陀螺仪数据必加 10Hz 低通滤波'),
    ('高低/轨向通道 (独立)', '陀螺仪 Pitch/Yaw 配合里程计积分算纵向和横向平顺度', '剔除 >2G 的硬冲击伪影'),
    ('缺陷ROI截流 (独立)', '边缘量化YOLO仅裁剪病害框，云端大模型复核，激光查截面磨耗', '剔除 90% 正常钢轨背景图')
]
for item in data3:
    r = table3.add_row().cells
    r[0].text, r[1].text, r[2].text = item

add_heading(doc, '4. 边缘-云端均衡通信 (抗弱网传输)', 2)
table4 = doc.add_table(rows=1, cols=3, style='Table Grid')
hdr4 = table4.rows[0].cells
hdr4[0].text, hdr4[1].text, hdr4[2].text = '核心子项', '设计规范', '一致性对齐参数'
data4 = [
    ('链路探针', '高频 TCP/UDP 包测算 RTT 及丢包率，反推 SIM 卡上行带宽', '实时带宽估算 (Bandwidth Estimation)'),
    ('四级 QoS 队列', 'P0急停, P1状态, P2缺陷切片, P3全量大图/点云', 'P0 强制预留 10% 带宽'),
    ('隧道写锁与 I/O 保底', '进隧道无网时，截断 P2/P3 写入本地，出隧道游标断点续传', '强制采用企业级/工业级高 TBW 不掉速固态硬盘')
]
for item in data4:
    r = table4.add_row().cells
    r[0].text, r[1].text, r[2].text = item

add_heading(doc, '5. 机械电气与容错系统 (物理防爆盾)', 2)
table5 = doc.add_table(rows=1, cols=3, style='Table Grid')
hdr5 = table5.rows[0].cells
hdr5[0].text, hdr5[1].text, hdr5[2].text = '核心子项', '设计规范', '一致性对齐参数'
data5 = [
    ('定制多通道电源板 (PMB)', '48V总入，分路隔离(强电/净电/通信)，阻断电机反电动势，防级联短路', '各通道独立 PTC 熔断保护与固态开关控制'),
    ('光学防震减灾', '无减震底盘下，相机镜头全胶水死锁，加装高分子阻尼基座', '杜绝定焦环离焦松动'),
    ('遮光罩热力学对策', '针对密闭腔体6相机发热，强制对流或半导体制冷', '防环境突破 55℃ 引发 CMOS 热噪声')
]
for item in data5:
    r = table5.add_row().cells
    r[0].text, r[1].text, r[2].text = item

add_heading(doc, '6. 技术文档与规范 / 7. 学术知识产权', 2)
doc.add_paragraph('所有文档采用纯 Word 表格化。核心提炼：无减震底盘多源解耦测量法、巨型帧零拷贝并发架构、主动探针流控机制等高质量论文创新点。')

for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = u'微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
                    run.font.size = Pt(9.5)

doc.save(r'D:\铁路线路智能检测机器人\04-项目文档\设计文档\铁路线路智能检测机器人_7大核心工作板块_V2.1对齐版.docx')
print("7CORE_V2.1_ALIGNED_DONE")
