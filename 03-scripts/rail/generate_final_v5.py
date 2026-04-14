import os
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn

output_dir = r"D:\铁路线路智能检测机器人\04-项目文档\设计文档"
os.makedirs(output_dir, exist_ok=True)

doc = Document()

# Set default font
style = doc.styles['Normal']
style.font.name = u'微软雅黑'
style._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
style.font.size = Pt(10.5)

def add_h(text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = u'微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    return h

add_h('铁路线路智能检测机器人', 0)
add_h('七大核心工作板块（深度论证完整版）', 1)

doc.add_paragraph('本文档为经过极限技术论证、结合物理工况（如遮光罩、编码器定位、10网口、48V隔离）优化的终极架构指导文件。')

data = [
    ("1. 硬件控制与多源感知", "物理拓扑与隔离", "主板LAN1独占EtherCAT(控制电机)；LAN2独占直连SIM卡路由器；两张PCIe扩展卡承接6相机+2线激光；低速总线走多串口。"),
    ("1. 硬件控制与多源感知", "全天候光学闭环", "物理遮光罩隔绝环境光，固定相机参数；微秒级频闪同步（相机曝光+LED爆闪）消除拖影。"),
    ("1. 硬件控制与多源感知", "绝对自主定位", "弃用GPS，依据4个电机绝对值编码器计算脉冲里程，配合人工标定与防滑转补偿。"),
    
    ("2. 软件与高并发架构", "CPU内核隔离防中断", "将特定CPU核心独占绑定给EtherCAT，防止8个千兆网口巨型帧引发的网络中断抢占资源导致电机失控。"),
    ("2. 软件与高并发架构", "零拷贝与数据重组", "C#开辟未托管内存池防GC卡顿。所有数据以“里程脉冲序号(FrameID)”为主键强制对齐重组。"),
    ("2. 软件与高并发架构", "物理看门狗", "探针发现路由器死机连不上外网时，通过GPIO继电器对其进行物理断电重启。"),
    
    ("3. 数据处理与AI核心", "边缘AI截流减负", "部署轻量化YOLO量化模型，仅截取“螺栓”“疑似裂纹”的ROI微小切片，剔除90%背景，极大降低传输压力。"),
    ("3. 数据处理与AI核心", "多源传感器联合解算", "建立外参标定矩阵。将8个测距仪、2个线激光与陀螺仪横滚角融合计算，消除车体晃动，输出真实高低差。"),
    ("3. 数据处理与AI核心", "云端大模型复核", "接收ROI切片，利用大算力进行长尾病害（擦伤、鳞纹）分类与尺寸测算。"),
    
    ("4. 均衡通信架构", "主动链路探针", "向云端高频发送测速包计算延迟与丢包率，穿透路由器感知真实4G/5G信号带宽。"),
    ("4. 均衡通信架构", "QoS四级队列与防掉速", "P0(告警/急停)>P1(坐标状态)>P2(ROI切片)>P3(大图点云)。弱网触发大文件写锁落盘；选用企业级SSD防缓存耗尽掉速。"),
    
    ("5. 机械电气数字化资产", "48V隔离与EMC", "宽压隔离DC-DC供电防电机反电动势冲击；强电动力与千兆网线物理隔离20cm并星型接地。"),
    ("5. 机械电气数字化资产", "遮光罩热力学设计", "针对罩内6相机+2激光+频闪灯的高温，引入强制风道或半导体制冷，防热噪声。"),
    
    ("6. 技术文档规范", "绝对排版铁律", "所有方案、协议、测试报告必须输出为Word(.docx)，核心参数与逻辑关系呈现为专业表格。"),
    ("6. 技术文档规范", "协议自动化映射", "通过脚本从代码注释提取接口变动，实时更新Word文档。"),
    
    ("7. 学术论文转化", "专利交底书锚点", "围绕“10网口隔离架构”、“遮光罩频闪同步”、“探针边缘截流调度”布局核心发明专利。"),
    ("7. 学术论文转化", "SCI/EI论点", "基于跑车真实数据，撰写边缘ROI截流带宽优化、多源空间补偿算法相关论文。")
]

table = doc.add_table(rows=1, cols=3, style='Table Grid')
hdr = table.rows[0].cells
hdr[0].text = '板块名称'
hdr[1].text = '核心节点'
hdr[2].text = '工程落地防线'

for item in data:
    row = table.add_row().cells
    row[0].text = item[0]
    row[1].text = item[1]
    row[2].text = item[2]

for row in table.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.name = u'微软雅黑'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
                r.font.size = Pt(10)

doc.save(r'D:\铁路线路智能检测机器人\04-项目文档\设计文档\铁路线路智能检测机器人_7大核心工作板块_深度论证完整版.docx')
print("FINAL DOCX GENERATED")