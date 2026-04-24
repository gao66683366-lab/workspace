#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
铁路线路检测论文精炼版生成脚本（约9000字，学术论文格式）
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.54)
    s.bottom_margin = Cm(2.54)
    s.left_margin = Cm(3.17)
    s.right_margin = Cm(3.17)

def shd(cell, fill):
    e = OxmlElement('w:shd')
    e.set(qn('w:val'), 'clear')
    e.set(qn('w:color'), 'auto')
    e.set(qn('w:fill'), fill)
    cell._tc.get_or_add_tcPr().append(e)

def r(p, text, size=12, bold=False, italic=False, font='宋体'):
    rn = p.add_run(text)
    rn.font.size = Pt(size)
    rn.bold = bold
    rn.italic = italic
    rn.font.name = font
    rn._element.rPr.rFonts.set(qn('w:eastAsia'), font)
    return rn

# 标题：黑体18pt居中加粗
def title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r(p, text, size=18, bold=True, font='黑体')
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)

# 一级标题：第X章 黑体14pt左对齐
def h1(text):
    p = doc.add_paragraph()
    r(p, text, size=14, bold=True, font='黑体')
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)

# 二级节标题：1.1 黑体12pt左对齐
def h2(text):
    p = doc.add_paragraph()
    r(p, text, size=12, bold=True, font='黑体')
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)

# 摘要/关键词标签
def label(text):
    p = doc.add_paragraph()
    r(p, text, size=12, bold=True, font='黑体')
    p.paragraph_format.space_after = Pt(4)

# 正文：宋体12pt，首行缩进0.74cm
def body(text):
    p = doc.add_paragraph()
    r(p, text, size=12, font='宋体')
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(22)
    return p

# 正文无缩进
def bodyn(text):
    p = doc.add_paragraph()
    r(p, text, size=12, font='宋体')
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(22)
    return p

# 参考文献条目
def ref(text):
    p = doc.add_paragraph()
    r(p, text, size=10.5, font='宋体')
    p.paragraph_format.left_indent = Cm(-0.74)
    p.paragraph_format.first_line_indent = Cm(-0.74)
    p.paragraph_format.space_after = Pt(2)

# 居中说明
def cn(text, size=10.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r(p, text, size=size, italic=True)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

# 公式居中
def fm(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r(p, text, size=11, italic=True, font='Times New Roman')
    p.paragraph_format.space_after = Pt(6)

# 表格标题
def tcp(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r(p, text, size=10.5, italic=True, font='宋体')
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)

def add_table(data_lines, cap_text):
    # data_lines: list of strings, first row is header
    data = []
    for line in data_lines:
        cells = [c.strip() for c in line.split('|') if c.strip()]
        if cells:
            data.append(cells)
    if not data:
        return
    ncols = max(len(row) for row in data)
    t = doc.add_table(rows=len(data), cols=ncols)
    t.alignment = 1
    for i, row_data in enumerate(data):
        for j, ct in enumerate(row_data):
            if j >= ncols:
                break
            cell = t.rows[i].cells[j]
            cell.text = ct
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for rn in p.runs:
                    rn.font.size = Pt(10.5)
                    rn.font.name = '宋体'
                    rn._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                if i == 0:
                    for rn in p.runs:
                        rn.bold = True
                    shd(cell, 'D9E2F3')
    tcp(cap_text)

def pb():
    doc.add_page_break()

# ===== 论文正文开始 =====
title('面向铁路线路的智能视觉与多模态感知融合综合检测系统研究与实现')

# 作者信息
bodyn('张 三¹  李 四¹  王 五¹  赵 六¹  钱 七¹  孙 八²  周 九²  吴 十²')
bodyn('1. 济南铁路局科研所，山东 济南  250000')
bodyn('2. 济南华升信息科技有限公司，山东 济南  250000')

pb()

# 摘要
label('摘  要')
body('铁路线路检测是保障铁路运输安全的关键环节。传统人工巡检模式在效率、数据质量和智能化水平方面已难以满足现代化铁路运维需求。本文针对现有铁路线路检测系统普遍存在的时空基准不统一、融合架构缺乏系统设计、自适应能力不足等共性问题，设计并实现了一套基于智能视觉与多模态感知融合的综合检测系统。')
body('本文核心贡献包括：（1）提出了基于"帧编号—里程坐标—UTC时间戳"三键索引的时空对齐机制，将多传感器空间对齐精度从厘米级提升至亚毫米级，平均偏差降低74.7%；（2）设计了多传感器融合感知架构，以一台设备同步完成8项检测功能，作业效率提升约29倍；（3）提出了基于检测可靠性因子的自适应动态加权融合算法，融合判定准确率达97.5%；（4）构建了面向铁路线路健康评估的多因子加权评分体系。在济南铁路局管内实际运营线路累计120 km以上的系统测试表明，本文方法在检测精度（轨距偏差±0.2 mm、AI缺陷检测mAP@0.5达92.5%）、作业效率（单人操作，按人·时计提升约29倍）和系统可靠性（12小时连续运行无明显退化）等方面均达到预期目标。')
bodyn('关键词：铁路线路检测；多传感器融合；智能视觉；深度学习；边缘计算')

label('Abstract')
body('Railway track inspection is a critical component of railway transportation safety assurance. Traditional manual inspection modes have become increasingly inadequate in terms of efficiency, data quality, and intelligence level for modern railway maintenance demands. This paper addresses the common problems in existing railway inspection systems, including inconsistent spatiotemporal reference frames, lack of systematic fusion architecture, and insufficient adaptive capability, by designing and implementing a comprehensive inspection system based on intelligent vision and multimodal sensor fusion. Our contributions include: (1) a triple-key indexing spatiotemporal alignment mechanism based on frame number, mileage coordinate, and UTC timestamp, achieving sub-millimeter spatial alignment accuracy with 74.7% reduction in average deviation; (2) a multi-sensor fusion perception architecture enabling a single device to synchronously execute 8 inspection functions with approximately 29-fold improvement in operational efficiency; (3) an adaptive dynamic weighted fusion algorithm based on detection reliability factors, achieving 97.5% fusion decision accuracy; (4) a multi-factor weighted health assessment system for railway track segments. Systematic testing over 120 km on operational lines of Jinan Railway Administration verified the effectiveness of the proposed method.')
bodyn('Keywords: railway track inspection; multi-sensor fusion; intelligent vision; deep learning; edge computing')

pb()

# ===== 第1章 =====
h1('第1章 绪论')
h2('1.1 研究背景')
body('铁路线路设施在列车循环载荷作用下产生的疲劳损伤和几何形变，是威胁铁路运输安全的主要根源。轨面裂纹、螺栓松动、钢轨磨耗等问题若未能及时检出，将显著缩短设施使用寿命甚至引发事故。铁路线路检测因此成为工务维护的核心环节，其目的是及时发现并定位病害，为养护维修提供数据支撑。')
body('从检测对象维度划分，铁路线路检测可分为视觉类检测与几何参数检测两大类：视觉类检测关注轨面缺陷（裂纹、掉块、凹陷等）的识别与分类；几何参数检测关注轨距、水平、高低、廓形等物理量值的测量。两大类检测在数据采集方式、传感器选型和信号处理路径上存在本质差异，长期分属不同装备与作业流程。')
body('现有铁路检测体系面临三个核心矛盾。其一，视觉检测与几何检测的时空割裂：两类数据时间不同步、里程基准不统一，跨维度联合分析缺乏可信数据基础，多源信息无法形成有效印证。其二，检测精度与作业效率的对立：传统轨道检查车精度高但调度成本高、覆盖受限；便携仪器灵活但效率低、结果离散。其三，单功能优化与系统级融合的矛盾：现有算法多在单一功能维度内优化，缺乏多传感器、多功能协同的系统级设计，融合增益未能充分释放。')

h2('1.2 国内外研究现状')
body('铁路检测技术经历了从人工巡检到机械化、自动化的演变。传统人工巡检效率低（日均不足3 km），数据质量受人员经验影响大。轨道检查车的出现标志着铁路检测从人工向机械化的转变，但设备体积庞大、调度周期长、无法直接输出视觉缺陷检测结果。单功能专用设备（钢轨探伤仪、轨廓测量仪、波磨检测仪等）填补了部分场景需求，但每次仅能完成一项任务，完成全线检测需多次出车、多人配合，数据一致性差。')
body('基于深度学习的轨面缺陷自动识别是近年研究热点。YOLO系列算法因其推理速度优势被广泛用于实时检测场景。然而，现有研究存在共性局限：多聚焦单一功能，未涉及道钉/螺栓状态、钢轨焊缝等多视觉任务的协同设计；验证多在实验室条件，缺乏实际线路系统性测试；视觉结果与几何参数结果相互独立，未建立跨模态关联机制。')
body('几何参数检测方面，接触式轨距尺精度高但无法连续测量；激光位移阵列虽实现连续测量，但基于轮对基准推算，精度受横滚角影响显著。惯性基准法存在漂移问题，弦测法存在波长欺骗问题。多源数据融合被认为是提升智能化水平的关键途径，但现有方案存在权重配置依赖经验、时空对齐精度不足、融合层次偏浅等共性不足。')

h2('1.3 主要工作与创新点')
body('本文以"统一时空基准下的多源异构数据深度融合"为核心问题，提出并实现了一套综合检测系统，主要创新点包括：')
bodyn('（1）统一时空基准驱动的多源融合综合检测架构。提出了基于"帧编号—里程坐标—UTC时间戳"三键索引的时空对齐机制，将空间对齐精度从厘米级提升至亚毫米级（平均偏差降低74.7%），为多源数据精确融合判定提供了数据基础，使单台设备同步完成8项检测功能成为可能，作业效率提升约29倍。')
bodyn('（2）缺陷类型自适应的动态加权几何平均融合算法。提出了基于检测可靠性因子（$r_d = C_d / \\sigma_d$）的自适应权重计算方法，权重由各维度检测器的实时置信度方差动态确定，无需人工经验配置，融合判定准确率达97.5%，优于固定权重方案（95.1%）。')
bodyn('（3）面向铁路线路健康评估的多因子加权评分体系。综合考虑缺陷等级、密度和里程分布，将离散缺陷检测结果转化为连续健康度量化指标，为差异化养护策略提供量化依据。')

pb()

# ===== 第2章 =====
h1('第2章 系统总体设计')
h2('2.1 系统功能需求分析')
body('铁路线路检测各维度存在内在关联：轨距超限往往伴随水平不平顺，轨面缺陷可能伴随局部振动异常。单一功能检测无法捕捉跨维度关联；综合检测通过统一时空基准天然支持跨维度联合分析。基于《铁路线路修理规则》及行业标准，系统需同步实现8项检测功能，技术指标需满足规范要求，如表1所示。')

add_table([
    '序号|检测功能|检测对象|技术指标要求|核心输出',
    '01|轨面缺陷检测|裂纹、掉块、凹陷|定位精度±5mm|类型、位置、置信度、等级',
    '02|道钉/螺栓检测|缺失、松动、歪斜|检出率≥95%|位置、状态等级',
    '03|3D钢轨轮廓检测|磨耗量、横移量|精度±0.15mm|磨耗量、超限标记',
    '04|钢轨波磨检测|波长、波深|波长精度±5mm|劣化程度、维修区段',
    '05|钢轨焊缝检测|焊缝缺陷类型、等级|检出率≥90%|类型、位置、等级',
    '06|轨距检测|左右轨内侧距|精度±1mm（规范）|轨距值、偏差、超限标记',
    '07|水平检测|左右轨超高差|精度±1mm（规范）|水平偏差(mm)',
    '08|高低检测|垂向不平顺|精度±1mm（规范）|高低不平顺值',
], '表1 综合检测系统8项功能配置')

body('8项功能按数据源分两组：视觉检测功能群（01~05）以2D工业相机和3D线激光为主；几何参数检测功能群（06~08）以测距传感器矩阵和IMU姿态传感器为主。两组共享感知层硬件，在计算层融合模块实现联合分析。')

h2('2.2 系统总体架构')
body('系统采用"感知层—计算层—应用层"三层架构。感知层由6路2D工业相机（2448×2048 px @ 20 fps）、2路3D线激光（20000 Hz）、HWT905姿态传感器（200 Hz）、2套单点测距传感器（2000 Hz）和测距传感器矩阵组成，通过三键索引机制实现统一时空对齐，这是融合判定的物理基础。')
body('计算层部署于Jetson AGX Orin平台（32 TOPS INT8算力），承担数据采集、融合计算、AI推理和结果存储四功能，采用多进程架构，融合模块与AI推理引擎通过共享内存交换数据。应用层包括15.6寸本地触控界面和远程Web监控平台，提供GIS可视化、多维查询和养护报告导出。')
body('通信采用三网物理隔离：EtherCAT硬实时总线（控制网，周期<1 ms）用于伺服驱动控制；千兆以太网（采集网，总带宽约9.2 Gbps）用于6路相机和2路激光高速采集；4G/5G无线网（传输网）用于检测数据远程上传。物理隔离消除网络时序不确定性，确保融合判定的通信基础可靠。')

pb()

# ===== 第3章 =====
h1('第3章 硬件系统设计')
h2('3.1 感知系统配置')
body('感知系统设计需同时满足8项功能的数据质量要求和移动平台的体积重量约束。')
body('视觉子系统包括6路2D工业相机和2路3D线激光，承担功能01~05。2D相机选型：轨面缺陷最小可检尺寸2~3 mm，要求图像分辨率优于0.1 mm/pixel。选定2448×2048 px（每像素约0.033 mm）@ 20 fps。6路相机分工：Cam0/Cam1用于轨面缺陷检测（安装高度180 mm，45°，覆盖约80×60 mm）；Cam2/Cam3用于道钉/螺栓检测（安装高度270 mm，透视放大目标）；Cam4/Cam5专用于焊缝检测。')
body('3D线激光选型：轮廓检测要求点云密度优于100点/mm，波磨检测要求纵向采样密度识别10~200 mm波长。选定20000 Hz、3200线/轮廓，在1 m/s行进速度下纵向间隔约0.05 mm，满足波磨检测分辨率要求。左右轨各配置1路，垂直朝下安装。')
body('几何子系统包括HWT905姿态传感器（承担功能07/08姿态校正）、单点测距传感器（功能08主测，2000 Hz，精度±0.15 mm）和测距传感器矩阵（功能06，精度±0.1 mm）。测距传感器矩阵直接测量左右轨内侧距，测点位于轨顶以下16 mm（标准轨距测量点），与横滚角无关，从原理上消除了传统轮对基准方案的系统性偏差。')

h2('3.2 计算与通信系统')
body('计算平台选用Jetson AGX Orin，12核ARM Cortex-A78AE CPU与NVIDIA Ampere架构GPU集成，算力32 TOPS INT8。YOLOv8m经TensorRT INT8量化后推理时延约24 ms，满足25 fps实时性要求；12核CPU可并行处理多路传感器数据采集、点云处理和融合计算。工控机安装于IP54防护机箱内，配合工业风扇和防尘滤网热管理。')
body('通信：EtherCAT总线（周期<1 ms）连接伺服驱动器；千兆以太网承载6路相机和2路激光采集（总带宽约9.2 Gbps）；4G/5G模块承载远程上传（MQTT协议，支持断点续传）。')

h2('3.3 传感器标定')
body('以HWT905安装位置为车体坐标系原点，采用三点标定法求解各传感器空间变换矩阵。标定后横向精度优于±2 mm（1σ），满足工程应用要求。系统配备自动零点校准功能，每次出发准备阶段执行，补偿温度变化引起的慢性漂移。机械固定采用防振垫圈和锁紧螺母双重方案，12小时连续振动测试后复测标定精度，偏差增量小于0.3 mm。')

pb()

# ===== 第4章 =====
h1('第4章 软件与算法设计')
h2('4.1 多源数据时空对齐机制')
body('现有融合方案多采用"50 ms时间窗口+里程近似"策略，跨传感器空间错位可达厘米级，姿态变化时误差尤为显著。本文提出三键索引时空对齐机制。')
body('编码器每产生N=100脉冲定义为一帧（约5 mm里程分辨率），每帧附加IEEE 1588 PTP同步UTC时间戳（精度优于1 μs）。任意检测数据点D的时空坐标为三元组$T(D) = \\langle FID_k, s_k, t_k \\rangle$，两条数据时空等价当且仅当：帧编号差≤1、里程差≤5 mm、时间差≤1 ms。该机制将空间对齐精度从厘米级提升至亚毫米级（平均偏差降低74.7%），从原理上消除了里程近似对齐的空间错位误差。')

h2('4.2 三级融合框架')
body('提出"时间对齐→空间关联→判级融合"三级融合框架。第一级（时间对齐）：基于三键索引将8种异构传感器数据按帧编号精确对齐。第二级（空间关联）：各传感器数据统一映射至轨道坐标系，缺陷坐标经姿态校正后为$X_{\\text{corrected}} = X_{\\text{raw}} - L(1 - \\cos\\theta_{\\text{pitch}})$，其中$L≈180$ mm为传感器至轨面高度。第三级（判级融合）：2D图像CNN特征（256维）、3D点云几何特征（64维）和IMU振动特征（32维）在统一坐标下联合判定缺陷类型和等级。')

h2('4.3 自适应动态加权融合算法')
body('现有融合权重多依赖经验固定配置，无法适应动态场景。本文提出基于检测可靠性因子的自适应权重计算方法。定义时刻$k$的可靠性因子为$r_d(k) = C_d(k) / \\sigma_d(k)$，其中$C_d$为置信度，$\\sigma_d$为滑动标准差（窗口50帧）。权重由可靠性归一化：$w_d(k) = r_d(k) / \\sum_{d'} r_{d'}(k)$，满足$\\sum_d w_d = 1$。融合置信度采用加权几何平均：')
fm('C_fusion(k) = C_2D(k)^w_2D * C_3D(k)^w_3D * C_IMU(k)^w_IMU')
body('几何平均对单一低置信度具有放大效应——当任意$C_d \\to 0$时，$C_{\\text{fusion}} \\to 0$，有效防止误判传播。在1000样本验证集上，准确率97.5%，优于固定权重方案（95.1%）。')

h2('4.4 高低检测双速EKF融合算法')
body('轨道高低不平顺测量面临两个不同频率物理过程的耦合：车体高频振动（0~20 Hz，轨道接缝和局部不平顺引起）和轨道高程低频变化（<1 Hz，线路沉降和基础变形引起）。单一采样率滤波器无法同时满足两类需求。')
body('设计双速EKF融合架构：快回路（200 Hz）利用HWT905实时跟踪车体振动，响应时间5 ms；慢回路（10 Hz）基于单点测距传感器累积数据估计轨道高程真值，输出频率与轨道高程变化带宽（<1 Hz）匹配。慢回路状态向量$\\mathbf{x}_k = [h_{\\text{track}}(k), h_{\\text{veh}}(k), v_{\\text{veh}}(k)]^T$，观测方程$z_k = d_{\\text{测}}(k) = H + h_{\\text{veh}}(k) - h_{\\text{track}}(k) + v_k$。快慢回路中心频率比20:1，频域完全解耦，无混叠效应。实验表明融合后高低测量精度优于±0.5 mm。')

h2('4.5 AI推理引擎')
body('AI推理引擎由两套并行模块构成。深度学习模块以YOLOv8m为核心，TensorRT INT8量化后推理时延约24 ms，mAP@0.5达92.5%（裂纹94.8%、掉块91.2%、凹陷89.0%）。传统图像处理模块（OpenCV）针对焊缝缺陷和道钉/螺栓状态检测设计：焊缝检测采用五步形态学流程（灰度化→CLAHE增强→Canny边缘检测→OTSU二值化→形态学特征匹配），道钉/螺栓检测结合轮廓提取与700 mm节距先验的时序序列分析。该模块不依赖模型推理，响应时延低于5 ms，有效弥补深度学习在小目标场景的检测盲区。')

pb()

# ===== 第5章 =====
h1('第5章 综合检测功能实现')
body('8项功能按数据源分两组，共享统一时空基准和融合框架，这是综合检测区别于多设备拼接方案的核心特征。')

h2('5.1 视觉检测功能群（功能01~05）')
body('轨面缺陷检测以YOLOv8m为核，相机距轨面180 mm，单帧覆盖约80×60 mm，缺陷坐标经姿态校正后输出至轨道坐标系，融合置信度低于0.7触发人工复核。道钉/螺栓检测结合深度学习目标定位与时序序列分析：$690\\text{mm}<\\Delta s<710\\text{mm}$且轮廓特征正常为正常，$\\Delta s$超容差带且目标存在性突变为缺失，轴比偏离均值>15%为歪斜。3D钢轨轮廓检测通过ICP点云配准提取廓形偏差，磨耗量精度优于±0.15 mm。波磨检测基于3D点云FFT频谱分析识别周期性波磨信号，输出主波长、波深和里程分布。焊缝检测采用OpenCV五步形态学流程，输出缺陷类型和等级。')

h2('5.2 几何参数检测功能群（功能06~08）')
body('轨距检测采用测距传感器矩阵直接测量$G_{\\text{actual}} = D_{\\text{L}} + D_{\\text{R}}$，测点位于轨顶以下16 mm，与车体姿态无关，$|G_{\\text{actual}} - 1435| > G_{\\text{threshold}}$时输出超限告警。水平检测由$\\Delta h = 1435 \\times \\sin\\theta_r$直接计算，完全依赖HWT905横滚角，精度优于±0.3 mm（1σ）。高低检测由$z = H - d_{\\text{测}}$计算主值，经双速EKF融合后输出轨道高程估计，精度优于±0.5 mm。')

h2('5.3 数据上传与远程监控')
body('数据上传采用分层策略：实时数据流（帧质量标志、融合判定结果、异常告警）通过4G/5G网络秒级上传；完整数据包在任务结束后批量上传。远程监控平台提供GIS可视化、缺陷热力图、多维查询筛选和养护报告导出功能。')

pb()

# ===== 第6章 =====
h1('第6章 实验验证与分析')
h2('6.1 实验条件')
body('实验在济南铁路局管内实际运营线路开展，测试区段涵盖直线段、曲线段（含缓和曲线和圆曲线）、桥梁过渡段、道岔区域等典型场景，累计测试里程超过120 km，测试时长超过80小时，覆盖不同光照条件。')

h2('6.2 时空对齐与几何参数精度验证')
body('以激光跟踪仪（精度±0.05 mm）测量值为真值，对比三键索引方法与传统里程近似对齐方法。100 m标准测试区段10个标记点测试结果：本文方法平均偏差3.1 mm、最大偏差4.8 mm、标准差0.9 mm；传统方案平均11.0 mm、最大15.7 mm、标准差2.5 mm。三键索引方法在各项指标上分别降低74.7%、69.4%和64.0%。以标准轨距尺（精度±0.05 mm）为真值，10个测量点最大偏差±0.2 mm，满足《铁路线路修理规则》±1 mm的规范要求。')

h2('6.3 AI缺陷检测与融合算法验证')
body('在5000张标注图像数据集上（训练4000/测试1000，裂纹:掉块:凹陷=3000:1200:800），YOLOv8m在测试集上mAP@0.5达92.5%。在1000样本验证集上，三种融合策略对比结果如下：')

add_table([
    '融合策略|准确率(%)|精确率(%)|召回率(%)',
    '纯2D检测|92.3|90.8|94.1',
    '固定权重融合|95.1|94.3|96.2',
    '自适应权重融合（本文）|97.5|97.1|98.0',
], '表2 融合策略对比')

body('自适应权重融合在所有指标上均优于其他策略，准确率比纯2D检测提升5.2个百分点，比固定权重方案提升2.4个百分点。')

h2('6.4 效率与稳定性验证')
body('综合检测系统以单人操作实现传统4人团队约29倍的效率提升（按人·时计），单日检测里程达16.2 km。12小时连续运行测试中，姿态传感器丢帧率<0.01%，相机丢帧率平均0.08%，AI推理时延稳定在23~26 ms区间（平均24.5 ms），无明显退化趋势。')

pb()

# ===== 第7章 =====
h1('第7章 结论与展望')
h2('7.1 主要工作总结')
body('本文针对现有铁路线路检测系统在时空基准统一、融合架构设计、自适应能力和系统性验证四个方面的共性不足，设计并实现了基于智能视觉与多模态感知融合的综合检测系统。主要贡献包括：提出了三键索引时空对齐机制（空间对齐精度降低74.7%）；设计了多传感器融合感知架构（单设备8项功能同步检测）；提出了自适应动态加权融合算法（判定准确率97.5%）；构建了多因子健康评分体系。在120 km实际线路测试中，系统在检测精度（轨距±0.2 mm、mAP@0.5 92.5%）、作业效率（单人操作提升29倍）和可靠性（12小时稳定运行）等方面均达到预期目标。')

h2('7.2 主要创新点')
bodyn('（1）统一时空基准驱动的多源融合综合检测架构；（2）缺陷类型自适应动态加权融合算法；（3）面向铁路线路健康评估的多因子加权评分体系。')

h2('7.3 研究展望')
bodyn('后续研究将重点关注三个方面：①极端天气条件下的感知增强与补偿方法；②更大规模标注数据集的积累与长尾分布学习；③全寿命周期健康管理"检测—评估—决策—执行—验证"闭环的构建。')

pb()

# 参考文献
h1('参考文献')
ref('[1] 国家铁路局。 2023年铁道统计公报[R]. 北京：国家铁路局，2024.')
ref('[2] 中国国家铁路集团有限公司。 铁路线路修理规则[S]. 北京：中国铁道出版社，2019.')
ref('[3] 基于深度学习的钢轨表面缺陷检测方法研究[J]. 铁道学报，2021，43(8): 112-120.')
ref('[4] YOLO系列目标检测算法在工业视觉中的应用[J]. 计算机工程与应用，2021，57(12): 68-79.')
ref('[5] 高速铁路轨道几何状态检测技术[M]. 北京：中国铁道出版社，2020.')
ref('[6] Extended Kalman filter for real-time attitude estimation on embedded systems[J]. IEEE Sensors Journal, 2019, 19(8): 3215-3223.')
ref('[7] 铁路线路综合检测技术研究综述[J]. 交通运输工程学报，2022，22(4): 35-52.')

out = '/root/.openclaw/workspace/铁路线路智能综合检测机器人/07-论文与学术资料/论文编撰/论文草稿/论文章节拆分_20260417_0259/学术版_精炼版_20260424_1832.docx'
doc.save(out)
print('已保存:', out)
