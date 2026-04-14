import docx
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

output_dir = r'D:\铁路线路智能检测机器人\04-项目文档\设计文档'
os.makedirs(output_dir, exist_ok=True)
doc = docx.Document()

doc.styles['Normal'].font.name = u'微软雅黑'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
doc.styles['Normal'].font.size = Pt(10.5)

def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = u'微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    return h

add_heading(doc, '铁路线路智能检测机器人', 0)
add_heading(doc, '底层硬件与系统生死线深度论证报告 V1.0', 1)

doc.add_paragraph('本文档旨在排除实验室环境的理想化假设，针对项目在真实铁路线路恶劣工况下可能面临的致命失效风险，进行严苛的工程级理论推演与极限论证，并给出强制落地对策。')

# 数据定义
topics = [
    ('1. 系统内核与网络中断风暴 (控制失效风险)', 
     '8个千兆网口(6相机+2激光)同时以9014Bytes巨型帧并发，主板LAN1以1ms周期跑EtherCAT。',
     'Windows非实时OS。海量巨型帧引发极高DPC(延迟过程调用)与硬件中断。若网络抢占CPU，EtherCAT将丢失1ms通信周期，驱动器报Watchdog Error导致机器抱死急停或翻车。',
     '强制内核隔离：将Windows系统与网络收发限制在Core 0-5；使用C# ProcessorAffinity将EtherCAT主站线程死锁绑定在Core 6和7，实现物理级算力隔离。'),
     
    ('2. 热力学与遮光罩盲区 (热宕机风险)',
     '夏季铁轨暴晒，密闭铝制遮光罩内6台相机、2台激光、高频LED爆闪总发热约80W。',
     '无对流腔体内温度会迅速突破70℃。工业相机CMOS超55℃产生严重热噪声(雪花图)，超65℃触发主板保护关机，系统彻底失明。',
     '必须修改图纸：1. 相机必须用高导热硅脂硬连接铝支架作为热沉；2. 遮光罩需开百叶窗式进出气口并加装IP67工业风扇强制对流，极热环境需评估TEC半导体制冷。'),
     
    ('3. 固态硬盘大掉速防爆 (I/O阻塞风险)',
     '进入5km无信号长隧道，触发写锁。6相机按10fps产生约300MB/s数据，须全量落盘。',
     '消费级SSD(TLC/QLC)的SLC Cache(十几GB)耗尽后，真实写入速度会断崖下跌至50-100MB/s。由于硬盘速度 < 采集速度，C#内存池会瞬间挤爆引发 OutOfMemory 崩溃。',
     '1. 杜绝消费级，必须采购工业/企业级长效NVMe SSD(稳态写入>1000MB/s，带PLP断电保护)。2. 采集端设高水位线，内存占满80%时强制相机降频自救。'),
     
    ('4. 48V强弱电干扰 (EMC与烧毁风险)',
     '工控机、路由器、传感器与4个伺服电机并联在同一48V蓄电池母线上。',
     '电机下坡制动产生强反电动势(Back-EMF)，母线瞬间飙升至60-70V；高频PWM斩波带来严重EMI辐射。将直接烧毁工控机主板或导致千兆网卡频繁丢包断连。',
     '强制电气隔离：工控机与传感器电源必须经过宽压输入、带电气隔离的工业级DC-DC转换器；触发信号必经高速光耦；强弱电走线须保持20cm以上安全距离。'),
     
    ('5. 纯轮式里程计累计漂移 (坐标失效风险)',
     '完全依赖伺服电机编码器计算脉冲里程，不使用GPS/北斗。',
     '车轮长期运行产生磨耗(如直径磨损1mm)，或遇到油污打滑。运行10公里后，机械累计误差可能超50米。发现病害但无法在物理世界精确定位。',
     '1. IMU融合防滑：融合陀螺仪加速度与4轮转速作卡尔曼滤波，剔除打滑轮数据；2. 地标强制纠偏：底盘加装RFID读取器，每经既有里程标定桩时清零累计误差。')
]

for title, limit, fail, action in topics:
    add_heading(doc, title, 2)
    t = doc.add_table(rows=1, cols=2, style='Table Grid')
    
    row1 = t.add_row().cells
    row1[0].text = '极限工况'
    row1[1].text = limit
    
    row2 = t.add_row().cells
    row2[0].text = '失效推演'
    row2[1].text = fail
    
    row3 = t.add_row().cells
    row3[0].text = '强制落地对策'
    row3[1].text = action

for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = u'微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
                    run.font.size = Pt(10)

doc.save(r'D:\铁路线路智能检测机器人\04-项目文档\设计文档\铁路线路智能检测机器人_底层硬件与系统生死线深度论证报告_V1.0.docx')
print('REPORT V1.0 GENERATED')