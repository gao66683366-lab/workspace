#!/usr/bin/env python3
"""生成专利五书：一种铁路线路智能视觉与多模态感知融合综合检测方法"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os, shutil

PATENT_NAME = '一种铁路线路智能视觉与多模态感知融合综合检测方法'
DATE = '2026年05月28日'
src = '/root/.openclaw/media/inbound/202605-论文发明专利说明书---56b8e1c8-e704-46d5-b670-0a2c7df09978.docx'
dst_dir = '/root/.openclaw/workspace/铁路线路智能综合检测机器人/08-专利资料/'
os.makedirs(dst_dir, exist_ok=True)

fig_paths = {
    '图1': dst_dir + '图1_系统总体架构图.png',
    '图2': dst_dir + '图2_三键索引时空对齐机制示意图.png',
    '图3': dst_dir + '图3_自适应动态加权融合算法流程图.png',
    '图4': dst_dir + '图4_快慢双速EKF融合架构图.png',
    '图5': dst_dir + '图5_几何参数检测原理图.png',
}

doc = Document(src)
paras = doc.paragraphs
tables = doc.tables

def sf(run, size=12, bold=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = 'SimSun'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def center(doc, text, size=12, bold=False, before=0, after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    sf(r, size, bold)
    return p

def left(doc, text, size=12, bold=False, indent=False, before=0, after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if indent:
        p.paragraph_format.first_line_indent = Pt(22.05)
    r = p.add_run(text)
    sf(r, size, bold)
    return p

def para(doc, p_obj, size=12, bold=False, indent=False, before=0, after=0):
    """Copy paragraph content from source paragraph object to target doc."""
    np = doc.add_paragraph()
    np.alignment = p_obj.alignment if p_obj.alignment else WD_ALIGN_PARAGRAPH.JUSTIFY
    np.paragraph_format.space_before = Pt(before)
    np.paragraph_format.space_after = Pt(after)
    if indent:
        np.paragraph_format.first_line_indent = Pt(22.05)
    for run in p_obj.runs:
        nr = np.add_run(run.text)
        nr.font.size = Pt(size)
        nr.font.bold = run.font.bold
        nr.font.name = 'SimSun'
        nr._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return np

def fig(doc, path, caption, w=15):
    if not os.path.exists(path):
        print(f'  WARNING: {path} not found, skipping')
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run()
    r.add_picture(path, width=Cm(w))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(0)
    cap.paragraph_format.space_after = Pt(12)
    cr = cap.add_run(caption)
    sf(cr, 10)

def table_to_doc(doc, t):
    nt = doc.add_table(rows=len(t.rows), cols=len(t.columns))
    nt.style = 'Table Grid'
    for i, row in enumerate(t.rows):
        for j, cell in enumerate(row.cells):
            c = nt.cell(i, j)
            c.text = cell.text
            for tp in c.paragraphs:
                for tr in tp.runs:
                    tr.font.size = Pt(10.5)
                    tr.font.name = 'SimSun'
                    tr._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return nt

# Index key paragraphs
# [16] 摘要正文, [17] 摘要续
# [24] 装置权利要求1, [27] 方法权利要求2
# [32] 权利要求3, [35]权利4, [38]权利5, [41]权利6
# [44]技术领域, [47]背景技术, [49]背景技术续, [52]发明内容
# [54]5.1, [56]5.2, [57]5.2.1, [60]5.2.2, [62]5.2.3, [64]5.2.4, [65]5.2.5, [67]5.2.6, [69]5.2.7, [71]5.2.8?, [73]5.3
# Actually let me re-examine the actual indices from the source doc

print("Paragraph text summary:")
for i in [16,17,19,21,23,24,26,27,30,32,35,38,41,44,46,47,49,52,54,56,57,60,61,62,63,64,65,67,68,69,70,71,72,73,75,78,79,80,81,82,83,84,85,88,90,91,93,94,95,96,97,98,99,100,101,102,103,104,105]:
    if i < len(paras):
        t = paras[i].text[:80] if paras[i].text else '---EMPTY---'
        print(f"  [{i}] {t}")