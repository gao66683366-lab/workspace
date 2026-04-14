import os
import docx
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

output_dir = r'D:\铁路线路智能检测机器人\04-项目文档\设计文档'
os.makedirs(output_dir, exist_ok=True)
doc = docx.Document()

doc.styles['Normal'].font.name = u'微软雅黑'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
doc.styles['Normal'].font.size = Pt(10.5)

def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = u'微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    return h

add_heading(doc, '铁路线路检测机器人 —— 四步流转硬核技术白皮书 (工业仪器级)', 0)

add_heading(doc, '第一步：核心数据采集 (底层硬件级时空定格)', 2)
t1 = doc.add_table(rows=1, cols=3, style='Table Grid')
h1 = t1.rows[0].cells
h1[0].text, h1[1].text, h1[2].text = '核心技术点', '算法与参数约束', '解决的工程痛点'
d1 = [
    ('突破OS纳秒级时钟锁', '启用EtherCAT DC模式(Sync0)。驱动器硬件Latch捕获脉冲瞬间，将纳秒级绝对时间戳硬写入PDO寄存器。', '消除Windows非实时调度带来的15ms(对应41mm)时间戳漂移错位。'),
    ('光电纳秒级硬同步', '10km/h极速下，I/O板卡利用高速数字光耦(6N137, 延迟<50ns)触发方波，电子快门与LED频闪强制在200μs内闭合。', '消除百微秒级高速运动拖影与普通光耦的通道相位差。'),
    ('零拷贝(Zero-Copy)接收', '废弃TCP/IP协议栈，使用unsafe代码块(VirtualAlloc)申请4K对齐物理锁定内存。配合AVX向量指令集极速搬运像素。', '杜绝海量巨型帧引发DPC硬件中断风暴及C#垃圾回收(GC)卡顿死机。')
]
for row_data in d1:
    row = t1.add_row().cells
    row[0].text, row[1].text, row[2].text = row_data

add_heading(doc, '第二步：数据清洗与筛查 (信号层去伪存真)', 2)
t2 = doc.add_table(rows=1, cols=3, style='Table Grid')
h2 = t2.rows[0].cells
h2[0].text, h2[1].text, h2[2].text = '核心技术点', '算法与参数约束', '解决的工程痛点'
d2 = [
    ('扭矩交叉比对剔除打滑', '实时监控4电机ActualTorque。若某轮扭矩瞬降>=80%且转速激增，触发SLIP_FLAG，积分器丢弃该轮脉冲取其余均值。', '彻底解决无减震刚性底盘“三条腿板凳”效应导致的局部车轮空转里程飞车。'),
    ('陀螺仪低通滤波与相位补偿', '前置10Hz二阶巴特沃斯低通滤波器。同时计算其50ms群延迟，在C#缓冲队列中进行时间轴向后50ms的历史相位反向寻址对齐。', '滤除无减震带来的高频机械碎震，并解决滤波带来的空间坐标滞后错位(防138mm错位)。'),
    ('SIMD硬件级ROI截流', '固化相机焦距。调用CPU SIMD(256-bit寄存器)直接计算内存指针偏移，瞬间切割512x512像素的缺陷区域。', '绕过OpenCV冗余检查，50μs内抛弃90%无效背景，极大降低前端总线负荷。')
]
for row_data in d2:
    row = t2.add_row().cells
    row[0].text, row[1].text, row[2].text = row_data

add_heading(doc, '第三步：核心算法处理 (独立降维与异构计算)', 2)
t3 = doc.add_table(rows=1, cols=3, style='Table Grid')
h3 = t3.rows[0].cells
h3[0].text, h3[1].text, h3[2].text = '核心技术点', '算法与参数约束', '解决的工程痛点'
d3 = [
    ('四大几何通道绝对解耦', '轨距=测距仪相加; 水平=Roll角乘基准; 高低=Pitch积分; 轨向=Yaw积分。各司其职，绝不交叉干涉。', '避免伪科学算法中多传感器强行补偿造成的级联误差(一损俱损)。'),
    ('ICP抗噪轮廓配准', '升级为Point-to-Plane(点到面)ICP算法，目标函数引入Huber Loss鲁棒核函数，强行压制突变噪点权重。', '解决钢轨生锈、轨面异物导致的3D点云匹配发散，确保磨耗测算精准。'),
    ('GPU异步流水线(Pinned Mem)', '调用cudaHostAlloc分配页锁定内存，ROI数据通过PCIe异步推入GPU，TensorRT执行INT8量化模型推理。', '实现CPU算几何与GPU算视觉的完全异步并行，互不阻塞主线程。')
]
for row_data in d3:
    row = t3.add_row().cells
    row[0].text, row[1].text, row[2].text = row_data

add_heading(doc, '第四步：轻量化上传 (协议级的弱网流控机制)', 2)
t4 = doc.add_table(rows=1, cols=3, style='Table Grid')
h4 = t4.rows[0].cells
h4[0].text, h4[1].text, h4[2].text = '核心技术点', '算法与参数约束', '解决的工程痛点'
d4 = [
    ('内存对齐的二进制裸协议', '废除JSON，定义C语言风格Pack=1的StructLayout。包体含FrameID、里程、轨距、病害ID等，总计压缩至25 Bytes。', '彻底消除字符串序列化产生的内存碎片与流量冗余，2G弱网下亦能毫秒级告警。'),
    ('SAEA对象池与探针调度', '预分配1000个SocketAsyncEventArgs入池复用。后台探针监控RTT，动态调节Token Bucket令牌桶下发流量。', '规避频繁创建Socket对象的系统开销，极其敏锐地适应SIM卡基站切换与掉包。'),
    ('底层工业级落盘写锁', '断网时触发SQLite写锁，开启WAL预写式日志和synchronous=OFF。依赖带PLP电容的工业级NVMe全盘不掉速写入。', '杜绝消费级固态硬盘SLC缓存耗尽后掉速到50MB/s，从而反向憋死C#采集内存池的灾难。')
]
for row_data in d4:
    row = t4.add_row().cells
    row[0].text, row[1].text, row[2].text = row_data

for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = u'微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
                    run.font.size = Pt(10)

file_path = r'D:\铁路线路智能检测机器人\04-项目文档\设计文档\铁路线路智能检测机器人_四步流转硬核技术白皮书_工业仪器级.docx'
doc.save(file_path)
print(file_path)