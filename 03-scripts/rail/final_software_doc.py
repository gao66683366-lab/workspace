# -*- coding: utf-8 -*-
"""
软件系统架构设计文档 - 完整专业版
全表格化设计
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import sys

print("开始生成软件架构设计文档...")

doc = Document()

# 样式设置
style = doc.styles['Normal']
style.font.name = '宋体'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.font.size = Pt(12)

def create_table(headers, data, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)
    
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = header
        hdr_cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in hdr_cells[idx].paragraphs[0].runs:
            run.font.bold = True
    
    for row_data in data:
        row = table.add_row()
        for idx, cell_data in enumerate(row_data):
            row.cells[idx].text = str(cell_data)
    
    doc.add_paragraph()
    return table

# 封面
print("生成封面...")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('\n\n\n\n\n')
title = p.add_run('铁路线路智能检测机器人')
title.font.name = '黑体'
title._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
title.font.size = Pt(26)
title.font.bold = True
p.add_run('\n\n')
subtitle = p.add_run('软件系统架构设计文档')
subtitle.font.name = '黑体'
subtitle._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
subtitle.font.size = Pt(20)
p.add_run('\n\n\n\n\n\n\n\n\n\n')
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info_run = info.add_run('编制日期：2026年3月5日\n版本：V1.0\n密级：内部')
info_run.font.size = Pt(12)
doc.add_page_break()

# 修订历史
print("生成修订历史...")
doc.add_heading('文档修订历史', 1)
create_table(
    ['版本', '日期', '修订内容', '修订人', '审核人'],
    [['V1.0', '2026-03-05', '初始版本，完成软件系统详细设计', '小测', '道']],
    [2, 3, 6, 2, 2]
)

doc.add_heading('目录', 1)
doc.add_paragraph('【此处应插入自动目录，在Word中通过"引用→目录→自动目录"功能生成】')
doc.add_page_break()

# 第1章
print("第1章：概述...")
doc.add_heading('1. 概述', 1)
doc.add_heading('1.1 系统简介', 2)
doc.add_paragraph(
    '本软件系统是铁路线路智能检测机器人的核心控制系统，采用C# .NET 8.0开发，'
    '运行于工控上位机，负责设备控制、数据采集、实时处理、本地AI推理、人机交互等功能。'
    '系统采用模块化架构设计，各模块职责清晰，便于开发、测试和维护。'
)

doc.add_heading('1.2 设计目标', 2)
create_table(
    ['目标类别', '具体指标', '说明'],
    [
        ['实时性', '图像处理延迟<100ms', 'AI推理响应时间满足实时检测要求'],
        ['可靠性', '连续运行8小时无崩溃', '内存泄漏控制，异常处理完善'],
        ['准确性', 'AI识别准确率≥95%', '缺陷检测漏检率<5%'],
        ['可维护性', '模块化设计', '单个模块可独立更换升级'],
        ['可扩展性', '支持功能扩展', '预留接口，支持新传感器接入'],
        ['易用性', '界面友好', '操作简单，10分钟培训即可上手']
    ],
    [3, 4, 8]
)

doc.add_heading('1.3 开发环境', 2)
create_table(
    ['类别', '名称', '版本', '用途'],
    [
        ['开发语言', 'C#', '.NET 8.0', '主要开发语言'],
        ['IDE', 'Visual Studio', '2022', '集成开发环境'],
        ['UI框架', 'WPF', '.NET 8.0', '用户界面开发'],
        ['图像处理', 'OpenCvSharp', '4.x', '图像预处理'],
        ['AI推理', 'ONNX Runtime', '1.17+', 'AI模型推理'],
        ['通信库', 'NModbus4', '3.0+', 'Modbus通信'],
        ['数据库', 'SQLite / LiteDB', '最新稳定版', '本地数据存储'],
        ['日志', 'Serilog', '3.x', '日志记录']
    ],
    [3, 3.5, 2.5, 6.5]
)

doc.add_page_break()

# 第2章
print("第2章：系统架构...")
doc.add_heading('2. 系统架构', 1)
doc.add_heading('2.1 整体架构', 2)
create_table(
    ['层次', '模块', '主要功能', '技术选型'],
    [
        ['表示层', 'UI界面模块', '实时显示、操作控制、参数设置', 'WPF + MVVM'],
        ['', '报表模块', '检测报告生成、数据导出', 'Word/Excel自动化'],
        ['业务逻辑层', '检测流程控制', '任务调度、状态管理', '状态机模式'],
        ['', '数据处理分析', 'AI推理、参数计算、异常判定', 'ONNX Runtime'],
        ['', '运动控制', '速度控制、位置跟踪', 'EtherCAT库'],
        ['数据访问层', '数据存储模块', '数据库读写、文件管理', 'SQLite + FileSystem'],
        ['', '云端通信模块', '数据上传、模型下载', 'HTTP/WebSocket'],
        ['设备驱动层', '硬件抽象层', '传感器驱动、设备接口', 'SDK封装']
    ],
    [2.5, 3, 6, 4]
)

doc.add_heading('2.2 模块划分', 2)
create_table(
    ['模块名称', '子模块', '核心类', '依赖项'],
    [
        ['运动控制模块', 'EtherCAT通信\n伺服控制\n轨迹规划', 'MotionController\nServoDriver\nPathPlanner', 'EtherCAT主站库'],
        ['视觉采集模块', '相机管理\n图像采集\n图像预处理', 'CameraManager\nImageCapture\nImageProcessor', 'Basler Pylon SDK\nOpenCvSharp'],
        ['3D测量模块', '激光控制\n点云采集\n轮廓提取', 'LaserScanner\nPointCloudProcessor', '激光厂商SDK'],
        ['传感器模块', 'Modbus通信\n测距驱动\n陀螺仪驱动', 'ModbusManager\nSensorDriver', 'NModbus4'],
        ['数据处理模块', 'AI推理引擎\n参数计算\n异常检测', 'AIInference\nGeometryCalculator\nDefectDetector', 'ONNX Runtime'],
        ['数据存储模块', '数据库管理\n文件管理\n日志管理', 'DatabaseManager\nFileManager\nLogger', 'SQLite\nSerilog'],
        ['通信模块', 'HTTP客户端\nWebSocket\n数据同步', 'CloudClient\nDataSyncer', 'RestSharp'],
        ['UI模块', '主窗口\n实时显示\n参数设置', 'MainWindow\nRealtimeView\nSettingsView', 'WPF']
    ],
    [3, 3.5, 3.5, 5.5]
)

doc.add_page_break()

# 第3章 - 核心模块详细设计
print("第3章：核心模块设计...")
doc.add_heading('3. 核心模块详细设计', 1)

doc.add_heading('3.1 运动控制模块', 2)
doc.add_heading('3.1.1 模块职责', 3)
create_table(
    ['功能', '输入', '输出', '性能要求'],
    [
        ['速度控制', '目标速度（km/h）', '电机转速指令', '响应时间<50ms'],
        ['位置跟踪', '编码器脉冲', '当前位置（m）', '更新频率1kHz'],
        ['轨迹规划', '起点、终点、速度', '运动曲线', '平滑度±5%'],
        ['急停控制', '急停信号', '停止指令', '响应时间<10ms']
    ],
    [3, 3.5, 3.5, 5.5]
)

doc.add_heading('3.1.2 关键类设计', 3)
create_table(
    ['类名', '属性', '方法', '说明'],
    [
        ['MotionController', 'Speed\nPosition\nState', 'Start()\nStop()\nSetSpeed()\nEmergencyStop()', 'EtherCAT主站控制'],
        ['ServoDriver', 'AxisId\nEncoder\nCurrent', 'Enable()\nDisable()\nMove()\nGetPosition()', '单轴伺服控制'],
        ['PathPlanner', 'StartPos\nEndPos\nProfile', 'PlanPath()\nInterpolate()', 'S曲线轨迹规划']
    ],
    [3, 4, 5, 3.5]
)

doc.add_heading('3.2 视觉采集模块', 2)
doc.add_heading('3.2.1 相机管理策略', 3)
create_table(
    ['相机类型', '采集模式', '触发方式', '帧率', '数据处理'],
    [
        ['轨面相机×2', '连续采集', '自由运行', '30fps', '实时AI推理'],
        ['螺栓相机×4', '触发采集', '编码器触发', '按需', '检测后存储'],
        ['采集线程', '独立线程×6', '异步采集', '-', '避免阻塞主线程'],
        ['图像队列', '环形缓冲区', '生产者-消费者', '-', '平滑数据流']
    ],
    [3, 3, 3, 2, 4.5]
)

doc.add_heading('3.2.2 图像预处理流程', 3)
create_table(
    ['步骤', '算法', '参数', '耗时（ms）'],
    [
        ['1. 图像读取', '从SDK获取', '-', '<5'],
        ['2. 格式转换', 'BGR→灰度', '-', '<2'],
        ['3. 去噪', '高斯滤波', '核大小5×5', '<3'],
        ['4. 增强', '直方图均衡化', '-', '<3'],
        ['5. ROI提取', '感兴趣区域', '根据位置', '<1'],
        ['6. 缓存', '写入队列', '-', '<1'],
        ['总计', '-', '-', '<15']
    ],
    [2, 3.5, 4, 3]
)

doc.add_heading('3.3 AI推理模块', 2)
doc.add_heading('3.3.1 模型部署方案', 3)
create_table(
    ['模型', '输入', '输出', '推理时间', '部署格式'],
    [
        ['螺栓检测', '640×640 RGB', 'Boxes + Classes', '<30ms', 'YOLOv8.onnx'],
        ['轨面缺陷分类', '224×224 灰度', 'Class + Confidence', '<20ms', 'ResNet50.onnx'],
        ['轮廓分割', '512×512 灰度', 'Segmentation Mask', '<50ms', 'UNet.onnx']
    ],
    [3, 3, 3.5, 3, 3]
)

doc.add_heading('3.3.2 推理优化策略', 3)
create_table(
    ['优化项', '方法', '效果', '实施难度'],
    [
        ['模型量化', 'INT8量化', '速度提升2-3倍', '中'],
        ['GPU加速', 'CUDA推理', '速度提升5-10倍', '低'],
        ['批处理', '批量推理', '吞吐量提升30%', '低'],
        ['异步推理', '多线程并行', 'CPU利用率提升50%', '中'],
        ['模型剪枝', '移除冗余层', '模型缩小30%', '高']
    ],
    [3, 3.5, 3.5, 3.5]
)

doc.add_page_break()

# 第4章 - 数据流设计
print("第4章：数据流设计...")
doc.add_heading('4. 数据流设计', 1)

doc.add_heading('4.1 实时数据流', 2)
create_table(
    ['数据源', '数据类型', '频率', '流向', '处理方式'],
    [
        ['工业相机×6', '图像（5MP）', '30fps', '采集→预处理→AI', '多线程管道'],
        ['3D激光×2', '点云（XYZ）', '2kHz', '采集→轮廓提取→计算', '实时处理'],
        ['测距传感器×8', '距离值', '10Hz', '采集→滤波→计算', '卡尔曼滤波'],
        ['陀螺仪×2', '姿态角', '100Hz', '采集→解算→计算', '四元数解算'],
        ['AI推理结果', '检测框/类别', '30fps', '推理→判定→存储', '异常报警']
    ],
    [3, 2.5, 2, 3.5, 4.5]
)

doc.add_heading('4.2 数据存储策略', 2)
create_table(
    ['数据类型', '存储格式', '存储位置', '保留时间', '压缩方式'],
    [
        ['原始图像', 'JPEG', 'NVMe SSD', '7天', 'JPEG质量85'],
        ['检测结果', 'SQLite', 'NVMe SSD', '永久', '无'],
        ['点云数据', 'PCD格式', 'NVMe SSD', '7天', 'LZ4压缩'],
        ['传感器数据', 'CSV/SQLite', 'NVMe SSD', '永久', '无'],
        ['日志文件', 'TXT', 'HDD', '30天', 'GZip压缩']
    ],
    [2.5, 2.5, 3, 2.5, 3.5]
)

doc.add_page_break()

# 第5章 - 数据库设计
print("第5章：数据库设计...")
doc.add_heading('5. 数据库设计', 1)

doc.add_heading('5.1 主要数据表', 2)
create_table(
    ['表名', '主要字段', '索引', '说明'],
    [
        ['DetectionRecord', 'ID, Timestamp, Position, TrackGauge, Level, Alignment', 'Timestamp, Position', '检测记录主表'],
        ['BoltDetection', 'ID, RecordID, CameraID, BoltID, Status, ImagePath', 'RecordID', '螺栓检测结果'],
        ['RailDefect', 'ID, RecordID, CameraID, DefectType, Severity, ImagePath', 'RecordID, DefectType', '轨面缺陷'],
        ['ProfileData', 'ID, RecordID, LaserID, ProfilePoints', 'RecordID', '轮廓数据'],
        ['SystemLog', 'ID, Timestamp, Level, Module, Message', 'Timestamp, Level', '系统日志'],
        ['ConfigParams', 'Key, Value, Description', 'Key', '配置参数']
    ],
    [3, 6, 3, 3.5]
)

doc.add_heading('5.2 检测记录表结构', 2)
create_table(
    ['字段名', '类型', '约束', '说明'],
    [
        ['ID', 'INTEGER', 'PRIMARY KEY', '自增主键'],
        ['Timestamp', 'TEXT', 'NOT NULL', 'ISO8601格式时间'],
        ['Position', 'REAL', 'NOT NULL', '位置（米）'],
        ['Speed', 'REAL', '', '检测速度（km/h）'],
        ['TrackGauge', 'REAL', '', '轨距（mm）'],
        ['Level', 'REAL', '', '水平（mm）'],
        ['Alignment', 'REAL', '', '高低（mm）'],
        ['Direction', 'REAL', '', '轨向（mm）'],
        ['Status', 'TEXT', '', '状态（正常/异常）']
    ],
    [3, 2.5, 3, 7]
)

doc.add_page_break()

# 第6章 - 通信协议
print("第6章：通信协议...")
doc.add_heading('6. 通信协议设计', 1)

doc.add_heading('6.1 云端通信', 2)
create_table(
    ['接口', '方法', '路径', '参数', '响应'],
    [
        ['数据上传', 'POST', '/api/data/upload', 'JSON数据包', '{"status":"ok"}'],
        ['模型下载', 'GET', '/api/model/download', 'modelId', '模型文件流'],
        ['心跳', 'POST', '/api/heartbeat', 'deviceId', '{"alive":true}'],
        ['配置获取', 'GET', '/api/config', 'deviceId', '配置JSON']
    ],
    [3, 2, 4, 4, 3.5]
)

doc.add_heading('6.2 Modbus通信参数', 2)
create_table(
    ['参数', '值', '说明'],
    [
        ['波特率', '115200', '高速通信'],
        ['数据位', '8', '标准配置'],
        ['停止位', '1', '标准配置'],
        ['校验位', 'None / Even', '根据设备'],
        ['功能码', '0x03（读）、0x10（写）', '读保持寄存器、写多个寄存器'],
        ['设备地址', '1-247', '每个设备唯一地址'],
        ['超时时间', '1000ms', '响应超时']
    ],
    [3, 3.5, 9]
)

doc.add_page_break()

# 第7章 - 异常处理
print("第7章：异常处理...")
doc.add_heading('7. 异常处理与容错', 1)

doc.add_heading('7.1 异常分类', 2)
create_table(
    ['异常类型', '触发条件', '处理策略', '记录级别'],
    [
        ['通信异常', 'EtherCAT/Modbus中断', '自动重连，3次失败报警', 'Error'],
        ['相机异常', '采集失败、超时', '跳过当前帧，连续10次报警', 'Warning'],
        ['AI推理异常', '模型加载失败、推理错误', '降级为规则判定', 'Error'],
        ['存储异常', '磁盘空间不足', '停止采集，清理旧数据', 'Critical'],
        ['传感器异常', '数据超限、无响应', '使用上一次有效值', 'Warning'],
        ['系统异常', '内存不足、CPU过热', '降速运行或停机', 'Critical']
    ],
    [2.5, 3, 5.5, 3.5]
)

doc.add_heading('7.2 容错机制', 2)
create_table(
    ['机制', '实现方式', '效果'],
    [
        ['数据冗余', '关键数据双备份', '防止单点故障'],
        ['断线续传', '队列+重试机制', '网络中断不丢数据'],
        ['自动恢复', '监控线程+重启', '无人值守运行'],
        ['降级运行', '功能分级+优先级', '部分故障不影响核心功能'],
        ['看门狗', '定时心跳检测', '死锁自动重启']
    ],
    [3, 6, 6.5]
)

doc.add_page_break()

# 第8章 - 性能优化
print("第8章：性能优化...")
doc.add_heading('8. 性能优化', 1)

doc.add_heading('8.1 多线程设计', 2)
create_table(
    ['线程名称', '优先级', '职责', 'CPU占用'],
    [
        ['主UI线程', '正常', '界面响应、用户交互', '<5%'],
        ['运动控制线程', '实时', 'EtherCAT通信、伺服控制', '5-10%'],
        ['相机采集线程×6', '高', '图像采集', '10-15%'],
        ['AI推理线程×2', '正常', 'AI推理（CPU/GPU）', '20-40%'],
        ['数据处理线程', '正常', '参数计算、数据分析', '10-15%'],
        ['存储线程', '低', '数据库写入、文件保存', '5-10%'],
        ['通信线程', '低', '云端数据上传', '<5%']
    ],
    [3.5, 2, 4.5, 3]
)

doc.add_heading('8.2 内存管理', 2)
create_table(
    ['策略', '实现', '效果'],
    [
        ['对象池', '图像缓冲区复用', '减少GC压力'],
        ['及时释放', 'using语句、Dispose模式', '避免内存泄漏'],
        ['图像压缩', 'JPEG压缩后存储', '降低内存占用70%'],
        ['分页加载', '历史数据按需加载', '启动时间缩短50%'],
        ['大对象池', 'LOH优化', '减少内存碎片']
    ],
    [3, 6, 6.5]
)

doc.add_page_break()

# 第9章 - 测试方案
print("第9章：测试方案...")
doc.add_heading('9. 测试方案', 1)

doc.add_heading('9.1 测试类型', 2)
create_table(
    ['测试类型', '测试内容', '工具', '通过标准'],
    [
        ['单元测试', '核心算法、工具类', 'xUnit / NUnit', '代码覆盖率≥80%'],
        ['集成测试', '模块间接口', 'xUnit + Moq', '所有接口测试通过'],
        ['性能测试', 'CPU/内存/响应时间', 'BenchmarkDotNet', '满足性能指标'],
        ['压力测试', '长时间运行稳定性', '实际运行8小时', '无崩溃、无内存泄漏'],
        ['功能测试', '完整业务流程', '手动+自动化', '功能100%可用']
    ],
    [3, 4, 3.5, 5]
)

doc.add_heading('9.2 关键测试用例', 2)
create_table(
    ['模块', '测试场景', '预期结果', '优先级'],
    [
        ['运动控制', '紧急停止响应时间', '<10ms', 'P0'],
        ['视觉采集', '6路相机并发采集', '无丢帧，帧率稳定', 'P0'],
        ['AI推理', '批量图像推理性能', '<30ms/张', 'P0'],
        ['数据存储', '磁盘满时降级处理', '自动清理，不崩溃', 'P1'],
        ['通信异常', '断网恢复', '自动重连，数据续传', 'P1']
    ],
    [2.5, 4.5, 4, 3.5]
)

doc.add_page_break()

# 第10章 - 部署与维护
print("第10章：部署维护...")
doc.add_heading('10. 部署与维护', 1)

doc.add_heading('10.1 部署清单', 2)
create_table(
    ['组件', '路径', '配置', '说明'],
    [
        ['主程序', 'C:\\Program Files\\RailInspector\\', 'app.config', '主可执行文件'],
        ['AI模型', 'C:\\Program Files\\RailInspector\\Models\\', '*.onnx', 'ONNX模型文件'],
        ['配置文件', 'C:\\ProgramData\\RailInspector\\', 'config.json', '系统配置'],
        ['数据库', 'D:\\RailData\\Database\\', 'local.db', 'SQLite数据库'],
        ['图像数据', 'D:\\RailData\\Images\\', '按日期分文件夹', '原始图像'],
        ['日志', 'D:\\RailData\\Logs\\', '按日期分文件', '运行日志']
    ],
    [2.5, 5, 3.5, 4.5]
)

doc.add_heading('10.2 日常维护', 2)
create_table(
    ['维护项', '频率', '操作', '负责人'],
    [
        ['日志检查', '每日', '查看Error/Warning日志', '操作员'],
        ['磁盘空间', '每日', '确保可用空间>100GB', '操作员'],
        ['数据备份', '每周', '导出检测数据到云端', '操作员'],
        ['软件更新', '每月', '检查并安装更新', '技术员'],
        ['性能评估', '每月', '检查CPU/内存/响应时间', '技术员'],
        ['模型更新', '按需', '下载新模型并测试', '技术员']
    ],
    [2.5, 2, 6.5, 3.5]
)

doc.add_page_break()

# 附录
print("生成附录...")
doc.add_heading('附录A：技术栈清单', 1)
create_table(
    ['类别', '名称', '版本', '开源协议', '用途'],
    [
        ['框架', '.NET', '8.0', 'MIT', '开发框架'],
        ['UI', 'WPF', '.NET 8.0', 'MIT', '界面开发'],
        ['图像处理', 'OpenCvSharp', '4.9+', 'Apache 2.0', '图像处理'],
        ['AI推理', 'ONNX Runtime', '1.17+', 'MIT', 'AI推理引擎'],
        ['通信', 'NModbus4', '3.0+', 'MIT', 'Modbus通信'],
        ['数据库', 'SQLite', '3.45+', 'Public Domain', '本地数据库'],
        ['日志', 'Serilog', '3.1+', 'Apache 2.0', '日志记录'],
        ['HTTP', 'RestSharp', '110+', 'Apache 2.0', 'HTTP客户端'],
        ['JSON', 'Newtonsoft.Json', '13.0+', 'MIT', 'JSON序列化']
    ],
    [2.5, 3, 2.5, 3, 4.5]
)

doc.add_heading('附录B：缩略语', 1)
create_table(
    ['缩略语', '英文全称', '中文说明'],
    [
        ['AI', 'Artificial Intelligence', '人工智能'],
        ['API', 'Application Programming Interface', '应用程序编程接口'],
        ['CPU', 'Central Processing Unit', '中央处理器'],
        ['GPU', 'Graphics Processing Unit', '图形处理器'],
        ['HTTP', 'HyperText Transfer Protocol', '超文本传输协议'],
        ['MVVM', 'Model-View-ViewModel', '模型-视图-视图模型'],
        ['ONNX', 'Open Neural Network Exchange', '开放神经网络交换格式'],
        ['REST', 'Representational State Transfer', '表述性状态转移'],
        ['SDK', 'Software Development Kit', '软件开发工具包'],
        ['SQL', 'Structured Query Language', '结构化查询语言'],
        ['UI', 'User Interface', '用户界面'],
        ['WPF', 'Windows Presentation Foundation', 'Windows呈现基础']
    ],
    [3, 6, 6.5]
)

# 保存文档
output_path = 'D:/铁路线路智能检测机器人/04-项目文档/设计文档/软件系统架构设计文档 V1.0.docx'
doc.save(output_path)
print(f"\n文档已保存: {output_path}")
print(f"文档大小: {len(doc.element.xml) / 1024:.2f} KB")
print("\n软件架构设计文档生成完成！")
