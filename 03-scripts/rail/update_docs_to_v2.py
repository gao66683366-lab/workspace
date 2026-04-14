import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

out_dir = r"D:\铁路线路智能检测机器人\04-项目文档\设计文档"
os.makedirs(out_dir, exist_ok=True)

def set_font_and_style(doc):
    doc.styles['Normal'].font.name = u'微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    doc.styles['Normal'].font.size = Pt(10.5)

# ================= 1. 硬件系统架构设计文档 V2.0 =================
doc_hw = Document()
set_font_and_style(doc_hw)
h1 = doc_hw.add_heading('铁路线路智能检测机器人', 0)
h2 = doc_hw.add_heading('硬件系统架构设计文档 V2.0 (纯工程技术版)', 1)

doc_hw.add_paragraph('【物理基线】：最高时速 10km/h、120mm小轮径刚性底盘、48V主供电、10网口全隔离。')
doc_hw.add_heading('一、 硬件网络与电气物理拓扑', level=2)

table_hw = doc_hw.add_table(rows=1, cols=3, style='Table Grid')
cells = table_hw.rows[0].cells
cells[0].text, cells[1].text, cells[2].text = '硬件模块', '核心组件配置', '工程级防线设计'
data_hw = [
    ('定制电源管理板(PMB)', '48V转多路隔离输出(强/弱电/网络独立)', '带独立PTC熔断与工控机引脚控制硬重启，彻底隔绝伺服反电动势。'),
    ('网络通道硬隔离', '主板2网口(EtherCAT+SIM路由器) + PCIe扩展(8口)', '10个千兆口纯物理隔离，视觉口强制开启 9014 Bytes 巨型帧。'),
    ('高精光学总成', '物理遮光罩 + 50mm柔性裙边 + 纯单色频闪光源', '隔绝外界光，相机曝光写死<200μs，微秒级硬同步消除10km/h拖影。'),
    ('刚性抗震防护', '定焦防震工业镜头 + 航空螺纹胶死锁 + 聚氨酯垫', '应对无减震底盘过道岔时的百G级瞬间高频硬冲击，防止镜片离焦。')
]
for row_data in data_hw:
    row = table_hw.add_row().cells
    row[0].text, row[1].text, row[2].text = row_data

for row in table_hw.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.name = u'微软雅黑'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

doc_hw.save(os.path.join(out_dir, '机器人硬件系统架构设计文档_V2.0_纯工程版.docx'))

# ================= 2. 软件系统架构设计文档 V2.0 =================
doc_sw = Document()
set_font_and_style(doc_sw)
doc_sw.add_heading('铁路线路智能检测机器人', 0)
doc_sw.add_heading('软件系统架构设计文档 V2.0 (纯工程技术版)', 1)

doc_sw.add_paragraph('【核心逻辑】：测量原理绝对解耦、零拷贝内存并发、CPU内核级隔离。')
doc_sw.add_heading('一、 核心软件控制与流转架构', level=2)

table_sw = doc_sw.add_table(rows=1, cols=3, style='Table Grid')
cells = table_sw.rows[0].cells
cells[0].text, cells[1].text, cells[2].text = '软件模块', '关键技术实现', '工程落地防死锁逻辑'
data_sw = [
    ('操作系统内核层', 'Thread Affinity (CPU核心独占绑定)', '将EtherCAT主站进程强制绑定独立CPU核，防止被视觉网络中断抢占导致急停。'),
    ('高并发采集引擎', '非托管内存池(Zero-Copy) + FrameID 重组', '底层SDK直写内存防GC卡顿。以电机里程脉冲序号为主键，强行对齐所有传感器数据。'),
    ('算法层(绝对解耦)', '四通道独立测量解算与数字滤波', '轨距、水平、高低互不干涉。针对刚性底盘，陀螺仪前置10Hz低通滤波剔除碎震。'),
    ('端云均衡流控', 'SIM卡主动探针 + I/O 防掉速保底', '进隧道触发写锁。大文件全量写入工业级企业盘(防消费级SSD掉速憋死内存池)。')
]
for row_data in data_sw:
    row = table_sw.add_row().cells
    row[0].text, row[1].text, row[2].text = row_data

for row in table_sw.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.name = u'微软雅黑'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

doc_sw.save(os.path.join(out_dir, '软件系统架构设计文档_V2.0_纯工程版.docx'))

# ================= 3. 综合技术方案 V2.0 =================
doc_all = Document()
set_font_and_style(doc_all)
doc_all.add_heading('铁路线路智能检测机器人', 0)
doc_all.add_heading('综合技术方案 V2.0 (纯工程技术版)', 1)

doc_all.add_paragraph('【重大更迭声明】：全面剥离不切实际的“伪科学”交叉补偿，转向符合铁路物理定律的工业级重构方案。')
doc_all.add_heading('一、 技术演进与架构升级对照', level=2)

table_all = doc_all.add_table(rows=1, cols=3, style='Table Grid')
cells = table_all.rows[0].cells
cells[0].text, cells[1].text, cells[2].text = '系统特性', 'V1.0 旧方案盲区', 'V2.0 纯工程版核心升版内容'
data_all = [
    ('定位与里程系统', '过度依赖易断联的GPS/北斗信号。', '纯绝对值编码器定位。配合输出扭矩监控剔除打滑轮，结合发车初始桩号标定。'),
    ('几何形位测量', '将测距仪与陀螺仪强行绞合进行补偿计算。', '各司其职，物理解耦。测距仪专精轨距，陀螺仪专精水平/高低，任意损坏互不连累。'),
    ('光学与视觉环境', '依赖耗费算力的图像算法处理环境光剧变。', '机械遮光罩 + 恒定参数写死 + 频闪灯微秒级硬同步，降维打击光照问题。'),
    ('电气与容错系统', '简单的DC-DC模块供电，易受电机干扰。', '全定制电源管理板(PMB)，强弱电物理劈开，带独立保护与底层软硬件看门狗联动。')
]
for row_data in data_all:
    row = table_all.add_row().cells
    row[0].text, row[1].text, row[2].text = row_data

for row in table_all.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.name = u'微软雅黑'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

doc_all.save(os.path.join(out_dir, '铁路线路智能检测机器人技术方案_V2.0_纯工程版.docx'))
print("DOCS V2 SUCCESS")
