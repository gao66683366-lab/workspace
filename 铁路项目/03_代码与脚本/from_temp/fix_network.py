# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

doc = Document('D:\\铁路线路智能检测机器人\\04-项目文档\\设计文档\\机器人硬件系统架构设计文档 V1.0.docx')

# 查找"3.5.2 通信接口要求"章节
for i, para in enumerate(doc.paragraphs):
    if '3.5.2 通信接口要求' in para.text:
        # 删除旧的接口说明
        delete_count = 0
        for j in range(i+1, len(doc.paragraphs)):
            if doc.paragraphs[j].style.name.startswith('Heading'):
                break
            delete_count += 1
        
        # 从后向前删除段落
        for j in range(delete_count):
            p = doc.paragraphs[i + delete_count - j]
            p._element.getparent().remove(p._element)
        
        # 插入新内容
        insert_pos = i + 1
        
        # 添加说明段落
        p1 = doc.add_paragraph('工控机通信接口配置方案：', style='List Bullet')
        p1._element.getparent().insert(insert_pos, p1._element)
        insert_pos += 1
        
        # 网口需求分析
        p2 = doc.add_paragraph('网口需求统计：', style='List Bullet 2')
        p2._element.getparent().insert(insert_pos, p2._element)
        insert_pos += 1
        
        p3 = doc.add_paragraph('工业相机×6 = 6个千兆网口', style='List Bullet 3')
        p3._element.getparent().insert(insert_pos, p3._element)
        insert_pos += 1
        
        p4 = doc.add_paragraph('3D线激光×2 = 2个千兆网口', style='List Bullet 3')
        p4._element.getparent().insert(insert_pos, p4._element)
        insert_pos += 1
        
        p5 = doc.add_paragraph('运动控制器（EtherCAT主站）×1 = 1个千兆网口', style='List Bullet 3')
        p5._element.getparent().insert(insert_pos, p5._element)
        insert_pos += 1
        
        p6 = doc.add_paragraph('云端通信×1 = 1个千兆网口', style='List Bullet 3')
        p6._element.getparent().insert(insert_pos, p6._element)
        insert_pos += 1
        
        p7 = doc.add_paragraph('备用×2 = 2个千兆网口', style='List Bullet 3')
        p7._element.getparent().insert(insert_pos, p7._element)
        insert_pos += 1
        
        p8 = doc.add_paragraph('合计：至少12个千兆网口', style='List Bullet 3')
        p8.runs[0].bold = True
        p8._element.getparent().insert(insert_pos, p8._element)
        insert_pos += 1
        
        # 解决方案
        p9 = doc.add_paragraph('解决方案：', style='List Bullet 2')
        p9._element.getparent().insert(insert_pos, p9._element)
        insert_pos += 1
        
        p10 = doc.add_paragraph('方案一：工控机主板集成2-4个千兆网口 + 安装多口千兆网卡（如Intel i350-T4四口网卡×2张）', style='List Bullet 3')
        p10._element.getparent().insert(insert_pos, p10._element)
        insert_pos += 1
        
        p11 = doc.add_paragraph('方案二：工控机集成2个网口 + 工业级千兆以太网交换机（≥16口，支持GigE Vision）', style='List Bullet 3')
        p11._element.getparent().insert(insert_pos, p11._element)
        insert_pos += 1
        
        p12 = doc.add_paragraph('推荐方案：方案二（交换机方案）', style='List Bullet 2')
        p12.runs[0].bold = True
        p12._element.getparent().insert(insert_pos, p12._element)
        insert_pos += 1
        
        p13 = doc.add_paragraph('优势：扩展性好、布线简洁、降低工控机PCI-E槽位占用', style='List Bullet 3')
        p13._element.getparent().insert(insert_pos, p13._element)
        insert_pos += 1
        
        p14 = doc.add_paragraph('交换机要求：工业级、非网管型或网管型、支持VLAN划分、支持巨型帧', style='List Bullet 3')
        p14._element.getparent().insert(insert_pos, p14._element)
        insert_pos += 1
        
        # 其他接口
        p15 = doc.add_paragraph('串口：≥2个RS485接口（用于Modbus通信）', style='List Bullet')
        p15._element.getparent().insert(insert_pos, p15._element)
        insert_pos += 1
        
        p16 = doc.add_paragraph('USB：≥4个USB 3.0接口', style='List Bullet')
        p16._element.getparent().insert(insert_pos, p16._element)
        insert_pos += 1
        
        p17 = doc.add_paragraph('显示：HDMI或VGA接口', style='List Bullet')
        p17._element.getparent().insert(insert_pos, p17._element)
        insert_pos += 1
        
        p18 = doc.add_paragraph('扩展槽：≥2个PCI-E插槽（用于安装采集卡、通信卡等）', style='List Bullet')
        p18._element.getparent().insert(insert_pos, p18._element)
        
        break

doc.save('D:\\铁路线路智能检测机器人\\04-项目文档\\设计文档\\机器人硬件系统架构设计文档 V1.0.docx')
print('[OK] 已修正网口配置方案')
