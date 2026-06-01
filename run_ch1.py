#!/usr/bin/env python3
import sys, os
sys.path.insert(0, '/root/.openclaw/workspace/word_tools')

try:
    from docx import Document
    from docx_gen import h1, h2, h3, body, add_table

    folder = '/root/.openclaw/workspace/列车底部构件检测机器人_技术文档章节拆分_20260526_1546/'
    doc = Document()

    # Only chapter 1
    path = os.path.join(folder, '列车底部构件检测机器人_技术文档_V1.0_第01章_项目概述与核心定位_20260526_1546.md')
    with open(path, encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    print(f'Lines: {len(lines)}')

    in_code = False
    code_lines = []
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.strip().startswith('```'):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                in_code = False
            i += 1
            continue
        if in_code:
            i += 1
            continue
        if line.startswith('# '):
            h1(doc, line[2:].strip())
            i += 1
        elif line.startswith('## '):
            h2(doc, line[3:].strip())
            i += 1
        elif line.startswith('### '):
            h3(doc, line[4:].strip())
            i += 1
        elif line.strip().startswith('|'):
            parts = [p.strip() for p in line.split('|')]
            interior = parts[1:-1]
            is_sep = all(p == '---' for p in interior)
            if len(parts) > 2 and not is_sep:
                table_rows = []
                while i < len(lines) and lines[i].strip().startswith('|'):
                    row_line = lines[i]
                    row_parts = [p.strip() for p in row_line.split('|')]
                    row_interior = row_parts[1:-1]
                    row_is_sep = all(p == '---' for p in row_interior)
                    if row_interior and not row_is_sep:
                        table_rows.append([p for p in row_line.split('|') if p.strip()])
                    i += 1
                if len(table_rows) > 1:
                    header = table_rows[0]
                    data_rows = table_rows[1:]
                    grouped = {}
                    for row in data_rows:
                        key = row[0].strip()
                        grouped.setdefault(key, []).append(row)
                    merged = [header]
                    for key, rows in grouped.items():
                        for ri, row in enumerate(rows):
                            if ri == 0:
                                merged.append(row)
                            else:
                                merged.append(['merge:' + key] + row[1:])
                    add_table(doc, merged)
                else:
                    add_table(doc, table_rows)
            else:
                i += 1
        elif line.strip() and not line.strip().startswith('<!--') and not line.strip().startswith('***'):
            body(doc, line.strip())
            i += 1
        else:
            i += 1

    out = '/tmp/test_ch1.docx'
    doc.save(out)
    print(f'Saved to {out}')
except Exception as e:
    print(f'Error: {e}')
    import traceback
    traceback.print_exc()