# -*- coding: utf-8 -*-
"""
铁路线路智能检测机器人技术方案
整合硬件架构和软件架构的完整技术方案文档
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import datetime

print("开始生成技术方案文档...")

doc = Document()

# 样式设置
style = doc.styles['Normal']
style.font.name = '宋体'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5

# 设置标题样式
for i in range(1, 4):
    heading = doc.styles[f'Heading {i}']
    heading.font.name = '黑体'
    heading._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    heading.font.color.rgb = RGBColor(0, 0, 0)
    heading.font.bold = True
    if i == 1:
        heading.font.size = Pt(18)
    elif i == 2:
        heading.font.size = Pt(15)
    else:
        heading.font.size = Pt(13)

def create_table(headers, data, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)
    
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = header
        hdr_cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in hdr_cells[idx].paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(11)
    
    for row_data in data:
        row = table.add_row()
        for idx, cell_data in enumerate(row_data):
            row.cells[idx].text = str(cell_data)
            row.cells[idx].paragraphs[0].runs[0].font.size = Pt(10.5)
    
    doc.add_paragraph()
    return table

# ==================== 封面 ====================
print("1/20 封面...")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('\n\n\n\n\n\n')

title = p.add_run('铁路线路智能检测机器人')
title.font.name = '黑体'
title._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
title.font.size = Pt(32)
title.font.bold = True
title.font.color.rgb = RGBColor(0, 51, 102)

p.add_run('\n\n')

subtitle = p.add_run('技术方案')
subtitle.font.name = '黑体'
subtitle._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
subtitle.font.size = Pt(24)
subtitle.font.color.rgb = RGBColor(0, 51, 102)

p.add_run('\n\n\n\n\n\n\n\n\n')

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info_text = f'''
编制单位：[单位名称]
编制日期：{datetime.date.today().strftime("%Y年%m月%d日")}
版本号：V1.0
密级：内部
'''
info_run = info.add_run(info_text)
info_run.font.size = Pt(14)

doc.add_page_break()

# ==================== 编制说明 ====================
print("2/20 编制说明...")
doc.add_heading('编制说明', 1)

doc.add_paragraph('本技术方案是针对铁路线路智能检测机器人项目编制的综合技术文档，详细阐述了系统的总体架构、硬件设计、软件设计、AI算法、系统集成、实施计划等内容。')

doc.add_heading('编制依据', 2)
create_table(
    ['序号', '文件名称', '文件编号', '版本'],
    [
        ['1', '铁路轨道几何尺寸测量仪', 'TB/T 3147-2012', '2012版'],
        ['2', '铁路线路修理规则', 'TG/GW102-2019', '2019版'],
        ['3', '工业机器人 安全要求', 'GB 11291.1-2011', '2011版'],
        ['4', '电磁兼容 试验和测量技术', 'GB/T 17626-2008', '2008版'],
        ['5', 'GigE Vision工业相机标准', 'GigE Vision 2.0', 'V2.0'],
        ['6', 'EtherCAT通信协议', 'IEC 61158', '最新版']
    ],
    [1.5, 5, 4, 2]
)

doc.add_heading('文档修订历史', 2)
create_table(
    ['版本', '日期', '修订内容', '编制', '审核', '批准'],
    [
        ['V1.0', '2026-03-05', '初始版本，完成技术方案编制', '小测', '道', '']
    ],
    [1.5, 2.5, 5, 1.5, 1.5, 1.5]
)

doc.add_heading('目录', 1)
doc.add_paragraph('【此处应插入自动目录】')

doc.add_page_break()

# ==================== 第1章 项目概述 ====================
print("3/20 第1章...")
doc.add_heading('1. 项目概述', 1)

doc.add_heading('1.1 项目背景', 2)
doc.add_paragraph(
    '铁路作为国家重要的基础设施和综合交通运输体系的骨干，其安全运营至关重要。'
    '铁路线路的几何参数（轨距、水平、高低、轨向）、钢轨状态（磨损、裂纹、剥落）、'
    '紧固件状态等直接影响行车安全。传统的人工巡检方式存在效率低、劳动强度大、'
    '主观性强、数据不全面等问题，难以满足现代铁路高密度、高速度运营的需求。'
)
doc.add_paragraph(
    '随着人工智能、机器视觉、传感器技术的快速发展，利用智能机器人进行铁路线路自动化检测'
    '成为可能。本项目旨在研发一款集机械、电子、光学、计算机技术于一体的智能检测机器人，'
    '实现对铁路线路状态的全面、精确、高效检测，为铁路维护决策提供科学依据，'
    '保障铁路运输安全。'
)

doc.add_heading('1.2 项目目标', 2)
create_table(
    ['目标类别', '具体指标', '实现路径'],
    [
        ['检测精度', '轨距/水平/高低：±0.5mm\n轨向：±1mm\n钢轨磨耗：±0.1mm', '高精度传感器+算法优化'],
        ['检测效率', '检测速度≥5km/h\n单日检测≥40km', '自动化运行+实时处理'],
        ['智能化', 'AI识别准确率≥95%\n自动生成检测报告', '深度学习+专家系统'],
        ['可靠性', '连续运行≥8小时\n故障率<0.1%', '工业级设计+冗余保护'],
        ['实用性', '操作培训≤10分钟\n维护便捷', '人性化设计+模块化']
    ],
    [2.5, 5.5, 7.5]
)

doc.add_heading('1.3 项目范围', 2)
doc.add_paragraph('本项目包括以下内容：')
create_table(
    ['序号', '内容', '说明'],
    [
        ['1', '硬件系统开发', '机械结构、运动控制、传感器集成、工控机配置'],
        ['2', '软件系统开发', 'C# .NET 8.0上位机软件、数据处理、AI推理'],
        ['3', 'AI算法开发', '螺栓检测、轨面缺陷识别、轮廓分析算法'],
        ['4', '云端平台开发', 'Python后端、数据管理、模型训练平台'],
        ['5', '系统集成测试', '硬件调试、软件测试、联调联试'],
        ['6', '现场试运行', '实际线路测试、数据采集、性能验证'],
        ['7', '技术文档编制', '设计文档、用户手册、维护手册']
    ],
    [1.5, 3, 11]
)

doc.add_page_break()

# ==================== 第2章 总体技术方案 ====================
print("4/20 第2章...")
doc.add_heading('2. 总体技术方案', 1)

doc.add_heading('2.1 系统定位', 2)
doc.add_paragraph(
    '本系统是一款软硬件一体化、智能化、模块化的铁路线路检测装备，'
    '通过搭载工业相机、3D线激光、测距传感器、陀螺仪等多种传感器，'
    '在铁路轨道上自主运行，实时采集轨道状态数据，'
    '结合深度学习算法进行在线分析与判断，及时发现安全隐患，'
    '为铁路线路维护提供科学依据。'
)

doc.add_heading('2.2 总体架构', 2)
doc.add_paragraph('系统采用"云-边-端"三层架构：')

create_table(
    ['层级', '组成', '核心功能', '技术特点', '部署位置'],
    [
        ['云端层\nCloud', '云端服务器\n数据中心', '• 海量数据存储\n• AI模型训练\n• 大数据分析\n• Web管理平台\n• 远程诊断', '• Python开发\n• PostgreSQL数据库\n• PyTorch训练框架\n• 微服务架构', '云端机房\n或私有云'],
        ['边缘层\nEdge', '工控上位机\n机器人本体', '• 实时数据采集\n• 本地AI推理\n• 运动控制\n• 人机交互\n• 本地存储', '• C# .NET 8.0\n• Windows系统\n• ONNX Runtime\n• 多线程并行', '机器人\n工控机箱'],
        ['设备层\nDevice', '传感器\n执行器', '• 图像采集（相机×6）\n• 3D扫描（激光×2）\n• 距离测量（传感器×8）\n• 姿态检测（陀螺仪×2）\n• 运动执行（伺服电机×4）', '• EtherCAT总线\n• GigE Vision\n• Modbus RS485\n• 工业级设计', '机器人\n各安装位']
    ],
    [2, 2.5, 4.5, 4, 2.5]
)

doc.add_heading('2.3 技术路线', 2)
create_table(
    ['技术领域', '选用技术', '技术优势', '替代方案'],
    [
        ['运动控制', 'EtherCAT + 伺服系统', '实时性好、精度高、同步性强', 'CANopen'],
        ['视觉检测', '工业相机 + GigE Vision', '高分辨率、传输距离远', 'USB3 Vision'],
        ['3D测量', '线激光三角测量', '精度高、速度快', '双目视觉'],
        ['几何测量', '激光测距 + MEMS陀螺仪', '非接触、精度高', '接触式传感器'],
        ['AI算法', 'YOLOv8 + ResNet + UNet', '准确率高、速度快', 'Faster R-CNN'],
        ['AI推理', 'ONNX Runtime + GPU', '跨平台、速度快', 'TensorRT'],
        ['数据存储', 'SQLite + 文件系统', '轻量级、无需服务器', 'SQL Server'],
        ['通信协议', 'HTTPS + WebSocket', '安全、实时', 'MQTT']
    ],
    [2.5, 3.5, 4.5, 3]
)

doc.add_page_break()

# 继续生成... 由于内容太多，分批保存
print("5/20 第3章 硬件系统方案...")

doc.add_heading('3. 硬件系统方案', 1)

doc.add_heading('3.1 硬件总体设计', 2)
doc.add_paragraph('硬件系统由五大子系统构成，各子系统功能明确、接口清晰：')

create_table(
    ['子系统', '核心设备', '主要参数', '功能', '技术指标'],
    [
        ['运动控制\n子系统', '伺服电机×4\n运动控制器\n编码器×4', '功率200W\n转速3000rpm\n通信周期1ms', '驱动机器人运动\n速度0-10km/h可调\n位置精度±1mm', 'EtherCAT总线\n响应时间<10ms\n同步精度±1μs'],
        ['视觉检测\n子系统', '工业相机×6\nLED光源×6\n镜头×6', '分辨率5MP/2MP\n帧率30fps\n接口GigE', '轨面检测×2\n螺栓检测×4\n实时AI推理', '检测准确率≥95%\n漏检率<5%\n误报率<10%'],
        ['3D测量\n子系统', '3D线激光×2\n采集卡', '扫描频率2kHz\n分辨率1024点/线\n精度±0.05mm', '钢轨轮廓扫描\n磨耗测量\n几何分析', '测量精度±0.05mm\n线宽≥100mm\n测量范围±50mm'],
        ['几何参数\n测量子系统', '测距传感器×8\n陀螺仪×2', '测距精度±0.1mm\n姿态精度0.01°\nModbus RS485', '轨距测量\n水平/高低测量\n轨向测量', '轨距±0.5mm\n水平±0.5mm\n高低±0.5mm\n轨向±1mm'],
        ['控制处理\n子系统', '工控机\n存储设备\n通信模块', 'CPU i7-12700\n内存32GB\n存储6.5TB\n48V供电', '数据采集控制\nAI推理\n数据存储\n云端通信', 'AI推理<30ms\n存储速度>6GB/s\n网口×10\n串口×3']
    ],
    [2, 2.5, 3, 3.5, 4.5]
)

doc.add_heading('3.2 工控机详细配置', 2)
create_table(
    ['类别', '配置项', '具体规格', '数量', '作用'],
    [
        ['处理器', 'CPU', 'Intel Core i7-12700\n12核20线程\n2.1-4.9GHz', '1', '数据处理\n多任务并行'],
        ['', '主板', 'Intel B760芯片组\nATX板型', '1', '系统平台'],
        ['内存', 'RAM', 'DDR4-3200 32GB\n双通道 2×16GB', '1套', '数据缓存\nAI推理'],
        ['存储', '系统盘', 'NVMe SSD 512GB\nPCIe 4.0 M.2', '1', '系统+软件'],
        ['', '数据盘', 'NVMe SSD 2TB×2\nRAID 0阵列', '2', '实时数据存储'],
        ['', '备份盘', 'SATA SSD 4TB', '1', '数据备份'],
        ['显卡', 'GPU', 'NVIDIA RTX A2000 8GB\n或RTX 4060', '1', 'AI推理加速'],
        ['网络', '板载网口', 'Intel千兆网卡', '2', 'EtherCAT+云端'],
        ['', '扩展网卡', 'Intel I350-T4 四口网卡', '2张', '相机+激光'],
        ['', '交换机（备选）', '16口千兆工业交换机', '1', '网络扩展'],
        ['串口', '扩展卡', 'PCIe转RS485 四口', '1', 'Modbus通信'],
        ['电源', 'DC-ATX', '48V转ATX 500W\n多重保护', '1', '系统供电'],
        ['散热', '散热器', '塔式风冷+机箱风扇', '若干', '散热降温'],
        ['机箱', '工控机箱', '4U机架式 IP54防护', '1', '防护+安装']
    ],
    [2, 2, 4.5, 1.5, 5.5]
)

doc.add_page_break()

# 保存进度并继续...
doc.save('D:/铁路线路智能检测机器人/04-项目文档/设计文档/铁路线路智能检测机器人技术方案 V1.0.docx')
print("[Progress] 前5章已生成，继续...")
