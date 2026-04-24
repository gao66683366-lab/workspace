#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
学术版论文Word生成脚本 - 将7章学术版md文件合并为一个docx
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os

doc = Document()

# 页面设置
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

def add_shading(cell, fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        tag = 'w:{}'.format(edge)
        element = OxmlElement(tag)
        element.set(qn('w:val'), kwargs.get(edge, 'single'))
        element.set(qn('w:sz'), '4')
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), '000000')
        tcBorders.append(element)
    tcPr.append(tcBorders)

def h1_center(text):
    """一级标题：第X章 黑体三号16pt居中"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16)
    r.font.name = '黑体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(10)
    return p

def h1_main(text):
    """论文标题：黑体18pt居中"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(18)
    r.font.name = '黑体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)
    return p

def h2_sec(text):
    """二级节标题：黑体四号14pt左对齐"""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    r.font.name = '黑体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p

def h3_sub(text):
    """三级节标题：黑体小四12pt左对齐"""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.name = '黑体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    return p

def body_text(text):
    """正文：宋体小四12pt首行缩进"""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.name = '宋体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(24)
    return p

def body_text_noindent(text):
    """正文无首行缩进"""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.name = '宋体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(24)
    return p

def ref_text(text):
    """参考文献条目：宋体五号10.5pt"""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.name = '宋体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.left_indent = Cm(-0.74)
    p.paragraph_format.first_line_indent = Cm(-0.74)
    p.paragraph_format.space_after = Pt(3)
    return p

def center_text(text, size=Pt(12), bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = size
    r.bold = bold
    p.paragraph_format.space_after = Pt(4)
    return p

def table_caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.name = '宋体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    r.italic = True
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_page_break():
    doc.add_page_break()

def parse_inline(text):
    """处理行内格式：*斜体*、**粗体**、$行内公式$"""
    parts = []
    # 分割粗体
    segments = re.split(r'(\*\*[^*]+\*\*|\$[^\$]+\$|\*[^*]+\*)', text)
    for seg in segments:
        if seg.startswith('**') and seg.endswith('**'):
            p = doc.add_paragraph()
            r = p.add_run(seg[2:-2])
            r.bold = True
            r.font.size = Pt(12)
            r.font.name = '宋体'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            p.paragraph_format.space_after = Pt(0)
        elif seg.startswith('$') and seg.endswith('$'):
            p = doc.add_paragraph()
            r = p.add_run(seg[1:-1])
            r.font.size = Pt(11)
            r.font.name = 'Times New Roman'
            r.italic = True
            p.paragraph_format.space_after = Pt(0)
        elif seg.startswith('*') and seg.endswith('*'):
            p = doc.add_paragraph()
            r = p.add_run(seg[1:-1])
            r.italic = True
            r.font.size = Pt(12)
            r.font.name = '宋体'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            p.paragraph_format.space_after = Pt(0)
        elif seg.strip():
            p = doc.add_paragraph()
            r = p.add_run(seg)
            r.font.size = Pt(12)
            r.font.name = '宋体'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            p.paragraph_format.space_after = Pt(0)
    return doc.paragraphs[-1]

def parse_table(lines):
    """解析md表格"""
    if len(lines) < 2:
        return None
    # 计算列数
    header = lines[0].strip('|').split('|')
    ncols = len([h for h in header if h.strip()])
    # 创建表格
    table = doc.add_table(rows=len(lines), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, line in enumerate(lines):
        cells = line.strip('|').split('|')
        row = table.rows[i]
        for j, cell_text in enumerate(cells):
            if j >= ncols:
                continue
            cell = row.cells[j]
            cell.text = cell_text.strip()
            # 设置字体
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10.5)
                    r.font.name = '宋体'
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # 表头加粗
            if i == 0:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True
                add_shading(cell, 'D9E2F3')
    return table

def process_file(filepath):
    """处理单个md文件，返回是否成功"""
    if not os.path.exists(filepath):
        print(f"文件不存在: {filepath}")
        return False
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()
    lines = content.split('\n')
    in_table = False
    table_lines = []
    para_buffer = []
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        # 跳过论文大标题（文件开头）
        if stripped == '# 面向铁路线路的智能视觉与多模态感知融合综合检测系统研究与实现':
            i += 1
            continue
        # 一级标题（第X章）
        m = re.match(r'^## 第([一二三四五六七八九十]+)章\s+(.+)$', stripped)
        if m:
            if para_buffer:
                for pb in para_buffer:
                    if pb.strip():
                        body_text(pb)
                para_buffer = []
            h1_center(stripped)
            i += 1
            continue
        # 二级节标题（数字编号）
        m = re.match(r'^### (\d+\.\d+)\s+(.+)$', stripped)
        if m:
            if para_buffer:
                for pb in para_buffer:
                    if pb.strip():
                        body_text(pb)
                para_buffer = []
            h2_sec(stripped)
            i += 1
            continue
        # 三级节标题（数字编号数字）
        m = re.match(r'^#### (\d+\.\d+\.\d+)\s+(.+)$', stripped)
        if m:
            if para_buffer:
                for pb in para_buffer:
                    if pb.strip():
                        body_text(pb)
                para_buffer = []
            h3_sub(stripped)
            i += 1
            continue
        # 表格开始
        if stripped.startswith('|') and not in_table:
            in_table = True
            table_lines = []
            # 检查是否是分隔行
            if '---' in stripped or '---' in line:
                i += 1
                continue
            table_lines.append(stripped)
            i += 1
            continue
        # 表格内容
        if in_table:
            if stripped.startswith('|'):
                if '---' in stripped or ':-:' in stripped or ':-' in stripped or '-:' in stripped:
                    i += 1
                    continue
                table_lines.append(stripped)
                i += 1
                continue
            else:
                # 表格结束
                in_table = False
                if para_buffer:
                    for pb in para_buffer:
                        if pb.strip():
                            body_text(pb)
                    para_buffer = []
                parse_table(table_lines)
                table_caption(f'（续表）')
                continue
        # 参考文献
        if stripped.startswith('### 参考文献') or stripped == '### 参考文献':
            if para_buffer:
                for pb in para_buffer:
                    if pb.strip():
                        body_text(pb)
                para_buffer = []
            h2_sec('参考文献')
            i += 1
            continue
        # 参考文献条目
        m = re.match(r'^\[(\d+)\]\s+(.+)', stripped)
        if m:
            if para_buffer:
                for pb in para_buffer:
                    if pb.strip():
                        ref_text(pb)
                para_buffer = []
            ref_text(stripped)
            i += 1
            continue
        # 参考文献之后的空行区域跳过
        if stripped.startswith('---'):
            i += 1
            continue
        # 空行
        if not stripped:
            if para_buffer:
                for pb in para_buffer:
                    if pb.strip():
                        body_text(pb)
                para_buffer = []
            i += 1
            continue
        # 普通正文段落
        para_buffer.append(stripped)
        i += 1
    # 处理剩余buffer
    if para_buffer:
        for pb in para_buffer:
            if pb.strip():
                body_text(pb)
    return True

# ============ 主程序 ============
base_dir = '/root/.openclaw/workspace/铁路线路智能综合检测机器人/07-论文与学术资料/论文编撰/论文草稿/论文章节拆分_20260417_0259'
files = [
    ('学术版_第1章_绪论_20260424_1747.md', '第1章 绪论'),
    ('学术版_第2章_系统总体设计_20260424_1752.md', '第2章 系统总体设计'),
    ('学术版_第3章_硬件系统设计_20260424_1753.md', '第3章 硬件系统设计'),
    ('学术版_第4章_软件与AI算法设计_20260424_1754.md', '第4章 软件与AI算法设计'),
    ('学术版_第5章_综合检测功能实现_20260424_1755.md', '第5章 综合检测功能实现'),
    ('学术版_第6章_实验验证与分析_20260424_1757.md', '第6章 实验验证与分析'),
    ('学术版_第7章_结论与展望_20260424_1758.md', '第7章 结论与展望'),
]

# 添加论文标题页
h1_main('面向铁路线路的智能视觉与多模态感知融合综合检测系统')
h1_main('研究与实现')

add_page_break()

# 依次处理每章
for filename, chapter in files:
    filepath = os.path.join(base_dir, filename)
    print(f'处理: {chapter} ({filename})')
    if process_file(filepath):
        print(f'  ✓ 完成')
    else:
        print(f'  ✗ 失败')
    add_page_break()

# 保存
output_path = os.path.join(base_dir, '学术版_完整1-7章_20260424_1802.docx')
doc.save(output_path)
print(f'\n已保存: {output_path}')
