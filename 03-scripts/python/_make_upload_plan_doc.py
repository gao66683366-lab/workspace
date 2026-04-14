from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from pathlib import Path

out = Path(r'C:\Users\DELL\.openclaw\workspace\02-docx\drafts\数据均衡上传方案_V1.0_2026-03-08.docx')
out.parent.mkdir(parents=True, exist_ok=True)

doc = Document()
style = doc.styles['Normal']
style.font.name = '宋体'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.font.size = Pt(11)

def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16)
    r.font.name = '黑体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

def add_h1(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    r.font.name = '黑体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

def add_h2(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.name = '黑体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

def add_para(text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(22)
    p.paragraph_format.line_spacing = 1.5
    p.add_run(text)

def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing = 1.5
    p.add_run(text)

add_title('数据均衡上传方案')
sub = doc.add_paragraph()
sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
sub.add_run('版本：V1.0    日期：2026-03-08')

add_h1('1. 文档目的')
add_para('本方案用于明确铁路线路智能检测机器人在本地采集、大容量缓存、链路调度、断点续传、分级上传与云端接收等环节的整体实现思路，形成一套能够在工业路由器带 SIM 卡接入公网条件下稳定运行的数据均衡上传机制。')

add_h1('2. 项目已知约束')
for t in [
    '机器人平台为铁路线路智能检测机器人，采用 4 个轮子在铁轨上运行。',
    '工控机主板原生 LAN 1 独占 EtherCAT 总线，主板原生 LAN 2 独占连接外置工业路由器。',
    'PCIe 扩展网口连接 6 个 2D 工业相机与 2 个 3D 线激光，采集侧各设备独立子网隔离。',
    '低速传感器采用串口/RS485 或 EtherCAT I/O，不占用采集网口。',
    '采集基准按里程脉冲（Frame ID）进行硬同步对齐，并开启巨型帧（9014 Bytes）。',
    '机器人运行假定工况为 3 km/h，每天运行 2 小时。'
]: add_bullet(t)

add_h1('3. 方案目标')
for t in [
    '保证采集业务与上传业务网络隔离，互不抢占关键采集链路。',
    '将上传过程从“实时硬推”改为“本地缓存 + 分级调度 + 均衡上传”。',
    '在公网波动、弱网、断网条件下保持数据不丢失、可恢复、可追溯。',
    '使上传策略可按任务优先级、文件类型、时间窗口和带宽状态动态调整。'
]: add_bullet(t)

add_h1('4. 总体方案')
add_para('整体采用“采集域、缓存域、上传域、云端接收域”四层结构。采集域负责将多源设备数据安全落盘；缓存域负责按任务批次、时间片和文件类型进行组织；上传域负责通过工业路由器执行带宽感知的均衡传输；云端接收域负责完成校验、入库、重传确认与状态回写。')

add_h2('4.1 采集域与上传域隔离原则')
for t in [
    '采集网与外网通信物理隔离，上传只能经主板 LAN 2 连接工业路由器。',
    '上传进程不得直接占用相机/激光采集链路，不得影响 EtherCAT 控制实时性。',
    '本地落盘优先于上传，所有上传数据均以本地稳定文件为唯一源头。'
]: add_bullet(t)

add_h2('4.2 本地缓存组织')
for t in [
    '按日期、任务批次、设备类型组织目录。',
    '按“原始数据 / 索引数据 / 摘要结果 / 异常片段”分层存储。',
    '每个上传单元生成唯一任务 ID、文件清单、大小、校验值与上传状态。'
]: add_bullet(t)

add_h2('4.3 均衡上传机制')
for t in [
    '采用后台守护式上传，不与采集主线程耦合。',
    '按优先级分队列：告警摘要 > 缺陷结果 > 索引元数据 > 原始大文件。',
    '按带宽状态进行速率控制，弱网时优先上传小文件和高价值结果。',
    '大文件采用分片上传与断点续传，避免单次失败导致整体重传。'
]: add_bullet(t)

add_h1('5. 上传优先级策略')
add_para('建议将上传内容划分为四级，以保证有限链路下优先交付最有价值的数据。')

table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '级别'
hdr[1].text = '数据类型'
hdr[2].text = '上传时机'
hdr[3].text = '说明'
rows = [
    ('P1','告警摘要/关键结果','优先实时或准实时','体量小、价值高，优先上云'),
    ('P2','结构化索引与元数据','网络稳定即上传','支撑检索、统计与任务追踪'),
    ('P3','缺陷相关片段/截图/短片段','空闲带宽时上传','用于快速复核与远程判断'),
    ('P4','原始全量数据','非高峰或任务结束后上传','体量最大，允许延后完成'),
]
for r in rows:
    cells = table.add_row().cells
    for i,v in enumerate(r):
        cells[i].text = v

add_h1('6. 断点续传与可靠性设计')
for t in [
    '每个分片记录偏移量、分片号、校验值与确认状态。',
    '上传中断后仅重传失败分片，不重传已确认分片。',
    '上传成功后保留本地状态记录，待云端回执确认后再进入归档或清理阶段。',
    '对关键结果文件采用双重校验：文件级校验 + 分片级校验。'
]: add_bullet(t)

add_h1('7. 建议的数据目录结构')
add_para('建议在工控机本地建立统一的数据组织规则，便于检索、重传和追责。')
for t in [
    '任务目录：按日期/线路/批次建立主目录。',
    '原始目录：保存相机、线激光、姿态、测距等原始采集数据。',
    '索引目录：保存 Frame ID、时间戳、里程、设备映射等结构化索引。',
    '结果目录：保存检测结果、统计结果、异常记录与摘要。',
    '上传状态目录：保存待传、传输中、已完成、失败待重传等状态文件。'
]: add_bullet(t)

add_h1('8. 云端接收建议')
for t in [
    '云端接口应支持分片合并、断点续传、重复分片去重和状态回执。',
    '云端应先接收元数据和摘要结果，再逐步接收原始数据。',
    '云端入库后返回任务级确认信息，用于本地状态闭环。'
]: add_bullet(t)

add_h1('9. 风险与控制措施')
risks = [
    ('公网带宽波动大','采用优先级队列与限速调度'),
    ('长时间断网','本地缓存保底，网络恢复后自动续传'),
    ('原始数据体量过大','按分片上传并延后低优先级原始数据'),
    ('上传占用系统资源','上传进程限 CPU/IO/带宽，不干扰采集控制'),
]
rt = doc.add_table(rows=1, cols=2)
rt.style = 'Table Grid'
rt.rows[0].cells[0].text = '风险'
rt.rows[0].cells[1].text = '控制措施'
for a,b in risks:
    c = rt.add_row().cells
    c[0].text = a
    c[1].text = b

add_h1('10. 结论与下一步')
add_para('本方案建议采用“本地缓存优先、结果优先上传、原始数据分级延后、断点续传保障”的总体策略，以满足铁路线路智能检测机器人在移动公网环境下的数据可靠传输需求。下一步建议继续细化带宽估算、数据量预算、分片大小、云端接口字段和状态机定义，并形成可直接实施的软件设计说明。')

doc.save(str(out))
print(out)
