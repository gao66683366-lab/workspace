import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

output_dir = r"D:\铁路线路智能检测机器人\04-项目文档\设计文档"
os.makedirs(output_dir, exist_ok=True)

doc = docx.Document()

# 样式设置
doc.styles['Normal'].font.name = u'微软雅黑'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
doc.styles['Normal'].font.size = Pt(10.5)

def add_h(text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = u'微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    return h

add_h('铁路线路智能检测机器人', 0)
add_h('四步流转硬核技术白皮书 V3.0 (工业仪器级)', 1)

doc.add_paragraph('本文档为系统的最高工程执行标准，所有参数、协议及底层逻辑必须严格按此规范在 C#（上位机）与 Python（云端）中实现，绝不允许任何技术降级。')

add_h('第一步：核心数据采集 (底层硬件级时空定格)', 2)
doc.add_paragraph('目标：突破 Windows 非实时限制，利用 10 网口隔离获取微秒级时空对齐的绝对纯净切片。')
t1 = doc.add_table(rows=1, cols=3, style='Table Grid')
h1 = t1.rows[0].cells
h1[0].text, h1[1].text, h1[2].text = '网络设备', '物理拓扑', '底层驱动与 C# 接收代码规范'
data1 = [
    ('运动控制 (伺服电机)', '主板 LAN 1', '【生死红线】彻底卸载 Windows 标准 IPv4/IPv6 协议栈，替换为 EtherCAT Real-time NDIS 驱动。C# 绕过 Socket，通过 TwinCAT ADS 直读底层寄存器，获取绝对值编码器脉冲，生成 Frame ID。'),
    ('2D 工业相机 (6个)', 'PCIe-A & B', '开启 MTU 9014 巨型帧。C# 必须调用相机原厂 GigE Filter Driver，绕过 NDIS 层直接将图像 DMA 映射至 `unsafe` 非托管内存池。必须加入 GVSP 丢包重组与 PacketResend 逻辑。'),
    ('3D 线激光 (2个)', 'PCIe-B', '开辟专属 16-bit 深度图内存池。并行接收深度包，零拷贝传递给 PCL 点云库解析，与 2D 图像流水线绝对物理与线程隔离。')
]
for d in data1:
    row = t1.add_row().cells
    row[0].text, row[1].text, row[2].text = d
doc.add_paragraph('光学触发：基于 Frame ID，I/O 板卡输出 TTL 方波，相机 200μs 电子快门与频闪 LED 纳秒级同步。')

add_h('第二步：数据压缩、清洗与筛查 (信号层去伪存真)', 2)
doc.add_paragraph('目标：截断 90% 无效背景与机械噪声，防止 C# 内存环形队列爆满。')
t2 = doc.add_table(rows=1, cols=3, style='Table Grid')
h2 = t2.rows[0].cells
h2[0].text, h2[1].text, h2[2].text = '清洗维度', '异常源', 'C# 算法处理逻辑'
data2 = [
    ('里程脉冲清洗', '刚性底盘导致车轮悬空打滑', '实时读取 EtherCAT 的 4 轮 ActualTorque，判定扭矩突降轮为打滑轮，瞬间剔除该轮高频脉冲，取接地轮求均值。'),
    ('姿态高频清洗', '10km/h 带来的金属碎震', '串口接入的陀螺仪数据必须前置 10Hz 二阶巴特沃斯低通滤波器。并利用 C# 队列进行时间回溯（相位延迟补偿），确保与相机图像空间对齐。'),
    ('视觉 ROI 截流', '90% 无用钢轨碎石背景', '利用镜头点胶死锁特性，C# 调用 SIMD 硬件指令集，在毫秒级内直接裁剪 512x512 的扣件/裂纹 ROI 切片，全景大图直接抛弃入库。')
]
for d in data2:
    row = t2.add_row().cells
    row[0].text, row[1].text, row[2].text = d

add_h('第三步：核心算法处理 (独立矩阵的降维解算)', 2)
doc.add_paragraph('目标：坚决贯彻“各司其职”，杜绝伪科学补偿。')
t3 = doc.add_table(rows=1, cols=2, style='Table Grid')
h3 = t3.rows[0].cells
h3[0].text, h3[1].text = '算法通道', '解算模型与输出'
data3 = [
    ('轨距独立计算', '8 个测距传感器独立求和平均，不参杂任何姿态补偿，直出轨距。'),
    ('水平与高低计算', '陀螺仪 Roll 角算超高；Pitch 角结合里程计积分算纵向平顺度。'),
    ('磨耗 3D 匹配', '3D 激光点云采用 Point-to-Plane ICP（引入 Huber Loss 抗噪），配准标称 CAD 轨型。'),
    ('边缘 AI 确诊', 'GPU 加载 TensorRT INT8 量化 YOLO 模型，对 ROI 切片推理病害。')
]
for d in data3:
    row = t3.add_row().cells
    row[0].text, row[1].text = d

add_h('第四步：轻量化数据上传 (协议级的弱网流控机制)', 2)
doc.add_paragraph('目标：工业路由器(LAN 2)带宽极致压榨，确保秒级告警。')
t4 = doc.add_table(rows=1, cols=2, style='Table Grid')
h4 = t4.rows[0].cells
h4[0].text, h4[1].text = '核心流控技术', '工程落地规范'
data4 = [
    ('二进制极限压缩', '针对 P0 极高优告警，废弃 JSON，C# 采用 `StructLayout` 内存对齐，封装为仅 25 Bytes 的二进制结构体裸传。'),
    ('令牌桶探针调度', '基于 UDP RTT 延迟，动态调整外网 QoS 队列：网好传 ROI 切片，网差只传 25 Bytes 定位符。'),
    ('固态降速防爆', '隧道彻底断网时，写入本地工业级 NVMe SSD。SQLite 开启 WAL 模式防 I/O 阻塞。')
]
for d in data4:
    row = t4.add_row().cells
    row[0].text, row[1].text = d

for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = u'微软雅黑'
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
                    r.font.size = Pt(10)

file_path = r'D:\铁路线路智能检测机器人\04-项目文档\设计文档\铁路线路智能检测机器人_四步流转硬核技术白皮书_V3.0.docx'
doc.save(file_path)
print("SUCCESS")
