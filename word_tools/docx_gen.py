"""
docx_gen.py — 标准化 Word 文档生成工具 v6
永久工具 v6：V3.0架构图+更多类型自动识别

用法:
  python docx_gen.py <md文件夹/> <输出.docx> [--cover "标题" "副标题" "日期" "单位"]
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, re, sys, argparse

PAGE_W_CM = 21.0
PAGE_H_CM = 29.7
MARGIN_CM = 2.54
TEXT_W_CM = PAGE_W_CM - 2 * MARGIN_CM

LH_BODY   = Pt(22)
LH_CODE   = Pt(15)
LH_TABLE  = Pt(17)
INDENT_BODY  = Cm(0.74)
INDENT_TABLE = Cm(0.2)
HDR_FILL  = 'D9E2F3'

def sf(run, cn, size, bold=False, italic=False, color=None):
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), cn)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def pfmt(p, before, after, lh, indent=None,
         align=WD_ALIGN_PARAGRAPH.LEFT):
    f = p.paragraph_format
    f.space_before = before
    f.space_after  = after
    f.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    f.line_spacing = lh
    if indent is not None:
        f.first_line_indent = indent
    f.alignment = align
    f.keep_with_next = True

def shade_cell(cell, hex_fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)
    tcPr.append(shd)

def set_border(cell, color='000000', sz='6'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        t = OxmlElement(f'w:{edge}')
        t.set(qn('w:val'), 'single')
        t.set(qn('w:sz'), sz)
        t.set(qn('w:space'), '0')
        t.set(qn('w:color'), color)
        borders.append(t)
    tcPr.append(borders)

def set_table_border(tbl, color='000000', sz='6'):
    tblPr = tbl._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        t = OxmlElement(f'w:{edge}')
        t.set(qn('w:val'), 'single')
        t.set(qn('w:sz'), sz)
        t.set(qn('w:space'), '0')
        t.set(qn('w:color'), color)
        tblBorders.append(t)
    tblPr.append(tblBorders)

def col_widths(rows_data):
    if not rows_data:
        return []
    ncols = max(len(r) for r in rows_data)
    weights = [0.0] * ncols
    for ri, row in enumerate(rows_data):
        for ci, val in enumerate(row):
            if ci < ncols:
                txt = str(val)
                score = sum(1 if '一' <= c <= '鿿' else 0.55 for c in txt) + len(txt) * 0.2
                if ri == 0:
                    score *= 2
                weights[ci] = max(weights[ci], score)
    total = sum(weights)
    if total == 0:
        return [TEXT_W_CM / ncols] * ncols
    return [TEXT_W_CM * w / total for w in weights]

def fmt_text(p, text, font_cn, font_size):
    import re
    # Full-line bold: line starts with ** but does NOT end with ** (e.g. "**1. 数据离散")
    # or line is entirely **bold** wrapped (e.g. "**2. 数据离散**")
    stripped = text.strip()
    if stripped.startswith('**') and '**' in stripped[2:]:
        # Find the last ** and strip both
        last_star = stripped.rfind('**')
        if last_star > 2:
            content = stripped[2:last_star]
            if content:
                r = p.add_run(content)
                sf(r, font_cn, font_size, bold=True)
            return
    # Split on well-formed **bold** or `code`
    pattern = r'(\*\*.+?\*\*|`.+?`)'
    for seg in re.split(pattern, text):
        if not seg:
            continue
        elif seg.startswith('**') and seg.endswith('**'):
            r = p.add_run(seg[2:-2])
            sf(r, font_cn, font_size, bold=True)
        elif seg.startswith('`') and seg.endswith('`'):
            r = p.add_run(seg[1:-1])
            sf(r, font_cn, font_size, italic=True)
        else:
            r = p.add_run(seg)
            sf(r, font_cn, font_size)

def add_page_numbers(doc):
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        for r in p.runs:
            r.text = ''
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    def page_field(run, instr):
        fc1 = OxmlElement('w:fldChar')
        fc1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fc1)
        it = OxmlElement('w:instrText')
        it.set(qn('xml:space'), 'preserve')
        it.text = f' {instr} '
        run._r.append(it)
        fc2 = OxmlElement('w:fldChar')
        fc2.set(qn('w:fldCharType'), 'separate')
        run._r.append(fc2)
        t = OxmlElement('w:t')
        t.text = '1'
        run._r.append(t)
        fc3 = OxmlElement('w:fldChar')
        fc3.set(qn('w:fldCharType'), 'end')
        run._r.append(fc3)
    r = p.add_run()
    r.font.size = Pt(10)
    r.font.name = 'Times New Roman'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    page_field(r, 'PAGE')
    r2 = p.add_run(' / ')
    sf(r2, '宋体', 10)
    r3 = p.add_run()
    r3.font.size = Pt(10)
    r3.font.name = 'Times New Roman'
    r3._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    page_field(r3, 'NUMPAGES')

def add_cover(doc, title, subtitle='', date='', unit=''):
    section = doc.sections[0]
    section.top_margin = Cm(5)
    for _ in range(6):
        p = doc.add_paragraph()
        pfmt(p, Pt(0), Pt(0), LH_BODY)
    p_title = doc.add_paragraph()
    pfmt(p_title, Pt(0), Pt(14), Pt(40), align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p_title.add_run(title)
    sf(r, '黑体', 28, bold=True)
    if subtitle:
        p_sub = doc.add_paragraph()
        pfmt(p_sub, Pt(4), Pt(10), Pt(24), align=WD_ALIGN_PARAGRAPH.CENTER)
        r2 = p_sub.add_run(subtitle)
        sf(r2, '黑体', 18)
    p_line = doc.add_paragraph()
    pfmt(p_line, Pt(14), Pt(14), Pt(18), align=WD_ALIGN_PARAGRAPH.CENTER)
    r3 = p_line.add_run('─' * 44)
    sf(r3, '宋体', 10, color=(120, 120, 120))
    if date:
        p_date = doc.add_paragraph()
        pfmt(p_date, Pt(8), Pt(6), Pt(15), align=WD_ALIGN_PARAGRAPH.CENTER)
        r4 = p_date.add_run(date)
        sf(r4, '宋体', 13)
    if unit:
        p_unit = doc.add_paragraph()
        pfmt(p_unit, Pt(42), Pt(4), Pt(13), align=WD_ALIGN_PARAGRAPH.CENTER)
        r5 = p_unit.add_run(unit)
        sf(r5, '宋体', 11, color=(80, 80, 80))
    section.top_margin = Cm(MARGIN_CM)
    doc.add_page_break()

def empty(doc):
    p = doc.add_paragraph()
    pfmt(p, Pt(0), Pt(0), LH_BODY)

def h1(doc, text):
    p = doc.add_paragraph()
    pfmt(p, Pt(20), Pt(8), Pt(30), Cm(0), WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run(text)
    sf(r, '黑体', 16, bold=True)

def h2(doc, text):
    p = doc.add_paragraph()
    pfmt(p, Pt(14), Pt(6), Pt(22), Cm(0))
    r = p.add_run(text)
    sf(r, '黑体', 14, bold=True)

def h3(doc, text):
    p = doc.add_paragraph()
    pfmt(p, Pt(10), Pt(4), Pt(22), Cm(0))
    r = p.add_run(text)
    sf(r, '黑体', 12, bold=True)

def body(doc, text):
    p = doc.add_paragraph()
    pfmt(p, Pt(2), Pt(2), LH_BODY, INDENT_BODY)
    fmt_text(p, text, '宋体', 12)

def add_table(doc, rows_data):
    if not rows_data:
        return
    # Fix: detect flat list (single row passed as [col1, col2, ...]) and wrap
    if rows_data and isinstance(rows_data[0], str):
        rows_data = [rows_data]
    ncols = max(len(r) for r in rows_data)
    tbl = doc.add_table(rows=len(rows_data), cols=ncols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
    widths = col_widths(rows_data)
    set_table_border(tbl, '000000', '6')
    for ri, row_vals in enumerate(rows_data):
        row = tbl.rows[ri]
        is_head = (ri == 0)
        for ci, val in enumerate(row_vals):
            if ci >= ncols:
                break
            cell = row.cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_border(cell)
            cell.width = Cm(widths[ci])
            for p in cell.paragraphs:
                p.clear()
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_head else WD_ALIGN_PARAGRAPH.LEFT
            pfmt(p, Pt(2), Pt(2), LH_TABLE, Cm(0) if is_head else INDENT_TABLE)
            fmt_text(p, str(val), '黑体' if is_head else '宋体', 10.5)
            if is_head:
                shade_cell(cell, HDR_FILL)
                for r in p.runs:
                    r.font.bold = True
    return tbl

def add_image_to_doc(doc, image_path, width=Cm(15.1)):
    p = doc.add_paragraph()
    pfmt(p, Pt(4), Pt(4), Pt(12))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    from PIL import Image
    import io
    if image_path.endswith('.png'):
        img = Image.open(image_path)
        if img.mode == 'RGBA':
            bg = Image.new('RGB', img.size, (255, 255, 255))
            bg.paste(img, mask=img.split()[3])
            img = bg
        else:
            img = img.convert('RGB')
        buf = io.BytesIO()
        img.save(buf, 'JPEG', quality=90)
        buf.seek(0)
        run.add_picture(buf, width=width)
    else:
        run.add_picture(image_path, width=width)

def _get_cjk_font():
    import matplotlib.font_manager as fm
    for f in fm.fontManager.ttflist:
        if 'Noto Serif CJK' in f.name:
            return f.fname
    return None

def _get_prop():
    fp = _get_cjk_font()
    if fp:
        import matplotlib.font_manager as fm
        fm.fontManager.addfont(fp)
        return fm.FontProperties(fname=fp)
    return None

def make_architecture_matplotlib(lines, text_w=TEXT_W_CM):
    all_text = '\n'.join(lines)
    if '感矩层' in all_text and '运动层' in all_text and '计算层' in all_text and '交互层' in all_text and '供电层' not in all_text:
        return _make_v1_hw_arch_png(lines, text_w)
    if all(x in all_text for x in ['传感层', '运动层', '计算层', '交互层', '供电层']):
        return _make_v3_hw_arch_png(lines, text_w)
    if '三网物理障离' in all_text or all(x in all_text for x in ['控制网络', '采集网络', '传输网络']):
        return _make_network_arch_png(lines, text_w)
    if '应用层' in all_text and 'AI算法层' in all_text:
        return _make_sw_arch_png(lines, text_w)
    if '主显示区' in all_text:
        return _make_hmi_arch_png(lines, text_w)
    if '实时采集' in all_text and '存储管理' in all_text:
        return _make_data_flow_png(lines, text_w)
    if '六大核心价值' in all_text or ('全自动' in all_text and '实时检测' in all_text and '大数据' in all_text):
        return _make_value_tree_png(lines, text_w)
    return _make_generic_arch_png(lines, text_w)



def _make_v1_hw_arch_png(lines, text_w):
    import matplotlib; matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    from matplotlib.patches import FancyBboxPatch
    prop = _get_prop()
    fig, ax = plt.subplots(1,1, figsize=(text_w*0.55, 3.5))
    ax.set_xlim(0, text_w); ax.set_ylim(0, 3.5); ax.axis('off')
    fig.patch.set_facecolor('white')
    outer = FancyBboxPatch((0.3,0.4), text_w-0.6, 2.7, boxstyle='round,pad=0.06', lw=1.2, edgecolor='#444444', facecolor='#F7F7F7', zorder=0)
    ax.add_patch(outer)
    cols = [
        {'name':'感知层','color':'#D9E2F3','items':['摄像模块\n(2×工业相机)','照明模块\n(频闪补光)','传感融合\n(IMU+里程计)']},
        {'name':'运动层','color':'#E2F0D9','items':['轨底平台\n(4驱动轮+1电缸)','电动云台\n(双云台俯仰)','线缆收放\n(自动拖链)']},
        {'name':'计算层','color':'#F2D9D9','items':['主控单元\n(Intel i7工控机)','AI加速\n(GPU扩展)','通信单元\n(4G/wifi)']},
        {'name':'交互层','color':'#E2D9F2','items':['手持终端\n(Android)','急停装置\n(硬件保护)','状态指示\n(LED声光)']},
    ]
    col_w = (text_w-0.8)/4; row_y = [2.85, 1.85, 0.85]
    for ci, col in enumerate(cols):
        x = 0.5 + ci*col_w
        ax.add_patch(FancyBboxPatch((x+0.05,2.95), col_w-0.1, 0.45, boxstyle='round,pad=0.03', lw=0.8, edgecolor='#666666', facecolor=col['color'], zorder=2))
        ax.text(x+col_w/2, 3.175, col['name'], ha='center', va='center', fontsize=9, fontweight='bold', fontproperties=prop)
        ax.plot([x+col_w,x+col_w],[0.5,3.3], color='#BBBBBB', lw=0.5, ls='--', zorder=1)
        for ii, item in enumerate(col['items']):
            ax.add_patch(FancyBboxPatch((x+0.05,row_y[2-ii]-0.35), col_w-0.1, 0.55, boxstyle='round,pad=0.04', lw=0.6, edgecolor='#BBBBBB', facecolor='white', zorder=2))
            ax.text(x+col_w/2, row_y[2-ii]-0.08, item, ha='center', va='center', fontsize=7.5, fontproperties=prop or None)
    ax.text(text_w/2, 3.42, '列车底部构件检测机器人 — 硬件系统总体架构', ha='center', va='center', fontsize=10, fontweight='bold', fontproperties=prop)
    out = f'/tmp/arch_v1_{hash(tuple(lines))%100000}.png'
    plt.savefig(out, dpi=150, bbox_inches='tight', facecolor='white', edgecolor='none')
    plt.close(); return out

def _make_v3_hw_arch_png(lines, text_w):
    import matplotlib; matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    from matplotlib.patches import FancyBboxPatch
    prop = _get_prop()
    fig, ax = plt.subplots(1,1, figsize=(text_w*0.55, 3.8))
    ax.set_xlim(0, text_w); ax.set_ylim(0, 3.8); ax.axis('off')
    fig.patch.set_facecolor('white')
    outer = FancyBboxPatch((0.3,0.3), text_w-0.6, 3.1, boxstyle='round,pad=0.06', lw=1.2, edgecolor='#444444', facecolor='#F7F7F7', zorder=0)
    ax.add_patch(outer)
    layers = [
        {'name':'传感层','color':'#D9E2F3','items':['2D相机+3D激光','姿态传感器\nHWT905','编码器+单点\n测距仪矩阵']},
        {'name':'运动层','color':'#E2F0D9','items':['4轮行走机构','伺服驱动','轮速编码器']},
        {'name':'计算层','color':'#F2D9D9','items':['工控机\n边缘AI推理','数据融合单元','通信模块']},
        {'name':'交互层','color':'#E2D9F2','items':['触摸屏','遥控器','物理按钮']},
        {'name':'供电层','color':'#FFF3CD','items':['48V蓄电池组','智能电源管理','三级保护回路']},
    ]
    n=5; layer_w=(text_w-0.8)/n; row_y=[3.15,2.05,1.05]
    for ci, layer in enumerate(layers):
        x = 0.5 + ci*layer_w
        ax.add_patch(FancyBboxPatch((x+0.05,3.15), layer_w-0.1, 0.4, boxstyle='round,pad=0.03', lw=0.8, edgecolor='#666666', facecolor=layer['color'], zorder=2))
        ax.text(x+layer_w/2, 3.35, layer['name'], ha='center', va='center', fontsize=8.5, fontweight='bold', fontproperties=prop)
        ax.plot([x+layer_w,x+layer_w],[0.4,3.45], color='#BBBBBB', lw=0.5, ls='--', zorder=1)
        for ii, item in enumerate(layer['items']):
            ax.add_patch(FancyBboxPatch((x+0.05,row_y[2-ii]-0.28), layer_w-0.1, 0.5, boxstyle='round,pad=0.04', lw=0.6, edgecolor='#BBBBBB', facecolor='white', zorder=2))
            ax.text(x+layer_w/2, row_y[2-ii]-0.03, item, ha='center', va='center', fontsize=7, fontproperties=prop or None)
    ax.text(text_w/2, 3.65, '铁路线路智能综合检测机器人 — 硬件系统总体架构', ha='center', va='center', fontsize=10, fontweight='bold', fontproperties=prop)
    out = f'/tmp/arch_v3_{hash(tuple(lines))%100000}.png'
    plt.savefig(out, dpi=150, bbox_inches='tight', facecolor='white', edgecolor='none')
    plt.close(); return out

def _make_network_arch_png(lines, text_w):
    import matplotlib; matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    from matplotlib.patches import FancyBboxPatch
    prop = _get_prop()
    fig, ax = plt.subplots(1,1, figsize=(text_w*0.55, 2.8))
    ax.set_xlim(0, text_w); ax.set_ylim(0, 2.8); ax.axis('off')
    fig.patch.set_facecolor('white')
    outer = FancyBboxPatch((0.3,0.4), text_w-0.6, 2.0, boxstyle='round,pad=0.06', lw=1.2, edgecolor='#444444', facecolor='#F7F7F7', zorder=0)
    ax.add_patch(outer)
    nets = [
        {'name':'控制网络\n(EtherCAT)','color':'#D9E2F3','items':['伺服驱动\n毫秒级响应','步进电机\n闭环控制']},
        {'name':'采集网络\n(千兆以太网)','color':'#E2F0D9','items':['相机数据\n高速传输','AI推理\n实时分析']},
        {'name':'传输网络\n(4G/5G)','color':'#F2D9D9','items':['远程监控\n数据上报','指令下发\n状态回传']},
    ]
    col_w = (text_w-0.8)/3; row_y = [2.1, 1.2, 0.7]
    for ci, net in enumerate(nets):
        x = 0.5 + ci*col_w
        ax.add_patch(FancyBboxPatch((x+0.08,1.85), col_w-0.16, 0.55, boxstyle='round,pad=0.03', lw=0.8, edgecolor='#666666', facecolor=net['color'], zorder=2))
        ax.text(x+col_w/2, 2.12, net['name'], ha='center', va='center', fontsize=8.5, fontweight='bold', fontproperties=prop)
        ax.plot([x+col_w,x+col_w],[0.5,2.3], color='#BBBBBB', lw=0.5, ls='--', zorder=1)
        for ii, item in enumerate(net['items']):
            ax.add_patch(FancyBboxPatch((x+0.08,row_y[2-ii]-0.28), col_w-0.16, 0.5, boxstyle='round,pad=0.04', lw=0.6, edgecolor='#BBBBBB', facecolor='white', zorder=2))
            ax.text(x+col_w/2, row_y[2-ii]-0.03, item, ha='center', va='center', fontsize=7.5, fontproperties=prop)
    ax.text(text_w/2, 2.62, '三网物理隔离架构', ha='center', va='center', fontsize=10, fontweight='bold', fontproperties=prop)
    out = f'/tmp/arch_net_{hash(tuple(lines))%100000}.png'
    plt.savefig(out, dpi=150, bbox_inches='tight', facecolor='white', edgecolor='none')
    plt.close(); return out

def _make_sw_arch_png(lines, text_w):
    import matplotlib; matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    from matplotlib.patches import FancyBboxPatch
    prop = _get_prop()
    fig, ax = plt.subplots(1,1, figsize=(text_w*0.55, 3.0))
    ax.set_xlim(0, text_w); ax.set_ylim(0, 3.0); ax.axis('off')
    fig.patch.set_facecolor('white')
    outer = FancyBboxPatch((0.3,0.4), text_w-0.6, 2.3, boxstyle='round,pad=0.06', lw=1.2, edgecolor='#444444', facecolor='#F7F7F7', zorder=0)
    ax.add_patch(outer)
    layers = [
        {'name':'应用层','color':'#D9E2F3','items':['人机交互','任务管理','报告生成','系统对接']},
        {'name':'AI算法层','color':'#E2F0D9','items':['目标检测','缺陷识别','数据融合','决策推理']},
        {'name':'数据层','color':'#F2D9D9','items':['实时采集','存储管理','预处理','安全传输']},
        {'name':'基础层','color':'#E2D9F2','items':['操作系统','设备驱动','通信中间件','系统监控']},
    ]
    col_w = (text_w-0.8)/4; row_y = [2.45, 1.7, 0.85]
    for ci, layer in enumerate(layers):
        x = 0.5 + ci*col_w
        ax.add_patch(FancyBboxPatch((x+0.05,2.5), col_w-0.1, 0.4, boxstyle='round,pad=0.03', lw=0.8, edgecolor='#666666', facecolor=layer['color'], zorder=2))
        ax.text(x+col_w/2, 2.7, layer['name'], ha='center', va='center', fontsize=9, fontweight='bold', fontproperties=prop)
        ax.plot([x+col_w,x+col_w],[0.5,2.8], color='#BBBBBB', lw=0.5, ls='--', zorder=1)
        for ii, item in enumerate(layer['items']):
            ax.add_patch(FancyBboxPatch((x+0.05,row_y[2-ii]-0.28), col_w-0.1, 0.48, boxstyle='round,pad=0.04', lw=0.6, edgecolor='#BBBBBB', facecolor='white', zorder=2))
            ax.text(x+col_w/2, row_y[2-ii]-0.04, item, ha='center', va='center', fontsize=7.5, fontproperties=prop)
    ax.text(text_w/2, 2.92, '软件系统总体架构', ha='center', va='center', fontsize=10, fontweight='bold', fontproperties=prop)
    out = f'/tmp/arch_sw_{hash(tuple(lines))%100000}.png'
    plt.savefig(out, dpi=150, bbox_inches='tight', facecolor='white', edgecolor='none')
    plt.close(); return out

def _make_hmi_arch_png(lines, text_w):
    import matplotlib; matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    from matplotlib.patches import FancyBboxPatch
    prop = _get_prop()
    fig, ax = plt.subplots(1,1, figsize=(text_w*0.55, 3.2))
    ax.set_xlim(0, text_w); ax.set_ylim(0, 3.2); ax.axis('off')
    fig.patch.set_facecolor('white')
    outer = FancyBboxPatch((0.3,0.4), text_w-0.6, 2.5, boxstyle='round,pad=0.06', lw=1.2, edgecolor='#444444', facecolor='#F7F7F7', zorder=0)
    ax.add_patch(outer)
    main_box = FancyBboxPatch((0.5,0.5), text_w*0.55, 1.6, boxstyle='round,pad=0.06', lw=1.0, edgecolor='#888888', facecolor='#D9E2F3', zorder=2)
    ax.add_patch(main_box)
    ax.text(text_w*0.30+0.25, 1.75, '主显示区（>=4K）', ha='center', va='center', fontsize=8, fontweight='bold', fontproperties=prop)
    ax.text(text_w*0.30+0.25, 1.55, '实时视频 / 局部放大 / 画中画 / 冻结', ha='center', va='center', fontsize=7, fontproperties=prop)
    left_box = FancyBboxPatch((0.5,2.2), text_w*0.25, 0.55, boxstyle='round,pad=0.04', lw=0.8, edgecolor='#888888', facecolor='#E2F0D9', zorder=2)
    ax.add_patch(left_box)
    ax.text(text_w*0.125+0.25, 2.47, '状态栏\n(电量/里程/模式)', ha='center', va='center', fontsize=7.5, fontproperties=prop)
    right_box = FancyBboxPatch((text_w*0.30+0.55,2.2), text_w*0.25, 0.55, boxstyle='round,pad=0.04', lw=0.8, edgecolor='#888888', facecolor='#F2D9D9', zorder=2)
    ax.add_patch(right_box)
    ax.text(text_w*0.425+0.55, 2.47, '控制栏\n(快捷操作/告警)', ha='center', va='center', fontsize=7.5, fontproperties=prop)
    bottom_box = FancyBboxPatch((0.5,0.5), text_w*0.55, 0.55, boxstyle='round,pad=0.04', lw=0.8, edgecolor='#888888', facecolor='#E2D9F2', zorder=2)
    ax.add_patch(bottom_box)
    ax.text(text_w*0.30+0.25, 0.77, '任务列表 / 检测结果 / 导出操作', ha='center', va='center', fontsize=7.5, fontproperties=prop)
    ax.text(text_w/2, 3.02, 'HMI 界面布局', ha='center', va='center', fontsize=10, fontweight='bold', fontproperties=prop)
    out = f'/tmp/arch_hmi_{hash(tuple(lines))%100000}.png'
    plt.savefig(out, dpi=150, bbox_inches='tight', facecolor='white', edgecolor='none')
    plt.close(); return out

def _make_data_flow_png(lines, text_w):
    import matplotlib; matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    from matplotlib.patches import FancyBboxPatch
    prop = _get_prop()
    fig, ax = plt.subplots(1,1, figsize=(text_w*0.55, 2.6))
    ax.set_xlim(0, text_w); ax.set_ylim(0, 2.6); ax.axis('off')
    fig.patch.set_facecolor('white')
    outer = FancyBboxPatch((0.3,0.4), text_w-0.6, 1.9, boxstyle='round,pad=0.06', lw=1.2, edgecolor='#444444', facecolor='#F7F7F7', zorder=0)
    ax.add_patch(outer)
    nodes = [
        {'name':'传感器数据\n(2D/3D/IMU)','color':'#D9E2F3'},
        {'name':'实时采集\n边缘预处理','color':'#E2F0D9'},
        {'name':'AI缺陷识别\n(<1s响应)','color':'#FFF3CD'},
        {'name':'本地SSD存储\n(>=1TB)','color':'#F2D9D9'},
        {'name':'4G/5G上传\n云端分析','color':'#E2D9F2'},
    ]
    n=len(nodes); node_w=(text_w-0.8)/n
    for ci, node in enumerate(nodes):
        x = 0.5 + ci*node_w + node_w/2; y = 1.45
        ax.add_patch(FancyBboxPatch((x-node_w/2+0.05,y-0.45), node_w-0.1, 0.9, boxstyle='round,pad=0.05', lw=0.8, edgecolor='#888888', facecolor=node['color'], zorder=2))
        ax.text(x, y, node['name'], ha='center', va='center', fontsize=7.5, fontweight='bold', fontproperties=prop)
        if ci < n-1:
            dx = (node_w-0.1)/2+0.03
            ax.annotate('', xy=(x+dx,y), xytext=(x+dx-0.05,y), arrowprops=dict(arrowstyle='->',color='#888888',lw=1.2), zorder=3)
    ax.text(text_w/2, 2.42, '数据采集与实时处理流程', ha='center', va='center', fontsize=10, fontweight='bold', fontproperties=prop)
    out = f'/tmp/arch_df_{hash(tuple(lines))%100000}.png'
    plt.savefig(out, dpi=150, bbox_inches='tight', facecolor='white', edgecolor='none')
    plt.close(); return out

def _make_value_tree_png(lines, text_w):
    import matplotlib; matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    from matplotlib.patches import FancyBboxPatch
    prop = _get_prop()
    fig, ax = plt.subplots(1,1, figsize=(text_w*0.55, 3.2))
    ax.set_xlim(0, text_w); ax.set_ylim(0, 3.2); ax.axis('off')
    fig.patch.set_facecolor('white')
    outer = FancyBboxPatch((0.3,0.4), text_w-0.6, 2.5, boxstyle='round,pad=0.06', lw=1.2, edgecolor='#444444', facecolor='#F7F7F7', zorder=0)
    ax.add_patch(outer)
    values = [
        {'name':'全自动','sub':'一键出发·自动巡航·自动返航','color':'#D9E2F3'},
        {'name':'实时检测','sub':'边缘AI推理·缺陷秒级报警','color':'#E2F0D9'},
        {'name':'全说明','sub':'AI依据可溯源·决策透明可审计','color':'#FFF3CD'},
        {'name':'大数据','sub':'实时汇聚·智能存储·长期积累','color':'#F2D9D9'},
        {'name':'无人干预','sub':'自主巡检·自动生成报告','color':'#E2D9F2'},
    ]
    n=len(values); col_w=(text_w-0.8)/n
    for ci, v in enumerate(values):
        x = 0.5 + ci*col_w
        ax.add_patch(FancyBboxPatch((x+0.05,2.5), col_w-0.1, 0.5, boxstyle='round,pad=0.04', lw=0.8, edgecolor='#666666', facecolor=v['color'], zorder=2))
        ax.text(x+col_w/2, 2.75, v['name'], ha='center', va='center', fontsize=9, fontweight='bold', fontproperties=prop)
        sub_box = FancyBboxPatch((x+0.05,1.1), col_w-0.1, 1.2, boxstyle='round,pad=0.04', lw=0.6, edgecolor='#BBBBBB', facecolor='white', zorder=2)
        ax.add_patch(sub_box)
        ax.text(x+col_w/2, 1.7, v['sub'], ha='center', va='center', fontsize=6.5, fontproperties=prop)
        ax.plot([x+col_w/2,x+col_w/2],[1.1,2.4], color='#CCCCCC', lw=0.8, zorder=1)
    ax.text(text_w/2, 3.02, '六大核心价值', ha='center', va='center', fontsize=10, fontweight='bold', fontproperties=prop)
    out = f'/tmp/arch_val_{hash(tuple(lines))%100000}.png'
    plt.savefig(out, dpi=150, bbox_inches='tight', facecolor='white', edgecolor='none')
    plt.close(); return out

def _make_generic_arch_png(lines, text_w):
    import matplotlib; matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    from matplotlib.patches import FancyBboxPatch
    prop = _get_prop()
    fig, ax = plt.subplots(1,1, figsize=(text_w*0.55, 2.5))
    ax.set_xlim(0, text_w); ax.set_ylim(0, 2.5); ax.axis('off')
    fig.patch.set_facecolor('white')
    box_chars = set('┌┬┐└┴┘├┼─│ ')
    diagram_lines = [l for l in lines if any(c in box_chars for c in l)]
    if not diagram_lines:
        diagram_lines = lines[:8]
    ax.text(text_w/2, 1.3, '\n'.join(diagram_lines[:6]), ha='center', va='center', fontsize=8, fontproperties=prop, family='monospace')
    out = f'/tmp/arch_gen_{hash(tuple(lines))%100000}.png'
    plt.savefig(out, dpi=150, bbox_inches='tight', facecolor='white', edgecolor='none')
    plt.close(); return out


def parse_md(folder, doc):
    files = sorted([f for f in os.listdir(folder) if f.endswith('.md')])
    for fname in files:
        path = os.path.join(folder, fname)
        with open(path, encoding='utf-8') as f:
            content = f.read()
        in_code = False; code_lines = []
        for line in content.split('\n'):
            if line.strip().startswith('```'):
                if not in_code:
                    in_code = True; code_lines = []
                else:
                    in_code = False
                    box_chars = set('┌┬┐└┴┘├┼─│┃ ')
                    if sum(1 for l in code_lines for c in l if c in box_chars) >= 6:
                        try:
                            img_path = make_architecture_matplotlib(code_lines)
                            add_image_to_doc(doc, img_path)
                        except Exception:
                            pass
                continue
            if in_code:
                code_lines.append(line)
                continue
            if line.startswith('# '):
                h1(doc, line[2:].strip())
            elif line.startswith('## '):
                h2(doc, line[3:].strip())
            elif line.startswith('### '):
                h3(doc, line[4:].strip())
            elif line.strip().startswith('|'):
                parts = [p.strip() for p in line.split('|')]
                interior = parts[1:-1]
                is_sep = all(p == '---' for p in interior)
                if len(parts) > 2 and not is_sep:
                    add_table(doc, [p for p in line.split('|') if p.strip()])
            elif line.strip() and not line.strip().startswith('<!--') and not line.strip().startswith('***'):
                body(doc, line.strip())

if __name__ == '__main__':
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument('folder', help='MD文件夹路径')
    parser.add_argument('output', help='输出DOCX路径')
    parser.add_argument('--cover', nargs=4, default=[], help='封面: 标题 副标题 日期 单位')
    args = parser.parse_args()
    doc = Document()
    if args.cover:
        add_cover(doc, *args.cover)
    parse_md(args.folder, doc)
    add_page_numbers(doc)
    doc.save(args.output)
    print(f'Saved: {args.output}')
