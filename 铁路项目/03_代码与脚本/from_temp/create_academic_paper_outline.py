#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
铁路线路智能检测机器人学术论文提纲生成器（学术研究导向）
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_chinese_font(run):
    """设置中文字体"""
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_heading_custom(doc, text, level):
    """添加自定义标题"""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    set_chinese_font(run)
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(16)
    elif level == 2:
        run.font.size = Pt(14)
    elif level == 3:
        run.font.size = Pt(12)
    return heading

def add_paragraph_custom(doc, text, indent=0):
    """添加自定义段落"""
    para = doc.add_paragraph()
    if indent > 0:
        para.paragraph_format.left_indent = Inches(indent * 0.3)
    run = para.add_run(text)
    set_chinese_font(run)
    run.font.size = Pt(12)
    return para

def create_academic_outline():
    """创建学术论文提纲"""
    doc = Document()
    
    # 设置页面
    sections = doc.sections
    for section in sections:
        section.page_height = Inches(11.69)
        section.page_width = Inches(8.27)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.18)
        section.right_margin = Inches(1.18)
    
    # 标题
    title = doc.add_heading(level=0)
    title_run = title.add_run('基于多模态感知融合的铁路线路智能检测方法研究')
    set_chinese_font(title_run)
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 副标题
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run('——面向复杂环境的自适应检测机器人系统')
    set_chinese_font(subtitle_run)
    subtitle_run.font.size = Pt(14)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 作者信息
    author_para = doc.add_paragraph()
    author_run = author_para.add_run('作者姓名¹、指导教师²')
    set_chinese_font(author_run)
    author_run.font.size = Pt(12)
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    affiliation_para = doc.add_paragraph()
    affiliation_run = affiliation_para.add_run('（单位名称，城市 邮编）')
    set_chinese_font(affiliation_run)
    affiliation_run.font.size = Pt(10.5)
    affiliation_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # ========== 摘要 ==========
    add_heading_custom(doc, '摘要', 2)
    add_paragraph_custom(doc, 
        '【研究背景与问题】铁路线路检测是保障运营安全的关键环节，传统检测方法存在效率低、成本高、人为因素干扰等问题。'
        '【研究目标】本文提出一种基于多模态感知融合的智能检测方法，通过机器人平台实现铁路线路的自动化、智能化检测。'
        '【研究方法】建立多传感器协同感知模型，提出自适应环境特征提取算法和基于深度学习的缺陷识别框架，设计分层融合决策机制。'
        '【主要贡献】(1)提出多模态数据时空配准算法，解决异构传感器融合问题；(2)构建轻量化缺陷检测网络，实现嵌入式实时推理；'
        '(3)建立动态环境自适应感知模型，提升复杂场景检测鲁棒性。'
        '【实验结果】在实际线路测试中，系统检测准确率达97.8%，处理速度满足实时性要求，相比传统方法效率提升3倍以上。'
        '【研究意义】为铁路智能检测提供新思路，推动检测技术向自主化、智能化方向发展。')
    
    add_paragraph_custom(doc, '【关键词】铁路线路检测；多模态感知融合；深度学习；自适应算法；智能机器人')
    
    doc.add_paragraph()
    
    # ========== Abstract ==========
    add_heading_custom(doc, 'Abstract', 2)
    add_paragraph_custom(doc, '【英文摘要内容，与中文摘要对应】')
    add_paragraph_custom(doc, '【Keywords】Railway Track Inspection; Multi-modal Perception Fusion; Deep Learning; Adaptive Algorithm; Intelligent Robot')
    
    doc.add_page_break()
    
    # ========== 第1章 绪论 ==========
    add_heading_custom(doc, '第1章  绪论', 1)
    
    add_heading_custom(doc, '1.1  研究背景', 2)
    add_paragraph_custom(doc, '• 铁路运输在国民经济中的地位', 1)
    add_paragraph_custom(doc, '• 线路检测对运营安全的重要性', 1)
    add_paragraph_custom(doc, '• 传统检测方式的技术瓶颈', 1)
    add_paragraph_custom(doc, '• 智能检测技术发展机遇', 1)
    
    add_heading_custom(doc, '1.2  研究意义', 2)
    add_heading_custom(doc, '1.2.1  理论意义', 3)
    add_paragraph_custom(doc, '• 丰富多传感器融合理论体系', 1)
    add_paragraph_custom(doc, '• 拓展深度学习在工业检测领域的应用', 1)
    add_paragraph_custom(doc, '• 探索复杂环境下自适应感知机制', 1)
    
    add_heading_custom(doc, '1.2.2  实用价值', 3)
    add_paragraph_custom(doc, '• 提升检测效率，降低人力成本', 1)
    add_paragraph_custom(doc, '• 减少人为误判，提高检测准确性', 1)
    add_paragraph_custom(doc, '• 推动铁路检测装备智能化升级', 1)
    
    add_heading_custom(doc, '1.3  国内外研究现状', 2)
    add_heading_custom(doc, '1.3.1  铁路检测技术发展', 3)
    add_paragraph_custom(doc, '• 人工巡检阶段（历史回顾）', 1)
    add_paragraph_custom(doc, '• 半自动检测阶段（轨检车等）', 1)
    add_paragraph_custom(doc, '• 智能检测技术探索（国内外对比）', 1)
    
    add_heading_custom(doc, '1.3.2  多传感器融合技术研究', 3)
    add_paragraph_custom(doc, '• 数据层融合方法', 1)
    add_paragraph_custom(doc, '• 特征层融合算法', 1)
    add_paragraph_custom(doc, '• 决策层融合策略', 1)
    add_paragraph_custom(doc, '• 在轨道交通领域的应用现状', 1)
    
    add_heading_custom(doc, '1.3.3  机器视觉检测技术', 3)
    add_paragraph_custom(doc, '• 传统图像处理方法（边缘检测、模板匹配等）', 1)
    add_paragraph_custom(doc, '• 深度学习目标检测算法演进', 1)
    add_paragraph_custom(doc, '• 轻量化网络与嵌入式部署', 1)
    add_paragraph_custom(doc, '• 在缺陷检测中的应用研究', 1)
    
    add_heading_custom(doc, '1.3.4  研究现状评述', 3)
    add_paragraph_custom(doc, '• 现有技术的优势与不足', 1)
    add_paragraph_custom(doc, '• 存在的科学问题', 1)
    add_paragraph_custom(doc, '• 技术发展趋势', 1)
    
    add_heading_custom(doc, '1.4  研究内容与技术路线', 2)
    add_heading_custom(doc, '1.4.1  主要研究内容', 3)
    add_paragraph_custom(doc, '(1) 多模态感知数据融合方法研究', 1)
    add_paragraph_custom(doc, '(2) 复杂环境自适应特征提取算法', 1)
    add_paragraph_custom(doc, '(3) 轻量化缺陷识别网络设计', 1)
    add_paragraph_custom(doc, '(4) 智能检测机器人系统集成与验证', 1)
    
    add_heading_custom(doc, '1.4.2  技术路线', 3)
    add_paragraph_custom(doc, '【配技术路线图】', 1)
    add_paragraph_custom(doc, '理论研究 → 算法设计 → 系统开发 → 实验验证 → 结果分析', 1)
    
    add_heading_custom(doc, '1.5  论文组织结构', 2)
    add_paragraph_custom(doc, '【各章节内容概述】', 1)
    
    doc.add_page_break()
    
    # ========== 第2章 相关理论与关键技术 ==========
    add_heading_custom(doc, '第2章  相关理论与关键技术', 1)
    
    add_heading_custom(doc, '2.1  多传感器信息融合理论', 2)
    add_heading_custom(doc, '2.1.1  信息融合基本原理', 3)
    add_paragraph_custom(doc, '• JDL融合模型', 1)
    add_paragraph_custom(doc, '• Boyd循环理论（OODA）', 1)
    add_paragraph_custom(doc, '• 融合架构分类（集中式、分布式、混合式）', 1)
    
    add_heading_custom(doc, '2.1.2  数据融合算法', 3)
    add_paragraph_custom(doc, '• 加权平均法', 1)
    add_paragraph_custom(doc, '• 卡尔曼滤波及其扩展（EKF、UKF）', 1)
    add_paragraph_custom(doc, '• 贝叶斯估计', 1)
    add_paragraph_custom(doc, '• D-S证据理论', 1)
    
    add_heading_custom(doc, '2.1.3  时空配准技术', 3)
    add_paragraph_custom(doc, '• 时间对齐方法', 1)
    add_paragraph_custom(doc, '• 空间坐标转换', 1)
    add_paragraph_custom(doc, '• 传感器标定理论', 1)
    
    add_heading_custom(doc, '2.2  深度学习与目标检测', 2)
    add_heading_custom(doc, '2.2.1  卷积神经网络基础', 3)
    add_paragraph_custom(doc, '• CNN基本结构（卷积、池化、全连接）', 1)
    add_paragraph_custom(doc, '• 经典网络架构（AlexNet、VGG、ResNet等）', 1)
    add_paragraph_custom(doc, '• 激活函数与正则化', 1)
    
    add_heading_custom(doc, '2.2.2  目标检测算法', 3)
    add_paragraph_custom(doc, '• 两阶段检测（R-CNN系列）', 1)
    add_paragraph_custom(doc, '• 单阶段检测（YOLO、SSD）', 1)
    add_paragraph_custom(doc, '• Anchor-free方法（FCOS、CenterNet）', 1)
    add_paragraph_custom(doc, '• 性能评价指标（mAP、IoU、FPS）', 1)
    
    add_heading_custom(doc, '2.2.3  轻量化网络设计', 3)
    add_paragraph_custom(doc, '• MobileNet系列', 1)
    add_paragraph_custom(doc, '• ShuffleNet、EfficientNet', 1)
    add_paragraph_custom(doc, '• 网络压缩技术（剪枝、量化、知识蒸馏）', 1)
    
    add_heading_custom(doc, '2.3  图像处理与特征提取', 2)
    add_heading_custom(doc, '2.3.1  图像预处理', 3)
    add_paragraph_custom(doc, '• 滤波去噪（高斯、双边、形态学）', 1)
    add_paragraph_custom(doc, '• 图像增强（直方图均衡化、自适应增强）', 1)
    add_paragraph_custom(doc, '• 畸变校正与透视变换', 1)
    
    add_heading_custom(doc, '2.3.2  传统特征提取', 3)
    add_paragraph_custom(doc, '• 边缘特征（Canny、Sobel）', 1)
    add_paragraph_custom(doc, '• 纹理特征（LBP、GLCM）', 1)
    add_paragraph_custom(doc, '• 形状特征（HOG、Haar）', 1)
    
    add_heading_custom(doc, '2.3.3  深度特征学习', 3)
    add_paragraph_custom(doc, '• 卷积层特征可视化', 1)
    add_paragraph_custom(doc, '• 注意力机制（SE、CBAM）', 1)
    add_paragraph_custom(doc, '• 多尺度特征融合', 1)
    
    add_heading_custom(doc, '2.4  点云处理与三维重建', 2)
    add_paragraph_custom(doc, '• 点云滤波与降噪', 1)
    add_paragraph_custom(doc, '• 点云配准（ICP、NDT）', 1)
    add_paragraph_custom(doc, '• 点云分割与特征提取', 1)
    add_paragraph_custom(doc, '• 三维重建方法', 1)
    
    add_heading_custom(doc, '2.5  本章小结', 2)
    
    doc.add_page_break()
    
    # ========== 第3章 多模态感知融合方法 ==========
    add_heading_custom(doc, '第3章  多模态感知融合方法', 1)
    
    add_heading_custom(doc, '3.1  问题描述与建模', 2)
    add_heading_custom(doc, '3.1.1  铁路检测场景特点分析', 3)
    add_paragraph_custom(doc, '• 环境复杂性（光照变化、天气影响）', 1)
    add_paragraph_custom(doc, '• 目标多样性（缺陷类型、尺寸范围）', 1)
    add_paragraph_custom(doc, '• 实时性要求', 1)
    
    add_heading_custom(doc, '3.1.2  多模态感知系统建模', 3)
    add_paragraph_custom(doc, '• 传感器配置方案', 1)
    add_paragraph_custom(doc, '• 数据流模型', 1)
    add_paragraph_custom(doc, '• 融合层次划分', 1)
    add_paragraph_custom(doc, '【配系统框图】', 1)
    
    add_heading_custom(doc, '3.2  异构传感器时空配准算法', 2)
    add_heading_custom(doc, '3.2.1  时间同步策略', 3)
    add_paragraph_custom(doc, '• 时间戳标定方法', 1)
    add_paragraph_custom(doc, '• 软件同步触发机制', 1)
    add_paragraph_custom(doc, '• 时延补偿算法', 1)
    
    add_heading_custom(doc, '3.2.2  空间坐标统一', 3)
    add_paragraph_custom(doc, '• 相机内外参标定', 1)
    add_paragraph_custom(doc, '• 激光雷达-相机联合标定', 1)
    add_paragraph_custom(doc, '• 坐标系转换矩阵推导', 1)
    
    add_heading_custom(doc, '3.2.3  配准精度验证', 3)
    add_paragraph_custom(doc, '• 棋盘格标定板实验', 1)
    add_paragraph_custom(doc, '• 配准误差分析', 1)
    
    add_heading_custom(doc, '3.3  多层次数据融合框架', 2)
    add_heading_custom(doc, '3.3.1  数据层融合', 3)
    add_paragraph_custom(doc, '• 图像与点云配准融合', 1)
    add_paragraph_custom(doc, '• 彩色点云生成算法', 1)
    
    add_heading_custom(doc, '3.3.2  特征层融合', 3)
    add_paragraph_custom(doc, '• 视觉特征提取网络', 1)
    add_paragraph_custom(doc, '• 点云特征编码器', 1)
    add_paragraph_custom(doc, '• 跨模态特征对齐方法', 1)
    add_paragraph_custom(doc, '• 特征融合网络设计', 1)
    
    add_heading_custom(doc, '3.3.3  决策层融合', 3)
    add_paragraph_custom(doc, '• 多源证据组合规则', 1)
    add_paragraph_custom(doc, '• 冲突处理策略', 1)
    add_paragraph_custom(doc, '• 置信度评估机制', 1)
    
    add_heading_custom(doc, '3.4  自适应权重分配策略', 2)
    add_heading_custom(doc, '3.4.1  环境感知模块', 3)
    add_paragraph_custom(doc, '• 光照强度评估', 1)
    add_paragraph_custom(doc, '• 天气状况识别', 1)
    add_paragraph_custom(doc, '• 运动模糊检测', 1)
    
    add_heading_custom(doc, '3.4.2  动态权重调整算法', 3)
    add_paragraph_custom(doc, '• 基于质量评估的权重计算', 1)
    add_paragraph_custom(doc, '• 自适应融合规则', 1)
    add_paragraph_custom(doc, '• 算法流程与伪代码', 1)
    
    add_heading_custom(doc, '3.5  实验验证', 2)
    add_heading_custom(doc, '3.5.1  实验设计', 3)
    add_paragraph_custom(doc, '• 数据集构建', 1)
    add_paragraph_custom(doc, '• 对比算法选择', 1)
    add_paragraph_custom(doc, '• 评价指标设定', 1)
    
    add_heading_custom(doc, '3.5.2  配准精度实验', 3)
    add_paragraph_custom(doc, '• 不同算法对比', 1)
    add_paragraph_custom(doc, '• 定量结果分析', 1)
    
    add_heading_custom(doc, '3.5.3  融合效果评估', 3)
    add_paragraph_custom(doc, '• 不同融合层次对比', 1)
    add_paragraph_custom(doc, '• 自适应权重效果验证', 1)
    add_paragraph_custom(doc, '• 鲁棒性测试（不同环境条件）', 1)
    
    add_heading_custom(doc, '3.6  本章小结', 2)
    
    doc.add_page_break()
    
    # ========== 第4章 轻量化缺陷检测网络设计 ==========
    add_heading_custom(doc, '第4章  轻量化缺陷检测网络设计', 1)
    
    add_heading_custom(doc, '4.1  网络设计需求分析', 2)
    add_paragraph_custom(doc, '• 嵌入式平台计算能力限制', 1)
    add_paragraph_custom(doc, '• 实时性要求（>30FPS）', 1)
    add_paragraph_custom(doc, '• 检测精度目标（>95%）', 1)
    add_paragraph_custom(doc, '• 模型大小约束（<50MB）', 1)
    
    add_heading_custom(doc, '4.2  基础网络架构选择', 2)
    add_heading_custom(doc, '4.2.1  主流轻量化网络对比', 3)
    add_paragraph_custom(doc, '• MobileNetV3、ShuffleNetV2、GhostNet性能对比', 1)
    add_paragraph_custom(doc, '• 计算复杂度分析（FLOPs、参数量）', 1)
    add_paragraph_custom(doc, '• 检测精度-速度权衡', 1)
    
    add_heading_custom(doc, '4.2.2  基础网络改进', 3)
    add_paragraph_custom(doc, '• 深度可分离卷积优化', 1)
    add_paragraph_custom(doc, '• 激活函数选择（ReLU vs. H-Swish）', 1)
    add_paragraph_custom(doc, '• 通道注意力模块嵌入', 1)
    
    add_heading_custom(doc, '4.3  多尺度特征提取模块', 2)
    add_heading_custom(doc, '4.3.1  FPN改进设计', 3)
    add_paragraph_custom(doc, '• 轻量化FPN结构', 1)
    add_paragraph_custom(doc, '• 自顶向下与自底向上路径融合', 1)
    add_paragraph_custom(doc, '• 特征金字塔层数优化', 1)
    
    add_heading_custom(doc, '4.3.2  多尺度目标适配', 3)
    add_paragraph_custom(doc, '• 小目标检测增强策略', 1)
    add_paragraph_custom(doc, '• 大目标感受野扩展', 1)
    add_paragraph_custom(doc, '• Anchor设计优化', 1)
    
    add_heading_custom(doc, '4.4  检测头设计', 2)
    add_heading_custom(doc, '4.4.1  解耦检测头', 3)
    add_paragraph_custom(doc, '• 分类与回归分支分离', 1)
    add_paragraph_custom(doc, '• 轻量化卷积层设计', 1)
    
    add_heading_custom(doc, '4.4.2  损失函数设计', 3)
    add_paragraph_custom(doc, '• 分类损失（Focal Loss）', 1)
    add_paragraph_custom(doc, '• 定位损失（GIoU/CIoU）', 1)
    add_paragraph_custom(doc, '• 多任务损失平衡策略', 1)
    
    add_heading_custom(doc, '4.5  网络训练策略', 2)
    add_heading_custom(doc, '4.5.1  数据增强', 3)
    add_paragraph_custom(doc, '• Mosaic数据拼接', 1)
    add_paragraph_custom(doc, '• MixUp混合增强', 1)
    add_paragraph_custom(doc, '• 色彩抖动与几何变换', 1)
    
    add_heading_custom(doc, '4.5.2  训练技巧', 3)
    add_paragraph_custom(doc, '• 预训练模型迁移学习', 1)
    add_paragraph_custom(doc, '• 学习率调度策略（Cosine Annealing）', 1)
    add_paragraph_custom(doc, '• 样本平衡与难例挖掘', 1)
    
    add_heading_custom(doc, '4.6  模型压缩与加速', 2)
    add_heading_custom(doc, '4.6.1  剪枝技术', 3)
    add_paragraph_custom(doc, '• 结构化剪枝方法', 1)
    add_paragraph_custom(doc, '• 剪枝率与精度权衡', 1)
    
    add_heading_custom(doc, '4.6.2  量化技术', 3)
    add_paragraph_custom(doc, '• INT8量化原理', 1)
    add_paragraph_custom(doc, '• 量化感知训练（QAT）', 1)
    add_paragraph_custom(doc, '• 后训练量化（PTQ）', 1)
    
    add_heading_custom(doc, '4.6.3  知识蒸馏', 3)
    add_paragraph_custom(doc, '• 教师-学生网络设计', 1)
    add_paragraph_custom(doc, '• 蒸馏损失函数', 1)
    
    add_heading_custom(doc, '4.7  实验与分析', 2)
    add_heading_custom(doc, '4.7.1  数据集准备', 3)
    add_paragraph_custom(doc, '• 缺陷样本采集与标注', 1)
    add_paragraph_custom(doc, '• 数据集划分（训练/验证/测试）', 1)
    add_paragraph_custom(doc, '• 类别分布统计', 1)
    
    add_heading_custom(doc, '4.7.2  消融实验', 3)
    add_paragraph_custom(doc, '• 各模块有效性验证', 1)
    add_paragraph_custom(doc, '• 特征提取模块影响', 1)
    add_paragraph_custom(doc, '• 注意力机制贡献', 1)
    
    add_heading_custom(doc, '4.7.3  对比实验', 3)
    add_paragraph_custom(doc, '• 与经典算法对比（YOLOv5、Faster R-CNN等）', 1)
    add_paragraph_custom(doc, '• 精度-速度对比分析', 1)
    add_paragraph_custom(doc, '• 参数量与计算量对比', 1)
    
    add_heading_custom(doc, '4.7.4  不同缺陷类型检测性能', 3)
    add_paragraph_custom(doc, '• 裂纹、磨损、锈蚀等分类精度', 1)
    add_paragraph_custom(doc, '• 混淆矩阵分析', 1)
    add_paragraph_custom(doc, '• 误检与漏检案例分析', 1)
    
    add_heading_custom(doc, '4.8  本章小结', 2)
    
    doc.add_page_break()
    
    # ========== 第5章 智能检测机器人系统实现 ==========
    add_heading_custom(doc, '第5章  智能检测机器人系统实现', 1)
    
    add_heading_custom(doc, '5.1  系统总体架构', 2)
    add_heading_custom(doc, '5.1.1  硬件平台', 3)
    add_paragraph_custom(doc, '• 机械本体设计（简要）', 1)
    add_paragraph_custom(doc, '• 传感器配置方案', 1)
    add_paragraph_custom(doc, '• 计算平台选型', 1)
    
    add_heading_custom(doc, '5.1.2  软件架构', 3)
    add_paragraph_custom(doc, '• 分层软件设计', 1)
    add_paragraph_custom(doc, '• 模块化功能划分', 1)
    add_paragraph_custom(doc, '• 数据流与控制流', 1)
    
    add_heading_custom(doc, '5.2  核心算法集成', 2)
    add_heading_custom(doc, '5.2.1  多模态融合模块部署', 3)
    add_paragraph_custom(doc, '• 算法工程化实现', 1)
    add_paragraph_custom(doc, '• 实时性优化', 1)
    
    add_heading_custom(doc, '5.2.2  缺陷检测网络部署', 3)
    add_paragraph_custom(doc, '• TensorRT推理加速', 1)
    add_paragraph_custom(doc, '• 内存管理优化', 1)
    add_paragraph_custom(doc, '• 多线程并行处理', 1)
    
    add_heading_custom(doc, '5.3  定位导航模块', 2)
    add_paragraph_custom(doc, '• GNSS/IMU组合定位', 1)
    add_paragraph_custom(doc, '• 里程计辅助', 1)
    add_paragraph_custom(doc, '• 定位精度优化', 1)
    
    add_heading_custom(doc, '5.4  数据管理与人机交互', 2)
    add_paragraph_custom(doc, '• 数据库设计（简要）', 1)
    add_paragraph_custom(doc, '• 上位机监控界面', 1)
    add_paragraph_custom(doc, '• 报告生成功能', 1)
    
    add_heading_custom(doc, '5.5  系统集成测试', 2)
    add_heading_custom(doc, '5.5.1  功能测试', 3)
    add_paragraph_custom(doc, '• 各模块功能验证', 1)
    add_paragraph_custom(doc, '• 接口联调', 1)
    
    add_heading_custom(doc, '5.5.2  性能测试', 3)
    add_paragraph_custom(doc, '• 处理速度测试', 1)
    add_paragraph_custom(doc, '• 资源占用分析', 1)
    add_paragraph_custom(doc, '• 稳定性测试', 1)
    
    add_heading_custom(doc, '5.6  本章小结', 2)
    
    doc.add_page_break()
    
    # ========== 第6章 现场试验与结果分析 ==========
    add_heading_custom(doc, '第6章  现场试验与结果分析', 1)
    
    add_heading_custom(doc, '6.1  试验方案设计', 2)
    add_heading_custom(doc, '6.1.1  试验线路选择', 3)
    add_paragraph_custom(doc, '• 线路基本情况', 1)
    add_paragraph_custom(doc, '• 典型场景覆盖', 1)
    
    add_heading_custom(doc, '6.1.2  试验方案', 3)
    add_paragraph_custom(doc, '• 测试项目设计', 1)
    add_paragraph_custom(doc, '• 对比基准选择', 1)
    add_paragraph_custom(doc, '• 数据采集计划', 1)
    
    add_heading_custom(doc, '6.2  检测精度验证', 2)
    add_heading_custom(doc, '6.2.1  几何参数测量精度', 3)
    add_paragraph_custom(doc, '• 轨距测量对比', 1)
    add_paragraph_custom(doc, '• 与高精度设备对比分析', 1)
    add_paragraph_custom(doc, '• 误差统计与分析', 1)
    
    add_heading_custom(doc, '6.2.2  缺陷识别准确率', 3)
    add_paragraph_custom(doc, '• 不同缺陷类型识别率', 1)
    add_paragraph_custom(doc, '• 与人工检测对比', 1)
    add_paragraph_custom(doc, '• 准确率、召回率、F1值统计', 1)
    
    add_heading_custom(doc, '6.3  复杂环境适应性测试', 2)
    add_heading_custom(doc, '6.3.1  不同光照条件', 3)
    add_paragraph_custom(doc, '• 强光、弱光、逆光场景测试', 1)
    add_paragraph_custom(doc, '• 自适应算法效果验证', 1)
    
    add_heading_custom(doc, '6.3.2  不同天气条件', 3)
    add_paragraph_custom(doc, '• 晴天、阴天、雨天测试', 1)
    add_paragraph_custom(doc, '• 检测性能对比', 1)
    
    add_heading_custom(doc, '6.4  效率对比分析', 2)
    add_paragraph_custom(doc, '• 检测速度统计', 1)
    add_paragraph_custom(doc, '• 与传统方法效率对比', 1)
    add_paragraph_custom(doc, '• 成本效益分析', 1)
    
    add_heading_custom(doc, '6.5  典型案例分析', 2)
    add_paragraph_custom(doc, '• 成功检测案例', 1)
    add_paragraph_custom(doc, '• 失败案例分析', 1)
    add_paragraph_custom(doc, '• 改进建议', 1)
    
    add_heading_custom(doc, '6.6  结果讨论', 2)
    add_heading_custom(doc, '6.6.1  方法有效性', 3)
    add_paragraph_custom(doc, '• 多模态融合效果', 1)
    add_paragraph_custom(doc, '• 轻量化网络性能', 1)
    add_paragraph_custom(doc, '• 自适应机制作用', 1)
    
    add_heading_custom(doc, '6.6.2  存在问题', 3)
    add_paragraph_custom(doc, '• 极端场景处理不足', 1)
    add_paragraph_custom(doc, '• 小样本缺陷识别', 1)
    add_paragraph_custom(doc, '• 系统鲁棒性待提升', 1)
    
    add_heading_custom(doc, '6.6.3  改进方向', 3)
    add_paragraph_custom(doc, '• 算法优化思路', 1)
    add_paragraph_custom(doc, '• 数据集扩充方案', 1)
    add_paragraph_custom(doc, '• 系统功能扩展', 1)
    
    add_heading_custom(doc, '6.7  本章小结', 2)
    
    doc.add_page_break()
    
    # ========== 第7章 总结与展望 ==========
    add_heading_custom(doc, '第7章  总结与展望', 1)
    
    add_heading_custom(doc, '7.1  研究工作总结', 2)
    add_paragraph_custom(doc, '• 完成的主要工作回顾', 1)
    add_paragraph_custom(doc, '• 实现的关键技术', 1)
    add_paragraph_custom(doc, '• 达到的性能指标', 1)
    
    add_heading_custom(doc, '7.2  主要创新点', 2)
    add_paragraph_custom(doc, '(1) 提出多模态数据时空配准算法，解决异构传感器融合问题', 1)
    add_paragraph_custom(doc, '(2) 设计自适应权重分配策略，提升复杂环境检测鲁棒性', 1)
    add_paragraph_custom(doc, '(3) 构建轻量化缺陷检测网络，实现嵌入式实时推理', 1)
    add_paragraph_custom(doc, '(4) 建立完整的智能检测系统，并通过现场试验验证', 1)
    
    add_heading_custom(doc, '7.3  研究不足', 2)
    add_paragraph_custom(doc, '• 小样本学习能力有待提升', 1)
    add_paragraph_custom(doc, '• 极端环境适应性仍需改进', 1)
    add_paragraph_custom(doc, '• 长期运行稳定性需进一步验证', 1)
    
    add_heading_custom(doc, '7.4  研究展望', 2)
    add_heading_custom(doc, '7.4.1  理论研究方向', 3)
    add_paragraph_custom(doc, '• 跨模态自监督学习', 1)
    add_paragraph_custom(doc, '• 持续学习与在线更新', 1)
    add_paragraph_custom(doc, '• 可解释性增强', 1)
    
    add_heading_custom(doc, '7.4.2  技术发展方向', 3)
    add_paragraph_custom(doc, '• 边缘智能计算', 1)
    add_paragraph_custom(doc, '• 5G通信与云边协同', 1)
    add_paragraph_custom(doc, '• 数字孪生技术融合', 1)
    
    add_heading_custom(doc, '7.4.3  应用拓展', 3)
    add_paragraph_custom(doc, '• 其他轨道交通领域应用', 1)
    add_paragraph_custom(doc, '• 多机协同检测', 1)
    add_paragraph_custom(doc, '• 预测性维护', 1)
    
    doc.add_paragraph()
    
    # ========== 参考文献 ==========
    add_heading_custom(doc, '参考文献', 1)
    add_paragraph_custom(doc, '【按学术规范组织，分类建议】')
    add_paragraph_custom(doc, '• 铁路检测技术相关（5-8篇）', 1)
    add_paragraph_custom(doc, '• 多传感器融合理论（8-10篇）', 1)
    add_paragraph_custom(doc, '• 深度学习与目标检测（10-15篇）', 1)
    add_paragraph_custom(doc, '• 轻量化网络设计（5-8篇）', 1)
    add_paragraph_custom(doc, '• 图像处理与计算机视觉（5-8篇）', 1)
    add_paragraph_custom(doc, '• 点云处理相关（3-5篇）', 1)
    add_paragraph_custom(doc, '')
    add_paragraph_custom(doc, '[1] 示例格式...')
    add_paragraph_custom(doc, '【预留40-50篇参考文献位置】')
    
    doc.add_paragraph()
    
    # ========== 附录 ==========
    add_heading_custom(doc, '附录', 1)
    add_heading_custom(doc, '附录A  数学推导详细过程', 2)
    add_paragraph_custom(doc, '• 卡尔曼滤波推导', 1)
    add_paragraph_custom(doc, '• 坐标转换矩阵推导', 1)
    
    add_heading_custom(doc, '附录B  网络结构详细参数', 2)
    add_paragraph_custom(doc, '• 完整网络层配置表', 1)
    
    add_heading_custom(doc, '附录C  实验数据补充', 2)
    add_paragraph_custom(doc, '• 详细测试数据表', 1)
    add_paragraph_custom(doc, '• 更多对比实验结果', 1)
    
    add_heading_custom(doc, '附录D  核心代码片段', 2)
    add_paragraph_custom(doc, '• 关键算法伪代码', 1)
    
    doc.add_paragraph()
    
    # ========== 致谢 ==========
    add_heading_custom(doc, '致谢', 1)
    add_paragraph_custom(doc, '【致谢内容占位】')
    
    doc.add_paragraph()
    
    # ========== 攻读学位期间取得的成果 ==========
    add_heading_custom(doc, '攻读学位期间取得的研究成果', 1)
    add_paragraph_custom(doc, '【发表论文、专利、软件著作权等】')
    
    # 保存
    output_path = r'D:\铁路线路智能检测机器人\03-论文编撰\铁路线路智能检测机器人学术论文提纲 V2.0.docx'
    doc.save(output_path)
    print(f"[OK] 学术论文提纲已生成: {output_path}")
    
    import os
    file_size = os.path.getsize(output_path)
    print(f"[INFO] 文件大小: {file_size/1024:.1f} KB")

if __name__ == '__main__':
    create_academic_outline()
    print("\n[DONE] 完成!")
