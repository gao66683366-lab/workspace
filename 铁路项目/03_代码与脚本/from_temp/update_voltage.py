# -*- coding: utf-8 -*-
from docx import Document

# 打开现有文档
doc = Document('D:\\铁路线路智能检测机器人\\04-项目文档\\设计文档\\机器人硬件系统架构设计文档 V1.0.docx')

# 查找"3.5.3 电源系统"章节并更新
found = False
for i, para in enumerate(doc.paragraphs):
    if '3.5.3 电源系统' in para.text:
        found = True
        # 找到后续段落并更新
        for j in range(i+1, min(i+10, len(doc.paragraphs))):
            if '输入电压' in doc.paragraphs[j].text:
                # 更新这一段
                doc.paragraphs[j].text = '  • 输入电压：DC 48V（电池供电）'
                break
            if doc.paragraphs[j].style.name.startswith('Heading'):
                break
        break

if found:
    # 同时更新技术指标表中的相关内容
    for table in doc.tables:
        for row in table.rows:
            cells_text = [cell.text for cell in row.cells]
            if '输入电压' in ' '.join(cells_text):
                for cell in row.cells:
                    if '24V' in cell.text or 'AC 220V' in cell.text:
                        cell.text = 'DC 48V'
    
    doc.save('D:\\铁路线路智能检测机器人\\04-项目文档\\设计文档\\机器人硬件系统架构设计文档 V1.0.docx')
    print('[OK] 已更新供电电压为 DC 48V')
else:
    print('[WARN] 未找到电源系统章节')
