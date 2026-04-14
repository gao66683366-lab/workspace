# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

doc = Document('D:\\铁路线路智能检测机器人\\04-项目文档\\设计文档\\机器人硬件系统架构设计文档 V1.0.docx')

# 查找"3.5.1 工控机配置"章节
for i, para in enumerate(doc.paragraphs):
    if '3.5.1 工控机配置' in para.text:
        # 删除旧内容
        delete_count = 0
        for j in range(i+1, len(doc.paragraphs)):
            if doc.paragraphs[j].style.name.startswith('Heading'):
                break
            delete_count += 1
        
        for j in range(delete_count):
            p = doc.paragraphs[i + delete_count - j]
            p._element.getparent().remove(p._element)
        
        # 插入新的详细配置
        insert_pos = i + 1
        
        # 负载分析
        p1 = doc.add_paragraph('一、负载分析与需求计算', style='List Bullet')
        p1.runs[0].bold = True
        p1._element.getparent().insert(insert_pos, p1._element)
        insert_pos += 1
        
        p2 = doc.add_paragraph('1. 数据吞吐量计算', style='List Bullet 2')
        p2._element.getparent().insert(insert_pos, p2._element)
        insert_pos += 1
        
        p3 = doc.add_paragraph('工业相机×6：假设5MP分辨率，8位灰度，30fps', style='List Bullet 3')
        p3._element.getparent().insert(insert_pos, p3._element)
        insert_pos += 1
        
        p4 = doc.add_paragraph('  单相机数据率 = 5M × 1字节 × 30fps = 150MB/s', style='List Bullet 3')
        p4._element.getparent().insert(insert_pos, p4._element)
        insert_pos += 1
        
        p5 = doc.add_paragraph('  6个相机总数据率 = 900MB/s ≈ 7.2Gbps', style='List Bullet 3')
        p5._element.getparent().insert(insert_pos, p5._element)
        insert_pos += 1
        
        p6 = doc.add_paragraph('3D线激光×2：假设2kHz扫描频率，每线1024点', style='List Bullet 3')
        p6._element.getparent().insert(insert_pos, p6._element)
        insert_pos += 1
        
        p7 = doc.add_paragraph('  单激光数据率 = 1024点 × 3坐标 × 4字节 × 2000Hz = 24.6MB/s', style='List Bullet 3')
        p7._element.getparent().insert(insert_pos, p7._element)
        insert_pos += 1
        
        p8 = doc.add_paragraph('  2个激光总数据率 = 49.2MB/s ≈ 394Mbps', style='List Bullet 3')
        p8._element.getparent().insert(insert_pos, p8._element)
        insert_pos += 1
        
        p9 = doc.add_paragraph('总网络带宽需求：≈8Gbps（需要千兆交换机 + 多网卡）', style='List Bullet 3')
        p9.runs[0].bold = True
        p9._element.getparent().insert(insert_pos, p9._element)
        insert_pos += 1
        
        p10 = doc.add_paragraph('2. 存储容量需求（按8小时工作计算）', style='List Bullet 2')
        p10._element.getparent().insert(insert_pos, p10._element)
        insert_pos += 1
        
        p11 = doc.add_paragraph('相机数据：900MB/s × 3600s/h × 8h = 25.9TB', style='List Bullet 3')
        p11._element.getparent().insert(insert_pos, p11._element)
        insert_pos += 1
        
        p12 = doc.add_paragraph('3D数据：49.2MB/s × 3600s/h × 8h = 1.4TB', style='List Bullet 3')
        p12._element.getparent().insert(insert_pos, p12._element)
        insert_pos += 1
        
        p13 = doc.add_paragraph('实际存储：采用压缩（JPEG压缩率约10:1）+ 间隔采样', style='List Bullet 3')
        p13._element.getparent().insert(insert_pos, p13._element)
        insert_pos += 1
        
        p14 = doc.add_paragraph('预估每日数据量：2-5TB（需配置大容量存储）', style='List Bullet 3')
        p14.runs[0].bold = True
        p14._element.getparent().insert(insert_pos, p14._element)
        insert_pos += 1
        
        p15 = doc.add_paragraph('3. 计算性能需求', style='List Bullet 2')
        p15._element.getparent().insert(insert_pos, p15._element)
        insert_pos += 1
        
        p16 = doc.add_paragraph('实时图像处理：6路图像预处理（去噪、增强）', style='List Bullet 3')
        p16._element.getparent().insert(insert_pos, p16._element)
        insert_pos += 1
        
        p17 = doc.add_paragraph('AI推理：螺栓检测 + 轨面缺陷检测，推理时间要求<100ms', style='List Bullet 3')
        p17._element.getparent().insert(insert_pos, p17._element)
        insert_pos += 1
        
        p18 = doc.add_paragraph('3D点云处理：实时轮廓提取与计算', style='List Bullet 3')
        p18._element.getparent().insert(insert_pos, p18._element)
        insert_pos += 1
        
        p19 = doc.add_paragraph('需求：高性能CPU + GPU加速（推荐）', style='List Bullet 3')
        p19.runs[0].bold = True
        p19._element.getparent().insert(insert_pos, p19._element)
        insert_pos += 1
        
        # 配置方案
        p20 = doc.add_paragraph('', style='Normal')
        p20._element.getparent().insert(insert_pos, p20._element)
        insert_pos += 1
        
        p21 = doc.add_paragraph('二、工控机硬件配置方案', style='List Bullet')
        p21.runs[0].bold = True
        p21._element.getparent().insert(insert_pos, p21._element)
        insert_pos += 1
        
        # CPU
        p22 = doc.add_paragraph('1. 处理器（CPU）', style='List Bullet 2')
        p22._element.getparent().insert(insert_pos, p22._element)
        insert_pos += 1
        
        p23 = doc.add_paragraph('型号：Intel Core i7-12700 或更高（12代/13代/14代酷睿）', style='List Bullet 3')
        p23._element.getparent().insert(insert_pos, p23._element)
        insert_pos += 1
        
        p24 = doc.add_paragraph('核心数：12核心20线程（8性能核+4能效核）', style='List Bullet 3')
        p24._element.getparent().insert(insert_pos, p24._element)
        insert_pos += 1
        
        p25 = doc.add_paragraph('主频：基频2.1GHz，睿频最高4.9GHz', style='List Bullet 3')
        p25._element.getparent().insert(insert_pos, p25._element)
        insert_pos += 1
        
        p26 = doc.add_paragraph('缓存：25MB L3缓存', style='List Bullet 3')
        p26._element.getparent().insert(insert_pos, p26._element)
        insert_pos += 1
        
        p27 = doc.add_paragraph('TDP：65W（标准版）或 125W（性能版）', style='List Bullet 3')
        p27._element.getparent().insert(insert_pos, p27._element)
        insert_pos += 1
        
        p28 = doc.add_paragraph('备选：AMD Ryzen 7 7700（8核16线程）', style='List Bullet 3')
        p28._element.getparent().insert(insert_pos, p28._element)
        insert_pos += 1
        
        # 内存
        p29 = doc.add_paragraph('2. 内存（RAM）', style='List Bullet 2')
        p29._element.getparent().insert(insert_pos, p29._element)
        insert_pos += 1
        
        p30 = doc.add_paragraph('容量：32GB DDR4-3200 或 DDR5-4800', style='List Bullet 3')
        p30._element.getparent().insert(insert_pos, p30._element)
        insert_pos += 1
        
        p31 = doc.add_paragraph('配置：2×16GB 双通道（可扩展至64GB）', style='List Bullet 3')
        p31._element.getparent().insert(insert_pos, p31._element)
        insert_pos += 1
        
        p32 = doc.add_paragraph('类型：ECC内存（可选，提高可靠性）', style='List Bullet 3')
        p32._element.getparent().insert(insert_pos, p32._element)
        insert_pos += 1
        
        p33 = doc.add_paragraph('理由：图像缓存 + AI推理 + 系统运行需要大内存', style='List Bullet 3')
        p33._element.getparent().insert(insert_pos, p33._element)
        insert_pos += 1
        
        # 存储
        p34 = doc.add_paragraph('3. 存储系统', style='List Bullet 2')
        p34._element.getparent().insert(insert_pos, p34._element)
        insert_pos += 1
        
        p35 = doc.add_paragraph('系统盘：NVMe SSD 512GB（M.2接口，PCIe 4.0）', style='List Bullet 3')
        p35._element.getparent().insert(insert_pos, p35._element)
        insert_pos += 1
        
        p36 = doc.add_paragraph('  - 用途：操作系统 + 应用软件 + AI模型', style='List Bullet 3')
        p36._element.getparent().insert(insert_pos, p36._element)
        insert_pos += 1
        
        p37 = doc.add_paragraph('  - 性能要求：顺序读写≥3500MB/s', style='List Bullet 3')
        p37._element.getparent().insert(insert_pos, p37._element)
        insert_pos += 1
        
        p38 = doc.add_paragraph('数据盘：NVMe SSD 2TB×2（RAID 0阵列，提高写入速度）', style='List Bullet 3')
        p38._element.getparent().insert(insert_pos, p38._element)
        insert_pos += 1
        
        p39 = doc.add_paragraph('  - 用途：实时数据存储（图像、点云、检测结果）', style='List Bullet 3')
        p39._element.getparent().insert(insert_pos, p39._element)
        insert_pos += 1
        
        p40 = doc.add_paragraph('  - 阵列性能：写入速度≥6000MB/s', style='List Bullet 3')
        p40._element.getparent().insert(insert_pos, p40._element)
        insert_pos += 1
        
        p41 = doc.add_paragraph('备份盘：SATA SSD 4TB（数据镜像备份）', style='List Bullet 3')
        p41._element.getparent().insert(insert_pos, p41._element)
        insert_pos += 1
        
        p42 = doc.add_paragraph('总存储容量：6.5TB（512GB + 4TB + 4TB备份）', style='List Bullet 3')
        p42.runs[0].bold = True
        p42._element.getparent().insert(insert_pos, p42._element)
        insert_pos += 1
        
        # GPU
        p43 = doc.add_paragraph('4. 显卡（GPU）- AI推理加速', style='List Bullet 2')
        p43._element.getparent().insert(insert_pos, p43._element)
        insert_pos += 1
        
        p44 = doc.add_paragraph('推荐配置：NVIDIA RTX A2000（专业卡）或 RTX 4060（消费级）', style='List Bullet 3')
        p44._element.getparent().insert(insert_pos, p44._element)
        insert_pos += 1
        
        p45 = doc.add_paragraph('显存：≥8GB GDDR6', style='List Bullet 3')
        p45._element.getparent().insert(insert_pos, p45._element)
        insert_pos += 1
        
        p46 = doc.add_paragraph('CUDA核心：≥3000个', style='List Bullet 3')
        p46._element.getparent().insert(insert_pos, p46._element)
        insert_pos += 1
        
        p47 = doc.add_paragraph('功耗：≤130W', style='List Bullet 3')
        p47._element.getparent().insert(insert_pos, p47._element)
        insert_pos += 1
        
        p48 = doc.add_paragraph('用途：ONNX Runtime / TensorRT推理加速，AI推理速度提升5-10倍', style='List Bullet 3')
        p48._element.getparent().insert(insert_pos, p48._element)
        insert_pos += 1
        
        p49 = doc.add_paragraph('备选：仅CPU推理（成本更低，但推理速度较慢）', style='List Bullet 3')
        p49._element.getparent().insert(insert_pos, p49._element)
        insert_pos += 1
        
        # 主板
        p50 = doc.add_paragraph('5. 主板', style='List Bullet 2')
        p50._element.getparent().insert(insert_pos, p50._element)
        insert_pos += 1
        
        p51 = doc.add_paragraph('芯片组：Intel B760 / Z790 或 AMD B650', style='List Bullet 3')
        p51._element.getparent().insert(insert_pos, p51._element)
        insert_pos += 1
        
        p52 = doc.add_paragraph('板型：ATX或M-ATX（根据工控机箱尺寸）', style='List Bullet 3')
        p52._element.getparent().insert(insert_pos, p52._element)
        insert_pos += 1
        
        p53 = doc.add_paragraph('网口：板载2个Intel千兆网口（I219-V或I225-V）', style='List Bullet 3')
        p53._element.getparent().insert(insert_pos, p53._element)
        insert_pos += 1
        
        p54 = doc.add_paragraph('PCIe插槽：', style='List Bullet 3')
        p54._element.getparent().insert(insert_pos, p54._element)
        insert_pos += 1
        
        p55 = doc.add_paragraph('  - 1×PCIe 4.0 x16（用于显卡）', style='List Bullet 3')
        p55._element.getparent().insert(insert_pos, p55._element)
        insert_pos += 1
        
        p56 = doc.add_paragraph('  - 2×PCIe 3.0 x4（用于多口网卡、串口卡）', style='List Bullet 3')
        p56._element.getparent().insert(insert_pos, p56._element)
        insert_pos += 1
        
        p57 = doc.add_paragraph('M.2插槽：≥2个（用于NVMe SSD）', style='List Bullet 3')
        p57._element.getparent().insert(insert_pos, p57._element)
        insert_pos += 1
        
        p58 = doc.add_paragraph('USB接口：≥6个USB 3.0/3.2', style='List Bullet 3')
        p58._element.getparent().insert(insert_pos, p58._element)
        insert_pos += 1
        
        # 网卡
        p59 = doc.add_paragraph('6. 扩展网卡', style='List Bullet 2')
        p59._element.getparent().insert(insert_pos, p59._element)
        insert_pos += 1
        
        p60 = doc.add_paragraph('型号：Intel I350-T4（四口千兆网卡）× 2张', style='List Bullet 3')
        p60._element.getparent().insert(insert_pos, p60._element)
        insert_pos += 1
        
        p61 = doc.add_paragraph('总网口：板载2个 + 扩展8个 = 10个千兆网口', style='List Bullet 3')
        p61._element.getparent().insert(insert_pos, p61._element)
        insert_pos += 1
        
        p62 = doc.add_paragraph('分配：相机6口 + 3D激光2口 + EtherCAT 1口 + 云端1口', style='List Bullet 3')
        p62._element.getparent().insert(insert_pos, p62._element)
        insert_pos += 1
        
        p63 = doc.add_paragraph('备选方案：改用千兆交换机（16口），工控机只需2-4个网口', style='List Bullet 3')
        p63._element.getparent().insert(insert_pos, p63._element)
        insert_pos += 1
        
        # 串口卡
        p64 = doc.add_paragraph('7. 串口扩展卡', style='List Bullet 2')
        p64._element.getparent().insert(insert_pos, p64._element)
        insert_pos += 1
        
        p65 = doc.add_paragraph('型号：PCI-E转4口RS485/RS232卡', style='List Bullet 3')
        p65._element.getparent().insert(insert_pos, p65._element)
        insert_pos += 1
        
        p66 = doc.add_paragraph('配置：4个独立RS485接口（支持自动流控）', style='List Bullet 3')
        p66._element.getparent().insert(insert_pos, p66._element)
        insert_pos += 1
        
        p67 = doc.add_paragraph('分配：测距传感器1口 + 陀螺仪1口 + 备用2口', style='List Bullet 3')
        p67._element.getparent().insert(insert_pos, p67._element)
        insert_pos += 1
        
        # 电源
        p68 = doc.add_paragraph('8. 电源系统', style='List Bullet 2')
        p68._element.getparent().insert(insert_pos, p68._element)
        insert_pos += 1
        
        p69 = doc.add_paragraph('输入：DC 48V（来自电池或外部电源）', style='List Bullet 3')
        p69._element.getparent().insert(insert_pos, p69._element)
        insert_pos += 1
        
        p70 = doc.add_paragraph('DC-ATX电源模块：48V转ATX标准（12V/5V/3.3V）', style='List Bullet 3')
        p70._element.getparent().insert(insert_pos, p70._element)
        insert_pos += 1
        
        p71 = doc.add_paragraph('功率计算：', style='List Bullet 3')
        p71._element.getparent().insert(insert_pos, p71._element)
        insert_pos += 1
        
        p72 = doc.add_paragraph('  - CPU：125W（满载）', style='List Bullet 3')
        p72._element.getparent().insert(insert_pos, p72._element)
        insert_pos += 1
        
        p73 = doc.add_paragraph('  - GPU：130W', style='List Bullet 3')
        p73._element.getparent().insert(insert_pos, p73._element)
        insert_pos += 1
        
        p74 = doc.add_paragraph('  - 主板+内存+SSD：50W', style='List Bullet 3')
        p74._element.getparent().insert(insert_pos, p74._element)
        insert_pos += 1
        
        p75 = doc.add_paragraph('  - 扩展卡（网卡+串口卡）：30W', style='List Bullet 3')
        p75._element.getparent().insert(insert_pos, p75._element)
        insert_pos += 1
        
        p76 = doc.add_paragraph('  - 风扇+其他：15W', style='List Bullet 3')
        p76._element.getparent().insert(insert_pos, p76._element)
        insert_pos += 1
        
        p77 = doc.add_paragraph('总功耗：≈350W（峰值）', style='List Bullet 3')
        p77.runs[0].bold = True
        p77._element.getparent().insert(insert_pos, p77._element)
        insert_pos += 1
        
        p78 = doc.add_paragraph('电源规格：DC-ATX 500W（留有40%余量）', style='List Bullet 3')
        p78.runs[0].bold = True
        p78._element.getparent().insert(insert_pos, p78._element)
        insert_pos += 1
        
        p79 = doc.add_paragraph('保护功能：过压、欠压、过流、过温、短路保护', style='List Bullet 3')
        p79._element.getparent().insert(insert_pos, p79._element)
        insert_pos += 1
        
        # 散热
        p80 = doc.add_paragraph('9. 散热系统', style='List Bullet 2')
        p80._element.getparent().insert(insert_pos, p80._element)
        insert_pos += 1
        
        p81 = doc.add_paragraph('CPU散热器：塔式风冷散热器（TDP≥150W）', style='List Bullet 3')
        p81._element.getparent().insert(insert_pos, p81._element)
        insert_pos += 1
        
        p82 = doc.add_paragraph('机箱风扇：前进后出，2×12cm风扇', style='List Bullet 3')
        p82._element.getparent().insert(insert_pos, p82._element)
        insert_pos += 1
        
        p83 = doc.add_paragraph('防尘措施：进风口配置防尘网（可清洗）', style='List Bullet 3')
        p83._element.getparent().insert(insert_pos, p83._element)
        insert_pos += 1
        
        # 操作系统
        p84 = doc.add_paragraph('10. 操作系统', style='List Bullet 2')
        p84._element.getparent().insert(insert_pos, p84._element)
        insert_pos += 1
        
        p85 = doc.add_paragraph('推荐：Windows 10/11 专业版 64位', style='List Bullet 3')
        p85._element.getparent().insert(insert_pos, p85._element)
        insert_pos += 1
        
        p86 = doc.add_paragraph('理由：C# .NET 8.0开发，相机SDK兼容性好', style='List Bullet 3')
        p86._element.getparent().insert(insert_pos, p86._element)
        insert_pos += 1
        
        p87 = doc.add_paragraph('备选：Ubuntu 22.04 LTS（开源，成本低）', style='List Bullet 3')
        p87._element.getparent().insert(insert_pos, p87._element)
        insert_pos += 1
        
        # 机箱
        p88 = doc.add_paragraph('11. 工控机箱', style='List Bullet 2')
        p88._element.getparent().insert(insert_pos, p88._element)
        insert_pos += 1
        
        p89 = doc.add_paragraph('类型：4U/5U 19英寸标准机架式工控机箱', style='List Bullet 3')
        p89._element.getparent().insert(insert_pos, p89._element)
        insert_pos += 1
        
        p90 = doc.add_paragraph('材质：全钢结构，表面喷塑处理', style='List Bullet 3')
        p90._element.getparent().insert(insert_pos, p90._element)
        insert_pos += 1
        
        p91 = doc.add_paragraph('防护等级：IP54（防尘防溅水）', style='List Bullet 3')
        p91._element.getparent().insert(insert_pos, p91._element)
        insert_pos += 1
        
        p92 = doc.add_paragraph('散热：侧面或后部散热孔 + 风扇', style='List Bullet 3')
        p92._element.getparent().insert(insert_pos, p92._element)
        insert_pos += 1
        
        p93 = doc.add_paragraph('扩展：支持ATX主板 + 全高显卡 + 多个3.5"/2.5"硬盘位', style='List Bullet 3')
        p93._element.getparent().insert(insert_pos, p93._element)
        insert_pos += 1
        
        # 配置汇总
        p94 = doc.add_paragraph('', style='Normal')
        p94._element.getparent().insert(insert_pos, p94._element)
        insert_pos += 1
        
        p95 = doc.add_paragraph('三、配置汇总表', style='List Bullet')
        p95.runs[0].bold = True
        p95._element.getparent().insert(insert_pos, p95._element)
        
        break

doc.save('D:\\铁路线路智能检测机器人\\04-项目文档\\设计文档\\机器人硬件系统架构设计文档 V1.0.docx')
print('[OK] 工控机详细配置方案已更新')

