import os
import docx
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 确保目录存在
output_dir = r"D:\铁路线路智能检测机器人\04-项目文档\设计文档"
os.makedirs(output_dir, exist_ok=True)

doc = docx.Document()

# 设置默认字体
style = doc.styles['Normal']
style.font.name = u'微软雅黑'
style._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
style.font.size = Pt(10.5)

def add_title(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = u'微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

add_title(doc, '铁路线路智能检测机器人 - 7大核心工作板块 (深度论证优化版)', 0)

doc.add_paragraph('本文档基于“做扎实、多方论证、可落地实施”的工程原则，对系统7大核心工作板块进行了极限工况下的技术论证与逻辑重构，排除了纸面架构在真实物理环境中的潜在死穴。')

# 板块 1
add_title(doc, '1. 硬件控制与多源感知集成 (底层感知域)', 2)
doc.add_paragraph('【核心配置】：10网口物理隔离、遮光罩恒定环境、绝对值编码器里程计。')
table1 = doc.add_table(rows=1, cols=3, style='Table Grid')
hdr1 = table1.rows[0].cells
hdr1[0].text = '关键落地设计'
hdr1[1].text = '工程论证与隐患防范'
hdr1[2].text = '实施标准'
data1 = [
    ('遮光罩+恒定频闪', '论证：封闭罩内容易导致相机过热(60℃宕机)。防范：必须论证并在罩内设计物理风道或加装主动散热风扇，确保长时运行。', '消除环境光干扰，相机曝光参数写死，降低AI训练难度。'),
    ('EtherCAT实时性隔离', '论证：Windows非实时系统在处理8网口巨型帧并发时，会导致EtherCAT(LAN1)丢包使电机急停。', '防范：必须在系统底层绑定CPU亲和性(Affinity)，划出独立CPU核心专门跑TwinCAT。'),
    ('多源联合标定', '论证：空间错位导致点云与图像无法叠合。', '基于定制标定板，出厂固化相机、激光、测距相对于车体中心的SE(3)外参矩阵。')
]
for item in data1:
    row = table1.add_row().cells
    row[0].text = item[0]
    row[1].text = item[1]
    row[2].text = item[2]

# 板块 2
add_title(doc, '2. 软件平台与高并发采集架构 (代码落地域)', 2)
doc.add_paragraph('【核心配置】：C# 全异步零拷贝引擎、巨型帧(9014 Bytes)。')
table2 = doc.add_table(rows=1, cols=2, style='Table Grid')
hdr2 = table2.rows[0].cells
hdr2[0].text = '关键落地设计'
hdr2[1].text = '工程论证与隐患防范'
data2 = [
    ('内存池与零拷贝', '论证：C# 频繁GC会导致系统周期性卡顿丢帧。防范：启动时预分配大容量非托管内存池，直接DMA映射。'),
    ('硬盘I/O极限写锁', '论证：进隧道断网时，缓存数据全量落盘。普通SSD的SLC Cache耗尽后掉速至100MB/s会导致全线崩溃。防范：必须选用企业级高TBW的NVMe硬盘。')
]
for item in data2:
    row = table2.add_row().cells
    row[0].text = item[0]
    row[1].text = item[1]

# 板块 3
add_title(doc, '3. 数据处理与核心 AI 算法 (系统大脑域)', 2)
doc.add_paragraph('【核心配置】：边缘端ROI截流预处理、3D/1D测距姿态融合解算。')
p3 = doc.add_paragraph()
p3.add_run('落地动作论证：\n').bold = True
p3.add_run('1. 边缘AI防洪：不在工控机跑复杂分割，只用轻量YOLO裁剪“螺栓/裂纹”ROI切片，剔除90%冗余背景，极大降低传输压力与发热。\n')
p3.add_run('2. 测距与姿态解算：8个测距传感器数据必须实时耦合2个陀螺仪的横滚角(Roll)与俯仰角(Pitch)，进行三角函数补偿，才能得出真实物理轨距。')

# 板块 4
add_title(doc, '4. 边缘-云端均衡通信架构 (数据流转域)', 2)
doc.add_paragraph('【核心配置】：LAN2独占带SIM卡工业路由器、主动链路探针、QoS四级队列。')
table4 = doc.add_table(rows=1, cols=3, style='Table Grid')
hdr4 = table4.rows[0].cells
hdr4[0].text = 'QoS级别'
hdr4[1].text = '数据类型'
hdr4[2].text = '调度论证与逻辑'
data4 = [
    ('P0极高', '告警指令/绝对坐标', '强保底，无论网多差必须发出，通过UDP/TCP双通道冗余。'),
    ('P1高', '电量/姿态/测距时序', '秒级发送，占用极小带宽，反映设备实时心跳。'),
    ('P2中', 'AI截取的局部缺陷图片', '根据探针测得的真实SIM卡带宽，动态调节发送频率。'),
    ('P3低', '原始大图/完整点云', '弱网直接“写锁”存入本地SQLite，网络极佳或驻车时通过文件指针分块(Chunk)断点续传。')
]
for item in data4:
    row = table4.add_row().cells
    row[0].text = item[0]
    row[1].text = item[1]
    row[2].text = item[2]

# 板块 5, 6, 7
add_title(doc, '5. 机械与电气系统数字化资产 (物理建档域)', 2)
doc.add_paragraph('论证难点：48V动力电机频繁加减速会产生巨大反电动势。防范：强弱电必须间隔20cm以上物理走线，相机硬触发必须采用高速光耦隔离，防止高压毛刺烧毁主板。图纸需严格受控。')

add_title(doc, '6. 技术文档与规范化管理 (工程管理域)', 2)
doc.add_paragraph('所有架构、接口、测试数据，全部采用 Word (.docx) 格式与专业表格排版，坚决摒弃非标格式，确保图纸与程序变动的全生命周期追溯。')

add_title(doc, '7. 学术论文与知识产权转化 (学术产出域)', 2)
doc.add_paragraph('论证转化：不写空泛理论。直接基于“遮光罩恒定频闪”、“EtherCAT与GigE带宽隔离”、“探针自适应均衡调度”等工程实战难点的解决过程，提炼国家发明专利交底书与高水平学术论文。')

# 设置所有表格样式
for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = u'微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
                    run.font.size = Pt(10)

file_path = r'D:\铁路线路智能检测机器人\04-项目文档\设计文档\铁路线路智能检测机器人_7大核心工作板块_深度论证优化版.docx'
doc.save(file_path)
print("Opt_V4.0 DOCX generated successfully")
