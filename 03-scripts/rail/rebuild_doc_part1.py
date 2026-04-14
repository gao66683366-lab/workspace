# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_table(doc, headers, data, title=None):
    """创建标准表格"""
    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(12)
    
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # 表头
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = header
        hdr_cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in hdr_cells[idx].paragraphs[0].runs:
            run.font.bold = True
    
    # 数据行
    for row_data in data:
        row = table.add_row()
        for idx, cell_data in enumerate(row_data):
            row.cells[idx].text = str(cell_data)
    
    doc.add_paragraph()  # 空行
    return table

# 打开文档
doc = Document('D:\\铁路线路智能检测机器人\\04-项目文档\\设计文档\\机器人硬件系统架构设计文档 V1.0.docx')

# 删除旧内容，重新生成
# 保留封面、修订历史、目录
# 从第1章开始重写

# 查找"1. 概述"
start_idx = -1
for i, para in enumerate(doc.paragraphs):
    if para.text == '1. 概述':
        start_idx = i
        break

if start_idx > 0:
    # 删除从"1. 概述"之后的所有内容
    to_delete = []
    for i in range(start_idx + 1, len(doc.paragraphs)):
        to_delete.append(doc.paragraphs[i])
    
    for p in to_delete:
        p._element.getparent().remove(p._element)
    
    # 同时删除所有表格
    for table in doc.tables:
        table._element.getparent().remove(table._element)

# 重新构建内容
doc.add_heading('1. 概述', 1)
doc.add_heading('1.1 系统简介', 2)
doc.add_paragraph(
    '铁路线路智能检测机器人是一款集机械、电子、光学、计算机技术于一体的智能检测装备，'
    '用于铁路轨道及相关设施的自动化巡检。系统采用多传感器融合技术，实现对钢轨轨面状态、'
    '几何参数、紧固件状态等关键指标的实时检测与分析。'
)

doc.add_heading('1.2 设计目标', 2)
create_table(doc, 
    ['目标类别', '具体指标', '说明'],
    [
        ['可靠性', '故障率<0.1%', '系统稳定运行，满足工业级要求'],
        ['精度', '毫米级', '检测精度达到铁路行业标准要求'],
        ['效率', '检测速度≥5km/h', '提升巡检效率，满足实际作业需求'],
        ['智能化', 'AI自动识别', '集成深度学习算法，实现缺陷自动分类'],
        ['模块化', '模块化设计', '便于维护、升级和功能扩展']
    ]
)

doc.add_heading('1.3 适用范围', 2)
doc.add_paragraph('本文档适用于铁路线路智能检测机器人硬件系统的设计、开发、测试、维护等工作。')

doc.add_page_break()

# 第2章：系统总体架构
doc.add_heading('2. 系统总体架构', 1)
doc.add_heading('2.1 系统定位', 2)
doc.add_paragraph(
    '本系统为软硬件一体化智能检测平台，通过搭载多种传感器在铁路轨道上自主运行，'
    '实时采集轨道状态数据，结合AI算法进行在线分析与判断，及时发现安全隐患。'
)

doc.add_heading('2.2 系统组成', 2)
create_table(doc,
    ['子系统名称', '主要功能', '核心设备', '备注'],
    [
        ['运动控制子系统', '机器人运动控制、速度调节、位置反馈', '伺服电机×4、运动控制器、编码器', 'EtherCAT总线'],
        ['视觉检测子系统', '轨面状态检测、螺栓识别', '工业相机×6、光源、镜头', 'GigE Vision'],
        ['三维测量子系统', '钢轨轮廓扫描、磨耗测量', '3D线激光×2、高速采集卡', '激光三角测量'],
        ['几何参数测量子系统', '轨距、水平、高低、轨向测量', '测距传感器×8、陀螺仪×2', 'Modbus RS485'],
        ['控制与处理子系统', '数据采集、处理、存储、通信', '工控机、存储设备、通信模块', '48V供电']
    ]
)

doc.add_heading('2.3 系统层级架构', 2)
create_table(doc,
    ['层级', '组成', '功能', '技术栈'],
    [
        ['云端层', '云端服务器', '数据存储、AI训练、大数据分析、Web管理', 'Python / PostgreSQL / PyTorch'],
        ['边缘层', '工控上位机', '实时采集、本地AI推理、运动控制、人机交互', 'C# .NET 8.0 / Windows 10'],
        ['设备层', '传感器与执行器', '数据采集、运动执行', 'EtherCAT / GigE / Modbus']
    ]
)

doc.add_page_break()

# 第3章：硬件详细设计
doc.add_heading('3. 硬件详细设计', 1)

# 3.1 运动控制子系统
doc.add_heading('3.1 运动控制子系统', 2)
doc.add_heading('3.1.1 系统组成', 3)
create_table(doc,
    ['设备名称', '数量', '型号规格', '主要参数', '功能说明'],
    [
        ['伺服电机', '4', '额定200W', '额定转速3000rpm\n额定扭矩0.64Nm\n编码器2500线', '驱动车轮运动'],
        ['伺服驱动器', '4', 'AC220V输入', '输出电流2A\n控制模式：位置/速度/扭矩', '控制伺服电机'],
        ['运动控制器', '1', 'EtherCAT主站', '4轴控制\n通信周期1ms\n同步精度±1μs', '协调多轴运动'],
        ['编码器', '4', '增量式', '分辨率2500线\nA/B相输出', '位置反馈']
    ]
)

doc.add_heading('3.1.2 EtherCAT通信参数', 3)
create_table(doc,
    ['参数名称', '参数值', '说明'],
    [
        ['通信周期', '1ms', '满足实时控制要求'],
        ['拓扑结构', '线型/星型', '支持热插拔'],
        ['同步精度', '±1μs', '确保多轴协调运动'],
        ['传输距离', '100m/段', '可通过中继扩展'],
        ['节点数量', '65535', '可扩展性强']
    ]
)

doc.add_heading('3.1.3 运动模式', 3)
create_table(doc,
    ['运动模式', '速度范围', '应用场景', '控制方式'],
    [
        ['匀速运动', '1-10 km/h', '常规检测作业', '恒定速度控制'],
        ['变速运动', '0.5-8 km/h', '复杂线路条件', '自适应速度调节'],
        ['点动模式', '0-0.1 km/h', '调试与精确定位', '手动控制'],
        ['回零模式', '0.2 km/h', '系统初始化', '自动寻找零点']
    ]
)

doc.add_heading('3.1.4 安全保护机制', 3)
create_table(doc,
    ['保护类型', '触发条件', '响应动作', '优先级'],
    [
        ['硬件急停', '急停按钮按下', '直接切断电机电源', '最高'],
        ['软件限位', '运动超出设定范围', '减速停止+报警', '高'],
        ['过载保护', '电流>额定值120%', '降速或停止', '高'],
        ['通信异常', 'EtherCAT中断>100ms', '紧急停止', '高'],
        ['编码器故障', '位置反馈异常', '停止并报警', '中']
    ]
)

doc.add_page_break()

# 3.2 视觉检测子系统
doc.add_heading('3.2 视觉检测子系统', 2)
doc.add_heading('3.2.1 相机配置方案', 3)
create_table(doc,
    ['相机编号', '检测对象', '安装位置', '分辨率', '视野范围', '帧率'],
    [
        ['CAM-1', '左侧轨面', '机器人前部，俯视45°', '5MP（2448×2048）', '宽500mm×长800mm', '30fps'],
        ['CAM-2', '右侧轨面', '机器人前部，俯视45°', '5MP（2448×2048）', '宽500mm×长800mm', '30fps'],
        ['CAM-3', '左侧螺栓（前）', '机器人左侧，侧视', '2MP（1920×1080）', '宽200mm×高300mm', '触发'],
        ['CAM-4', '左侧螺栓（后）', '机器人左侧，前视', '2MP（1920×1080）', '宽200mm×高300mm', '触发'],
        ['CAM-5', '右侧螺栓（前）', '机器人右侧，侧视', '2MP（1920×1080）', '宽200mm×高300mm', '触发'],
        ['CAM-6', '右侧螺栓（后）', '机器人右侧，前视', '2MP（1920×1080）', '宽200mm×高300mm', '触发']
    ]
)

doc.add_heading('3.2.2 相机技术规格', 3)
create_table(doc,
    ['技术参数', '轨面相机（×2）', '螺栓相机（×4）', '备注'],
    [
        ['传感器', 'CMOS彩色/黑白', 'CMOS彩色/黑白', '根据算法需求选择'],
        ['分辨率', '≥5MP', '≥2MP', '满足缺陷检测要求'],
        ['帧率', '30fps连续', '触发模式', '按轨枕间距触发'],
        ['接口', 'GigE Vision', 'GigE Vision', '千兆以太网'],
        ['曝光', '外触发/自由运行', '外触发', '与光源同步'],
        ['镜头', 'C口，焦距12-16mm', 'C口，焦距8-12mm', '根据视野计算'],
        ['防护', 'IP65+防护罩', 'IP65+防护罩', '户外防尘防水']
    ]
)

doc.add_heading('3.2.3 光源系统', 3)
create_table(doc,
    ['光源类型', '数量', '规格', '安装位置', '用途'],
    [
        ['LED条形光源', '2', '白光6000K，宽度600mm，功率30W', '轨面相机两侧', '轨面照明'],
        ['LED环形光源', '4', '白光6000K，直径100mm，功率15W', '螺栓相机周围', '螺栓照明'],
        ['光源控制器', '1', '6通道独立调光，触发同步', '工控机箱内', '亮度控制']
    ]
)

save_doc('3.2.4 图像采集策略', 3)
create_table(doc,
    ['相机类型', '采集方式', '触发源', '数据量（30fps）', '存储策略'],
    [
        ['轨面相机×2', '连续采集', '内部时钟', '900MB/s', 'JPEG压缩后存储'],
        ['螺栓相机×4', '位置触发', '编码器脉冲', '按需采集', '检测到异常才存储'],
        ['合计数据率', '-', '-', '≈1GB/s（压缩前）', '本地SSD + 云端备份']
    ]
)

doc.add_page_break()

# 继续添加其他章节...
# 由于内容过多，这里先保存
doc.save('D:\\铁路线路智能检测机器人\\04-项目文档\\设计文档\\机器人硬件系统架构设计文档 V1.0.docx')
print('[Progress] 前3章已转换为表格形式，继续处理...')
