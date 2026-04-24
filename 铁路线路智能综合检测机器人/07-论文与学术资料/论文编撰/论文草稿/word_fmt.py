"""
Word 学术论文排版工具
用法: python word_fmt.py <输入docx> <输出docx>

学术规范:
- 论文标题: 黑体18pt 居中
- 作者/单位: 宋体10.5pt 居中
- 摘要标题: 黑体14pt 加粗 居中
- Abstract: Times New Roman12pt 加粗 居中
- 第X章(一/二/三级): 黑体16pt 加粗 居中
- 1.X 节标题: 黑体13pt 加粗 左对齐
- 正文: 宋体12pt 首行缩进0.74cm 行距22pt
- 参考文献: 黑体14pt 加粗 居中

表格规范:
- 所有单元格: 1pt 黑色边框 (#000000)
- 表头: 蓝色底纹 (#D9E2F3) + 黑体 + 加粗 + 居中
- 数据行: 宋体10.5pt + 居中
- 列宽: 按内容合理分配（手动指定 dxa 单位）
"""

from docx import Document
from lxml import etree
import re

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

# ==================== 格式原语 ====================

def _rpr(e, fc, fe, sz, bold, italic):
    rfs = etree.SubElement(e, f'{{{W}}}rFonts')
    rfs.set(f'{{{W}}}ascii', fe or fc)
    rfs.set(f'{{{W}}}hAnsi', fe or fc)
    rfs.set(f'{{{W}}}eastAsia', fc)
    if bold: etree.SubElement(e, f'{{{W}}}b')
    if italic: etree.SubElement(e, f'{{{W}}}i')
    if sz:
        for tag in [f'{{{W}}}sz', f'{{{W}}}szCs']:
            s = etree.SubElement(e, tag); s.set(f'{{{W}}}val', str(int(sz * 2)))

def _ppr(e, align, sb, sa, indent, ls):
    sp = etree.SubElement(e, f'{{{W}}}spacing')
    if sb is not None: sp.set(f'{{{W}}}before', str(int(sb * 20)))
    if sa is not None: sp.set(f'{{{W}}}after', str(int(sa * 20)))
    if ls is not None:
        sp.set(f'{{{W}}}line', str(int(ls * 20)))
        sp.set(f'{{{W}}}lineRule', 'exact')
    if indent is not None:
        ind = etree.SubElement(e, f'{{{W}}}ind')
        ind.set(f'{{{W}}}firstLine', str(int(indent * 567)))
    if align:
        jc = etree.SubElement(e, f'{{{W}}}jc'); jc.set(f'{{{W}}}val', align)

def _tcpr(e, w_dxa, fill, vAlign='center'):
    """单元格属性: 宽度 + 底纹(可选) + 垂直对齐"""
    tcW = etree.SubElement(e, f'{{{W}}}tcW')
    tcW.set(f'{{{W}}}type', 'dxa'); tcW.set(f'{{{W}}}w', str(w_dxa))
    if fill:
        shd = etree.SubElement(e, f'{{{W}}}shd')
        shd.set(f'{{{W}}}val', 'clear')
        shd.set(f'{{{W}}}color', 'auto'); shd.set(f'{{{W}}}fill', fill)
    if vAlign:
        va = etree.SubElement(e, f'{{{W}}}vAlign'); va.set(f'{{{W}}}val', vAlign)
    # 单元格边距
    tcM = etree.SubElement(e, f'{{{W}}}tcMar')
    for side in ['top', 'bottom']:
        m = etree.SubElement(tcM, f'{{{W}}}{side}')
        m.set(f'{{{W}}}w', '60'); m.set(f'{{{W}}}type', 'dxa')
    for side in ['left', 'right']:
        m = etree.SubElement(tcM, f'{{{W}}}{side}')
        m.set(f'{{{W}}}w', '100'); m.set(f'{{{W}}}type', 'dxa')

# ==================== 表格格式化 ====================

BORDER_BLACK = ('single', 6, '000000')  # (线型, 宽度1/8pt, 颜色)

def format_table(tbl, col_widths_dxa, header_fill='D9E2F3', font_sz=10.5):
    """
    格式化整个表格:
    - col_widths_dxa: list of int, 每列宽度(twips), 例 [600, 1500, 1700]
    - header_fill: 表头底色 (默认浅蓝)
    - font_sz: 字体大小 pt
    """
    rows = list(tbl.rows)
    n_cols = len(col_widths_dxa)

    for ri, row in enumerate(rows):
        cells = list(row.cells)
        for ci in range(min(len(cells), n_cols)):
            cell = cells[ci]
            w = col_widths_dxa[ci]
            is_header = (ri == 0)
            fill = header_fill if is_header else None

            tc = cell._tc
            tcPr = tc.find(f'{{{W}}}tcPr')
            if tcPr is None:
                tcPr = etree.SubElement(tc, f'{{{W}}}tcPr')
            else:
                for child in list(tcPr): tcPr.remove(child)  # 清空原有属性

            # 添加单元格属性
            _tcpr(tcPr, w, fill)

            # 添加边框
            tcB = etree.SubElement(tcPr, f'{{{W}}}tcBorders')
            for side in ['top', 'left', 'bottom', 'right']:
                b = etree.SubElement(tcB, f'{{{W}}}{side}')
                b.set(f'{{{W}}}val', BORDER_BLACK[0])
                b.set(f'{{{W}}}sz', str(BORDER_BLACK[1]))
                b.set(f'{{{W}}}space', '0')
                b.set(f'{{{W}}}color', BORDER_BLACK[2])

            # 单元格内容
            p = cell.paragraphs[0]._element
            for child in list(p):
                if child.tag != f'{{{W}}}pPr': p.remove(child)
            pPr = p.find(f'{{{W}}}pPr')
            if pPr is None: pPr = etree.SubElement(p, f'{{{W}}}pPr')
            else:
                for child in list(pPr): pPr.remove(child)
            jc = etree.SubElement(pPr, f'{{{W}}}jc'); jc.set(f'{{{W}}}val', 'center')
            for r_el in p.findall(f'{{{W}}}r'): p.remove(r_el)

            # 新建 run
            r = etree.SubElement(p, f'{{{W}}}r')
            rp = etree.SubElement(r, f'{{{W}}}rPr')
            fc = '黑体' if is_header else '宋体'
            fe = 'SimHei' if is_header else 'SimSun'
            _rpr(rp, fc, fe, font_sz, True, False)
            t = etree.SubElement(r, f'{{{W}}}t')
            t.text = cell.paragraphs[0].text
            if t.text: t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

    # 表级边框 (备选，增强)
    tbl_el = tbl._tbl
    tblPr = tbl_el.find(f'{{{W}}}tblPr')
    if tblPr is None:
        tblPr = etree.SubElement(tbl_el, f'{{{W}}}tblPr')
    existing = tblPr.find(f'{{{W}}}tblBorders')
    if existing is not None: tblPr.remove(existing)
    tblB = etree.SubElement(tblPr, f'{{{W}}}tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = etree.SubElement(tblB, f'{{{W}}}{side}')
        b.set(f'{{{W}}}val', BORDER_BLACK[0])
        b.set(f'{{{W}}}sz', str(BORDER_BLACK[1]))
        b.set(f'{{{W}}}space', '0')
        b.set(f'{{{W}}}color', BORDER_BLACK[2])


def add_tbl_borders_all(doc):
    """给文档中所有表格加表级边框 (已在 format_table 中调用)"""
    for tbl in doc.tables:
        tbl_el = tbl._tbl
        tblPr = tbl_el.find(f'{{{W}}}tblPr')
        if tblPr is None:
            tblPr = etree.SubElement(tbl_el, f'{{{W}}}tblPr')
        existing = tblPr.find(f'{{{W}}}tblBorders')
        if existing is not None: tblPr.remove(existing)
        tblB = etree.SubElement(tblPr, f'{{{W}}}tblBorders')
        for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            b = etree.SubElement(tblB, f'{{{W}}}{side}')
            b.set(f'{{{W}}}val', BORDER_BLACK[0])
            b.set(f'{{{W}}}sz', str(BORDER_BLACK[1]))
            b.set(f'{{{W}}}space', '0')
            b.set(f'{{{W}}}color', BORDER_BLACK[2])


# ==================== 正文段落格式化 ====================

def format_paragraph(p_el, text, fc, fe, sz, bold, align, sb, sa, indent, ls):
    """清空段落, 用格式化文本重写"""
    for child in list(p_el):
        if child.tag != f'{{{W}}}pPr': p_el.remove(child)
    pPr = p_el.find(f'{{{W}}}pPr')
    if pPr is None: pPr = etree.SubElement(p_el, f'{{{W}}}pPr')
    else:
        for child in list(pPr): pPr.remove(child)
    _ppr(pPr, align, sb, sa, indent, ls)
    r = etree.SubElement(p_el, f'{{{W}}}r')
    rp = etree.SubElement(r, f'{{{W}}}rPr')
    _rpr(rp, fc, fe, sz, bold, False)
    t = etree.SubElement(r, f'{{{W}}}t')
    t.text = text
    if text: t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')


# ==================== 完整文档格式化 ====================

CHAPTER_PAT = re.compile(r'^第.+章')
SECTION_PAT = re.compile(r'^\d+\.\d+')

def format_doc(doc, table_col_widths=None):
    """
    对文档应用完整学术论文排版规范
    table_col_widths: list of col_widths lists, 按表格顺序对应
    """
    # 表格格式化
    if table_col_widths:
        for ti, tbl in enumerate(doc.tables):
            if ti < len(table_col_widths):
                format_table(tbl, table_col_widths[ti])

    # 段落格式化
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if not t: continue

        if i == 0:  # 论文标题
            format_paragraph(p._element, t, '黑体', 'SimHei', 18, True, 'CENTER', 24, 12, None, None)
            continue

        if i in [1, 2, 3]:  # 作者/单位
            format_paragraph(p._element, t, '宋体', 'SimSun', 10.5, False, 'CENTER', None, 6, None, None)
            continue

        if t == '摘  要':
            format_paragraph(p._element, t, '黑体', 'SimHei', 14, True, 'CENTER', None, 4, None, None)
            continue
        if t == 'Abstract':
            format_paragraph(p._element, t, '宋体', 'Times New Roman', 12, True, 'CENTER', None, 4, None, None)
            continue
        if t == '参考文献':
            format_paragraph(p._element, t, '黑体', 'SimHei', 14, True, 'CENTER', 14, 6, None, None)
            continue
        if t.startswith('关键词：') or t.startswith('Keywords:'):
            format_paragraph(p._element, t, '宋体', 'SimSun', 12, False, 'LEFT', None, 6, None, None)
            continue

        if CHAPTER_PAT.match(t):  # 第X章
            format_paragraph(p._element, t, '黑体', 'SimHei', 16, True, 'CENTER', 14, 6, None, None)
            continue

        if SECTION_PAT.match(t) and len(t) < 35:  # 1.X 节标题
            format_paragraph(p._element, t, '黑体', 'SimHei', 13, True, 'LEFT', 8, 4, None, None)
            continue

        # 正文
        format_paragraph(p._element, t, '宋体', 'SimSun', 12, False, None, None, 6, 0.74, 22)


if __name__ == '__main__':
    import sys
    if len(sys.argv) < 3:
        print("用法: python word_fmt.py <输入.docx> <输出.docx> [表1宽] [表2宽] ...")
        sys.exit(1)
    inp, out = sys.argv[1], sys.argv[2]

    doc = Document(inp)

    # 默认列宽 (twips): 适合 A4 页面(15cm宽)
    # 可通过命令行参数覆盖
    default_widths = {
        0: [1400, 1200, 1200, 2200],  # 表1: 传感器
        1: [600,  1500, 1700, 1100, 2000],  # 表2: 8项功能(5列)
        2: [2800, 1100, 1100, 1100],  # 表3: 融合对比
    }

    table_widths = []
    for ti in range(len(doc.tables)):
        w_arg = sys.argv[3 + ti] if 3 + ti < len(sys.argv) else None
        if w_arg:
            table_widths.append([int(x) for x in w_arg.split(',')])
        elif ti in default_widths:
            table_widths.append(default_widths[ti])
        else:
            # 默认等宽
            n_cols = len(doc.tables[ti].columns)
            total = 8640  # 约15cm
            table_widths.append([total // n_cols] * n_cols)

    format_doc(doc, table_widths)
    doc.save(out)
    print(f"已保存: {out}")
    print(f"总字符: {sum(len(p.text) for p in doc.paragraphs)}")
