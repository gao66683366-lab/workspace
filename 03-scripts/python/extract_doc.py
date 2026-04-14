import docx
doc = docx.Document(r'D:\铁路线路智能检测机器人\04-项目文档\设计文档\铁路线路智能检测机器人系统全景架构设计文档_V3.0_终极版.docx')
with open(r'D:\铁路线路智能检测机器人\07-临时文件\终极版_temp.txt', 'w', encoding='utf-8') as f:
    for p in doc.paragraphs:
        if p.text.strip():
            f.write(p.text + '\n')
    f.write('\n--- 表格内容 ---\n')
    for t in doc.tables:
        for row in t.rows:
            f.write(' | '.join([c.text.replace('\n', ' ') for c in row.cells]) + '\n')
        f.write('\n')
