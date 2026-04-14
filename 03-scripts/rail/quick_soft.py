# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# 样式设置
style = doc.styles['Normal']
style.font.name = '宋体'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.font.size = Pt(12)

# 封面
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('\n\n\n\n\n')
title_run = p.add_run('铁路线路智能检测机器人')
title_run.font.name = '黑体'
title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
title_run.font.size = Pt(24)
title_run.font.bold = True
p.add_run('\n\n')
subtitle_run = p.add_run('软件系统架构设计文档')
subtitle_run.font.name = '黑体'
subtitle_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
subtitle_run.font.size = Pt(18)
p.add_run('\n\n\n\n\n\n\n\n\n\n')
info_p = doc.add_paragraph()
info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
info_run = info_p.add_run('编制日期：2026年3月5日\n版本：V1.0')
info_run.font.size = Pt(12)
doc.add_page_break()

# 修订历史
doc.add_heading('文档修订历史', 1)
table = doc.add_table(rows=2, cols=5)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text, hdr[4].text = '版本', '日期', '修订内容', '修订人', '审核人'
row = table.rows[1].cells
row[0].text, row[1].text, row[2].text, row[3].text, row[4].text = 'V1.0', '2026-03-05', '初始版本', '小测', '道'
doc.add_paragraph()
doc.add_heading('目录', 1)
doc.add_paragraph('（此处应插入自动目录）')
doc.add_page_break()

doc.save('D:\\铁路线路智能检测机器人\\04-项目文档\\设计文档\\软件系统架构设计文档 V1.0.docx')
print('[OK] 软件架构文档已创建')
