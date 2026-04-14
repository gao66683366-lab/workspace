# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document('D:\\铁路线路智能检测机器人\\04-项目文档\\设计文档\\机器人硬件系统架构设计文档 V1.0.docx')

# 在"三、配置汇总表"后插入表格
for i, para in enumerate(doc.paragraphs):
    if '三、配置汇总表' in para.text:
        # 删除后续可能存在的内容
        delete_count = 0
        for j in range(i+1, len(doc.paragraphs)):
            if doc.paragraphs[j].style.name == 'Heading 3' or '3.5.3' in doc.paragraphs[j].text:
                break
            delete_count += 1
        
        for j in range(delete_count):
            if i + 1 < len(doc.paragraphs):
                p = doc.paragraphs[i + 1]
                p._element.getparent().remove(p._element)
        
        # 创建配置汇总表
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'
        
        # 设置列宽
        table.autofit = False
        table.allow_autofit = False
        widths = [Cm(2.5), Cm(3.5), Cm(4), Cm(4), Cm(2)]
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width
        
        # 表头
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '类别'
        hdr_cells[1].text = '组件'
        hdr_cells[2].text = '型号/规格'
        hdr_cells[3].text = '参数说明'
        hdr_cells[4].text = '数量'
        
        for cell in hdr_cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.bold = True
        
        # 数据行
        data = [
            ['核心\n处理器', 'CPU', 'Intel Core i7-12700', '12核20线程，2.1-4.9GHz，25MB缓存，TDP 125W', '1'],
            ['', '主板', 'Intel B760 / Z790', 'ATX板型，2×板载千兆网口，3×PCIe插槽，2×M.2插槽', '1'],
            ['内存\n存储', '内存', 'DDR4-3200 32GB', '2×16GB双通道，可扩展至64GB', '1'],
            ['', '系统盘', 'NVMe SSD 512GB', 'M.2接口，PCIe 4.0，读写≥3500MB/s', '1'],
            ['', '数据盘', 'NVMe SSD 2TB', 'M.2接口，RAID 0阵列，写入≥6000MB/s', '2'],
            ['', '备份盘', 'SATA SSD 4TB', '2.5英寸，数据镜像备份', '1'],
            ['计算\n加速', 'GPU', 'NVIDIA RTX A2000', '8GB GDDR6显存，3328个CUDA核心，功耗70W', '1'],
            ['网络\n通信', '扩展网卡', 'Intel I350-T4', '四口千兆网卡，PCIe x4接口', '2'],
            ['', '以太网交换机', '16口千兆交换机', '工业级，支持GigE Vision，非网管型（备选方案）', '1'],
            ['串口\n扩展', '串口卡', 'PCIe转RS485卡', '4口独立RS485，自动流控', '1'],
            ['电源\n系统', 'DC-ATX电源', '48V转ATX 500W', '输入DC 48V，输出12V/5V/3.3V，过压过流保护', '1'],
            ['散热\n系统', 'CPU散热器', '塔式风冷', 'TDP≥150W，4热管', '1'],
            ['', '机箱风扇', '12cm静音风扇', '前进后出风道', '2'],
            ['机箱\n防护', '工控机箱', '4U机架式', '19英寸标准，全钢结构，IP54防护', '1']
        ]
        
        for row_data in data:
            row = table.add_row()
            for idx, cell_data in enumerate(row_data):
                cell = row.cells[idx]
                cell.text = cell_data
                if idx == 4:  # 数量列居中
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 插入表格到文档
        table._element.getparent().insert(i + 1, table._element)
        
        # 添加接口配置表
        p_title = doc.add_paragraph()
        p_title.add_run('\n接口配置清单').bold = True
        p_title._element.getparent().insert(i + 2, p_title._element)
        
        table2 = doc.add_table(rows=1, cols=4)
        table2.style = 'Light Grid Accent 1'
        
        hdr2 = table2.rows[0].cells
        hdr2[0].text = '接口类型'
        hdr2[1].text = '数量'
        hdr2[2].text = '用途分配'
        hdr2[3].text = '备注'
        
        for cell in hdr2:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.bold = True
        
        data2 = [
            ['千兆网口', '10个', '相机6 + 3D激光2 + EtherCAT 1 + 云端1', '板载2 + I350-T4扩展8'],
            ['RS485串口', '4个', '测距传感器1 + 陀螺仪1 + 备用2', 'PCIe扩展卡提供'],
            ['USB 3.0', '6个', '鼠标/键盘/U盘/调试设备', '主板集成'],
            ['HDMI/VGA', '1个', '显示器连接', '主板或GPU输出'],
            ['PCIe插槽', '3个', 'GPU 1 + 网卡2', 'x16×1 + x4×2']
        ]
        
        for row_data in data2:
            row = table2.add_row()
            for idx, cell_data in enumerate(row_data):
                cell = row.cells[idx]
                cell.text = cell_data
                if idx == 1:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table2._element.getparent().insert(i + 3, table2._element)
        
        # 添加功耗分析表
        p_title2 = doc.add_paragraph()
        p_title2.add_run('\n功耗分析表').bold = True
        p_title2._element.getparent().insert(i + 4, p_title2._element)
        
        table3 = doc.add_table(rows=1, cols=4)
        table3.style = 'Light Grid Accent 1'
        
        hdr3 = table3.rows[0].cells
        hdr3[0].text = '组件'
        hdr3[1].text = '典型功耗'
        hdr3[2].text = '峰值功耗'
        hdr3[3].text = '占比'
        
        for cell in hdr3:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.bold = True
        
        data3 = [
            ['CPU', '65W', '125W', '36%'],
            ['GPU', '70W', '130W', '37%'],
            ['主板+内存', '30W', '40W', '11%'],
            ['SSD（3块）', '15W', '25W', '7%'],
            ['网卡+串口卡', '20W', '30W', '9%'],
            ['合计', '200W', '350W', '100%'],
            ['电源规格（留40%余量）', '', '500W', '']
        ]
        
        for row_data in data3:
            row = table3.add_row()
            for idx, cell_data in enumerate(row_data):
                cell = row.cells[idx]
                cell.text = cell_data
                if idx > 0:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                if '合计' in cell_data or '电源规格' in cell_data:
                    for run in cell.paragraphs[0].runs:
                        run.font.bold = True
        
        table3._element.getparent().insert(i + 5, table3._element)
        
        # 添加存储容量表
        p_title3 = doc.add_paragraph()
        p_title3.add_run('\n存储容量分配表').bold = True
        p_title3._element.getparent().insert(i + 6, p_title3._element)
        
        table4 = doc.add_table(rows=1, cols=5)
        table4.style = 'Light Grid Accent 1'
        
        hdr4 = table4.rows[0].cells
        hdr4[0].text = '存储设备'
        hdr4[1].text = '容量'
        hdr4[2].text = '用途'
        hdr4[3].text = '性能要求'
        hdr4[4].text = '冗余'
        
        for cell in hdr4:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.bold = True
        
        data4 = [
            ['系统盘', '512GB', 'Windows + 软件 + AI模型', '读写≥3500MB/s', '否'],
            ['数据盘1', '2TB', '实时数据存储（RAID 0）', '写入≥6000MB/s', 'RAID'],
            ['数据盘2', '2TB', '实时数据存储（RAID 0）', '写入≥6000MB/s', 'RAID'],
            ['备份盘', '4TB', '数据镜像备份', '读写≥500MB/s', '是'],
            ['合计', '8.5TB', '实际可用：6.5TB', '', '']
        ]
        
        for row_data in data4:
            row = table4.add_row()
            for idx, cell_data in enumerate(row_data):
                cell = row.cells[idx]
                cell.text = cell_data
                if idx == 1 or idx == 4:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                if '合计' in cell_data:
                    for run in cell.paragraphs[0].runs:
                        run.font.bold = True
        
        table4._element.getparent().insert(i + 7, table4._element)
        
        break

doc.save('D:\\铁路线路智能检测机器人\\04-项目文档\\设计文档\\机器人硬件系统架构设计文档 V1.0.docx')
print('[OK] 已重新排版为专业表格形式')
