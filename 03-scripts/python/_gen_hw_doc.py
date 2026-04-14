from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

path = r"D:\铁路线路智能检测机器人\02_原始技术文档\设计文档\硬件架构设计方案_定稿_2026-03-07.docx"

doc = Document()
style = doc.styles['Normal']
style.font.name = '宋体'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.font.size = Pt(11)

doc.add_heading('铁路线路智能检测机器人 硬件架构设计方案（定稿）', level=1)
doc.add_paragraph('版本：V1.0')
doc.add_paragraph('日期：2026-03-07')

# 1 总览表
doc.add_heading('1. 系统硬件总览', level=2)
t = doc.add_table(rows=1, cols=6)
t.style='Table Grid'
headers=['层级','子系统','核心设备','数量','接入方式','主要职责']
for i,h in enumerate(headers):
    t.rows[0].cells[i].text=h
rows=[
['运动执行层','行走执行','伺服电机','4','驱动器输出','驱动4个车轮在钢轨运行'],
['运动执行层','轮系机构','车轮','4','机械直连','小车运行与承载'],
['控制层','运动控制','运动控制器','1','网口直连工控机','下发运动控制指令、协调4轴'],
['控制层','伺服驱动','总线型驱动器','4','受运动控制器控制','驱动伺服电机'],
['感知层（高速）','视觉采集','2D工业相机','6','网口直连工控机','轨面/构件图像采集'],
['感知层（高速）','几何采集','3D线激光','2','网口直连工控机','轮廓几何数据采集'],
['感知层（低速）','距离感知','测距传感器','8','RS485串口直连工控机','距离/状态补充采集'],
['感知层（低速）','姿态感知','陀螺仪','2','RS485串口直连工控机','姿态与运动状态补充'],
['边缘计算层','数据中枢','工控机','1','网口+串口','统一接入、调度、缓存与处理']
]
for r in rows:
    c=t.add_row().cells
    for i,v in enumerate(r): c[i].text=v

# 2 运动控制链路
doc.add_heading('2. 运动控制链路', level=2)
t2=doc.add_table(rows=1, cols=5)
t2.style='Table Grid'
for i,h in enumerate(['链路序号','上游','下游','连接方式','说明']):
    t2.rows[0].cells[i].text=h
rows2=[
['1','工控机','运动控制器','网口直连','上位协同与控制交互'],
['2','运动控制器','总线型驱动器（4）','总线控制','统一控制4个驱动器'],
['3','驱动器（4）','伺服电机（4）','电机驱动连接','一驱一机'],
['4','伺服电机（4）','车轮（4）','机械传动','实现钢轨行走']
]
for r in rows2:
    c=t2.add_row().cells
    for i,v in enumerate(r): c[i].text=v

# 3 感知接入
doc.add_heading('3. 感知接入链路', level=2)
t3=doc.add_table(rows=1, cols=6)
t3.style='Table Grid'
for i,h in enumerate(['类别','设备','数量','到工控机链路','数据特征','用途']):
    t3.rows[0].cells[i].text=h
rows3=[
['高速感知','2D工业相机','6','网口直连','图像流','缺陷/纹理/标识检测'],
['高速感知','3D线激光','2','网口直连','轮廓/点云类数据','几何与形变检测'],
['低速感知','测距传感器','8','RS485串口直连','低速状态量','距离与近场状态补充'],
['低速感知','陀螺仪','2','RS485串口直连','姿态状态量','姿态补偿与工况判断']
]
for r in rows3:
    c=t3.add_row().cells
    for i,v in enumerate(r): c[i].text=v

# 4 接口资源
doc.add_heading('4. 工控机接口资源分配', level=2)
t4=doc.add_table(rows=1, cols=4)
t4.style='Table Grid'
for i,h in enumerate(['接口类型','对接对象','数量/规模','备注']):
    t4.rows[0].cells[i].text=h
rows4=[
['网口','运动控制器','1链路','工控机与运动控制器直连'],
['网口','2D工业相机','6链路','直连'],
['网口','3D线激光','2链路','直连'],
['串口（RS485）','测距传感器','8设备','串口总线接入'],
['串口（RS485）','陀螺仪','2设备','串口总线接入']
]
for r in rows4:
    c=t4.add_row().cells
    for i,v in enumerate(r): c[i].text=v

# 5 约束
doc.add_heading('5. 硬件架构约束', level=2)
t5=doc.add_table(rows=1, cols=2)
t5.style='Table Grid'
for i,h in enumerate(['约束项','要求']):
    t5.rows[0].cells[i].text=h
rows5=[
['行走结构','4伺服+4轮，钢轨运行'],
['控制结构','1运动控制器→4总线型驱动器'],
['控制互联','运动控制器与工控机网口直连'],
['高速采集','6个2D相机+2个3D线激光均网口直连工控机'],
['低速采集','8测距+2陀螺仪均RS485串口直连工控机']
]
for r in rows5:
    c=t5.add_row().cells
    c[0].text=r[0]; c[1].text=r[1]

doc.add_paragraph('\n说明：本文件仅覆盖硬件架构；供电方案另行章节说明。')
doc.save(path)
print(path)
