# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

def setup_document_styles(doc):
    """设置文档样式"""
    # 正文样式
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)
    
    # 标题样式
    for i in range(1, 4):
        heading = doc.styles[f'Heading {i}']
        heading.font.name = '黑体'
        heading._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        heading.font.color.rgb = RGBColor(0, 0, 0)
        if i == 1:
            heading.font.size = Pt(16)
            heading.font.bold = True
        elif i == 2:
            heading.font.size = Pt(14)
            heading.font.bold = True
        else:
            heading.font.size = Pt(13)
            heading.font.bold = True

def add_cover_page(doc, title, subtitle):
    """添加封面"""
    # 标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('\n\n\n\n\n')
    
    title_run = p.add_run(title)
    title_run.font.name = '黑体'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    
    p.add_run('\n\n')
    
    subtitle_run = p.add_run(subtitle)
    subtitle_run.font.name = '黑体'
    subtitle_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    subtitle_run.font.size = Pt(18)
    
    p.add_run('\n\n\n\n\n\n\n\n\n\n')
    
    # 文档信息
    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_p.add_run(f'编制日期：2026年3月5日\n版本：V1.0')
    info_run.font.size = Pt(12)
    
    doc.add_page_break()

def create_hardware_architecture_doc():
    """创建硬件系统架构文档"""
    doc = Document()
    setup_document_styles(doc)
    
    # 封面
    add_cover_page(doc, '铁路线路智能检测机器人', '硬件系统架构设计文档')
    
    # 文档修订历史
    doc.add_heading('文档修订历史', 1)
    table = doc.add_table(rows=2, cols=5)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '版本'
    hdr_cells[1].text = '日期'
    hdr_cells[2].text = '修订内容'
    hdr_cells[3].text = '修订人'
    hdr_cells[4].text = '审核人'
    
    row_cells = table.rows[1].cells
    row_cells[0].text = 'V1.0'
    row_cells[1].text = '2026-03-05'
    row_cells[2].text = '初始版本'
    row_cells[3].text = '小测'
    row_cells[4].text = '道'
    
    doc.add_paragraph()
    
    # 目录占位
    doc.add_heading('目录', 1)
    doc.add_paragraph('（此处应插入自动目录，在Word中通过"引用→目录"功能生成）')
    doc.add_page_break()
    
    # 1. 概述
    doc.add_heading('1. 概述', 1)
    
    doc.add_heading('1.1 系统简介', 2)
    doc.add_paragraph(
        '铁路线路智能检测机器人是一款集机械、电子、光学、计算机技术于一体的智能检测装备，'
        '用于铁路轨道及相关设施的自动化巡检。系统采用多传感器融合技术，实现对钢轨轨面状态、'
        '几何参数、紧固件状态等关键指标的实时检测与分析。'
    )
    
    doc.add_heading('1.2 设计目标', 2)
    doc.add_paragraph('本系统设计遵循以下目标：', style='List Bullet')
    doc.add_paragraph('高可靠性：系统稳定运行，故障率低于0.1%', style='List Bullet 2')
    doc.add_paragraph('高精度：检测精度达到毫米级，满足铁路行业标准要求', style='List Bullet 2')
    doc.add_paragraph('高效率：检测速度≥5km/h，提升巡检效率', style='List Bullet 2')
    doc.add_paragraph('智能化：集成AI算法，实现缺陷自动识别与分类', style='List Bullet 2')
    doc.add_paragraph('模块化：采用模块化设计，便于维护与升级', style='List Bullet 2')
    
    doc.add_heading('1.3 适用范围', 2)
    doc.add_paragraph('本文档适用于铁路线路智能检测机器人硬件系统的设计、开发、测试、维护等工作。')
    
    doc.add_page_break()
    
    # 2. 系统总体架构
    doc.add_heading('2. 系统总体架构', 1)
    
    doc.add_heading('2.1 系统定位', 2)
    doc.add_paragraph(
        '本系统为软硬件一体化智能检测平台，通过搭载多种传感器在铁路轨道上自主运行，'
        '实时采集轨道状态数据，结合AI算法进行在线分析与判断，及时发现安全隐患。'
    )
    
    doc.add_heading('2.2 系统组成', 2)
    doc.add_paragraph('系统由以下五大子系统构成：')
    
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '子系统名称'
    hdr_cells[1].text = '主要功能'
    hdr_cells[2].text = '核心设备'
    
    data = [
        ['运动控制子系统', '机器人运动控制、速度调节、位置反馈', '伺服电机×4、运动控制器、编码器'],
        ['视觉检测子系统', '轨面状态检测、螺栓识别', '工业相机×6、光源、镜头'],
        ['三维测量子系统', '钢轨轮廓扫描、磨耗测量', '3D线激光×2、高速采集卡'],
        ['几何参数测量子系统', '轨距、水平、高低、轨向测量', '测距传感器×8、陀螺仪×2'],
        ['控制与处理子系统', '数据采集、处理、存储、通信', '工控机、存储设备、通信模块']
    ]
    
    for i, row_data in enumerate(data, 1):
        row_cells = table.rows[i].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
    
    doc.add_paragraph()
    
    doc.add_heading('2.3 系统拓扑结构', 2)
    p = doc.add_paragraph()
    p.add_run('系统采用分层架构设计，分为三层：\n\n').bold = True
    p.add_run('【云端层】\n').bold = True
    p.add_run('- 云端服务器（Python）\n')
    p.add_run('  · 数据存储与管理\n')
    p.add_run('  · AI模型训练与优化\n')
    p.add_run('  · 大数据分析\n')
    p.add_run('  · Web管理平台\n\n')
    
    p.add_run('【边缘层】\n').bold = True
    p.add_run('- 工控上位机（C# .NET 8.0）\n')
    p.add_run('  · 实时数据采集\n')
    p.add_run('  · 本地AI推理\n')
    p.add_run('  · 运动控制\n')
    p.add_run('  · 人机交互\n\n')
    
    p.add_run('【设备层】\n').bold = True
    p.add_run('- 传感器与执行器\n')
    p.add_run('  · 伺服电机×4\n')
    p.add_run('  · 工业相机×6\n')
    p.add_run('  · 3D线激光×2\n')
    p.add_run('  · 测距传感器×8\n')
    p.add_run('  · 陀螺仪×2\n')
    
    doc.add_page_break()
    
    # 3. 硬件详细设计
    doc.add_heading('3. 硬件详细设计', 1)
    
    # 3.1 运动控制子系统
    doc.add_heading('3.1 运动控制子系统', 2)
    
    doc.add_heading('3.1.1 系统组成', 3)
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '设备名称'
    hdr_cells[1].text = '数量'
    hdr_cells[2].text = '主要参数'
    hdr_cells[3].text = '功能说明'
    
    data = [
        ['伺服电机', '4', '额定功率：200W\n额定转速：3000rpm\n额定扭矩：0.64Nm', '驱动车轮运动'],
        ['伺服驱动器', '4', '输入电压：AC220V\n输出电流：2A\n控制模式：位置/速度/扭矩', '控制伺服电机'],
        ['运动控制器', '1', '轴数：4轴\n通信：EtherCAT\n控制周期：1ms', '协调多轴运动'],
        ['编码器', '4', '分辨率：2500线\n输出：增量式', '位置反馈']
    ]
    
    for i, row_data in enumerate(data, 1):
        row_cells = table.rows[i].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
    
    doc.add_paragraph()
    
    doc.add_heading('3.1.2 控制方式', 3)
    doc.add_paragraph('采用EtherCAT工业以太网总线进行实时控制：', style='List Bullet')
    doc.add_paragraph('通信周期：1ms，满足实时控制要求', style='List Bullet 2')
    doc.add_paragraph('拓扑结构：星型或线型，支持热插拔', style='List Bullet 2')
    doc.add_paragraph('同步精度：±1μs，确保多轴协调运动', style='List Bullet 2')
    doc.add_paragraph('传输距离：单段最长100m', style='List Bullet 2')
    
    doc.add_heading('3.1.3 运动模式', 3)
    doc.add_paragraph('支持以下运动模式：', style='List Bullet')
    doc.add_paragraph('匀速运动：恒定速度巡检，适用于常规检测', style='List Bullet 2')
    doc.add_paragraph('变速运动：根据线路条件自动调速', style='List Bullet 2')
    doc.add_paragraph('点动模式：手动控制，用于调试与精确定位', style='List Bullet 2')
    doc.add_paragraph('回零模式：系统初始化，建立坐标系', style='List Bullet 2')
    
    doc.add_heading('3.1.4 安全保护', 3)
    doc.add_paragraph('集成多重安全保护机制：', style='List Bullet')
    doc.add_paragraph('硬件急停：物理按钮，直接切断电机电源', style='List Bullet 2')
    doc.add_paragraph('软件限位：软件监控运动范围，防止越界', style='List Bullet 2')
    doc.add_paragraph('过载保护：电流监测，防止电机过载', style='List Bullet 2')
    doc.add_paragraph('通信监测：EtherCAT通信中断检测与报警', style='List Bullet 2')
    
    # 3.2 视觉检测子系统
    doc.add_heading('3.2 视觉检测子系统', 2)
    
    doc.add_heading('3.2.1 相机配置方案', 3)
    
    table = doc.add_table(rows=4, cols=5)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '相机编号'
    hdr_cells[1].text = '检测对象'
    hdr_cells[2].text = '安装位置'
    hdr_cells[3].text = '分辨率建议'
    hdr_cells[4].text = '视野范围'
    
    data = [
        ['CAM1-2', '左右轨面', '机器人前部，俯视', '≥5MP', '宽度500mm×长度800mm'],
        ['CAM3-4', '左右侧螺栓', '机器人两侧，侧视', '≥2MP', '宽度200mm×高度300mm'],
        ['CAM5-6', '左右侧螺栓', '机器人两侧，前视', '≥2MP', '宽度200mm×高度300mm']
    ]
    
    for i, row_data in enumerate(data, 1):
        row_cells = table.rows[i].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
    
    doc.add_paragraph()
    
    doc.add_heading('3.2.2 相机技术参数', 3)
    doc.add_paragraph('工业相机应满足以下技术要求：', style='List Bullet')
    doc.add_paragraph('传感器：CMOS或CCD，黑白或彩色', style='List Bullet 2')
    doc.add_paragraph('分辨率：轨面相机≥5MP，螺栓相机≥2MP', style='List Bullet 2')
    doc.add_paragraph('帧率：≥30fps（匀速5km/h时）', style='List Bullet 2')
    doc.add_paragraph('接口：千兆以太网（GigE Vision）', style='List Bullet 2')
    doc.add_paragraph('曝光：支持外触发曝光，曝光时间可调', style='List Bullet 2')
    doc.add_paragraph('镜头：C口或CS口，焦距根据视野计算', style='List Bullet 2')
    doc.add_paragraph('防护等级：≥IP65，适应户外环境', style='List Bullet 2')
    
    doc.add_heading('3.2.3 光源系统', 3)
    doc.add_paragraph('配置LED光源系统以保证图像质量：', style='List Bullet')
    doc.add_paragraph('类型：LED条形光源或环形光源', style='List Bullet 2')
    doc.add_paragraph('色温：白光（6000-6500K）', style='List Bullet 2')
    doc.add_paragraph('亮度：可调节，应对不同光照条件', style='List Bullet 2')
    doc.add_paragraph('触发：与相机同步触发，减少曝光时间', style='List Bullet 2')
    
    doc.add_heading('3.2.4 图像采集策略', 3)
    doc.add_paragraph('采用位置触发与时间触发混合方式：', style='List Bullet')
    doc.add_paragraph('轨面相机：连续采集，帧率30fps', style='List Bullet 2')
    doc.add_paragraph('螺栓相机：位置触发，根据轨枕间距（0.6m）触发', style='List Bullet 2')
    doc.add_paragraph('触发源：编码器脉冲或运动控制器位置信号', style='List Bullet 2')
    
    # 3.3 三维测量子系统
    doc.add_heading('3.3 三维测量子系统', 2)
    
    doc.add_heading('3.3.1 3D线激光选型', 3)
    doc.add_paragraph('3D线激光传感器技术参数：', style='List Bullet')
    doc.add_paragraph('测量原理：激光三角测量法', style='List Bullet 2')
    doc.add_paragraph('激光波长：红光（650nm）或蓝光（450nm）', style='List Bullet 2')
    doc.add_paragraph('线宽：≥100mm，覆盖钢轨宽度', style='List Bullet 2')
    doc.add_paragraph('测量范围：Z向±50mm', style='List Bullet 2')
    doc.add_paragraph('测量精度：±0.05mm', style='List Bullet 2')
    doc.add_paragraph('扫描频率：≥2kHz', style='List Bullet 2')
    doc.add_paragraph('输出：点云数据，以太网接口', style='List Bullet 2')
    
    doc.add_heading('3.3.2 安装方案', 3)
    doc.add_paragraph('左右各安装1个3D线激光：', style='List Bullet')
    doc.add_paragraph('安装位置：机器人中部，激光线垂直于运动方向', style='List Bullet 2')
    doc.add_paragraph('安装高度：根据测量范围确定，建议150-200mm', style='List Bullet 2')
    doc.add_paragraph('入射角：30-45度，避免镜面反射', style='List Bullet 2')
    
    doc.add_heading('3.3.3 测量内容', 3)
    doc.add_paragraph('获取钢轨三维轮廓数据，用于：', style='List Bullet')
    doc.add_paragraph('轨面磨耗测量', style='List Bullet 2')
    doc.add_paragraph('轨侧磨耗测量', style='List Bullet 2')
    doc.add_paragraph('轨头宽度测量', style='List Bullet 2')
    doc.add_paragraph('轨顶面不平顺检测', style='List Bullet 2')
    
    # 3.4 几何参数测量子系统
    doc.add_heading('3.4 几何参数测量子系统', 2)
    
    doc.add_heading('3.4.1 测距传感器', 3)
    table = doc.add_table(rows=2, cols=5)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '参数'
    hdr_cells[1].text = '测量原理'
    hdr_cells[2].text = '测量范围'
    hdr_cells[3].text = '精度'
    hdr_cells[4].text = '通信接口'
    
    row_cells = table.rows[1].cells
    row_cells[0].text = '测距传感器×8'
    row_cells[1].text = '激光测距'
    row_cells[2].text = '0-200mm'
    row_cells[3].text = '±0.1mm'
    row_cells[4].text = 'Modbus RS485'
    
    doc.add_paragraph()
    doc.add_paragraph('测距传感器布置方案：', style='List Bullet')
    doc.add_paragraph('左右轨距测量：各布置2个传感器，测量轨距', style='List Bullet 2')
    doc.add_paragraph('高低测量：各布置2个传感器，测量轨面高程', style='List Bullet 2')
    
    doc.add_heading('3.4.2 陀螺仪', 3)
    table = doc.add_table(rows=2, cols=5)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '参数'
    hdr_cells[1].text = '测量轴数'
    hdr_cells[2].text = '量程'
    hdr_cells[3].text = '精度'
    hdr_cells[4].text = '通信接口'
    
    row_cells = table.rows[1].cells
    row_cells[0].text = '陀螺仪×2'
    row_cells[1].text = '6轴（3轴加速度+3轴角速度）'
    row_cells[2].text = '±2g / ±250°/s'
    row_cells[3].text = '0.01°'
    row_cells[4].text = 'Modbus RS485'
    
    doc.add_paragraph()
    doc.add_paragraph('陀螺仪测量内容：', style='List Bullet')
    doc.add_paragraph('横向倾斜角（水平）', style='List Bullet 2')
    doc.add_paragraph('纵向倾斜角（高低）', style='List Bullet 2')
    doc.add_paragraph('偏航角（轨向）', style='List Bullet 2')
    doc.add_paragraph('三轴加速度（振动监测）', style='List Bullet 2')
    
    doc.add_heading('3.4.3 测量精度要求', 3)
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '测量项目'
    hdr_cells[1].text = '测量精度'
    hdr_cells[2].text = '铁路标准要求'
    
    data = [
        ['轨距', '±0.5mm', '≤±2mm'],
        ['水平', '±0.5mm', '≤±3mm'],
        ['高低', '±0.5mm', '≤±4mm'],
        ['轨向', '±1mm', '≤±4mm']
    ]
    
    for i, row_data in enumerate(data, 1):
        row_cells = table.rows[i].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
    
    doc.add_paragraph()
    
    # 3.5 控制与处理子系统
    doc.add_heading('3.5 控制与处理子系统', 2)
    
    doc.add_heading('3.5.1 工控机配置', 3)
    doc.add_paragraph('工控机作为系统核心，需满足以下配置：', style='List Bullet')
    doc.add_paragraph('处理器：Intel Core i7 或更高，≥4核心', style='List Bullet 2')
    doc.add_paragraph('内存：≥16GB DDR4', style='List Bullet 2')
    doc.add_paragraph('存储：SSD ≥512GB（系统+软件） + HDD ≥2TB（数据）', style='List Bullet 2')
    doc.add_paragraph('显卡：独立显卡，用于AI推理加速（可选）', style='List Bullet 2')
    doc.add_paragraph('操作系统：Windows 10/11 专业版 或 Linux', style='List Bullet 2')
    
    doc.add_heading('3.5.2 通信接口要求', 3)
    doc.add_paragraph('工控机应具备以下接口：', style='List Bullet')
    doc.add_paragraph('以太网：≥2个千兆网口（1个用于相机，1个用于EtherCAT/云端）', style='List Bullet 2')
    doc.add_paragraph('串口：≥2个RS485接口（用于传感器通信）', style='List Bullet 2')
    doc.add_paragraph('USB：≥4个USB 3.0接口', style='List Bullet 2')
    doc.add_paragraph('显示：HDMI或VGA接口', style='List Bullet 2')
    
    doc.add_heading('3.5.3 电源系统', 3)
    doc.add_paragraph('系统供电方案：', style='List Bullet')
    doc.add_paragraph('输入电压：DC 24V（电池供电）或 AC 220V（外接电源）', style='List Bullet 2')
    doc.add_paragraph('电源模块：DC-DC或AC-DC模块，输出24V/12V/5V多路电压', style='List Bullet 2')
    doc.add_paragraph('功率预算：总功率≤500W', style='List Bullet 2')
    doc.add_paragraph('保护功能：过压、欠压、过流、短路保护', style='List Bullet 2')
    
    doc.add_page_break()
    
    # 4. 通信架构
    doc.add_heading('4. 通信架构', 1)
    
    doc.add_heading('4.1 总线拓扑', 2)
    
    table = doc.add_table(rows=6, cols=5)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '设备类型'
    hdr_cells[1].text = '数量'
    hdr_cells[2].text = '通信方式'
    hdr_cells[3].text = '协议/标准'
    hdr_cells[4].text = '带宽/速率'
    
    data = [
        ['伺服电机', '4', '工业以太网', 'EtherCAT', '100Mbps'],
        ['工业相机', '6', '以太网', 'GigE Vision', '1000Mbps'],
        ['3D线激光', '2', '以太网', 'TCP/IP', '1000Mbps'],
        ['测距传感器', '8', '串口总线', 'Modbus RTU RS485', '115200bps'],
        ['陀螺仪', '2', '串口总线', 'Modbus RTU RS485', '115200bps']
    ]
    
    for i, row_data in enumerate(data, 1):
        row_cells = table.rows[i].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
    
    doc.add_paragraph()
    
    doc.add_heading('4.2 EtherCAT总线', 2)
    doc.add_paragraph('EtherCAT用于运动控制，具有以下特点：', style='List Bullet')
    doc.add_paragraph('高实时性：周期时间≤1ms', style='List Bullet 2')
    doc.add_paragraph('高同步性：同步精度≤1μs', style='List Bullet 2')
    doc.add_paragraph('高可靠性：支持冗余拓扑', style='List Bullet 2')
    doc.add_paragraph('易扩展：支持最多65535个节点', style='List Bullet 2')
    
    doc.add_heading('4.3 GigE Vision', 2)
    doc.add_paragraph('工业相机采用GigE Vision标准，优势：', style='List Bullet')
    doc.add_paragraph('传输距离远：标准网线可达100m', style='List Bullet 2')
    doc.add_paragraph('带宽充足：1000Mbps满足多相机并发', style='List Bullet 2')
    doc.add_paragraph('成本低：使用标准以太网设备', style='List Bullet 2')
    doc.add_paragraph('兼容性好：符合工业相机标准', style='List Bullet 2')
    
    doc.add_heading('4.4 Modbus RS485', 2)
    doc.add_paragraph('测距传感器和陀螺仪采用Modbus RS485：', style='List Bullet')
    doc.add_paragraph('抗干扰强：差分信号，适合工业环境', style='List Bullet 2')
    doc.add_paragraph('传输距离：最长可达1200m', style='List Bullet 2')
    doc.add_paragraph('多设备：支持32个节点（可通过中继扩展）', style='List Bullet 2')
    doc.add_paragraph('成熟稳定：工业标准协议', style='List Bullet 2')
    
    doc.add_page_break()
    
    # 5. 机械结构
    doc.add_heading('5. 机械结构设计', 1)
    
    doc.add_heading('5.1 整机结构', 2)
    doc.add_paragraph('机器人整机由以下部分组成：', style='List Bullet')
    doc.add_paragraph('底盘：承载系统，安装车轮与驱动装置', style='List Bullet 2')
    doc.add_paragraph('主框架：铝型材或钣金结构，安装各传感器', style='List Bullet 2')
    doc.add_paragraph('防护罩：保护电控设备，防尘防水', style='List Bullet 2')
    doc.add_paragraph('操作面板：人机交互界面，状态指示灯', style='List Bullet 2')
    
    doc.add_heading('5.2 底盘设计', 2)
    doc.add_paragraph('底盘设计要点：', style='List Bullet')
    doc.add_paragraph('材质：高强度铝合金或钢材', style='List Bullet 2')
    doc.add_paragraph('车轮：金属轮，带橡胶缓冲层，适配轨道', style='List Bullet 2')
    doc.add_paragraph('驱动方式：四轮独立驱动或两轮驱动+两轮从动', style='List Bullet 2')
    doc.add_paragraph('承重：≥50kg（设备总重）', style='List Bullet 2')
    
    doc.add_heading('5.3 传感器安装', 2)
    doc.add_paragraph('传感器安装应满足以下要求：', style='List Bullet')
    doc.add_paragraph('稳定性：安装牢固，避免振动影响测量', style='List Bullet 2')
    doc.add_paragraph('可调性：支持位置微调，方便校准', style='List Bullet 2')
    doc.add_paragraph('防护性：加装防护罩，防止碰撞与污染', style='List Bullet 2')
    doc.add_paragraph('维护性：易拆卸，方便更换与维修', style='List Bullet 2')
    
    doc.add_heading('5.4 防护等级', 2)
    doc.add_paragraph('整机防护等级：≥IP54', style='List Bullet')
    doc.add_paragraph('电控箱防护等级：≥IP65', style='List Bullet')
    doc.add_paragraph('相机防护等级：≥IP65（带防护罩）', style='List Bullet')
    
    doc.add_page_break()
    
    # 6. 技术指标
    doc.add_heading('6. 系统技术指标', 1)
    
    table = doc.add_table(rows=16, cols=3)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '序号'
    hdr_cells[1].text = '技术参数'
    hdr_cells[2].text = '指标值'
    
    data = [
        ['1', '检测速度', '0-10 km/h，可调'],
        ['2', '工作速度', '≥5 km/h'],
        ['3', '轨距测量精度', '±0.5 mm'],
        ['4', '水平测量精度', '±0.5 mm'],
        ['5', '高低测量精度', '±0.5 mm'],
        ['6', '轨向测量精度', '±1 mm'],
        ['7', '钢轨磨耗测量精度', '±0.1 mm'],
        ['8', '轨面图像分辨率', '≥5 MP'],
        ['9', '螺栓图像分辨率', '≥2 MP'],
        ['10', '3D轮廓测量精度', '±0.05 mm'],
        ['11', '续航时间', '≥4小时（电池供电）'],
        ['12', '工作温度', '-20℃ ~ +50℃'],
        ['13', '存储温度', '-40℃ ~ +70℃'],
        ['14', '防护等级', 'IP54（整机）'],
        ['15', '整机重量', '≤50 kg']
    ]
    
    for i, row_data in enumerate(data, 1):
        row_cells = table.rows[i].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
    
    doc.add_page_break()
    
    # 7. 安全设计
    doc.add_heading('7. 安全设计', 1)
    
    doc.add_heading('7.1 电气安全', 2)
    doc.add_paragraph('系统设计遵循电气安全规范：', style='List Bullet')
    doc.add_paragraph('电源隔离：工控机、传感器、电机分别供电', style='List Bullet 2')
    doc.add_paragraph('过流保护：每路输出配置保险丝或断路器', style='List Bullet 2')
    doc.add_paragraph('接地保护：金属外壳可靠接地', style='List Bullet 2')
    doc.add_paragraph('漏电保护：配置漏电保护器', style='List Bullet 2')
    
    doc.add_heading('7.2 运动安全', 2)
    doc.add_paragraph('运动控制安全措施：', style='List Bullet')
    doc.add_paragraph('急停按钮：操作面板配置急停按钮', style='List Bullet 2')
    doc.add_paragraph('软件限位：软件监控运动范围', style='List Bullet 2')
    doc.add_paragraph('防撞检测：前置接近开关或碰撞传感器', style='List Bullet 2')
    doc.add_paragraph('故障停机：通信中断或传感器异常自动停机', style='List Bullet 2')
    
    doc.add_heading('7.3 数据安全', 2)
    doc.add_paragraph('数据采集与传输安全：', style='List Bullet')
    doc.add_paragraph('本地备份：数据实时存储到本地硬盘', style='List Bullet 2')
    doc.add_paragraph('断线续传：网络中断后自动重传未上传数据', style='List Bullet 2')
    doc.add_paragraph('数据加密：云端传输采用HTTPS/TLS加密', style='List Bullet 2')
    
    doc.add_page_break()
    
    # 8. 环境适应性
    doc.add_heading('8. 环境适应性', 1)
    
    doc.add_heading('8.1 温度适应', 2)
    doc.add_paragraph('系统适应宽温度范围：', style='List Bullet')
    doc.add_paragraph('工作温度：-20℃ ~ +50℃', style='List Bullet 2')
    doc.add_paragraph('存储温度：-40℃ ~ +70℃', style='List Bullet 2')
    doc.add_paragraph('热管理：工控机配置散热风扇', style='List Bullet 2')
    
    doc.add_heading('8.2 防尘防水', 2)
    doc.add_paragraph('防护措施：', style='List Bullet')
    doc.add_paragraph('整机防护等级：IP54', style='List Bullet 2')
    doc.add_paragraph('电控箱防护等级：IP65', style='List Bullet 2')
    doc.add_paragraph('相机保护：加装透明防护罩', style='List Bullet 2')
    doc.add_paragraph('接口保护：未使用接口加装防尘盖', style='List Bullet 2')
    
    doc.add_heading('8.3 抗振动', 2)
    doc.add_paragraph('振动防护：', style='List Bullet')
    doc.add_paragraph('减震设计：工控机、相机采用减震安装', style='List Bullet 2')
    doc.add_paragraph('固定方式：传感器采用锁紧螺钉固定', style='List Bullet 2')
    doc.add_paragraph('走线规范：线缆采用扎带固定，避免松动', style='List Bullet 2')
    
    doc.add_page_break()
    
    # 9. 扩展性设计
    doc.add_heading('9. 扩展性设计', 1)
    
    doc.add_heading('9.1 硬件扩展', 2)
    doc.add_paragraph('系统预留扩展接口：', style='List Bullet')
    doc.add_paragraph('相机扩展：网口预留2路，可增加2个相机', style='List Bullet 2')
    doc.add_paragraph('传感器扩展：RS485总线预留4个地址', style='List Bullet 2')
    doc.add_paragraph('USB扩展：预留2个USB 3.0接口', style='List Bullet 2')
    doc.add_paragraph('电源扩展：电源模块预留20%功率裕量', style='List Bullet 2')
    
    doc.add_heading('9.2 功能扩展', 2)
    doc.add_paragraph('可扩展的功能模块：', style='List Bullet')
    doc.add_paragraph('GPS定位：增加GPS模块，实现精确定位', style='List Bullet 2')
    doc.add_paragraph('超声检测：增加超声探头，检测内部缺陷', style='List Bullet 2')
    doc.add_paragraph('红外热成像：增加热成像相机，检测温度异常', style='List Bullet 2')
    doc.add_paragraph('环境监测：增加温湿度、噪声传感器', style='List Bullet 2')
    
    doc.add_page_break()
    
    # 10. 维护与保养
    doc.add_heading('10. 维护与保养', 1)
    
    doc.add_heading('10.1 日常维护', 2)
    doc.add_paragraph('每次使用后应进行以下检查：', style='List Bullet')
    doc.add_paragraph('清洁：清理传感器表面灰尘', style='List Bullet 2')
    doc.add_paragraph('紧固：检查螺钉是否松动', style='List Bullet 2')
    doc.add_paragraph('充电：电池供电模式下及时充电', style='List Bullet 2')
    doc.add_paragraph('数据备份：导出检测数据并备份', style='List Bullet 2')
    
    doc.add_heading('10.2 定期保养', 2)
    doc.add_paragraph('每月进行以下保养：', style='List Bullet')
    doc.add_paragraph('相机校准：检查相机成像质量，必要时重新标定', style='List Bullet 2')
    doc.add_paragraph('激光校准：检查3D激光测量精度', style='List Bullet 2')
    doc.add_paragraph('传感器校准：零位校准测距传感器和陀螺仪', style='List Bullet 2')
    doc.add_paragraph('软件更新：检查并更新系统软件', style='List Bullet 2')
    
    doc.add_heading('10.3 易损件更换', 2)
    doc.add_paragraph('以下部件为易损件，应定期检查并更换：', style='List Bullet')
    doc.add_paragraph('车轮：磨损严重时更换', style='List Bullet 2')
    doc.add_paragraph('光源：LED光源衰减时更换', style='List Bullet 2')
    doc.add_paragraph('风扇：工控机散热风扇异响时更换', style='List Bullet 2')
    doc.add_paragraph('电池：容量下降至80%以下时更换', style='List Bullet 2')
    
    doc.add_page_break()
    
    # 附录
    doc.add_heading('附录A：缩略语', 1)
    
    table = doc.add_table(rows=11, cols=2)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '缩略语'
    hdr_cells[1].text = '英文全称/中文说明'
    
    data = [
        ['AI', 'Artificial Intelligence / 人工智能'],
        ['CAM', 'Camera / 相机'],
        ['CMOS', 'Complementary Metal-Oxide-Semiconductor / 互补金属氧化物半导体'],
        ['CCD', 'Charge-Coupled Device / 电荷耦合器件'],
        ['EtherCAT', 'Ethernet for Control Automation Technology / 控制自动化技术以太网'],
        ['GigE', 'Gigabit Ethernet / 千兆以太网'],
        ['GPS', 'Global Positioning System / 全球定位系统'],
        ['IP', 'Ingress Protection / 防护等级'],
        ['LED', 'Light-Emitting Diode / 发光二极管'],
        ['RS485', 'Recommended Standard 485 / 推荐标准485（串行通信接口）']
    ]
    
    for i, row_data in enumerate(data, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = row_data[0]
        row_cells[1].text = row_data[1]
    
    # 保存文档
    doc.save('D:\\铁路线路智能检测机器人\\04-项目文档\\设计文档\\硬件系统架构设计文档.docx')
    print('[OK] 硬件系统架构设计文档.docx 已创建 (专业版)')

if __name__ == '__main__':
    create_hardware_architecture_doc()
