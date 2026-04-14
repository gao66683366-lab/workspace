# -*- coding: utf-8 -*-
"""
完整的硬件系统架构设计文档 - 全表格化版本
作者：小测
日期：2026-03-05
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def setup_styles(doc):
    """设置文档样式"""
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)
    
    for i in range(1, 4):
        heading = doc.styles[f'Heading {i}']
        heading.font.name = '黑体'
        heading._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        heading.font.color.rgb = RGBColor(0, 0, 0)
        if i == 1:
            heading.font.size = Pt(16)
        elif i == 2:
            heading.font.size = Pt(14)
        else:
            heading.font.size = Pt(13)

def add_cover(doc):
    """添加封面"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('\n\n\n\n\n')
    
    title = p.add_run('铁路线路智能检测机器人')
    title.font.name = '黑体'
    title._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    title.font.size = Pt(26)
    title.font.bold = True
    
    p.add_run('\n\n')
    
    subtitle = p.add_run('硬件系统架构设计文档')
    subtitle.font.name = '黑体'
    subtitle._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    subtitle.font.size = Pt(20)
    
    p.add_run('\n\n\n\n\n\n\n\n\n\n')
    
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info.add_run('编制日期：2026年3月5日\n版本：V1.0\n密级：内部')
    info_run.font.size = Pt(12)
    
    doc.add_page_break()

def create_table(doc, headers, data, col_widths=None):
    """创建表格"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)
    
    # 表头
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = header
        hdr_cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in hdr_cells[idx].paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(11)
    
    # 数据行
    for row_data in data:
        row = table.add_row()
        for idx, cell_data in enumerate(row_data):
            row.cells[idx].text = str(cell_data)
            row.cells[idx].paragraphs[0].runs[0].font.size = Pt(10.5)
    
    doc.add_paragraph()
    return table

# 创建文档
print('[1/12] 初始化文档...')
doc = Document()
setup_styles(doc)

# 封面
print('[2/12] 生成封面...')
add_cover(doc)

# 文档修订历史
print('[3/12] 生成修订历史...')
doc.add_heading('文档修订历史', 1)
create_table(doc, 
    ['版本', '日期', '修订内容', '修订人', '审核人'],
    [['V1.0', '2026-03-05', '初始版本，完成硬件系统详细设计', '小测', '道']],
    [2, 3, 6, 2, 2]
)

doc.add_heading('目录', 1)
doc.add_paragraph('【此处应插入自动目录，在Word中通过"引用→目录→自动目录"功能生成】')
doc.add_page_break()

# 第1章：概述
print('[4/12] 第1章：概述...')
doc.add_heading('1. 概述', 1)
doc.add_heading('1.1 系统简介', 2)
doc.add_paragraph(
    '铁路线路智能检测机器人是一款集机械、电子、光学、计算机技术于一体的智能检测装备，'
    '用于铁路轨道及相关设施的自动化巡检。系统采用多传感器融合技术，实现对钢轨轨面状态、'
    '几何参数、紧固件状态等关键指标的实时检测与分析。'
)

doc.add_heading('1.2 设计目标', 2)
create_table(doc,
    ['目标类别', '具体指标', '说明'],
    [
        ['可靠性', '系统故障率 < 0.1%', '满足工业级设备标准，连续运行时间≥8小时'],
        ['检测精度', '毫米级', '轨距/水平/高低±0.5mm，轨向±1mm，满足TB标准'],
        ['检测效率', '检测速度 ≥ 5km/h', '提升巡检效率，单日完成≥40km线路检测'],
        ['智能化', 'AI自动识别准确率≥95%', '集成深度学习算法，自动识别螺栓/轨面缺陷'],
        ['模块化', '模块化设计', '传感器、控制器可独立更换，维护时间<30分钟'],
        ['环境适应', '工作温度-20℃~+50℃', '适应户外全天候作业环境']
    ],
    [3, 4, 8]
)

doc.add_heading('1.3 适用范围', 2)
doc.add_paragraph('本文档适用于铁路线路智能检测机器人硬件系统的设计、开发、采购、测试、安装、维护等全生命周期工作。')

doc.add_heading('1.4 参考标准', 2)
create_table(doc,
    ['标准编号', '标准名称', '适用范围'],
    [
        ['TB/T 3147-2012', '铁路轨道几何尺寸测量仪', '几何参数测量精度要求'],
        ['GB/T 2423.1-2008', '电工电子产品环境试验', '环境适应性测试标准'],
        ['GB/T 17626-2008', '电磁兼容 试验和测量技术', 'EMC测试要求'],
        ['IEC 61000-6-2', '工业环境的抗扰度标准', '工业环境电磁兼容'],
        ['GigE Vision 2.0', '千兆以太网工业相机标准', '工业相机通信协议']
    ],
    [4, 6, 5]
)

doc.add_page_break()

# 第2章：系统总体架构
print('[5/12] 第2章：系统总体架构...')
doc.add_heading('2. 系统总体架构', 1)
doc.add_heading('2.1 系统定位', 2)
doc.add_paragraph(
    '本系统为软硬件一体化智能检测平台，通过搭载多种传感器在铁路轨道上自主运行，'
    '实时采集轨道状态数据，结合AI算法进行在线分析与判断，及时发现安全隐患。'
)

doc.add_heading('2.2 系统组成', 2)
create_table(doc,
    ['子系统', '主要功能', '核心设备', '通信方式', '数据类型'],
    [
        ['运动控制\n子系统', '机器人运动控制\n速度调节\n位置反馈', '伺服电机×4\n运动控制器\n编码器×4', 'EtherCAT\n100Mbps', '控制指令\n位置反馈'],
        ['视觉检测\n子系统', '轨面状态检测\n螺栓完好性识别', '工业相机×6\nLED光源×6\n镜头×6', 'GigE Vision\n1000Mbps', '图像数据\n900MB/s'],
        ['三维测量\n子系统', '钢轨轮廓扫描\n磨耗测量', '3D线激光×2\n采集卡', 'TCP/IP\n1000Mbps', '点云数据\n49MB/s'],
        ['几何参数\n测量子系统', '轨距/水平/高低\n轨向测量', '测距传感器×8\n陀螺仪×2', 'Modbus RS485\n115200bps', '传感器数据\n<1KB/s'],
        ['控制处理\n子系统', '数据采集处理\n本地AI推理\n人机交互', '工控机\n存储设备\n显示器', '内部总线', '所有数据流']
    ],
    [3, 3.5, 3.5, 3, 3]
)

doc.add_heading('2.3 系统层级架构', 2)
create_table(doc,
    ['层级', '组成', '主要功能', '开发语言/平台', '部署位置'],
    [
        ['云端层', '云端服务器', '海量数据存储\nAI模型训练\n大数据分析\nWeb管理平台', 'Python\nPostgreSQL\nPyTorch/TensorFlow', '云端机房\n或私有云'],
        ['边缘层', '工控上位机', '实时数据采集\n本地AI推理\n运动控制\n人机交互', 'C# .NET 8.0\nWindows 10/11\nONNX Runtime', '机器人本体\n工控机箱内'],
        ['设备层', '传感器\n执行器', '数据采集\n运动执行\n状态反馈', '嵌入式固件', '机器人\n各安装位置']
    ],
    [2.5, 3, 4.5, 4, 3]
)

doc.add_heading('2.4 数据流向', 2)
create_table(doc,
    ['数据源', '数据类型', '流向', '数据量', '处理方式'],
    [
        ['工业相机×6', '图像（JPEG）', '设备层→边缘层', '压缩后≈100MB/s', '本地AI推理+存储'],
        ['3D线激光×2', '点云（XYZ）', '设备层→边缘层', '≈50MB/s', '实时轮廓计算'],
        ['测距传感器×8', '距离值', '设备层→边缘层', '<100KB/s', '几何参数计算'],
        ['陀螺仪×2', '姿态角度', '设备层→边缘层', '<50KB/s', '姿态解算'],
        ['检测结果', '结构化数据', '边缘层→云端层', '≈1MB/分钟', '上传存储+分析'],
        ['AI模型', '模型文件', '云端层→边缘层', '按需下载', '模型更新']
    ],
    [3, 2.5, 3, 2.5, 4.5]
)

doc.add_page_break()

# 第3章继续...由于代码很长，分批生成
print('[6/12] 第3章：硬件详细设计（第1部分）...')
doc.add_heading('3. 硬件详细设计', 1)

# 3.1 运动控制子系统
doc.add_heading('3.1 运动控制子系统', 2)
doc.add_heading('3.1.1 系统组成', 3)
create_table(doc,
    ['设备', '数量', '型号/规格', '关键参数', '接口', '功能'],
    [
        ['伺服电机', '4', '额定200W\n无刷永磁', '额定转速：3000rpm\n额定扭矩：0.64Nm\n过载能力：200%', '编码器\nA/B/Z相', '驱动车轮\n提供动力'],
        ['伺服驱动器', '4', 'AC220V输入\n或DC48V', '输出电流：2A\n控制模式：位置/速度/扭矩\n响应频率：2kHz', 'EtherCAT\n编码器', '接收指令\n驱动电机'],
        ['运动控制器', '1', 'EtherCAT主站\n4轴控制', '控制周期：1ms\n同步精度：±1μs\n内存：512MB', 'EtherCAT\n以太网', '协调运动\n轨迹规划'],
        ['编码器', '4', '增量式\n2500线', '分辨率：2500脉冲/转\n输出：A/B/Z相\n电源：5V', '差分输出', '位置反馈\n速度检测']
    ],
    [2.5, 1.5, 2.5, 5, 2, 2.5]
)

doc.add_heading('3.1.2 EtherCAT通信参数', 3)
create_table(doc,
    ['参数项', '参数值', '说明', '技术优势'],
    [
        ['通信周期', '1ms', '每毫秒更新一次', '满足伺服控制实时性要求'],
        ['同步精度', '±1μs', '各轴同步误差', '保证多轴协调运动精度'],
        ['拓扑结构', '线型/星型/树型', '灵活组网', '支持热插拔，易扩展'],
        ['传输距离', '100m/段', '单段标准距离', '通过中继器可扩展至数公里'],
        ['节点数量', '最多65535', '可接入设备数', '扩展性极强'],
        ['传输速率', '100Mbps', '全双工', '带宽充足，延迟低']
    ],
    [3, 3, 4, 5.5]
)

# 继续后续章节...
# 保存进度
doc.save('D:\\铁路线路智能检测机器人\\04-项目文档\\设计文档\\机器人硬件系统架构设计文档 V1.0-NEW.docx')
print('[Progress] 已完成前3章，继续生成...')
