#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
铁路线路智能检测机器人专业论文提纲生成器
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_chinese_font(run):
    """设置中文字体"""
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_heading_custom(doc, text, level):
    """添加自定义标题"""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    set_chinese_font(run)
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(16)
    elif level == 2:
        run.font.size = Pt(14)
    elif level == 3:
        run.font.size = Pt(12)
    return heading

def add_paragraph_custom(doc, text, indent=0):
    """添加自定义段落"""
    para = doc.add_paragraph()
    if indent > 0:
        para.paragraph_format.left_indent = Inches(indent * 0.3)
    run = para.add_run(text)
    set_chinese_font(run)
    run.font.size = Pt(12)
    return para

def create_paper_outline():
    """创建论文提纲"""
    doc = Document()
    
    # 设置页面
    sections = doc.sections
    for section in sections:
        section.page_height = Inches(11.69)  # A4高度
        section.page_width = Inches(8.27)    # A4宽度
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.18)
        section.right_margin = Inches(1.18)
    
    # 标题
    title = doc.add_heading(level=0)
    title_run = title.add_run('基于智能视觉与多传感器融合的铁路线路检测机器人系统研究')
    set_chinese_font(title_run)
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 副标题/作者信息占位
    author_para = doc.add_paragraph()
    author_run = author_para.add_run('作者姓名¹、指导教师²')
    set_chinese_font(author_run)
    author_run.font.size = Pt(12)
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    affiliation_para = doc.add_paragraph()
    affiliation_run = affiliation_para.add_run('（单位名称，城市 邮编）')
    set_chinese_font(affiliation_run)
    affiliation_run.font.size = Pt(10.5)
    affiliation_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # 空行
    
    # ========== 摘要 ==========
    add_heading_custom(doc, '摘要', 2)
    add_paragraph_custom(doc, 
        '【摘要内容概述】本文针对铁路线路检测自动化、智能化需求，设计并实现了一套基于智能视觉与多传感器融合的铁路线路检测机器人系统。'
        '系统采用分层架构设计，集成高精度视觉检测、激光雷达测距、北斗/GPS定位等多源传感器，'
        '实现对铁路轨道几何参数、表面缺陷、周边环境的自动化检测与智能分析。'
        '论文详细阐述了系统硬件架构、软件设计、关键算法及工程实现，并通过实验验证了系统的有效性与可靠性。'
        '研究成果对提升铁路线路检测效率、保障运营安全具有重要意义。')
    
    add_paragraph_custom(doc, '【关键词】铁路线路检测；智能机器人；多传感器融合；机器视觉；故障诊断')
    
    doc.add_paragraph()
    
    # ========== 正文提纲 ==========
    add_heading_custom(doc, '1  引言', 1)
    
    add_heading_custom(doc, '1.1  研究背景与意义', 2)
    add_paragraph_custom(doc, '• 铁路运输安全重要性', 1)
    add_paragraph_custom(doc, '• 传统线路检测方式的局限性', 1)
    add_paragraph_custom(doc, '• 智能检测技术发展趋势', 1)
    add_paragraph_custom(doc, '• 本研究的实际应用价值', 1)
    
    add_heading_custom(doc, '1.2  国内外研究现状', 2)
    add_paragraph_custom(doc, '• 国外铁路检测机器人技术发展', 1)
    add_paragraph_custom(doc, '• 国内铁路检测装备现状', 1)
    add_paragraph_custom(doc, '• 多传感器融合技术应用', 1)
    add_paragraph_custom(doc, '• 机器视觉在轨道检测中的应用', 1)
    add_paragraph_custom(doc, '• 现有技术存在的问题与挑战', 1)
    
    add_heading_custom(doc, '1.3  论文主要研究内容', 2)
    add_paragraph_custom(doc, '• 系统总体架构设计', 1)
    add_paragraph_custom(doc, '• 硬件平台开发', 1)
    add_paragraph_custom(doc, '• 软件系统实现', 1)
    add_paragraph_custom(doc, '• 关键算法研究', 1)
    add_paragraph_custom(doc, '• 系统集成与测试', 1)
    
    add_heading_custom(doc, '1.4  论文组织结构', 2)
    add_paragraph_custom(doc, '【简述各章节内容安排】', 1)
    
    doc.add_page_break()
    
    # ========== 第2章 ==========
    add_heading_custom(doc, '2  系统总体设计', 1)
    
    add_heading_custom(doc, '2.1  系统需求分析', 2)
    add_heading_custom(doc, '2.1.1  功能需求', 3)
    add_paragraph_custom(doc, '• 轨道几何参数检测需求', 1)
    add_paragraph_custom(doc, '• 表面缺陷识别需求', 1)
    add_paragraph_custom(doc, '• 环境感知需求', 1)
    add_paragraph_custom(doc, '• 数据采集与存储需求', 1)
    add_paragraph_custom(doc, '• 实时分析与报警需求', 1)
    
    add_heading_custom(doc, '2.1.2  性能需求', 3)
    add_paragraph_custom(doc, '• 检测精度指标', 1)
    add_paragraph_custom(doc, '• 检测速度要求', 1)
    add_paragraph_custom(doc, '• 系统可靠性要求', 1)
    add_paragraph_custom(doc, '• 环境适应性要求', 1)
    
    add_heading_custom(doc, '2.2  系统总体架构', 2)
    add_heading_custom(doc, '2.2.1  分层架构设计', 3)
    add_paragraph_custom(doc, '• 感知层：传感器网络', 1)
    add_paragraph_custom(doc, '• 控制层：嵌入式控制系统', 1)
    add_paragraph_custom(doc, '• 应用层：数据处理与分析', 1)
    add_paragraph_custom(doc, '• 各层次功能与接口定义', 1)
    
    add_heading_custom(doc, '2.2.2  系统工作流程', 3)
    add_paragraph_custom(doc, '• 数据采集流程', 1)
    add_paragraph_custom(doc, '• 数据处理流程', 1)
    add_paragraph_custom(doc, '• 结果输出流程', 1)
    add_paragraph_custom(doc, '【配流程图】', 1)
    
    add_heading_custom(doc, '2.3  关键技术路线', 2)
    add_paragraph_custom(doc, '• 多传感器数据融合技术', 1)
    add_paragraph_custom(doc, '• 深度学习缺陷识别技术', 1)
    add_paragraph_custom(doc, '• 实时数据处理技术', 1)
    add_paragraph_custom(doc, '• 高精度定位技术', 1)
    
    doc.add_page_break()
    
    # ========== 第3章 ==========
    add_heading_custom(doc, '3  硬件系统设计与实现', 1)
    
    add_heading_custom(doc, '3.1  机械平台设计', 2)
    add_heading_custom(doc, '3.1.1  机器人本体结构', 3)
    add_paragraph_custom(doc, '• 底盘设计与材料选择', 1)
    add_paragraph_custom(doc, '• 行走机构设计', 1)
    add_paragraph_custom(doc, '• 载荷分析与强度校核', 1)
    
    add_heading_custom(doc, '3.1.2  传感器安装方案', 3)
    add_paragraph_custom(doc, '• 相机安装位置与角度', 1)
    add_paragraph_custom(doc, '• 激光雷达布置方案', 1)
    add_paragraph_custom(doc, '• 减震与防护设计', 1)
    
    add_heading_custom(doc, '3.2  供电系统设计', 2)
    add_paragraph_custom(doc, '• DC 48V 供电方案', 1)
    add_paragraph_custom(doc, '• 电源管理模块设计', 1)
    add_paragraph_custom(doc, '• 多路电源转换电路', 1)
    add_paragraph_custom(doc, '• 电源保护与监控', 1)
    
    add_heading_custom(doc, '3.3  控制系统硬件', 2)
    add_heading_custom(doc, '3.3.1  主控制器选型', 3)
    add_paragraph_custom(doc, '• 工控机配置方案', 1)
    add_paragraph_custom(doc, '• 嵌入式处理器选择', 1)
    add_paragraph_custom(doc, '• 性能与成本分析', 1)
    
    add_heading_custom(doc, '3.3.2  运动控制模块', 3)
    add_paragraph_custom(doc, '• 电机驱动器选型', 1)
    add_paragraph_custom(doc, '• 编码器反馈机制', 1)
    add_paragraph_custom(doc, '• 控制电路设计', 1)
    
    add_heading_custom(doc, '3.4  传感器系统', 2)
    add_heading_custom(doc, '3.4.1  视觉传感器', 3)
    add_paragraph_custom(doc, '• 工业相机选型（分辨率、帧率、接口）', 1)
    add_paragraph_custom(doc, '• 镜头参数设计', 1)
    add_paragraph_custom(doc, '• 照明系统设计', 1)
    
    add_heading_custom(doc, '3.4.2  激光雷达', 3)
    add_paragraph_custom(doc, '• 2D/3D 激光雷达选型', 1)
    add_paragraph_custom(doc, '• 测距精度与扫描频率', 1)
    add_paragraph_custom(doc, '• 数据接口与通信协议', 1)
    
    add_heading_custom(doc, '3.4.3  定位与导航传感器', 3)
    add_paragraph_custom(doc, '• 北斗/GPS 模块', 1)
    add_paragraph_custom(doc, '• IMU 惯性测量单元', 1)
    add_paragraph_custom(doc, '• 里程计编码器', 1)
    
    add_heading_custom(doc, '3.4.4  环境传感器', 3)
    add_paragraph_custom(doc, '• 温湿度传感器', 1)
    add_paragraph_custom(doc, '• 倾角传感器', 1)
    add_paragraph_custom(doc, '• 其他辅助传感器', 1)
    
    add_heading_custom(doc, '3.5  通信与数据接口', 2)
    add_paragraph_custom(doc, '• 4G/5G 无线通信模块', 1)
    add_paragraph_custom(doc, '• 工业以太网设计', 1)
    add_paragraph_custom(doc, '• CAN 总线通信', 1)
    add_paragraph_custom(doc, '• USB/串口扩展', 1)
    
    doc.add_page_break()
    
    # ========== 第4章 ==========
    add_heading_custom(doc, '4  软件系统设计与实现', 1)
    
    add_heading_custom(doc, '4.1  软件架构设计', 2)
    add_heading_custom(doc, '4.1.1  分层软件架构', 3)
    add_paragraph_custom(doc, '• 硬件抽象层（HAL）', 1)
    add_paragraph_custom(doc, '• 中间件层', 1)
    add_paragraph_custom(doc, '• 应用层', 1)
    add_paragraph_custom(doc, '【配软件架构图】', 1)
    
    add_heading_custom(doc, '4.1.2  模块化设计', 3)
    add_paragraph_custom(doc, '• 各功能模块划分', 1)
    add_paragraph_custom(doc, '• 模块间接口定义', 1)
    add_paragraph_custom(doc, '• 数据流设计', 1)
    
    add_heading_custom(doc, '4.2  数据采集模块', 2)
    add_heading_custom(doc, '4.2.1  图像采集', 3)
    add_paragraph_custom(doc, '• 相机驱动程序开发', 1)
    add_paragraph_custom(doc, '• 图像缓存管理', 1)
    add_paragraph_custom(doc, '• 多相机同步触发', 1)
    
    add_heading_custom(doc, '4.2.2  激光雷达数据采集', 3)
    add_paragraph_custom(doc, '• 点云数据获取', 1)
    add_paragraph_custom(doc, '• 数据预处理', 1)
    add_paragraph_custom(doc, '• 点云配准与拼接', 1)
    
    add_heading_custom(doc, '4.2.3  多传感器时间同步', 3)
    add_paragraph_custom(doc, '• 时间戳标定', 1)
    add_paragraph_custom(doc, '• 数据同步策略', 1)
    
    add_heading_custom(doc, '4.3  数据处理模块', 2)
    add_heading_custom(doc, '4.3.1  图像预处理', 3)
    add_paragraph_custom(doc, '• 图像去噪算法', 1)
    add_paragraph_custom(doc, '• 畸变校正', 1)
    add_paragraph_custom(doc, '• 图像增强技术', 1)
    
    add_heading_custom(doc, '4.3.2  特征提取', 3)
    add_paragraph_custom(doc, '• 轨道区域分割', 1)
    add_paragraph_custom(doc, '• 边缘检测算法', 1)
    add_paragraph_custom(doc, '• 关键点检测', 1)
    
    add_heading_custom(doc, '4.4  智能分析模块', 2)
    add_heading_custom(doc, '4.4.1  缺陷检测算法', 3)
    add_paragraph_custom(doc, '• 深度学习模型设计（YOLO/Faster R-CNN等）', 1)
    add_paragraph_custom(doc, '• 训练数据集构建', 1)
    add_paragraph_custom(doc, '• 模型训练与优化', 1)
    add_paragraph_custom(doc, '• 检测精度评估', 1)
    
    add_heading_custom(doc, '4.4.2  几何参数测量', 3)
    add_paragraph_custom(doc, '• 轨距测量算法', 1)
    add_paragraph_custom(doc, '• 高低、方向不平顺检测', 1)
    add_paragraph_custom(doc, '• 测量精度分析', 1)
    
    add_heading_custom(doc, '4.4.3  多传感器数据融合', 3)
    add_paragraph_custom(doc, '• 卡尔曼滤波融合算法', 1)
    add_paragraph_custom(doc, '• 贝叶斯决策融合', 1)
    add_paragraph_custom(doc, '• 融合结果可靠性评估', 1)
    
    add_heading_custom(doc, '4.5  定位导航模块', 2)
    add_paragraph_custom(doc, '• GNSS/IMU 组合定位算法', 1)
    add_paragraph_custom(doc, '• 里程计辅助定位', 1)
    add_paragraph_custom(doc, '• 定位精度优化', 1)
    
    add_heading_custom(doc, '4.6  故障诊断与预警模块', 2)
    add_paragraph_custom(doc, '• 故障模式识别', 1)
    add_paragraph_custom(doc, '• 风险评估算法', 1)
    add_paragraph_custom(doc, '• 多级报警机制', 1)
    
    add_heading_custom(doc, '4.7  数据管理模块', 2)
    add_paragraph_custom(doc, '• 数据库设计（MySQL/PostgreSQL）', 1)
    add_paragraph_custom(doc, '• 数据存储策略', 1)
    add_paragraph_custom(doc, '• 数据检索与查询', 1)
    
    add_heading_custom(doc, '4.8  人机交互界面', 2)
    add_paragraph_custom(doc, '• 上位机软件设计（Qt/Web）', 1)
    add_paragraph_custom(doc, '• 实时监控界面', 1)
    add_paragraph_custom(doc, '• 数据可视化', 1)
    add_paragraph_custom(doc, '• 报表生成功能', 1)
    
    doc.add_page_break()
    
    # ========== 第5章 ==========
    add_heading_custom(doc, '5  关键算法研究', 1)
    
    add_heading_custom(doc, '5.1  基于深度学习的缺陷识别算法', 2)
    add_heading_custom(doc, '5.1.1  算法原理', 3)
    add_paragraph_custom(doc, '• 卷积神经网络基础', 1)
    add_paragraph_custom(doc, '• 目标检测算法演进', 1)
    add_paragraph_custom(doc, '• 本文采用的网络结构', 1)
    
    add_heading_custom(doc, '5.1.2  模型训练', 3)
    add_paragraph_custom(doc, '• 数据标注与预处理', 1)
    add_paragraph_custom(doc, '• 数据增强策略', 1)
    add_paragraph_custom(doc, '• 超参数调优', 1)
    add_paragraph_custom(doc, '• 损失函数设计', 1)
    
    add_heading_custom(doc, '5.1.3  模型优化与部署', 3)
    add_paragraph_custom(doc, '• 模型压缩技术（剪枝、量化）', 1)
    add_paragraph_custom(doc, '• 推理加速（TensorRT/OpenVINO）', 1)
    add_paragraph_custom(doc, '• 嵌入式平台部署', 1)
    
    add_heading_custom(doc, '5.2  多传感器数据融合算法', 2)
    add_heading_custom(doc, '5.2.1  融合框架设计', 3)
    add_paragraph_custom(doc, '• 数据层融合', 1)
    add_paragraph_custom(doc, '• 特征层融合', 1)
    add_paragraph_custom(doc, '• 决策层融合', 1)
    
    add_heading_custom(doc, '5.2.2  卡尔曼滤波算法', 3)
    add_paragraph_custom(doc, '• 算法原理与推导', 1)
    add_paragraph_custom(doc, '• 状态估计模型', 1)
    add_paragraph_custom(doc, '• 实验验证', 1)
    
    add_heading_custom(doc, '5.3  轨道几何参数测量算法', 2)
    add_heading_custom(doc, '5.3.1  轨距测量', 3)
    add_paragraph_custom(doc, '• 图像坐标到实际坐标转换', 1)
    add_paragraph_custom(doc, '• 亚像素精度边缘检测', 1)
    add_paragraph_custom(doc, '• 误差分析与校正', 1)
    
    add_heading_custom(doc, '5.3.2  不平顺检测', 3)
    add_paragraph_custom(doc, '• 高低不平顺算法', 1)
    add_paragraph_custom(doc, '• 方向不平顺算法', 1)
    add_paragraph_custom(doc, '• 算法精度验证', 1)
    
    add_heading_custom(doc, '5.4  实时处理优化算法', 2)
    add_paragraph_custom(doc, '• 多线程并行处理', 1)
    add_paragraph_custom(doc, '• GPU 加速计算', 1)
    add_paragraph_custom(doc, '• 算法时间复杂度分析', 1)
    
    doc.add_page_break()
    
    # ========== 第6章 ==========
    add_heading_custom(doc, '6  系统集成与测试', 1)
    
    add_heading_custom(doc, '6.1  系统集成', 2)
    add_heading_custom(doc, '6.1.1  硬件集成', 3)
    add_paragraph_custom(doc, '• 模块装配与调试', 1)
    add_paragraph_custom(doc, '• 接口测试', 1)
    add_paragraph_custom(doc, '• 系统联调', 1)
    
    add_heading_custom(doc, '6.1.2  软件集成', 3)
    add_paragraph_custom(doc, '• 模块联合调试', 1)
    add_paragraph_custom(doc, '• 系统稳定性测试', 1)
    add_paragraph_custom(doc, '• 异常处理机制', 1)
    
    add_heading_custom(doc, '6.2  功能测试', 2)
    add_heading_custom(doc, '6.2.1  图像采集测试', 3)
    add_paragraph_custom(doc, '• 不同光照条件测试', 1)
    add_paragraph_custom(doc, '• 图像质量评估', 1)
    
    add_heading_custom(doc, '6.2.2  缺陷检测测试', 3)
    add_paragraph_custom(doc, '• 典型缺陷识别准确率', 1)
    add_paragraph_custom(doc, '• 误检率与漏检率', 1)
    add_paragraph_custom(doc, '• 不同缺陷类型测试', 1)
    
    add_heading_custom(doc, '6.2.3  几何参数测量测试', 3)
    add_paragraph_custom(doc, '• 轨距测量精度验证', 1)
    add_paragraph_custom(doc, '• 与标准设备对比', 1)
    
    add_heading_custom(doc, '6.3  性能测试', 2)
    add_paragraph_custom(doc, '• 检测速度测试', 1)
    add_paragraph_custom(doc, '• 数据处理实时性', 1)
    add_paragraph_custom(doc, '• 系统资源占用分析', 1)
    add_paragraph_custom(doc, '• 长时间运行稳定性', 1)
    
    add_heading_custom(doc, '6.4  环境适应性测试', 2)
    add_paragraph_custom(doc, '• 不同天气条件测试', 1)
    add_paragraph_custom(doc, '• 温度适应性测试', 1)
    add_paragraph_custom(doc, '• 抗振动与冲击测试', 1)
    
    add_heading_custom(doc, '6.5  现场试验', 2)
    add_heading_custom(doc, '6.5.1  试验方案', 3)
    add_paragraph_custom(doc, '• 试验线路选择', 1)
    add_paragraph_custom(doc, '• 试验方案设计', 1)
    add_paragraph_custom(doc, '• 数据采集计划', 1)
    
    add_heading_custom(doc, '6.5.2  试验结果', 3)
    add_paragraph_custom(doc, '• 检测数据统计', 1)
    add_paragraph_custom(doc, '• 典型案例分析', 1)
    add_paragraph_custom(doc, '• 与传统方法对比', 1)
    
    add_heading_custom(doc, '6.5.3  问题分析与改进', 3)
    add_paragraph_custom(doc, '• 试验中发现的问题', 1)
    add_paragraph_custom(doc, '• 改进措施', 1)
    add_paragraph_custom(doc, '• 效果验证', 1)
    
    doc.add_page_break()
    
    # ========== 第7章 ==========
    add_heading_custom(doc, '7  结果分析与讨论', 1)
    
    add_heading_custom(doc, '7.1  系统性能评估', 2)
    add_paragraph_custom(doc, '• 检测精度达标情况', 1)
    add_paragraph_custom(doc, '• 检测效率提升对比', 1)
    add_paragraph_custom(doc, '• 系统可靠性评估', 1)
    
    add_heading_custom(doc, '7.2  关键技术创新点', 2)
    add_paragraph_custom(doc, '• 多传感器融合方法创新', 1)
    add_paragraph_custom(doc, '• 深度学习模型优化', 1)
    add_paragraph_custom(doc, '• 实时处理技术突破', 1)
    
    add_heading_custom(doc, '7.3  系统优势分析', 2)
    add_paragraph_custom(doc, '• 与现有技术对比', 1)
    add_paragraph_custom(doc, '• 成本效益分析', 1)
    add_paragraph_custom(doc, '• 应用前景', 1)
    
    add_heading_custom(doc, '7.4  存在的问题与不足', 2)
    add_paragraph_custom(doc, '• 极端环境适应性', 1)
    add_paragraph_custom(doc, '• 复杂场景处理能力', 1)
    add_paragraph_custom(doc, '• 系统成本控制', 1)
    
    add_heading_custom(doc, '7.5  改进方向', 2)
    add_paragraph_custom(doc, '• 算法优化方向', 1)
    add_paragraph_custom(doc, '• 硬件升级方案', 1)
    add_paragraph_custom(doc, '• 功能扩展建议', 1)
    
    doc.add_page_break()
    
    # ========== 第8章 ==========
    add_heading_custom(doc, '8  结论与展望', 1)
    
    add_heading_custom(doc, '8.1  主要工作总结', 2)
    add_paragraph_custom(doc, '• 完成的主要工作', 1)
    add_paragraph_custom(doc, '• 实现的功能与指标', 1)
    add_paragraph_custom(doc, '• 创新点总结', 1)
    
    add_heading_custom(doc, '8.2  主要结论', 2)
    add_paragraph_custom(doc, '【总结研究成果与贡献】', 1)
    
    add_heading_custom(doc, '8.3  研究展望', 2)
    add_paragraph_custom(doc, '• 技术发展趋势', 1)
    add_paragraph_custom(doc, '• 下一步研究方向', 1)
    add_paragraph_custom(doc, '• 应用推广建议', 1)
    
    doc.add_paragraph()
    
    # ========== 参考文献 ==========
    add_heading_custom(doc, '参考文献', 1)
    add_paragraph_custom(doc, '[1] 作者. 文献标题[J]. 期刊名称, 年份, 卷(期): 页码.')
    add_paragraph_custom(doc, '[2] 作者. 文献标题[M]. 出版地: 出版社, 年份.')
    add_paragraph_custom(doc, '[3] ...')
    add_paragraph_custom(doc, '【预留20-30篇参考文献位置】')
    
    doc.add_paragraph()
    
    # ========== 附录 ==========
    add_heading_custom(doc, '附录', 1)
    add_heading_custom(doc, '附录A  系统技术参数表', 2)
    add_heading_custom(doc, '附录B  主要算法流程图', 2)
    add_heading_custom(doc, '附录C  测试数据详表', 2)
    add_heading_custom(doc, '附录D  主要程序代码', 2)
    
    doc.add_paragraph()
    
    # ========== 致谢 ==========
    add_heading_custom(doc, '致谢', 1)
    add_paragraph_custom(doc, '【致谢内容占位】')
    
    # 保存
    output_path = r'D:\铁路线路智能检测机器人\04-项目文档\设计文档\铁路线路智能检测机器人专业论文提纲 V1.0.docx'
    doc.save(output_path)
    print(f"[OK] 论文提纲已生成: {output_path}")
    
    import os
    file_size = os.path.getsize(output_path)
    print(f"[INFO] 文件大小: {file_size/1024:.1f} KB")

if __name__ == '__main__':
    create_paper_outline()
    print("\n[DONE] 完成!")
