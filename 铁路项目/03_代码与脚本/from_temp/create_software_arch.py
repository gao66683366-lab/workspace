# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# 设置默认字体
doc.styles['Normal'].font.name = '宋体'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
doc.styles['Normal'].font.size = Pt(10.5)

# 标题
title = doc.add_heading('铁路线路智能检测机器人 - 软件架构设计', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 整体架构
doc.add_heading('整体架构（三层架构）', 1)
doc.add_paragraph('本系统采用三层架构设计：云端服务器、工控上位机、硬件设备层。')

# 架构描述
p = doc.add_paragraph()
p.add_run('云端服务器 (Python)\n').bold = True
p.add_run('  - 数据存储与管理\n')
p.add_run('  - 深度学习模型训练\n')
p.add_run('  - 大数据分析与统计\n')
p.add_run('  - Web管理界面\n\n')
p.add_run('        ↑↓ HTTP/WebSocket\n\n')
p.add_run('工控上位机 (C# .NET 8.0)\n').bold = True
p.add_run('  应用层: 主控程序、人机交互界面、任务调度、数据上传下载\n')
p.add_run('  业务逻辑层: 检测流程控制、数据预处理、本地AI推理、轨道参数计算\n')
p.add_run('  硬件抽象层: 运动控制、视觉采集、传感器采集、数据存储\n\n')
p.add_run('        ↓ EtherCAT / 以太网 / Modbus RS485\n\n')
p.add_run('硬件设备层\n').bold = True
p.add_run('  伺服系统 | 工业相机×6 | 3D激光×2 | 测距×8 | 陀螺仪×2')

# 工控上位机软件模块
doc.add_heading('工控上位机软件模块（C# .NET 8.0）', 1)

doc.add_heading('1. 运动控制模块', 2)
doc.add_paragraph('职责：控制机器人在轨道上运行', style='List Bullet')
doc.add_paragraph('EtherCAT 通信接口', style='List Bullet 2')
doc.add_paragraph('运动控制器驱动', style='List Bullet 2')
doc.add_paragraph('速度、位置控制算法', style='List Bullet 2')
doc.add_paragraph('急停与安全保护', style='List Bullet 2')

doc.add_heading('2. 视觉采集模块', 2)
doc.add_paragraph('职责：采集并处理相机和3D激光数据', style='List Bullet')
doc.add_paragraph('工业相机 SDK 封装（6个相机）', style='List Bullet 2')
doc.add_paragraph('3D线激光 SDK 封装（2个激光）', style='List Bullet 2')
doc.add_paragraph('图像采集触发与同步', style='List Bullet 2')
doc.add_paragraph('图像预处理（去噪、增强）', style='List Bullet 2')
doc.add_paragraph('本地缓存管理', style='List Bullet 2')

doc.add_heading('3. 传感器采集模块', 2)
doc.add_paragraph('职责：采集测距和姿态数据', style='List Bullet')
doc.add_paragraph('Modbus RS485 通信', style='List Bullet 2')
doc.add_paragraph('测距传感器驱动（8个）', style='List Bullet 2')
doc.add_paragraph('陀螺仪驱动（2个）', style='List Bullet 2')
doc.add_paragraph('数据滤波与校准（卡尔曼滤波/中值滤波）', style='List Bullet 2')

doc.add_heading('4. 数据处理与分析模块', 2)
doc.add_paragraph('职责：实时数据分析与计算', style='List Bullet')
p = doc.add_paragraph('轨道几何参数计算：', style='List Bullet 2')
doc.add_paragraph('轨距（基于测距传感器）', style='List Bullet 3')
doc.add_paragraph('水平（基于陀螺仪）', style='List Bullet 3')
doc.add_paragraph('高低（基于陀螺仪）', style='List Bullet 3')
doc.add_paragraph('轨向（基于陀螺仪）', style='List Bullet 3')
doc.add_paragraph('视觉AI检测：', style='List Bullet 2')
doc.add_paragraph('螺栓完好性识别（目标检测+分类）', style='List Bullet 3')
doc.add_paragraph('轨面缺陷检测（磨损、鱼鳞纹、脱落）', style='List Bullet 3')
doc.add_paragraph('钢轨轮廓分析（基于3D激光）', style='List Bullet 3')
doc.add_paragraph('异常判定：参数超限报警、缺陷分级', style='List Bullet 2')

doc.add_heading('5. 数据存储模块', 2)
doc.add_paragraph('职责：本地数据管理', style='List Bullet')
doc.add_paragraph('SQLite / SQL Server Compact 本地数据库', style='List Bullet 2')
doc.add_paragraph('原始图像存储（按时间/位置索引）', style='List Bullet 2')
doc.add_paragraph('检测结果存储', style='List Bullet 2')
doc.add_paragraph('日志记录', style='List Bullet 2')

doc.add_heading('6. 通信模块', 2)
doc.add_paragraph('职责：与云端服务器通信', style='List Bullet')
doc.add_paragraph('HTTP/HTTPS API 调用', style='List Bullet 2')
doc.add_paragraph('WebSocket 实时数据推送', style='List Bullet 2')
doc.add_paragraph('数据上传队列（断线重传）', style='List Bullet 2')
doc.add_paragraph('模型下载与更新', style='List Bullet 2')

doc.add_heading('7. 人机交互界面', 2)
doc.add_paragraph('职责：操作与监控', style='List Bullet')
doc.add_paragraph('WPF / WinForms 界面', style='List Bullet 2')
doc.add_paragraph('实时显示：相机画面（6路）、3D点云、传感器数据、检测结果、运行状态', style='List Bullet 2')
doc.add_paragraph('操作控制：启动/停止、参数设置、手动/自动模式', style='List Bullet 2')

doc.add_heading('8. 任务调度模块', 2)
doc.add_paragraph('职责：协调各模块工作', style='List Bullet')
doc.add_paragraph('状态机管理（待机/检测/暂停/故障）', style='List Bullet 2')
doc.add_paragraph('多线程任务调度', style='List Bullet 2')
doc.add_paragraph('资源管理', style='List Bullet 2')
doc.add_paragraph('日志记录', style='List Bullet 2')

# 云端服务器软件
doc.add_heading('云端服务器软件（Python）', 1)

doc.add_heading('1. Web 管理平台', 2)
doc.add_paragraph('Flask / Django / FastAPI 框架', style='List Bullet')
doc.add_paragraph('用户管理与权限', style='List Bullet')
doc.add_paragraph('设备管理', style='List Bullet')
doc.add_paragraph('任务管理', style='List Bullet')

doc.add_heading('2. 数据管理', 2)
doc.add_paragraph('PostgreSQL / MySQL 数据库', style='List Bullet')
doc.add_paragraph('海量图像存储（对象存储 / NAS）', style='List Bullet')
doc.add_paragraph('检测数据统计分析', style='List Bullet')
doc.add_paragraph('报表生成', style='List Bullet')

doc.add_heading('3. AI 模型训练', 2)
doc.add_paragraph('PyTorch / TensorFlow 训练框架', style='List Bullet')
doc.add_paragraph('数据标注管理', style='List Bullet')
doc.add_paragraph('模型训练与优化', style='List Bullet')
doc.add_paragraph('模型版本管理', style='List Bullet')

doc.add_heading('4. 数据分析', 2)
doc.add_paragraph('缺陷趋势分析', style='List Bullet')
doc.add_paragraph('设备运行统计', style='List Bullet')
doc.add_paragraph('预测性维护', style='List Bullet')

# 关键技术栈
doc.add_heading('关键技术栈', 1)

doc.add_heading('工控上位机 (C# .NET 8.0)', 2)
table = doc.add_table(rows=10, cols=2)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '模块'
hdr_cells[1].text = '技术选型'

data = [
    ['EtherCAT 通信', 'TwinCAT / Acontis / SOEM'],
    ['工业相机 SDK', 'Basler Pylon / Hikvision MVS / Halcon'],
    ['3D 激光 SDK', '厂商 SDK'],
    ['Modbus 通信', 'NModbus4 / Modbus.Net'],
    ['AI 推理', 'ONNX Runtime / OpenVINO / TensorRT'],
    ['图像处理', 'OpenCvSharp / Emgu CV'],
    ['数据库', 'SQLite / LiteDB'],
    ['UI 框架', 'WPF (MVVM)'],
    ['日志', 'Serilog / NLog']
]

for i, row_data in enumerate(data, 1):
    row_cells = table.rows[i].cells
    row_cells[0].text = row_data[0]
    row_cells[1].text = row_data[1]

doc.add_heading('云端服务器 (Python)', 2)
table2 = doc.add_table(rows=7, cols=2)
table2.style = 'Light Grid Accent 1'
hdr_cells2 = table2.rows[0].cells
hdr_cells2[0].text = '模块'
hdr_cells2[1].text = '技术选型'

data2 = [
    ['Web 框架', 'FastAPI / Django'],
    ['数据库', 'PostgreSQL / MySQL'],
    ['对象存储', 'MinIO / OSS'],
    ['AI 框架', 'PyTorch / TensorFlow'],
    ['数据分析', 'Pandas / NumPy'],
    ['可视化', 'Matplotlib / Plotly']
]

for i, row_data in enumerate(data2, 1):
    row_cells = table2.rows[i].cells
    row_cells[0].text = row_data[0]
    row_cells[1].text = row_data[1]

# 软件工作流程
doc.add_heading('软件工作流程', 1)
doc.add_heading('典型检测流程', 2)

p = doc.add_paragraph()
p.add_run('1. 启动 → 系统初始化 → 硬件自检\n\n')
p.add_run('2. 开始检测:\n')
p.add_run('   a. 运动控制: 以设定速度运行\n')
p.add_run('   b. 位置触发: 按里程/时间触发采集\n')
p.add_run('   c. 数据采集:\n')
p.add_run('      - 相机拍摄 (6路)\n')
p.add_run('      - 3D激光扫描 (2路)\n')
p.add_run('      - 测距传感器读取 (8路)\n')
p.add_run('      - 陀螺仪读取 (2路)\n')
p.add_run('   d. 实时处理:\n')
p.add_run('      - AI推理 (螺栓/轨面)\n')
p.add_run('      - 几何参数计算\n')
p.add_run('      - 异常判定\n')
p.add_run('   e. 数据存储:\n')
p.add_run('      - 本地数据库\n')
p.add_run('      - 图像文件\n')
p.add_run('   f. 显示更新:\n')
p.add_run('      - UI界面刷新\n')
p.add_run('      - 异常报警\n\n')
p.add_run('3. 检测完成 → 数据上传云端 → 生成报告')

# 待确认与补充
doc.add_heading('待确认与补充', 1)
doc.add_paragraph('1. 工业相机具体型号与SDK？')
doc.add_paragraph('2. 3D线激光具体型号？')
doc.add_paragraph('3. AI模型部署方式（CPU / GPU / 专用AI加速卡）？')
doc.add_paragraph('4. 检测速度要求（km/h）？')
doc.add_paragraph('5. 数据量估算（每公里产生多少GB数据）？')
doc.add_paragraph('6. 云端服务器部署方式（私有云 / 公有云）？')

# 页脚
doc.add_paragraph()
p = doc.add_paragraph('创建时间: 2026-03-05')
p.runs[0].italic = True
p.runs[0].font.size = Pt(9)

doc.save('D:\\铁路线路智能检测机器人\\04-项目文档\\设计文档\\软件架构设计.docx')
print('[OK] 软件架构设计.docx 已创建')
