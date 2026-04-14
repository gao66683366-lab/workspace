# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = parse_xml(r'<w:tcBorders %s/>' % nsdecls('w'))
        tcPr.append(tcBorders)

def create_doc1():
    """创建机器人系统架构文档"""
    doc = Document()
    
    # 设置默认字体
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.styles['Normal'].font.size = Pt(10.5)
    
    # 标题
    title = doc.add_heading('铁路线路智能检测机器人系统架构', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 系统定位
    doc.add_heading('系统定位', 1)
    doc.add_paragraph('软硬件一体化智能检测机器人，用于铁路线路状态检测。')
    
    # 硬件组成
    doc.add_heading('硬件组成', 1)
    
    doc.add_heading('1. 底盘系统', 2)
    p = doc.add_paragraph('4个伺服电机 + 4个伺服驱动器', style='List Bullet')
    p.runs[0].bold = True
    doc.add_paragraph('控制4个车轮在铁路轨道上运行', style='List Bullet 2')
    doc.add_paragraph('由工控机（工业控制计算机）控制', style='List Bullet 2')
    
    doc.add_heading('2. 视觉检测系统', 2)
    p = doc.add_paragraph('6个工业相机', style='List Bullet')
    p.runs[0].bold = True
    doc.add_paragraph('检测对象：', style='List Bullet 2')
    doc.add_paragraph('两根钢轨轨面（磨损、鱼鳞纹、脱落等状态）', style='List Bullet 3')
    doc.add_paragraph('两根轨道旁的螺栓（每侧4颗，共8颗螺栓的完好状态）', style='List Bullet 3')
    
    doc.add_heading('3. 3D轮廓检测', 2)
    p = doc.add_paragraph('2个3D线激光', style='List Bullet')
    p.runs[0].bold = True
    doc.add_paragraph('用于两根钢轨的轮廓检测', style='List Bullet 2')
    
    doc.add_heading('4. 轨距检测', 2)
    p = doc.add_paragraph('8个测距传感器', style='List Bullet')
    p.runs[0].bold = True
    doc.add_paragraph('用于轨道轨距的检测', style='List Bullet 2')
    
    doc.add_heading('5. 姿态检测', 2)
    p = doc.add_paragraph('2个陀螺仪', style='List Bullet')
    p.runs[0].bold = True
    doc.add_paragraph('检测参数：', style='List Bullet 2')
    doc.add_paragraph('钢轨水平度', style='List Bullet 3')
    doc.add_paragraph('高低差', style='List Bullet 3')
    doc.add_paragraph('轨向（轨道方向）', style='List Bullet 3')
    doc.add_paragraph('其他空间姿态', style='List Bullet 3')
    
    # 控制架构
    doc.add_heading('控制架构', 1)
    
    doc.add_heading('主控制器', 2)
    doc.add_paragraph('工控机（工业控制计算机）', style='List Bullet').runs[0].bold = True
    doc.add_paragraph('统一控制所有硬件设备', style='List Bullet 2')
    
    doc.add_heading('运动控制层', 2)
    doc.add_paragraph('运动控制器 → 控制4个伺服电机', style='List Bullet').runs[0].bold = True
    doc.add_paragraph('通信方式：EtherCAT 总线（工业以太网）', style='List Bullet 2')
    doc.add_paragraph('驱动：4个伺服驱动器', style='List Bullet 2')
    
    # 软件架构
    doc.add_heading('软件架构', 1)
    
    doc.add_heading('上位机软件', 2)
    doc.add_paragraph('平台: 工控机', style='List Bullet')
    doc.add_paragraph('语言: C# .NET 8.0', style='List Bullet')
    doc.add_paragraph('功能: 设备控制、数据采集、本地处理', style='List Bullet')
    
    doc.add_heading('云端服务器', 2)
    doc.add_paragraph('语言: Python', style='List Bullet')
    doc.add_paragraph('功能: 数据处理、算法分析、远程管理', style='List Bullet')
    
    doc.add_heading('系统层级', 2)
    p = doc.add_paragraph()
    p.add_run('云端服务器 (Python)\n')
    p.add_run('    ↑\n')
    p.add_run('工控上位机 (C# .NET 8.0)\n')
    p.add_run('    ↓\n')
    p.add_run('运动控制器 (EtherCAT 总线)\n')
    p.add_run('    ↓\n')
    p.add_run('伺服驱动器 × 4 → 伺服电机 × 4')
    
    # 传感器通信架构
    doc.add_heading('传感器通信架构', 1)
    
    doc.add_heading('视觉与3D扫描', 2)
    doc.add_paragraph('6个工业相机 - 以太网连接（网线）', style='List Bullet')
    doc.add_paragraph('2个3D线激光 - 以太网连接（网线）', style='List Bullet')
    
    doc.add_heading('测距与姿态传感器', 2)
    doc.add_paragraph('8个测距传感器 - Modbus RS485 串口连接', style='List Bullet')
    doc.add_paragraph('2个陀螺仪 - Modbus RS485 串口连接', style='List Bullet')
    
    # 通信协议总览表格
    doc.add_heading('通信协议总览', 1)
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Light Grid Accent 1'
    
    # 表头
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '设备类型'
    hdr_cells[1].text = '数量'
    hdr_cells[2].text = '通信方式'
    hdr_cells[3].text = '协议'
    
    # 数据行
    data = [
        ['伺服电机', '4', '工业以太网', 'EtherCAT'],
        ['工业相机', '6', '以太网', 'GigE Vision / 厂商协议'],
        ['3D线激光', '2', '以太网', 'TCP/IP'],
        ['测距传感器', '8', '串口', 'Modbus RS485'],
        ['陀螺仪', '2', '串口', 'Modbus RS485']
    ]
    
    for i, row_data in enumerate(data, 1):
        row_cells = table.rows[i].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
            if j == 1:  # 数量列居中
                row_cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 页脚
    doc.add_paragraph()
    p = doc.add_paragraph('创建时间: 2026-03-05')
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(9)
    
    doc.save('D:\\铁路线路智能检测机器人\\04-项目文档\\设计文档\\机器人系统架构.docx')
    print('[OK] 机器人系统架构.docx 已创建')

def create_doc2():
    """创建项目目录说明文档"""
    doc = Document()
    
    # 设置默认字体
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.styles['Normal'].font.size = Pt(10.5)
    
    # 标题
    title = doc.add_heading('铁路线路智能检测机器人 - 项目目录说明', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 目录结构
    doc.add_heading('目录结构', 1)
    
    p = doc.add_paragraph()
    p.add_run('D:\\铁路线路智能检测机器人\\\n').bold = True
    p.add_run('│\n')
    p.add_run('├─ 01-代码/              # 所有代码文件\n')
    p.add_run('│  ├─ 上位机软件/         # C# .NET 8.0 工控软件\n')
    p.add_run('│  ├─ 云端服务器/         # Python 云端服务\n')
    p.add_run('│  ├─ 测试脚本/          # 单元测试、集成测试\n')
    p.add_run('│  └─ 工具脚本/          # 辅助工具、自动化脚本\n')
    p.add_run('│\n')
    p.add_run('├─ 02-技术资料/          # 技术文档和参考资料\n')
    p.add_run('│  ├─ 系统设计/          # 架构设计、详细设计\n')
    p.add_run('│  ├─ 硬件文档/          # 设备手册、接线图\n')
    p.add_run('│  ├─ API文档/          # 接口文档、SDK说明\n')
    p.add_run('│  ├─ 参考资料/          # 技术书籍、教程\n')
    p.add_run('│  └─ 标准规范/          # 行业标准、规范文件\n')
    p.add_run('│\n')
    p.add_run('├─ 03-论文编撰/          # 论文相关资料\n')
    p.add_run('│  ├─ 文献综述/          # 参考文献、综述材料\n')
    p.add_run('│  ├─ 数据分析/          # 实验数据、分析结果\n')
    p.add_run('│  ├─ 图表素材/          # 图片、表格、示意图\n')
    p.add_run('│  ├─ 论文草稿/          # 各版本草稿\n')
    p.add_run('│  └─ 发表材料/          # 投稿材料、修改意见\n')
    p.add_run('│\n')
    p.add_run('├─ 04-项目文档/          # 项目管理文档\n')
    p.add_run('│  ├─ 需求分析/          # 需求文档、用例\n')
    p.add_run('│  ├─ 设计文档/          # 系统架构、软件设计\n')
    p.add_run('│  ├─ 测试文档/          # 测试计划、测试报告\n')
    p.add_run('│  └─ 用户手册/          # 操作手册、维护手册\n')
    p.add_run('│\n')
    p.add_run('├─ 05-数据资料/          # 数据和模型\n')
    p.add_run('│  ├─ 测试数据/          # 测试用图像、传感器数据\n')
    p.add_run('│  ├─ 训练数据/          # AI模型训练数据集\n')
    p.add_run('│  └─ AI模型/           # 训练好的模型文件\n')
    p.add_run('│\n')
    p.add_run('├─ 06-会议记录/          # 会议纪要、讨论记录\n')
    p.add_run('│\n')
    p.add_run('└─ 07-临时文件/          # 临时工作文件')
    
    # 智能存储规则
    doc.add_heading('智能存储规则', 1)
    doc.add_paragraph('小测会根据对话内容自动判断并存储到对应文件夹。')
    
    doc.add_heading('代码相关 → 01-代码/', 2)
    doc.add_paragraph('C# 代码 → 上位机软件/', style='List Bullet')
    doc.add_paragraph('Python 代码 → 云端服务器/', style='List Bullet')
    doc.add_paragraph('测试代码 → 测试脚本/', style='List Bullet')
    doc.add_paragraph('工具脚本 → 工具脚本/', style='List Bullet')
    
    doc.add_heading('技术资料 → 02-技术资料/', 2)
    doc.add_paragraph('架构设计、技术方案 → 系统设计/', style='List Bullet')
    doc.add_paragraph('设备手册、规格书 → 硬件文档/', style='List Bullet')
    doc.add_paragraph('SDK、API说明 → API文档/', style='List Bullet')
    doc.add_paragraph('教程、参考书籍 → 参考资料/', style='List Bullet')
    doc.add_paragraph('国标、行标 → 标准规范/', style='List Bullet')
    
    doc.add_heading('论文编撰 → 03-论文编撰/', 2)
    doc.add_paragraph('参考文献、综述 → 文献综述/', style='List Bullet')
    doc.add_paragraph('实验数据、统计分析 → 数据分析/', style='List Bullet')
    doc.add_paragraph('图片、表格 → 图表素材/', style='List Bullet')
    doc.add_paragraph('论文各版本 → 论文草稿/', style='List Bullet')
    doc.add_paragraph('投稿、审稿意见 → 发表材料/', style='List Bullet')
    
    # 使用说明
    doc.add_heading('使用说明', 1)
    doc.add_paragraph('1. 自动存储：对话中产生的文件会自动分类存储')
    doc.add_paragraph('2. 手动管理：您也可以自己移动、整理文件')
    doc.add_paragraph('3. 命名规范：建议用日期+描述，如 2026-03-05_模块设计.docx')
    doc.add_paragraph('4. 定期清理：07-临时文件/ 可定期清理')
    
    # 页脚
    doc.add_paragraph()
    p = doc.add_paragraph('创建时间: 2026-03-05 | 维护者: 小测（智能体）')
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(9)
    
    doc.save('D:\\铁路线路智能检测机器人\\项目目录说明.docx')
    print('[OK] 项目目录说明.docx 已创建')

if __name__ == '__main__':
    create_doc1()
    create_doc2()
    print('\n[SUCCESS] 所有Word文档已创建完成！')
