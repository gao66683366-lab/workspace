import os
import docx
from docx.shared import Pt, Inches
from docx.oxml.ns import qn

doc = docx.Document()
doc.styles['Normal'].font.name = '微软雅黑'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
doc.styles['Normal'].font.size = Pt(10.5)

# Heading
h0 = doc.add_heading('铁路线路智能检测机器人', 0)
h1 = doc.add_heading('七大核心工作板块 (纯工程技术定稿版)', 1)

doc.add_paragraph('本文档基于小车真实物理约束（10km/h低速、120mm轮径、刚性底盘、物理遮光罩、定制电源管理板PMB、10网口硬隔离等）进行全盘深度重构与解耦，是指导后续所有软硬件开发的最高标准依据。')

sections = [
    ("1. 硬件控制与物理感知 (底层基石)", [
        ("1.1 物理拓扑 (10网口强制隔离)", "LAN1: 独占EtherCAT (1ms控制4电机)\nLAN2: 直连外置工业路由器(SIM卡)\nPCIe-A/B: 分别承载6相机+2线激光(巨型帧)\n低速总线: 测距/陀螺仪走串口或IO板卡，绝不抢占以太网。"),
        ("1.2 封闭光学与同步触发", "物理防线: 金属遮光罩+50mm柔性挡光裙边(防异物碰撞)\n参数锁死: 废弃软件调光，相机参数写死\n硬件触发: <200μs曝光时间配合单色频闪LED，消除10km/h运动拖影。"),
        ("1.3 测距基准 (120mm刚性轮)", "纯轮式里程计: EtherCAT高频读取4轮绝对值编码器\n防滑容错: 实时监测电机扭矩，自动剔除悬空/打滑轮的高频脉冲数据，取有效接触轮的均值作为里程。")
    ]),
    ("2. 软件平台与数据流转 (C#上位机)", [
        ("2.1 CPU内核物理隔离", "Thread Affinity机制：将Windows系统底层EtherCAT主站进程强绑至独立CPU核心，杜绝8个视觉网口的巨型帧中断风暴抢占资源导致电机宕机。"),
        ("2.2 零拷贝与Frame重组", "内存池: 开辟非托管连续内存池直写图像，杜绝C# GC卡顿\n数据对齐: 严格以编码器里程脉冲(Frame ID)为字典主键，集齐同物理截面的多源数据才拼装出栈，残缺帧强制丢弃。")
    ]),
    ("3. 数据处理与核心算法 (绝对解耦)", [
        ("3.1 轨道几何形位 (4独立通道)", "轨距: 8测距传感器求和，独立输出\n水平: 陀螺仪Roll角乘以基准，独立输出\n高低/轨向: 陀螺仪Pitch/Yaw角积分输出\n震动过滤: 针对刚性底盘的高频冲击，必须串联10Hz低通滤波器，过滤碎震噪音。"),
        ("3.2 截面磨耗 (3D激光)", "2个3D线激光提取钢轨截面点云，直接与标称钢轨CAD配准，计算垂磨与侧磨。"),
        ("3.3 AI表面病害截流", "边缘端部署YOLO量化模型，剔除90%无病害钢轨背景，仅裁剪包含螺栓与疑似裂纹的微小ROI图片供云端复核。")
    ]),
    ("4. 边缘-云端均衡通信 (恶劣网络对策)", [
        ("4.1 穿透路由探针", "C#程序高频发送测速心跳包，通过RTT延迟与丢包率动态推算SIM卡外置路由器的实际上行可用带宽。"),
        ("4.2 QoS队列与防掉速落盘", "队列优先级: P0(告警) > P1(状态) > P2(ROI切片) > P3(大图/点云)\n隧道生存机制: 弱网触发系统写锁，全量数据存入工业级高TBW固态硬盘(防消费级SSD掉速憋死内存)，网络恢复后按MD5游标断点续传。")
    ]),
    ("5. 机械电气与容错设计 (应对刚性冲击)", [
        ("5.1 定制电源管理板 (PMB)", "总线隔离: 48V总输入，光电/磁隔离多路输出(电机脏电与工控机/相机净电分离)，防电机反电动势击穿设备。\n智能保灾: 各通道独立PTC熔断；板载控制继电器，支持工控机指令对死机相机或路由器发起物理断电硬重启。"),
        ("5.2 光学防震与散热", "光学防震: 刚性连接过道岔极易导致离焦，相机镜头调校后必须全螺纹点胶死锁，基座增加高分子阻尼垫层。\n热力学设计: 密闭遮光罩内发热量巨大，必须加装强制对流风道(IP67工业风扇)。")
    ]),
    ("6. 技术文档与规范化管理", [
        ("6.1 格式铁律与自动化生成", "输出标准: 彻底放弃Markdown体系，所有接口定义、架构图、测试报告统一采用全表格化的Word (.docx)进行版本管理。\n自动化: 编写脚本读取源码数据字典，自动映射生成接口文档，确保开发与文档绝对对齐。")
    ]),
    ("7. 学术论文与知识产权", [
        ("7.1 工程向学术转化", "专利与论文锚点: \n1. 无机械减震四轮底盘下的多源传感器解耦测量与数字滤波补偿方法。\n2. 基于10网口物理隔离与巨型帧的零拷贝多模态数据采集架构。\n3. 基于定制PMB电源板隔离与弱网主动探针的高可靠巡检机器系统。")
    ])
]

for sec_title, items in sections:
    doc.add_heading(sec_title, level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '工程模块'
    hdr_cells[1].text = '硬核落地规范'
    
    for mod, desc in items:
        row = table.add_row().cells
        row[0].text = mod
        row[1].text = desc

# Set table fonts
for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    run.font.size = Pt(10.5)

filepath = r"D:\铁路线路智能检测机器人\04-项目文档\设计文档\铁路线路智能检测机器人_7大核心工作板块_纯工程定稿版.docx"
doc.save(filepath)
print(f"DONE:{filepath}")
