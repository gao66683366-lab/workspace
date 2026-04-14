# -*- coding: utf-8 -*-
import docx
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

output_dir = r"D:\铁路线路智能检测机器人\04-项目文档\设计文档"
os.makedirs(output_dir, exist_ok=True)

doc = docx.Document()

# 设置正文默认字体
doc.styles['Normal'].font.name = u'微软雅黑'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
doc.styles['Normal'].font.size = Pt(10.5)

def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = u'微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    return h

add_heading(doc, '铁路线路智能检测机器人', 0)
add_heading(doc, '系统全景架构设计文档 V3.0 (终极物理定调版)', 1)

doc.add_paragraph('本文档为项目系统架构的最终纲领（V3.0版）。全面融入了10网口硬约束、48V蓄电池直流供电、遮光罩恒定光源设计、纯编码器高精里程计定位，以及基于工业路由器SIM卡的QoS主动探针均衡传输方案。')

add_heading(doc, '1. 物理拓扑与硬件接口硬分配 (10网口铁律)', 2)
table1 = doc.add_table(rows=1, cols=4, style='Table Grid')
hdr_cells = table1.rows[0].cells
hdr_cells[0].text = '网络接口'
hdr_cells[1].text = '接口类型'
hdr_cells[2].text = '连接设备'
hdr_cells[3].text = '核心约束与说明'

data1 = [
    ('主板 LAN 1', '原生千兆网口', '4个伺服电机(驱动器)', '独占EtherCAT总线(1ms周期)。实时获取绝对值编码器脉冲，作为全车高精里程计(Odometry)核心数据源。'),
    ('主板 LAN 2', '原生千兆网口', '工业路由器(带SIM卡)', '全车唯一WAN口。C#程序通过主动链路探针(UDP Ping)探测真实带宽，执行QoS均衡上传策略。'),
    ('PCIe 扩展卡 A', '4口千兆 PoE', '4个工业相机', '巨型帧(9014 Bytes)独立子网配置。与扩展卡B共享底层同一微秒的I/O硬同步触发。'),
    ('PCIe 扩展卡 B', '4口千兆 PoE', '2个工业相机 + 2个3D线激光', '线激光与视觉同享纳秒级延时补偿，确保空间采集面对齐。'),
    ('多串口/EtherCAT I/O', 'RS485 / I/O', '8个测距传感器+2个陀螺仪', '走低速总线，绝不占用以太网口。结合陀螺仪数据进行三角融合解算真实轨距。')
]
for item in data1:
    row = table1.add_row().cells
    row[0].text = item[0]
    row[1].text = item[1]
    row[2].text = item[2]
    row[3].text = item[3]

add_heading(doc, '2. 机械/电气降维打击物理设计', 2)
table2 = doc.add_table(rows=1, cols=3, style='Table Grid')
hdr2 = table2.rows[0].cells
hdr2[0].text = '物理设计模块'
hdr2[1].text = '工程实现'
hdr2[2].text = '解决的痛点与红利'

data2 = [
    ('全天候恒定光照', '封闭遮光罩 + 内部高频频闪光源', '相机固定曝光参数，免除复杂动态调光算法；彻底屏蔽外界强光/阴影干扰，大幅降低AI识别难度。'),
    ('自主高精定位', '伺服电机编码器(脉冲)计算里程', '摆脱GPS信号依赖。配合人工起点标定、多轮防滑转均值滤波与磨耗补偿，实现隧道内毫米级连续定位。'),
    ('动力与续航监测', '48V蓄电池直流主电源', '系统统一供电。采集端读取BMS数据（电压/电流），在云端绘制数字孪生放电曲线。'),
    ('多模态联合标定', '定制3D-2D联合标定板出厂校准', '生成所有传感器相对于底盘中心的统一外参矩阵，保证点云与图片在三维空间的绝对重合。')
]
for item in data2:
    row = table2.add_row().cells
    row[0].text = item[0]
    row[1].text = item[1]
    row[2].text = item[2]

add_heading(doc, '3. 七大核心工作板块 (全景落地字典)', 2)
table3 = doc.add_table(rows=1, cols=3, style='Table Grid')
hdr3 = table3.rows[0].cells
hdr3[0].text = '版图分类'
hdr3[1].text = '核心落地动作'
hdr3[2].text = '关键技术点'

data3 = [
    ('1. 硬件控制与感知', '多模态硬同步触发、频闪光源控制', '依赖编码器里程脉冲等距触发，方波信号同时激活相机快门与频闪光源。'),
    ('2. 软件平台架构', 'C#高并发零拷贝采集、硬守护看门狗', '按Frame ID重组数据帧，防爆缓存。连续断网时GPIO继电器物理硬重启路由器。'),
    ('3. AI与核心算法', '边缘YOLO截流、测距姿态融合补偿', '裁剪出含有缺陷的微小ROI切片；8测距结合陀螺仪横滚角(Roll)输出高精轨距。'),
    ('4. 均衡通信架构', 'QoS四级队列与写锁落盘机制', '探针侦测到弱网立刻停止P2/P3大文件发送，写入SQLite缓存，待网络恢复后MD5分块续传。'),
    ('5. 资产与图纸管理', '机械刚性仿真与电气EMC防干扰', '强弱电物理隔离走线规范，相机支架防震高刚性保障。'),
    ('6. 技术文档管理', '全Word格式+表格化强制输出', '拒绝散乱文档，API变动与测试数据自动化抽取填充至文档模板。'),
    ('7. 学术论文与专利', '聚焦自适应均衡传输与多源融合', '提取遮光罩光照一致性对比数据、探针调度效率图表用于高水平SCI论文发表。')
]
for item in data3:
    row = table3.add_row().cells
    row[0].text = item[0]
    row[1].text = item[1]
    row[2].text = item[2]

for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = u'微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
                    run.font.size = Pt(10)

doc.save(os.path.join(output_dir, 'Architecture_V3.0.docx'))
print("V3.0 Success")
