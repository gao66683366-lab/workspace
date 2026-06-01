#!/usr/bin/env python3
import sys, os
sys.path.insert(0, '/root/.openclaw/workspace/word_tools')

# Just test add_table with merge data
from docx import Document
from docx_gen import add_table

doc = Document()
rows = [
    ['指标类别', '具体指标', '目标值'],
    ['运动性能', '行走速度', '0-3 m/s'],
    ['merge:运动性能', '定位精度', '±20mm'],
    ['merge:运动性能', '越障能力', '≤30mm'],
    ['成像性能', '图像分辨率', '4K'],
    ['merge:成像性能', '低照度灵敏度', '0.01 lux'],
]
add_table(doc, rows)
doc.save('/tmp/test_merge.docx')
print('Saved /tmp/test_merge.docx')