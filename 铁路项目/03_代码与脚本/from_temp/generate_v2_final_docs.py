import docx
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

output_dir = r"D:\铁路线路智能检测机器人\04-项目文档\设计文档"
os.makedirs(output_dir, exist_ok=True)

def set_font(doc):
    doc.styles['Normal'].font.name = u'微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    doc.styles['Normal'].font.size = Pt(10.5)

def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = u'微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    return h

def format_table(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = u'微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
                    run.font.size = Pt(10)

# ================= 1. 硬件系统架构 V2.0 =================
doc_hw = docx.Document()
set_font(doc_hw)
add_heading(doc_hw, '铁路线路智能检测机器人 - 硬件系统架构设计 V2.0', 1)
doc_hw.add_paragraph('【修订说明】：基于10km/h极速、120mm刚性轮径、完全解耦物理测量、以及定制电源管理板(PMB)的全新硬核基线重构。')

add_heading(doc_hw, '1. 核心网络与通信拓扑 (10网口强制隔离)', 2)
t1 = doc_hw.add_table(rows=1, cols=3, style='Table Grid')
t1.rows[0].cells[0].text, t1.rows[0].cells[1].text, t1.rows[0].cells[2].text = '网络通道', '物理接口', '承载设备与说明'
data_hw_1 = [
    ('运动控制(高实时)', '主板 LAN 1', '独占EtherCAT总线，1ms周期通信，接4个伺服驱动器。'),
    ('云端通信(外网)', '主板 LAN 2', '直连外置工业路由器(带SIM卡)，唯一上行通道。'),
    ('视觉感知通道 A', 'PCIe 扩展卡 (4口GigE)', '接4个工业相机(巨型帧，独立子网)。'),
    ('视觉感知通道 B', 'PCIe 扩展卡 (4口GigE)', '接2个工业相机 + 2个3D线激光(巨型帧，独立子网)。'),
    ('低速测量通道', '多串口卡/IO板', '接8个测距传感器 + 2个陀螺仪，彻底规避以太网拥堵。')
]
for row_data in data_hw_1:
    row = t1.add_row().cells
    row[0].text, row[1].text, row[2].text = row_data
format_table(t1)

add_heading(doc_hw, '2. 定制电源管理板 (PMB) 与电气隔离', 2)
t2 = doc_hw.add_table(rows=1, cols=3, style='Table Grid')
t2.rows[0].cells[0].text, t2.rows[0].cells[1].text, t2.rows[0].cells[2].text = '供电通道', '电压', '接入设备与保护机制'
data_hw_2 = [
    ('动力通道 A', '48V 直通', '4个伺服电机（承受反电动势，硬件过流保护）。'),
    ('主控通道 B', '24V/12V 稳压', '工控机专用（宽压隔离DC-DC，防电机干扰）。'),
    ('感知通道 C', '24V', '6相机+2激光（单路PTC自恢复保险丝短路隔离）。'),
    ('通信通道 D', '12V/5V', '工业路由器、传感器（支持工控机发指令继电器硬重启）。')
]
for row_data in data_hw_2:
    row = t2.add_row().cells
    row[0].text, row[1].text, row[2].text = row_data
format_table(t2)

add_heading(doc_hw, '3. 光学与机械抗震方案', 2)
doc_hw.add_paragraph('由于120mm小轮径无减震底盘，必须采取局部柔性与死锁：\n1. 光学：物理遮光罩+50mm柔性防尘裙边+LED微秒频闪，废除复杂软件测光。\n2. 紧固：所有相机镜头必须点胶/螺纹胶焊死，基座加装聚氨酯阻尼垫防高频离焦。')
doc_hw.save(os.path.join(output_dir, '硬件系统架构设计文档_V2.0_纯工程版.docx'))

# ================= 2. 软件系统架构 V2.0 =================
doc_sw = docx.Document()
set_font(doc_sw)
add_heading(doc_sw, '铁路线路智能检测机器人 - 软件系统架构设计 V2.0', 1)
doc_sw.add_paragraph('【修订说明】：彻底废除伪科学耦合计算，落实传感器“各司其职”模型，加入内核隔离与防掉速机制。')

add_heading(doc_sw, '1. 底层高并发采集与内核隔离', 2)
t3 = doc_sw.add_table(rows=1, cols=3, style='Table Grid')
t3.rows[0].cells[0].text, t3.rows[0].cells[1].text, t3.rows[0].cells[2].text = '机制名称', '实现方式', '解决的致命问题'
data_sw_1 = [
    ('CPU 内核绑定', 'Windows Thread Affinity', '将EtherCAT主站独占CPU核心，防巨型帧网络中断抢占导致电机失控。'),
    ('零拷贝内存池', 'C# 非托管内存 (Marshal)', '海量图像直写底层内存，彻底规避C#垃圾回收(GC)引发的严重卡顿丢帧。'),
    ('数据帧重组', '以Frame ID为主键对齐', '凑齐同一里程脉冲的6图+2点云才放行，残缺帧直接丢弃，绝不污染AI。')
]
for row_data in data_sw_1:
    row = t3.add_row().cells
    row[0].text, row[1].text, row[2].text = row_data
format_table(t3)

add_heading(doc_sw, '2. 轨道几何解算模型 (绝对解耦)', 2)
t4 = doc_sw.add_table(rows=1, cols=3, style='Table Grid')
t4.rows[0].cells[0].text, t4.rows[0].cells[1].text, t4.rows[0].cells[2].text = '几何指标', '独立数据源', '计算逻辑 (互不干涉)'
data_sw_2 = [
    ('轨距', '8个测距传感器', '两侧测值直接相加，输出绝对轨距。不引入姿态补偿。'),
    ('水平 (超高)', '陀螺仪 Roll (横滚角)', '读取Roll角乘以标称轨距1435mm，直接输出高低差。'),
    ('纵向高低', '陀螺仪 Pitch + 里程计', '俯仰角结合里程脉冲进行空间积分。'),
    ('轨向 (平顺度)', '陀螺仪 Yaw + 里程计', '航向角结合里程脉冲积分。需经过10Hz低通滤波器滤除刚性硬冲击。'),
    ('表面缺陷', '6相机 + 边缘YOLO', '仅作ROI截流裁剪，只留缺陷切片传云端复核。')
]
for row_data in data_sw_2:
    row = t4.add_row().cells
    row[0].text, row[1].text, row[2].text = row_data
format_table(t4)
doc_sw.save(os.path.join(output_dir, '软件系统架构设计文档_V2.0_纯工程版.docx'))

# ================= 3. 总体技术方案 V2.0 =================
doc_tech = docx.Document()
set_font(doc_tech)
add_heading(doc_tech, '铁路线路智能检测机器人 - 总体技术方案 V2.0', 1)
doc_tech.add_paragraph('本文档是本项目的最高技术纲领，基于 10km/h极速、120mm无减震轮式底盘、纯编码器无盲区定位与48V供电等真实物理边界编制。')

add_heading(doc_tech, '1. 系统总体拓扑', 2)
doc_tech.add_paragraph('以工控机为边缘大脑，以定制PMB电源板为电气护城河。底部通过10网口隔离接入视觉/激光/控制总线。上层通过工业路由器与云端握手。')

add_heading(doc_tech, '2. 弱网环境 QoS 均衡通信策略', 2)
t5 = doc_tech.add_table(rows=1, cols=3, style='Table Grid')
t5.rows[0].cells[0].text, t5.rows[0].cells[1].text, t5.rows[0].cells[2].text = '队列级别', '数据类型', '调度与保底策略'
data_tech = [
    ('P0 极高', '急停与系统硬件告警', '最高优先级，强行抢占带宽发送。'),
    ('P1 高', '几何数据/坐标/电量JSON', '秒级发送，保持前后端状态同步。'),
    ('P2 中', '边缘AI裁剪出的缺陷切片', '带宽探测良好时发送。'),
    ('P3 低', '原始全景大图与3D点云', '遇隧道网差时触发本地写锁(需工业级SSD防掉速)。网络恢复后MD5游标续传。')
]
for row_data in data_tech:
    row = t5.add_row().cells
    row[0].text, row[1].text, row[2].text = row_data
format_table(t5)
doc_tech.save(os.path.join(output_dir, '铁路线路智能检测机器人技术方案_V2.0_纯工程版.docx'))

print("ALL_V2_DOCS_GENERATED")
