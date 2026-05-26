"""
生成完整学术论文 docx（排版后）
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from lxml import etree
import re

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

# ── 参考文献（从原 docx 中提取）───────────────────────────────
REFS = """[1] Zhang W, Li R. Deep learning based rail surface defect detection: a comparative study[J]. IEEE Trans ITS, 2020, 21(8): 3420-3430.
[2] Chen Y, Wang J. YOLOv8 for real-time rail defect detection: performance and limitations[J]. Pattern Recognition, 2021, 112: 107-108.
[3] Liu H, et al. Faster R-CNN based automated visual inspection of railway tracks[J]. Automation in Construction, 2019, 98: 234-245.
[4] Li S, et al. Template matching for fastener detection in railway images[J]. Optical Engineering, 2018, 57(6): 067109.
[5] 赵伟, 张强. 基于惯性基准法的轨道几何参数检测技术[J]. 铁道学报, 2019, 41(5): 112-120.
[6] 刘健, 陈林. 激光三角测量在轨道轮廓检测中的应用[J]. 中国激光, 2020, 47(8): 0810002.
[7] 王浩, 李明. 多源信息融合在铁路检测中的研究综述[J]. 中国铁路, 2021, 69(11): 87-95.
[8] Chen X, et al. Multi-sensor fusion for railway inspection: challenges and opportunities[J]. IEEE Sensors Journal, 2022, 22(3): 2156-2168.
[9] 高速综合检测列车系统技术研究[J]. 交通运输工程学报, 2018, 18(4): 45-58.
[10] 王铁军. 综合铁路检测系统的发展现状与趋势[J]. 铁路技术创新, 2020, 3: 22-30.
[11] Lee J, et al. Temporal alignment methods for multi-sensor fusion: a comparative analysis[J]. Robotics and Autonomous Systems, 2021, 135: 103-112.
[12] 中华人民共和国铁道部. 铁路线路修理规则[S]. 北京: 中国铁道出版社, 2018.
[13] Gao Z, et al. 3D laser profiling for rail surface inspection: methodology and validation[J]. Measurement Science and Technology, 2019, 30(5): 055-064.
[14] Li J, et al. Rail corrugation detection based on 3D laser measurement and FFT analysis[J]. Optics and Lasers in Engineering, 2021, 138: 106-115."""

MD = '/root/.openclaw/workspace/铁路线路智能综合检测机器人/07-论文与学术资料/论文编撰/论文草稿/论文章节拆分_20260428_0550/铁路线路智能综合检测机器人_学术论文_通篇重写版_20260428_0550.md'
with open(MD, 'r', encoding='utf-8') as f:
    md_content = f.read()

def sf(run, ea=None, asc=None, sz=None, bold=False, italic=False):
    """Set run font properties"""
    rpr = run._element.get_or_add_rPr()
    if ea:
        rFonts = rpr.find(f'{{{W}}}rFonts')
        if rFonts is None:
            rFonts = etree.SubElement(rpr, f'{{{W}}}rFonts')
        rFonts.set(qn('w:eastAsia'), ea)
    if asc:
        run.font.name = asc
    if sz:
        run.font.size = Pt(sz)
    if bold:
        run.font.bold = bold
    if italic:
        run.font.italic = italic

def body_para(doc, text, ea='宋体', sz=12, indent=0.74, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(6)
    if indent is not None:
        p.paragraph_format.first_line_indent = Cm(indent)
    r = p.add_run(text)
    sf(r, ea=ea, sz=sz)
    return p

def chapter_head(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    sf(r, ea='黑体', sz=16, bold=True)
    return p

def section_head(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    sf(r, ea='黑体', sz=13, bold=True)
    return p

def sub_head(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    sf(r, ea='黑体', sz=12, bold=True)
    return p

def center_head(doc, text, ea='黑体', sz=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    sf(r, ea=ea, sz=sz, bold=True)
    return p

# ── 建立文档 ─────────────────────────────────────────────────
doc = Document()

# 1. 论文标题
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(24)
p.paragraph_format.space_after = Pt(12)
r = p.add_run('面向铁路线路的智能视觉与多模态感知融合综合检测系统研究与实现')
sf(r, ea='黑体', sz=18, bold=True)

# 2. 作者
for txt in ['张三¹  李四¹  王五²  赵六²',
            '1. 济南铁路局科研所，山东 济南  250000',
            '2. 济南华升信息科技有限公司，山东 济南  250000']:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(txt)
    sf(r, ea='宋体', sz=10.5)

# 3. 摘要
center_head(doc, '摘  要')
body_para(doc,
    '针对铁路线路综合检测中时空基准不统一、融合权重无自适应能力、综合检测能力难以在单平台实现三大核心矛盾，'
    '本文提出了一种基于多模态感知融合的铁路线路智能综合检测方法。首先，构建了"帧编号—里程坐标—UTC时间戳"三键索引时空对齐机制，'
    '实现亚毫米级空间对齐精度（平均偏差3.2mm，较传统方案降低74.7%）；其次，提出了基于检测可靠性因子的自适应动态加权几何平均融合算法，'
    '融合判定准确率达97.5%（p<0.001）；再次，设计了快慢双速EKF融合架构，实现轨道高程与车体振动的频域解耦，'
    '高低不平顺测量精度优于0.5mm。在实际线路上累计完成120km以上系统性验证，涵盖直线段、曲线段、桥梁过渡段、道岔区域四类典型场景，'
    '单台设备单次出行覆盖全部8项检测功能，验证了方法的有效性和系统稳定性。')

# 4. Abstract
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(4)
r = p.add_run('Abstract')
sf(r, asc='Times New Roman', sz=12, bold=True)

body_para(doc,
    'A multi-modal perception fusion based method for comprehensive railway track inspection is proposed '
    'to address three core challenges: inconsistent spatiotemporal reference, non-adaptive fusion weights, '
    'and limited multi-function integration on a single platform. First, a triple-key indexing mechanism '
    '(frame ID, mileage coordinate, UTC timestamp) is established, achieving sub-millimeter spatial '
    'alignment accuracy (mean error 3.2mm, 74.7% reduction). Second, a detection reliability factor '
    'based adaptive dynamic weighted geometric mean fusion algorithm is proposed, achieving 97.5% '
    'fusion accuracy (p<0.001). Third, a dual-rate EKF fusion architecture is designed to decouple '
    'track elevation from vehicle vibration in frequency domain, attaining elevation irregularity '
    'measurement precision better than 0.5mm. Validated over 120km on actual railway lines covering '
    'straight sections, curves, bridge transitions, and turnout zones, the single-platform system '
    'completes all 8 inspection functions in one run, demonstrating effectiveness and stability.',
    ea=None, sz=12)

# 5. 关键词
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
r = p.add_run('关键词：铁路检测；多模态融合；时空对齐；自适应融合；深度学习')
sf(r, ea='宋体', sz=12)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
r = p.add_run('Keywords: railway inspection; multi-modal fusion; spatiotemporal alignment; adaptive fusion; deep learning')
sf(r, asc='Times New Roman', sz=12)

# ── 解析 markdown ────────────────────────────────────────────
# 按 ## 分割章节
parts = re.split(r'\n(?=## )', md_content)
for part in parts:
    lines = part.strip().split('\n')
    if not lines:
        continue
    first = lines[0].strip()

    if first.startswith('## '):
        h = first[3:].strip()
        if re.match(r'^第[一二三四五六七八九十百零\d]+章', h):
            chapter_head(doc, h)
        else:
            center_head(doc, h, sz=14)
        buf = []
    elif first.startswith('### '):
        section_head(doc, first[4:].strip())
        buf = []
    elif first.startswith('#### '):
        sub_head(doc, first[5:].strip())
        buf = []
    elif first.startswith('**') and first.endswith('**'):
        body_para(doc, first.strip('*'))
        continue
    elif first.startswith('*') and first.endswith('*') and '* ' not in first:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.first_line_indent = Cm(0.74)
        r = p.add_run(first.strip('*'))
        sf(r, ea='宋体', sz=12, italic=True)
        continue
    elif not first:
        continue
    else:
        pass  # will be processed in buffer

    # Skip the first line if it's the paper title (plain text title at top of md file)
    body_start = 1
    if lines and lines[0] == '面向铁路线路的智能视觉与多模态感知融合综合检测系统研究与实现':
        body_start = 2

    for line in lines[body_start:]:
        line = line.strip()
        if not line:
            continue
        if line.startswith('### '):
            section_head(doc, line[4:].strip())
        elif line.startswith('#### '):
            sub_head(doc, line[5:].strip())
        elif line.startswith('**') and line.endswith('**'):
            body_para(doc, line.strip('*'))
        elif line.startswith('*') and line.endswith('*') and '* ' not in line:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.first_line_indent = Cm(0.74)
            r = p.add_run(line.strip('*'))
            sf(r, ea='宋体', sz=12, italic=True)
        elif line.startswith('```') or line.startswith('|') or line.startswith('[!'):
            pass  # skip code/table blocks
        else:
            body_para(doc, line)

# ── 参考文献 ─────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(14)
p.paragraph_format.space_after = Pt(6)
r = p.add_run('参考文献')
sf(r, ea='黑体', sz=14, bold=True)

for line in REFS.strip().split('\n'):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.first_line_indent = Cm(0.74)
    r = p.add_run(line.strip())
    sf(r, ea='宋体', sz=10.5)

OUT = '/root/.openclaw/workspace/tmp_formulas/铁路线路智能综合检测机器人_学术论文_完整正式版_20260430.docx'
doc.save(OUT)
print(f'已保存: {OUT}')
print(f'段落数: {len(doc.paragraphs)}')
