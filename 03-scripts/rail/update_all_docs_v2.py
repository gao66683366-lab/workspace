import docx
from docx.shared import Pt
from docx.oxml.ns import qn
import os

out_dir = r"D:\铁路线路智能检测机器人\04-项目文档\设计文档"
os.makedirs(out_dir, exist_ok=True)

def create_doc(title, desc, sections):
    doc = docx.Document()
    # 字体设置
    doc.styles['Normal'].font.name = u'微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    
    h0 = doc.add_heading(title, 0)
    for run in h0.runs: run.font.name = u'微软雅黑'; run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    
    doc.add_paragraph(desc)
    
    for sec_title, table_headers, table_data in sections:
        h = doc.add_heading(sec_title, 2)
        for run in h.runs: run.font.name = u'微软雅黑'; run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
        
        table = doc.add_table(rows=1, cols=len(table_headers), style='Table Grid')
        hdr_cells = table.rows[0].cells
        for i, th in enumerate(table_headers):
            hdr_cells[i].text = th
            
        for row_data in table_data:
            row_cells = table.add_row().cells
            for i, td in enumerate(row_data):
                row_cells[i].text = str(td)
                
        # 表格字体
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = u'微软雅黑'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
                        run.font.size = Pt(10)
    return doc

# 1. 硬件架构 V2.0
hw_sections = [
    ("1. 物理网络拓扑 (10网口)", ["接口", "用途", "强约束"], [
        ("主板 LAN 1", "EtherCAT主站", "1ms高频控制4个伺服电机，CPU内核隔离"),
        ("主板 LAN 2", "直连SIM卡路由器", "唯一的广域网出口，防电磁干扰"),
        ("PCIe卡 A", "4个工业相机", "巨型帧独立子网，物理隔离"),
        ("PCIe卡 B", "2相机 + 2线激光", "巨型帧独立子网，微秒级硬同步触发")
    ]),
    ("2. 定制多通道电源管理板 (PMB)", ["通道", "输出电压", "供电设备及保护机制"], [
        ("通道 A", "48V直通", "伺服电机 (承受反电动势，防浪涌)"),
        ("通道 B", "24V/12V稳压", "工控机主板 (宽压隔离，自恢复保险丝)"),
        ("通道 C", "24V独立隔离", "工业相机、线激光 (防电磁干扰)"),
        ("通道 D", "12V/5V", "路由器、传感器 (支持工控机指令级硬重启)")
    ]),
    ("3. 光学与机械抗震 (10km/h刚性底盘)", ["模块", "方案", "解决痛点"], [
        ("光学防震", "定焦镜头螺纹胶死锁 + 阻尼基座", "防刚性冲击导致离焦"),
        ("全天候光源", "物理遮光罩 + 柔性毛刷裙边 + LED频闪", "废除软件调光，200μs定格图像"),
        ("散热系统", "罩内防尘强制风道", "防高发热导致CMOS热噪声")
    ])
]
doc1 = create_doc("机器人硬件系统架构设计文档 V2.0", "核心更新：全面落实PMB电源管理板、10网口强制分配、刚性底盘光学防震与物理遮光恒定光源。", hw_sections)
doc1.save(os.path.join(out_dir, "硬件系统架构设计文档_V2.0_纯工程版.docx"))

# 2. 软件架构 V2.0
sw_sections = [
    ("1. 采集引擎底层机制", ["核心机制", "实现方式", "工程收益"], [
        ("CPU资源独占", "Windows Thread Affinity 绑定", "防止网卡中断风暴导致EtherCAT急停"),
        ("零拷贝内存池", "开辟非托管连续内存直写", "杜绝C#垃圾回收(GC)引发的严重掉帧"),
        ("时空对齐", "以里程编码器Frame ID为绝对主键", "解决多网口异步到达的数据重组难题")
    ]),
    ("2. 核心解耦算法 (各司其职)", ["几何参数", "数据源", "解耦算法逻辑"], [
        ("轨距", "8个测距传感器", "左右对向测距求和，不受车体姿态干扰"),
        ("水平(超高)", "陀螺仪 Roll 角", "独立读取横滚角 * 基准轨距"),
        ("高低与轨向", "陀螺仪 Pitch/Yaw + 编码器", "结合里程积分计算纵向与侧向平顺度"),
        ("表面病害", "6相机 + 边缘AI", "YOLO裁剪病害ROI，大图废弃防拥堵")
    ]),
    ("3. 边缘-云端网络流控", ["队列级别", "数据类型", "应对策略"], [
        ("P0-最高优", "急停与系统告警", "强行预留10%带宽，UDP/TCP双发"),
        ("P1-高优", "几何、姿态、电量时序数据", "秒级连续发送，保持监控不断联"),
        ("P3-低优大文件", "3D点云、原始大图", "遇弱网触发写锁落盘工业SSD，恢复后断点续传")
    ])
]
doc2 = create_doc("机器人软件系统架构设计文档 V2.0", "核心更新：摒弃伪科学耦合，确立各司其职的测量逻辑，强化CPU隔离与零拷贝机制。", sw_sections)
doc2.save(os.path.join(out_dir, "软件系统架构设计文档_V2.0_纯工程版.docx"))

# 3. 技术方案综合 V2.0
sol_sections = [
    ("1. 系统定位与核心策略", ["策略维度", "核心设计", "工程依据"], [
        ("里程定位策略", "纯伺服编码器防滑测距 + 人工标定", "摆脱隧道与边坡对GPS的遮挡，实现绝对无盲区"),
        ("光学检测策略", "恒定暗室 + 微秒频闪", "降维打击：用机械遮挡解决复杂的机器视觉光照干扰"),
        ("数据传输策略", "边缘截流减负 + 探针QoS调度", "适应铁路沿线4G/5G极其恶劣的信号波动")
    ])
]
doc3 = create_doc("铁路线路智能检测机器人技术方案 V2.0", "核心更新：对齐最终版物理硬件参数(120mm轮、10km/h极速)，确立项目最高工程宪法。", sol_sections)
doc3.save(os.path.join(out_dir, "铁路线路智能检测机器人技术方案_V2.0_纯工程版.docx"))

print("V2.0 Docs Successfully Generated!")
