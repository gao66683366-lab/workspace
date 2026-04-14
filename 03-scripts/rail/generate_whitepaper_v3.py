import docx
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

output_dir = r"D:\铁路线路智能检测机器人\04-项目文档\设计文档"
os.makedirs(output_dir, exist_ok=True)
doc_path = os.path.join(output_dir, "铁路线路智能检测机器人_四步流转硬核技术白皮书_工业仪器级_V3.0.docx")

doc = docx.Document()
doc.styles['Normal'].font.name = u'微软雅黑'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
doc.styles['Normal'].font.size = Pt(10.5)

def add_h(text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = u'微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    return h

add_h('铁路线路智能检测机器人', 0)
add_h('四步流转硬核技术白皮书 (工业仪器级 V3.0)', 1)
doc.add_paragraph('本文档基于10网口严格物理隔离、10km/h极速、120mm轮径刚性底盘、48V PMB隔离电源的客观约束，对工控机(C#)与云端(Python)的数据流转进行芯片级与代码级的深度约束设计。')

add_h('第一步：核心数据采集 (底层硬件级时空定格)', 2)
doc.add_paragraph('【执行目标】突破操作系统纳秒级枷锁，利用独立网卡实现高速协议栈替换与物理截流。')
table1 = doc.add_table(rows=1, cols=3, style='Table Grid')
hdr = table1.rows[0].cells
hdr[0].text, hdr[1].text, hdr[2].text = '网络端口', '协议与驱动重构(C#底层)', '数据接收与内存级操作'
data1 = [
    ('主板 LAN 1\n(运动控制)', '卸载Windows IPv4协议栈，替换为TwinCAT EtherCAT实时网卡驱动(RT-NDIS)。', '1ms周期收发PDO报文；启用DC分布式时钟(Sync0)截获驱动器硬件纳秒级时间戳，生成全局Frame ID。'),
    ('PCIe扩展卡 A\n(4个 GigE 2D相机)', '配置巨型帧(MTU 9014)，调用原厂GigE Vision Filter Driver绕过Windows内核。', 'C#使用unsafe块开辟非托管连续内存池；开启GVSP丢包监控，Block ID不连续时触发PacketResend。'),
    ('PCIe扩展卡 B\n(2个相机+2个3D激光)', '同上，3D线激光跑GigE或专用协议，同样依赖底层Filter Driver抓包。', '为3D线激光开辟专门的16-bit深度图(Depth Map)内存池，接收完整轮廓数据。'),
    ('数字I/O触发板', '接收编码器里程脉冲，通过高速光耦(如6N137)输出微秒级差分方波。', '固定相机电子快门200μs内，频闪LED同步过载爆闪，定格10km/h运动画面，彻底消除拖影。')
]
for item in data1:
    row = table1.add_row().cells
    row[0].text, row[1].text, row[2].text = item

add_h('第二步：数据清洗与筛查 (信号层去伪存真)', 2)
doc.add_paragraph('【执行目标】化解刚性震荡，彻底剥离物理噪声与图像无用背景。')
table2 = doc.add_table(rows=1, cols=3, style='Table Grid')
hdr2 = table2.rows[0].cells
hdr2[0].text, hdr2[1].text, hdr2[2].text = '清洗模块', '诱发原因与边界条件', 'C#底层算法与执行方案'
data2 = [
    ('里程脉冲清洗', '120mm小轮遇不平轨道导致某轮悬空打滑。', '监控4轴ActualTorque(实际扭矩)，瞬降即判定打滑。剔除该轮脉冲，取其余接触轮均值。'),
    ('姿态高频清洗', '无减震底盘导致陀螺仪掺杂轮轨硬冲击噪声。', '前置10Hz二阶巴特沃斯低通滤波器。注：需在C#队列中执行“相位延迟反向补偿”，对齐视觉Frame ID。'),
    ('图像空间清洗', '4K全分辨率图像仅有10%为有效病害区域。', '利用C# SIMD(AVX2)向量指令集，在非托管内存中直接对固定坐标(焦距已胶水死锁)进行极速ROI像素裁剪。'),
    ('时序残缺清洗', '串口(低速)与GigE(高速)的异步通讯错位。', '非对称时间窗：以Frame ID为锚点，低速数据在±5ms窗内就近吸附或插值；超过500ms未集齐帧直接冷血丢弃。')
]
for item in data2:
    row = table2.add_row().cells
    row[0].text, row[1].text, row[2].text = item

add_h('第三步：核心算法处理 (独立通道的降维解算)', 2)
doc.add_paragraph('【执行目标】各司其职，坚决摒弃传感器伪补偿，进行纯粹的工程定性。')
table3 = doc.add_table(rows=1, cols=3, style='Table Grid')
hdr3 = table3.rows[0].cells
hdr3[0].text, hdr3[1].text, hdr3[2].text = '算法通道', '解算模型与公式', '关键抗噪策略'
data3 = [
    ('几何四大通道', '轨距=D-(L+R); 水平=1435*sin(Roll); 高低/轨向=Pitch/Yaw空间积分。', '四大通道绝对独立解耦，互不干涉。防止级联失效。'),
    ('3D轮廓磨耗', '点云与60kg/m标准CAD模型进行ICP迭代配准。', '升级为Point-to-Plane ICP并引入Huber损失函数，强力压制钢轨表面铁锈/剥落造成的配准发散。'),
    ('AI病害定性', '边缘端基于TensorRT的INT8量化推理。', '使用Pinned Memory(页锁定内存)将ROI推入GPU异步流水线，不阻塞CPU几何解算。')
]
for item in data3:
    row = table3.add_row().cells
    row[0].text, row[1].text, row[2].text = item

add_h('第四步：轻量化数据上传 (协议级的弱网流控机制)', 2)
doc.add_paragraph('【执行目标】以几百字节的极限网络开销，穿透恶劣铁路隧道环境，实现秒级告警。')
table4 = doc.add_table(rows=1, cols=3, style='Table Grid')
hdr4 = table4.rows[0].cells
hdr4[0].text, hdr4[1].text, hdr4[2].text = '流控机制', '技术特征', '工业级落地参数'
data4 = [
    ('裸协议组装', '废弃庞大GC开销的JSON，改用C语言风格内存对齐结构体。', '[StructLayout(Pack=1)]，将病害坐标、分类、置信度压缩为约25 Bytes的二进制包，Span<byte>直发。'),
    ('主动探针调度', 'Token Bucket令牌桶流控，根据探针RTT动态调整。', '主板LAN2连接SIM路由器。P0告警(25B)>P1状态>P2图切片>P3原始大文件。'),
    ('极限落盘防爆', '隧道断网时触发C#写锁，数据转存本地工业级NVMe SSD。', '开启SQLite WAL模式与synchronous=OFF，防I/O阻塞引发内存池爆满；出隧道后凭游标断点续传。')
]
for item in data4:
    row = table4.add_row().cells
    row[0].text, row[1].text, row[2].text = item

# 表格字体统一
for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = u'微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
                    run.font.size = Pt(10)

doc.save(doc_path)
print(f"SUCCESS: {doc_path}")
