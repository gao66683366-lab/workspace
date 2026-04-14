import docx
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

output_dir = r"D:\铁路线路智能检测机器人\04-项目文档\设计文档"
os.makedirs(output_dir, exist_ok=True)
doc = docx.Document()
doc.styles['Normal'].font.name = u'微软雅黑'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
doc.styles['Normal'].font.size = Pt(10.5)

def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = u'微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    return h

add_heading(doc, '铁路线路智能检测机器人', 0)
add_heading(doc, '系统全景架构设计文档 V3.0 (终极版)', 1)

doc.add_paragraph('本文档基于最新物理硬件约束（10网口严格分配、48V蓄电池直流供电、工业路由器SIM卡回传）及核心物理设计（遮光罩恒定光源、编码器高精测距）进行最终定调版输出。')

add_heading(doc, '1. 硬件与网络物理拓扑架构 (10网口严格分配)', 2)
table1 = doc.add_table(rows=1, cols=4, style='Table Grid')
hdr_cells = table1.rows[0].cells
hdr_cells[0].text = '网络接口'
hdr_cells[1].text = '接口类型'
hdr_cells[2].text = '连接设备'
hdr_cells[3].text = '核心约束与说明'

data1 = [
    ('主板 LAN 1', '原生千兆网口', '4个伺服电机(驱动器)', '独占EtherCAT总线，保证1ms高实时运动控制，读取绝对值编码器脉冲。'),
    ('主板 LAN 2', '原生千兆网口', '工业路由器(带SIM卡)', '唯一对外WAN口，基于主动探针实现QoS均衡上传数据到云端。'),
    ('PCIe 扩展卡 A', '4口千兆 PoE', '4个工业相机', '巨型帧(9014 Bytes)独立子网，基于里程硬同步触发。'),
    ('PCIe 扩展卡 B', '4口千兆 PoE', '2个工业相机 + 2个3D线激光', '2个3D线激光占2个网口，与视觉同享微秒级硬同步触发。'),
    ('多串口/I/O卡', 'RS485 / I/O', '8个测距传感器+2个陀螺仪', '走串口或总线，绝不占用以太网口资源。')
]
for item in data1:
    row = table1.add_row().cells
    row[0].text = item[0]
    row[1].text = item[1]
    row[2].text = item[2]
    row[3].text = item[3]

add_heading(doc, '2. 供电、机械与物理光学系统', 2)
table2 = doc.add_table(rows=1, cols=3, style='Table Grid')
hdr2 = table2.rows[0].cells
hdr2[0].text = '系统模块'
hdr2[1].text = '核心组件'
hdr2[2].text = '设计要求'

data2 = [
    ('供电系统', '48V 蓄电池直流主电源', '提供系统总动力，配置DC-DC转换，支持BMS云端电量监测与续航预估。'),
    ('机械结构', '底盘与传感器支架', '满足轨道运行防滑设计，高刚性保障视觉与激光的空间标定不产生物理形变。'),
    ('光学系统', '全天候遮光罩+频闪光源', '物理隔绝外界环境光，相机参数写死(取消软件调光算法)，搭配硬件触发实现频闪定格高速运动。')
]
for item in data2:
    row = table2.add_row().cells
    row[0].text = item[0]
    row[1].text = item[1]
    row[2].text = item[2]

add_heading(doc, '3. 七大核心工作版图分解', 2)
table3 = doc.add_table(rows=1, cols=3, style='Table Grid')
hdr3 = table3.rows[0].cells
hdr3[0].text = '版图分类'
hdr3[1].text = '负责模块'
hdr3[2].text = '关键技术点'

data3 = [
    ('1. 硬件控制与感知', '多模态硬同步、EtherCAT闭环', '里程等距微秒级硬触发(结合遮光罩/频闪)，防滑转补偿。'),
    ('2. 软件平台与架构', 'C#高并发采集、Python云端微服务', '零拷贝内存池环形队列，看门狗物理硬重启外置工业路由器。'),
    ('3. AI与核心算法', '边缘AI截流裁剪、云端大模型复核', '3D点云与1D时序融合计算高低差，YOLO轻量化截流，多模态特征融合。8个测距传感器三角函数解算。'),
    ('4. 均衡通信架构', 'QoS四级队列、主动链路探针', '弱网探针探测，大文件写锁本地落盘(SQLite)与断点续传，P0-P3四级流控。'),
    ('5. 高精定位与资产', '编码器里程计、多轮融合防滑', '摆脱GPS，纯靠EtherCAT读取电机脉冲+起始点锚定进行高精病害定位。'),
    ('6. 技术文档管理', '规范化表格输出', '所有文档纯Word格式、全表格化排版，API与测试报告自动化生成。'),
    ('7. 学术论文转化', 'SCI/EI论文编撰、专利申请', '聚焦多模态硬同步、边缘QoS均衡调度、遮光物理降维打击算法的学术产出。')
]
for item in data3:
    row = table3.add_row().cells
    row[0].text = item[0]
    row[1].text = item[1]
    row[2].text = item[2]

for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = u'微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
                    run.font.size = Pt(10)

doc.save(r'D:\铁路线路智能检测机器人\04-项目文档\设计文档\铁路线路智能检测机器人系统全景架构设计文档_V3.0_终极版.docx')
print("V3.0 终极版架构文档生成成功")
