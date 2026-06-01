#!/usr/bin/env python3
"""生成专利申报五书合集版Word文档（含5张附图）"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.27)
section.page_height = Inches(11.69)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)

def set_font(run, name='宋体', size=12, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)

def add_center_heading(doc, text, size=16, bold=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, '黑体', size, bold)
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(10)

def add_heading(doc, text, size=13, bold=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font(run, '黑体', size, bold)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)

def add_body(doc, text, size=12):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_font(run, '宋体', size, False)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(22)

def add_figure(doc, fig_num, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'图{fig_num}：{caption}')
    set_font(run, '宋体', 10, False)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)

def insert_img(doc, path, width=Inches(5.5)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=width)

# ===================== 第一书：说明书摘要 =====================
add_center_heading(doc, '说　明　书　摘　要', 22)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('摘　　要')
set_font(run, '黑体', 14, True)
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after = Pt(10)

abstract = """本发明公开了一种铁路线路智能视觉与多模态感知融合综合检测方法及系统。系统由检测小车车体、6路2D工业相机、2路3D线激光传感器、HWT905姿态传感器、单点测距传感器、测距传感器矩阵、工控机及融合判定单元构成，单台设备单次出行覆盖轨面缺陷、道钉/螺栓状态、焊缝质量、钢轨廓形、波磨、轨距、水平、高低共8项检测功能。

本发明提出三键索引时空对齐机制，以帧编号—里程坐标—UTC时间戳为统一基准，将多传感器空间对齐精度从厘米级提升至亚毫米级（平均偏差3.2mm，降低74.7%）。提出基于检测可靠性因子的自适应动态加权几何平均融合算法，融合判定准确率达97.5%，传感器退化场景下仍保持96.1%。设计快慢双速EKF融合架构（频率比20:1），实现轨道高程与车体振动的频域解耦，高低不平顺测量精度优于0.5mm。轨距检测精度优于0.3mm，水平检测误差小于0.4mm，轨面缺陷mAP@0.5达92.5%。在济南铁路局管内实际线路120km以上验证，系统运行稳定，方法有效。"""
for para in abstract.strip().split('\n\n'):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(para.strip())
    set_font(run, '宋体', 12, False)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(22)

doc.add_page_break()
add_center_heading(doc, '摘　　要　　附　　图', 18)
insert_img(doc, '/root/.openclaw/media/tool-image-generation/fig1_PIL.png', Inches(5.5))
add_figure(doc, 1, '系统总体架构图')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run('图中标记说明：1—检测小车车体；2—2D工业相机；3—3D线激光传感器；4—HWT905姿态传感器；5—单点测距传感器；6—测距传感器矩阵；7—工控机；8—通信模块；9—编码器；10—供电模块；11—环形缓冲区；12—融合判定单元。')
set_font(run, '宋体', 9, False)

# ===================== 第二书：权利要求书 =====================
doc.add_page_break()
add_center_heading(doc, '权　利　要　求　书', 22)
add_heading(doc, '一、独立权利要求', 13, True)
add_heading(doc, '1. 装置权利要求', 12, True)
add_body(doc, '1. 一种铁路线路智能视觉与多模态感知融合综合检测系统，其特征在于，包括：')
items1 = [
    '检测小车车体（1），安装于铁路轨道上运行；',
    '视觉感知子系统，包括安装于所述车体的6路2D工业相机（2），分辨率2448×2048px、帧率20fps、千兆以太网接口，分工位覆盖轨面缺陷检测、道钉/螺栓状态检测和焊缝缺陷检测；',
    '廓形感知子系统，包括2路3D线激光传感器（3），扫描频率20000Hz、每轮廓3200点、IP67防护等级，用于钢轨廓形和波磨检测；',
    '姿态感知子系统，包括HWT905姿态传感器（4），采样率200Hz、三轴姿态角分辨率0.05°、支持IEEE 1588 PTP时间同步；',
    '单点测距传感器（5），共2个，采样率2000Hz，精度±0.15mm，用于高低检测；',
    '测距传感器矩阵（6），用于轨距检测，与车体横滚角完全解耦；',
    '编码器（9），用于里程测量和帧触发；',
    '工控机（7）及通信模块（8）；',
    '融合判定单元（12），实现多源数据的三级融合判定；',
    '其中，所述系统采用三键索引时空对齐机制，以帧编号FID、里程坐标s_k、UTC时间戳t_k为三键索引元组，为全部传感器数据建立统一时空基准；融合判定单元（12）采用基于检测可靠性因子的自适应动态加权几何平均融合算法，对多源数据进行三级融合判定；高低检测采用快慢双速EKF融合架构，以20:1频率比双回路实现轨道高程与车体振动的频域解耦。'
]
for item in items1:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(item)
    set_font(run, '宋体', 12, False)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = Pt(20)

add_heading(doc, '2. 方法权利要求', 12, True)
add_body(doc, '2. 一种铁路线路智能视觉与多模态感知融合综合检测方法，基于权利要求1所述的系统，其特征在于，包括以下步骤：')
steps = [
    '（1）三键索引时空对齐：编码器每积累N个脉冲（N=100）定义为一帧，生成帧编号FID_k；同时从PTP时钟获取UTC时间戳t_k，并将编码器脉冲累积值经标准轨段标定后映射为里程坐标s_k；以三键索引元组D(FID_k, s_k, t_k)作为各传感器数据的融合主键，在感知层完成帧边界对齐后写入环形缓冲区；',
    '（2）三级融合处理：第一级时间对齐，将8种异构传感器数据按帧编号精确对齐；第二级空间关联，经姿态角修正后统一映射至轨道坐标系；第三级判级融合，对多维数据进行缺陷类型和等级的联合判定；',
    '（3）自适应动态加权融合：计算各维度检测器的可靠性因子r_d(t)=C_d(t)/σ_d(t)，其中C_d(t)为维度置信度，σ_d(t)为置信度滑动标准差；归一化后得融合权重w_d(t)；采用加权几何平均计算融合置信度C_fuse=∏C_d(t)^w_d(t)；',
    '（4）快慢双速EKF高低检测：快回路以200Hz频率利用姿态传感器实时估计车体高频振动；慢回路以10Hz频率基于单点测距传感器累积数据进行EKF状态估计，输出轨道高程真值；实现高低不平顺测量精度优于0.5mm；',
    '（5）轨距及水平检测：测距传感器矩阵直接测量左右轨内侧距，轨距G=d_left+d_right与横滚角完全解耦；HWT905横滚角θ_r计算超高差Δh=1435×sin(θ_r)；',
    '（6）视觉检测与跨维度联合判定：采用YOLOv8m算法检测轨面缺陷，ZNCC模板匹配检测道钉/螺栓状态，OpenCV五步流程检测焊缝缺陷；所有检测结果通过三键索引在同一帧内精确关联，进行跨维度联合判定。'
]
for step in steps:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(step)
    set_font(run, '宋体', 12, False)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(20)

add_heading(doc, '二、从属权利要求', 13, True)
sub_items = [
    ('3. 关于时空对齐', '3. 根据权利要求1所述的系统，其特征在于，所述三键索引时空对齐机制中，UTC时间戳精度优于1μs，由IEEE 1588 PTP协议实现跨传感器时间基准统一；里程坐标通过标准轨段标定消除轮径磨损和打滑引入的里程累积误差；帧编号对应约5mm的里程分辨率。'),
    ('4. 关于融合算法', '4. 根据权利要求1或2所述的系统或方法，其特征在于，所述自适应动态加权几何平均融合算法中，融合置信度低于0.5的检测结果直接丢弃，介于0.5至0.7的进入人工复核队列，高于0.7的进入最终判定输出；几何平均的乘积特性使任一维度异常显著降低融合置信度，防止误判在多维度间的传播。'),
    ('5. 关于EKF架构', '5. 根据权利要求1或2所述的系统或方法，其特征在于，所述快慢双速EKF融合架构中，快慢回路中心频率比值为20:1，在频域上完全解耦；慢回路状态向量为x_k=[h_track, h_veh, v_veh]^T，其中h_track为轨道高程待估真值，h_veh为车体相对轨道高差，v_veh为车体垂向速度。'),
    ('6. 关于传感器配置', '6. 根据权利要求1所述的系统，其特征在于，所述6路2D工业相机中，2路用于轨面缺陷检测、2路用于道钉/螺栓状态检测、2路用于焊缝缺陷检测；轨面缺陷检测相机离轨面约180mm，道钉/螺栓检测相机离轨面约270mm；3D线激光传感器垂直朝下安装，用于扫描左右钢轨轨头轮廓。'),
    ('7. 关于容错降级', '7. 根据权利要求1所述的系统，其特征在于，所述系统还包括传感器可靠性评估模块，从数据质量、时序一致性、环境适应性三个维度计算传感器综合可靠性指数ρ；当ρ低于0.5时，系统进入降级运行模式，仅使用正常工作的传感器通道继续检测并标注降级标识；当传感器可靠性指数快速回升时，自动发起传感器重检流程。')
]
for title, content in sub_items:
    add_heading(doc, title, 12, True)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(content)
    set_font(run, '宋体', 12, False)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(20)

# ===================== 第三书：说明书 =====================
doc.add_page_break()
add_center_heading(doc, '说　　明　　书', 22)

sections_3 = [
    ('一、技术领域', '本发明涉及铁路线路检测技术领域，具体涉及一种基于智能视觉与多模态感知融合的铁路线路综合检测系统，以及基于该系统的时空对齐、多源数据融合、轨道几何参数测量方法。'),
    ('二、背景技术', '''铁路线路承受列车循环载荷，轨面裂纹、螺栓松动、钢轨磨耗等病害持续累积，及时准确发现线路病害是工务维护的核心任务。根据检测对象不同，铁路线路检测分为视觉类检测（轨面缺陷、道钉螺栓状态、焊缝质量、钢轨廓形、波磨）和几何参数类检测（轨距、水平、高低）两大类，两类检测在信号采集方式、数据特征和处理路径上存在本质差异，长期分立运行。

现有铁路线路综合检测面临以下技术瓶颈：

（1）多源数据时空基准不统一。视觉检测与几何参数检测设备各自独立时钟和里程基准，两类数据难以精确关联。

（2）融合机制停留于拼接层，自适应能力不足。现有方案多采用"各传感器独立检测后结果拼接"的两阶段方案，融合权重多基于工程经验固定配置，无法适应传感器性能退化。

（3）轨道几何参数检测受车体姿态扰动影响显著。传统基于轮距机械基准配合横滚角反推的方法，在曲线段车体横滚角较大时，微小测量误差通过三角函数关系放大为轨距显著偏差。

（4）钢轨廓形缺乏在线三维测量手段。传统接触式轮廓仪效率低且无法在线测量；3D线激光传感器虽然精度高，但与其他传感器数据缺乏统一时空基准，无法实现特征级融合。

（5）钢轨表面缺陷视觉检测精度有限。深度学习方法在受控环境下mAP@0.5可达89%以上，但与真实铁路环境（光照变化大、雨雪覆盖、振动模糊等）存在较大差距。

（6）道钉及螺栓完好状态检测效率低。传统人工敲击检查效率极低，基于深度学习目标检测在小目标、高密度场景下检测困难。'''),
]

for title, content in sections_3:
    add_heading(doc, title, 13, True)
    for para in content.strip().split('\n\n'):
        if para.strip():
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0.74)
            run = p.add_run(para.strip())
            set_font(run, '宋体', 12, False)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = Pt(22)

# 三、发明内容
add_heading(doc, '三、发明内容', 13, True)
add_heading(doc, '3.1 发明要解决的技术问题', 12, True)
add_body(doc, '首要问题：多源异构传感器数据时空基准不统一，导致视觉检测与几何参数检测无法精确融合，跨维度联合分析缺乏可信数据基础。')
add_body(doc, '次要问题：融合权重缺乏自适应机制，传感器性能退化时鲁棒性不足；轨道几何参数检测受姿态扰动影响精度；钢轨廓形、轨面缺陷、道钉螺栓缺乏统一的在线综合检测手段；现有综合检测设备体积大、成本高、维护复杂，难以在中小规模工务段推广。')

add_heading(doc, '3.2 技术方案', 12, True)
subsections = [
    ('3.2.1 三键索引时空对齐机制', '本发明提出"帧编号—里程坐标—UTC时间戳"三键索引机制。帧编号（FID）：编码器每产生N个脉冲（N=100）定义为一帧，对应约5mm里程分辨率。里程坐标（s_k）：编码器脉冲累积值经标准轨段标定后，建立帧编号与里程坐标的精确映射。UTC时间戳（t_k）：每帧数据附带IEEE 1588 PTP同步的UTC时间戳，精度优于1μs。三键索引元组D(FID_k, s_k, t_k)实现多传感器数据的硬件级精确对齐。'),
    ('3.2.2 自适应动态加权几何平均融合算法', '本发明提出基于检测可靠性因子的自适应权重计算方法。定义各维度检测器在时刻t的可靠性因子r_d(t) = C_d(t) / σ_d(t)，其中C_d为维度置信度，σ_d为该维度置信度的滑动标准差。各维度权重由可靠性归一化得到：w_d(t) = r_d(t) / Σr_d(t)。融合置信度采用加权几何平均计算：C_fuse = ∏C_d(t)^w_d(t)。几何平均的乘积特性使任一维度异常显著降低融合置信度，防止误判传播。'),
    ('3.2.3 快慢双速EKF融合架构', '本发明设计双速EKF融合架构，分别处理轨道高程（<1Hz）与车体振动（0~20Hz）两个不同频带的物理过程。快回路（200Hz）：利用HWT905姿态传感器实时估计并补偿车体的高频振动。慢回路（10Hz）：基于单点测距传感器累积数据，估计轨道高程真值。快慢回路中心频率比值20:1，在频域上完全解耦，无混叠效应。高低不平顺测量精度优于±0.5mm，振动噪声降低53.7%。'),
    ('3.2.4 测距传感器矩阵直接测量轨距方案', '轨距检测采用测距传感器矩阵直接测量方案：G = d_left + d_right直接即为轨距实测值，不依赖于横滚角姿态数据，与车体横滚角完全解耦，轨距检测精度优于±0.3mm。'),
    ('3.2.5 姿态传感器水平检测方案', '水平（超高差）检测基于HWT905姿态传感器的横滚角输出，计算超高差：Δh = 1435 × sin(θ_r)，误差小于0.4mm。'),
    ('3.2.6 基于YOLOv8m的轨面缺陷视觉检测方案', '本发明采用YOLOv8m目标检测算法实现钢轨表面缺陷的实时识别，mAP@0.5达92.5%。6路2D工业相机分工位配置，所有相机数据通过三键索引机制与3D线激光、姿态传感器、测距传感器数据在同一帧内精确关联，实现特征级融合。'),
    ('3.2.7 基于ZNCC模板匹配的道钉/螺栓完好状态检测方案', '本发明采用ZNCC（归一化零均值互相关）模板匹配方法实现道钉/螺栓的完好状态检测。缺失检出率97.2%，松动检出率94.8%。')
]

for title, content in subsections:
    add_heading(doc, title, 12, True)
    for para in content.strip().split('\n\n'):
        if para.strip():
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0.74)
            run = p.add_run(para.strip())
            set_font(run, '宋体', 12, False)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = Pt(22)

add_heading(doc, '3.3 有益效果', 12, True)
benefits = [
    '（1）多传感器空间对齐精度从厘米级提升至亚毫米级（平均偏差3.2mm，降低74.7%）。',
    '（2）3D线激光廓形测量精度优于±0.15mm，20000Hz扫描频率，3200点/轮廓。',
    '（3）轨面缺陷检测mAP@0.5达92.5%，道钉/螺栓缺失检出率97.2%、松动检出率94.8%。',
    '（4）自适应融合判定准确率97.5%，传感器退化场景下仍保持96.1%。',
    '（5）高低不平顺测量精度优于±0.5mm，振动噪声降低53.7%；轨距检测精度优于±0.3mm，水平检测误差小于0.4mm。',
    '（6）单台设备单次出行覆盖全部8项检测功能，适合中小规模工务段推广。'
]
for item in benefits:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(item)
    set_font(run, '宋体', 12, False)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(20)

add_heading(doc, '四、附图说明', 13, True)
add_body(doc, '图1：铁路线路综合检测系统总体架构图')
add_body(doc, '图2：三键索引时空对齐机制示意图')
add_body(doc, '图3：自适应动态加权融合算法流程图')
add_body(doc, '图4：快慢双速EKF融合架构图')
add_body(doc, '图5：几何参数检测原理图（轨距/水平/高低）')
add_body(doc, '图中标记说明：1—检测小车车体；2—2D工业相机；3—3D线激光传感器；4—HWT905姿态传感器；5—单点测距传感器；6—测距传感器矩阵；7—工控机；8—通信模块；9—编码器；10—供电模块；11—环形缓冲区；12—融合判定单元。')

add_heading(doc, '五、具体实施方式', 13, True)
impl_sections = [
    ('5.1 系统总体架构', '检测小车车体（1）安装于铁路轨道上运行。6路2D工业相机（2），分辨率2448×2048px，帧率20fps，用于采集轨面缺陷、道钉/螺栓、焊缝图像；2路3D线激光传感器（3），扫描频率20000Hz，每轮廓3200点，IP67防护，用于采集钢轨廓形和波磨数据；HWT905姿态传感器（4），采样率200Hz，PTP时间同步；2个单点测距传感器（5），采样率2000Hz，精度±0.15mm，用于高低检测；测距传感器矩阵（6），采样率约1000Hz，精度±0.1mm，用于轨距检测；工控机（7），防护等级IP54；通信模块（8），支持4G/5G无线传输；三网物理隔离架构（EtherCAT控制网+千兆采集网+无线传输网）。'),
    ('5.2 三键索引时空对齐与钢轨廓形测量实现', '编码器脉冲触发帧分割逻辑，每积累100个脉冲生成一帧FID_k，同时从PTP时钟获取UTC时间戳并更新里程坐标。各传感器数据在感知层被赋予帧编号后写入环形缓冲区（11）。三键索引元组作为融合主键，8种传感器的原始数据在硬件层面完成帧边界对齐。2路3D线激光传感器（3）安装于检测小车底部，垂直朝下对左右钢轨轨头进行扫描，20000Hz扫描频率，3200点/轮廓，完全满足廓形测量精度要求。'),
    ('5.3 自适应融合判定实现', '融合判定单元（12）对来自不同传感器的数据进行三级融合。第一级（时间对齐）：8种异构传感器原始数据按帧编号精确对齐。第二级（空间关联）：对帧对齐数据进行空间坐标变换，将各传感器观测数据统一映射到轨道坐标系下。第三级（判级融合）：在统一空间坐标下，对多维度数据进行缺陷类型和等级的联合判定。可靠性因子r_d(t)由各维度置信度及其滑动标准差计算，权重w_d(t)归一化后用于几何平均融合，置信度分档输出结果。'),
    ('5.4 双速EKF融合与几何参数检测实现', '快回路以200Hz频率运行，利用HWT905姿态传感器数据实时估计车体振动。慢回路以10Hz频率运行，基于单点测距传感器累积数据进行EKF状态估计。慢回路输出轨道高程估计值h_track。双速解耦使轨道高程变化（<1Hz）与车体振动（0~20Hz）在频域上完全分离，高低不平顺测量精度优于±0.5mm。测距传感器矩阵测量值之和G = d_left + d_right即为轨距实测值，与横滚角完全解耦，精度优于±0.3mm。HWT905姿态传感器实时输出横滚角θ_r，超高差计算Δh = 1435 × sin(θ_r)，误差小于0.4mm。'),
    ('5.5 系统容错与降级策略', '软故障容错：当传感器可靠性指数ρ低于阈值（0.5至0.8之间）时，融合算法自动降低该通道权重，对应区段检测结果标记为"待复核"状态。硬故障处理：当传感器完全失效时，切换至降级运行模式，仅使用正常传感器通道继续检测并标注降级标识，同时记录故障日志通知维护人员。系统持续监测ρ变化趋势，当ρ快速回升时自动发起重检流程。')
]
for title, content in impl_sections:
    add_heading(doc, title, 12, True)
    for para in content.strip().split('\n\n'):
        if para.strip():
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0.74)
            run = p.add_run(para.strip())
            set_font(run, '宋体', 12, False)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = Pt(22)

add_body(doc, '上面所述的实施例仅为本发明的优选实施方式进行描述，并非对本发明的范围进行限定，在不脱离本发明设计精神的前提下，本领域相关技术人员对本发明的各种变形和改进，均应落入本发明权利要求书所确定的保护范围内。')

# ===================== 第四书：说明书附图 =====================
doc.add_page_break()
add_center_heading(doc, '说　明　书　附　图', 22)

fig_list = [
    ('1', '系统总体架构图', '/root/.openclaw/media/tool-image-generation/fig1_PIL.png'),
    ('2', '三键索引时空对齐机制示意图', '/root/.openclaw/media/tool-image-generation/fig2_PIL.png'),
    ('3', '自适应动态加权融合算法流程图', '/root/.openclaw/media/tool-image-generation/fig3_PIL.png'),
    ('4', '快慢双速EKF融合架构图', '/root/.openclaw/media/tool-image-generation/fig4_PIL.png'),
    ('5', '几何参数检测原理图（轨距/水平/高低）', '/root/.openclaw/media/tool-image-generation/fig5_PIL.png'),
]

markers = [
    '图中标记：1—检测小车车体；2—2D工业相机；3—3D线激光传感器；4—HWT905姿态传感器；5—单点测距传感器；6—测距传感器矩阵；7—工控机；8—通信模块；9—编码器；10—供电模块；11—环形缓冲区；12—融合判定单元。',
    '',
    '',
    '',
    ''
]

for i, (fig_num, caption, img_path) in enumerate(fig_list):
    if i > 0:
        doc.add_page_break()
    insert_img(doc, img_path, Inches(5.5))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'图{fig_num}：{caption}')
    set_font(run, '宋体', 10, False)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    if markers[i]:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run2 = p2.add_run(markers[i])
        set_font(run2, '宋体', 9, False)
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(6)

out_path = '/root/.openclaw/workspace/铁路线路智能综合检测机器人/08-专利资料/专利申报五书_合集版_20260531.docx'
doc.save(out_path)
print(f'已保存：{out_path}')